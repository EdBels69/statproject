"""
Word Report Generator for Statistical Analysis Results.
Generates .docx files with properly formatted Table 1 and hypothesis test results.
"""

import io
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
import pandas as pd


# --- Interpretation Text Generator ---

METHOD_NAMES = {
    "t_test_ind": "независимый t-критерий Стьюдента",
    "t_test_welch": "t-критерий Уэлча",
    "mann_whitney": "критерий Манна-Уитни",
    "anova": "однофакторный дисперсионный анализ (ANOVA)",
    "anova_welch": "ANOVA Уэлча",
    "kruskal": "критерий Краскела-Уоллиса",
    "chi_square": "критерий χ²",
    "fisher_exact": "точный критерий Фишера",
    "pearson": "коэффициент корреляции Пирсона",
    "spearman": "коэффициент корреляции Спирмена",
}


def format_p_value(p: float) -> str:
    """Format p-value according to APA guidelines."""
    if p < 0.001:
        return "p<0.001"
    elif p < 0.01:
        return f"p={p:.3f}"
    else:
        return f"p={p:.2f}"


def generate_interpretation(result: Dict[str, Any], target: str = "показатель", group: str = "группы") -> str:
    """
    Generate human-readable interpretation of statistical test results.
    """
    method_id = result.get("method", "unknown")
    method_name = METHOD_NAMES.get(method_id, method_id)
    p_value = result.get("p_value", 1.0)
    significant = result.get("significant", False)
    
    p_formatted = format_p_value(p_value)
    
    if significant:
        conclusion = f"выявлены статистически значимые различия ({p_formatted})"
    else:
        conclusion = f"статистически значимых различий не выявлено ({p_formatted})"
    
    text = f"При сравнении {group} по показателю «{target}» с использованием {method_name} {conclusion}."
    
    # Add effect size if available
    effect_size = result.get("effect_size")
    if effect_size is not None:
        if abs(effect_size) < 0.2:
            effect_desc = "незначительный"
        elif abs(effect_size) < 0.5:
            effect_desc = "малый"
        elif abs(effect_size) < 0.8:
            effect_desc = "средний"
        else:
            effect_desc = "большой"
        text += f" Размер эффекта (d Коэна): {effect_size:.2f} ({effect_desc})."
    
    return text


def generate_method_justification(result: Dict[str, Any]) -> str:
    """
    Generate justification for why this method was chosen.
    """
    method_id = result.get("method", "")
    assumptions = result.get("assumptions", {})
    
    normality = assumptions.get("normality", {})
    homogeneity = assumptions.get("homogeneity", {})
    
    parts = []
    
    # Check normality
    all_normal = all(v.get("passed", False) for v in normality.values()) if normality else True
    if all_normal:
        parts.append("данные соответствуют нормальному распределению")
    else:
        failed_groups = [k for k, v in normality.items() if not v.get("passed", True)]
        if failed_groups:
            parts.append(f"нормальность не подтверждена для групп: {', '.join(failed_groups)}")
    
    # Check homogeneity
    if homogeneity:
        if homogeneity.get("passed", True):
            parts.append("дисперсии однородны")
        else:
            parts.append("дисперсии неоднородны")
    
    if parts:
        return f"Выбор метода обоснован: {'; '.join(parts)}."
    return ""


# --- Word Document Generator ---

def create_table_one(doc: Document, stats: Dict[str, Any], groups: List[str], target_name: str = "Показатель"):
    """
    Create a Table 1 (descriptive statistics) in APA style.
    """
    # Add header
    doc.add_heading("Таблица 1. Описательная статистика", level=2)
    
    # Calculate number of columns: Metric + groups + Overall
    num_cols = 2 + len(groups)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Показатель"
    for i, g in enumerate(groups):
        count = stats.get(str(g), {}).get("count", 0)
        header_cells[i + 1].text = f"Группа {g}\n(n={count})"
    header_cells[-1].text = f"Всего\n(n={stats.get('overall', {}).get('count', 0)})"
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Metrics to show
    metrics = [
        ("Среднее (M)", "mean"),
        ("Медиана (Me)", "median"),
        ("Станд. откл. (SD)", "std"),
        ("Мин – Макс", None),  # Special handling
        ("95% ДИ", None),  # Special handling
    ]
    
    for metric_name, metric_key in metrics:
        row = table.add_row().cells
        row[0].text = metric_name
        
        for i, g in enumerate(groups):
            g_stats = stats.get(str(g), {})
            
            if metric_key:
                val = g_stats.get(metric_key, 0)
                row[i + 1].text = f"{val:.2f}" if val else "—"
            elif metric_name == "Мин – Макс":
                min_v = g_stats.get("min", 0)
                max_v = g_stats.get("max", 0)
                row[i + 1].text = f"{min_v:.1f} – {max_v:.1f}"
            elif metric_name == "95% ДИ":
                ci_l = g_stats.get("ci_lower", 0)
                ci_u = g_stats.get("ci_upper", 0)
                row[i + 1].text = f"[{ci_l:.2f}; {ci_u:.2f}]"
        
        # Overall column (only mean/count for now)
        overall = stats.get("overall", {})
        if metric_key == "mean":
            row[-1].text = f"{overall.get('mean', 0):.2f}"
        else:
            row[-1].text = "—"
    
    doc.add_paragraph()  # Spacing


def create_hypothesis_section(doc: Document, result: Dict[str, Any], target_name: str = "Показатель"):
    """
    Add hypothesis test results section.
    """
    doc.add_heading("Результаты сравнительного анализа", level=2)
    
    method_id = result.get("method", "unknown")
    method_name = METHOD_NAMES.get(method_id, method_id)
    p_value = result.get("p_value", 1.0)
    stat_value = result.get("stat_value", 0)
    
    # Method and statistic
    para = doc.add_paragraph()
    para.add_run("Метод: ").bold = True
    para.add_run(f"{method_name}\n")
    para.add_run("Статистика: ").bold = True
    para.add_run(f"{stat_value:.3f}\n")
    para.add_run("p-value: ").bold = True
    para.add_run(format_p_value(p_value))
    
    # Interpretation
    doc.add_heading("Интерпретация", level=3)
    if result.get("narrative"):
        doc.add_paragraph(result["narrative"])
    else:
        interpretation = generate_interpretation(result, target=target_name)
        doc.add_paragraph(interpretation)
    
    # Method justification
    justification = generate_method_justification(result)
    if justification:
        doc.add_paragraph(justification).italic = True


def _add_plot_to_doc(doc: Document, result: Dict[str, Any]):
    """
    Generate and embed a plot into the document.
    """
    plot_data = result.get("plot_data", [])
    plot_stats = result.get("plot_stats", {})
    
    if not plot_data and not plot_stats:
        return

    try:
        plt.figure(figsize=(6, 4))
        sns.set_style("whitegrid")
        
        # Logic similar to reporting.py but adapted
        if plot_data:
            df_plot = pd.DataFrame(plot_data)
            
            # Determine Plot Type
            if "group" in df_plot.columns and "value" in df_plot.columns:
                # Boxplot + Strip
                sns.boxplot(x="group", y="value", data=df_plot, showfliers=False, color="lightblue", width=0.5)
                sns.stripplot(x="group", y="value", data=df_plot, color=".3", size=4, alpha=0.5)
                plt.title("Comparison")
                
            elif "diff" in df_plot.columns:
                 # Paired Difference Histogram (or Strip)
                 sns.histplot(data=df_plot, x="diff", kde=True, color="skyblue")
                 plt.axvline(x=0, color='r', linestyle='--')
                 plt.title("Distribution of Differences")

            elif "x" in df_plot.columns and "y" in df_plot.columns:
                # Scatter
                sns.scatterplot(x="x", y="y", data=df_plot)
                sns.regplot(x="x", y="y", data=df_plot, scatter=False, color="red")
        
        elif plot_stats:
             # Bar chart from stats
             groups = list(plot_stats.keys())
             means = [s["mean"] for s in plot_stats.values()]
             sems = [s["sem"] for s in plot_stats.values()]
             plt.bar(groups, means, yerr=sems, capsize=5, color="skyblue", alpha=0.8)
             plt.title("Means ± SEM")

        # Save to buffer
        buf = io.BytesIO()
        plt.tight_layout()
        plt.savefig(buf, format='png', dpi=100)
        plt.close()
        buf.seek(0)
        
        # Add to doc
        doc.add_picture(buf, width=Inches(5.5))
        doc.add_paragraph("Рис. Визуализация результатов", style='Caption')
        doc.add_paragraph() # Spacing

    except Exception as e:
        print(f"Word plot generation failed: {e}")


def create_paired_desc_table(doc: Document, res: Dict[str, Any]):
    """
    Table for descriptive_paired.
    """
    doc.add_heading("Описательная статистика разностей", level=2)
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = "Переменные"
    hdr[1].text = "Разность (M±SD)"
    hdr[2].text = "Мин – Макс"
    hdr[3].text = "n"
    
    for c in hdr:
        for p in c.paragraphs:
            p.runs[0].bold = True

    # Data
    row = table.rows[1].cells
    vars = res.get('variables', ['Var1', 'Var2'])
    row[0].text = f"{vars[0]} vs {vars[1]}"
    row[1].text = f"{res.get('mean_diff', 0):.2f} ± {res.get('std_diff', 0):.2f}"
    row[2].text = f"{res.get('min_diff', 0):.2f} – {res.get('max_diff', 0):.2f}"
    row[3].text = str(res.get('n', 0))
    
    doc.add_paragraph()


def create_paired_test_section(doc: Document, result: Dict[str, Any]):
    """
    Specific for compare_paired test results.
    """
    doc.add_heading("Результаты сравнения (Связанные выборки)", level=2)
    
    method_id = result.get("method", "unknown")
    method_name = METHOD_NAMES.get(method_id, method_id)
    p_value = result.get("p_value", 1.0)
    stat_val = result.get("stat_value", 0)
    eff_size = result.get("effect_size")
    
    para = doc.add_paragraph()
    para.add_run("Метод: ").bold = True
    para.add_run(f"{method_name}\n")
    para.add_run("Статистика: ").bold = True
    para.add_run(f"{stat_val:.3f}\n")
    para.add_run("p-value: ").bold = True
    para.add_run(format_p_value(p_value))
    
    if eff_size is not None:
        para.add_run("\nРазмер эффекта: ").bold = True
        para.add_run(f"{eff_size:.2f}")

    # Interpretation
    doc.add_heading("Интерпретация", level=3)
    if result.get("narrative"):
         doc.add_paragraph(result["narrative"])
    else:
         sig_text = "выявлены статистически значимые различия" if result.get("significant") else "значимых различий не обнаружено"
         doc.add_paragraph(f"Согласно критерию {method_name}, {sig_text} (p={format_p_value(p_value)}).")


def generate_word_report(run_data: Dict[str, Any], dataset_name: str = "Dataset") -> bytes:
    """
    Generate a Word document from analysis results.
    
    Args:
        run_data: The analysis run data containing results
        dataset_name: Name of the dataset for the report header
    
    Returns:
        bytes: The .docx file content
    """
    doc = Document()
    
    # Title
    title = doc.add_heading("Статистический отчёт", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dataset info
    doc.add_paragraph(f"Файл данных: {dataset_name}")
    doc.add_paragraph()
    
    # Process results
    results = run_data.get("results", {})
    
    for step_id, res in results.items():
        step_type = res.get("type", "")
        
        if step_type == "descriptive_compare":
            # Table 1
            stats = res.get("stats", {})
            groups = [g for g in stats.keys() if g != "overall"]
            target = res.get("target", "Показатель")
            create_table_one(doc, stats, groups, target)
            
        elif step_type == "hypothesis_test" or (step_type == "" and "p_value" in res):
            # Hypothesis test
            target = res.get("target", step_id)
            create_hypothesis_section(doc, res, target)
            _add_plot_to_doc(doc, res)

        elif step_type == "descriptive_paired":
            create_paired_desc_table(doc, res)
            
        elif step_type == "compare_paired":
            create_paired_test_section(doc, res)
            _add_plot_to_doc(doc, res)
            
        elif step_type == "correlation_matrix":
            doc.add_heading("Correlation Matrix", level=2)
            # Embed pre-calculated image
            img_str = res.get("plot_image")
            if img_str:
                try:
                    import base64
                    img_bytes = base64.b64decode(img_str)
                    buf = io.BytesIO(img_bytes)
                    doc.add_picture(buf, width=Inches(6.0))
                    doc.add_paragraph("Рис. Heatmap корреляций", style='Caption')
                except Exception as e:
                    doc.add_paragraph(f"[Error rendering image: {e}]")
            else:
                 doc.add_paragraph("No visualization available.")
                 
            # Add Table of correlations? (Optional, maybe too big. Heatmap is usually enough)

        elif step_type == "regression":
            doc.add_heading("Regression Analysis", level=2)
            doc.add_paragraph(f"Target: {res.get('target', 'Y')}")
            
            # 1. Fit Stats
            stats = res.get("fit_stats", {})
            doc.add_paragraph(f"N Obs: {stats.get('n_obs')} | R2: {stats.get('r_squared', 0):.3f} | AIC: {stats.get('aic', 0):.1f}")
            
            # 2. Plot
            img_b64 = res.get("plot_image")
            if img_b64:
                 try:
                    import base64
                    img_bytes = base64.b64decode(img_b64)
                    buf = io.BytesIO(img_bytes)
                    doc.add_picture(buf, width=Inches(5.0))
                 except: pass

            # 3. Coef Table
            coefs = res.get("coef_table", [])
            if coefs:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Variable'
                hdr_cells[1].text = 'Coef'
                hdr_cells[2].text = 'P-Value'
                hdr_cells[3].text = '95% CI'
                
                for row in coefs:
                    cells = table.add_row().cells
                    cells[0].text = str(row['variable'])
                    cells[1].text = f"{row['coef']:.3f}"
                    p = row['p_value']
                    cells[2].text = "<0.001" if p < 0.001 else f"{p:.4f}"
                    cells[3].text = f"[{row['ci_lower']:.2f}, {row['ci_upper']:.2f}]"

        elif step_type == "survival":
            doc.add_heading("Survival Analysis", level=2)
            
            # 1. Stats
            medians = res.get("median_survival", {})
            p_val = res.get("p_value")
            
            p_text = f"(Log-Rank P: {p_val:.4f})" if p_val is not None else ""
            doc.add_paragraph(f"Median Survival Time {p_text}")
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = "Group"
            table.rows[0].cells[1].text = "Median Time"
            
            for grp, val in medians.items():
                cells = table.add_row().cells
                cells[0].text = str(grp)
                cells[1].text = str(val)
                
            # 2. Plot
            img_b64 = res.get("plot_image")
            if img_b64:
                 try:
                    import base64
                    img_bytes = base64.b64decode(img_b64)
                    buf = io.BytesIO(img_bytes)
                    doc.add_picture(buf, width=Inches(6.0))
                    doc.add_paragraph("Kaplan-Meier Curve", style='Caption')
                 except: pass
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_selective_word_report(
    run_data: Dict[str, Any], 
    dataset_name: str = "Dataset",
    selected_vars: List[str] = None
) -> bytes:
    """
    Generate a Word document with a summary table for selected variables only.
    Creates a compact 'Table 1 Builder' style output.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading("Статистический отчёт", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dataset info
    doc.add_paragraph(f"Файл данных: {dataset_name}")
    if selected_vars:
        doc.add_paragraph(f"Выбранные показатели: {len(selected_vars)}")
    doc.add_paragraph()
    
    results = run_data.get("results", {})
    settings = run_data.get("export_settings", {})
    show_mean = settings.get("show_mean", True)
    show_median = settings.get("show_median", False)
    show_quartiles = settings.get("show_quartiles", False)
    
    # Collect data for summary table
    table_rows = []
    
    for var_name in (selected_vars or []):
        desc_key = f"desc_{var_name}"
        test_key = f"test_{var_name}"
        
        desc_data = results.get(desc_key, {}).get("data", {})
        test_data = results.get(test_key, {})
        
        groups = [g for g in desc_data.keys() if g != "overall"]
        
        if len(groups) >= 2:
            g_a = desc_data.get(groups[0], {})
            g_b = desc_data.get(groups[1], {})
            
            # Calculate delta
            delta_abs = None
            delta_pct = None
            if g_a.get("mean") and g_b.get("mean"):
                delta_abs = g_b["mean"] - g_a["mean"]
                delta_pct = (delta_abs / g_a["mean"]) * 100 if g_a["mean"] != 0 else None
            
            table_rows.append({
                "variable": var_name,
                "n": g_a.get("count", 0),
                "group_a": groups[0],
                "group_b": groups[1],
                "g_a_mean": g_a.get("mean"),
                "g_a_std": g_a.get("std"),
                "g_a_median": g_a.get("median"),
                "g_a_q1": g_a.get("q1"),
                "g_a_q3": g_a.get("q3"),
                "g_b_mean": g_b.get("mean"),
                "g_b_std": g_b.get("std"),
                "g_b_median": g_b.get("median"),
                "g_b_q1": g_b.get("q1"),
                "g_b_q3": g_b.get("q3"),
                "delta_abs": delta_abs,
                "delta_pct": delta_pct,
                "p_value": test_data.get("p_value"),
                "significant": test_data.get("significant", False),
                "conclusion": test_data.get("conclusion", "")
            })
    
    if not table_rows:
        doc.add_paragraph("Нет данных для отображения")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    # Create summary table
    doc.add_heading("Таблица 1. Сравнительный анализ показателей", level=2)
    
    # Determine columns based on settings
    cols = ["Показатель", "n"]
    if table_rows:
        cols.extend([table_rows[0]["group_a"], table_rows[0]["group_b"]])
    cols.extend(["Δ%", "Δабс", "p"])
    
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = 'Table Grid'
    
    # Header
    header_cells = table.rows[0].cells
    for i, col_name in enumerate(cols):
        header_cells[i].text = col_name
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    for row_data in table_rows:
        row = table.add_row().cells
        row[0].text = row_data["variable"]
        row[1].text = str(row_data["n"])
        
        # Group A value
        parts_a = []
        if show_mean and row_data["g_a_mean"] is not None:
            parts_a.append(f"{row_data['g_a_mean']:.1f}±{row_data['g_a_std']:.1f}")
        if show_median and row_data["g_a_median"] is not None:
            parts_a.append(f"Me: {row_data['g_a_median']:.1f}")
        if show_quartiles and row_data["g_a_q1"] is not None:
            parts_a.append(f"[{row_data['g_a_q1']:.1f}-{row_data['g_a_q3']:.1f}]")
        row[2].text = "\n".join(parts_a) if parts_a else "—"
        
        # Group B value
        parts_b = []
        if show_mean and row_data["g_b_mean"] is not None:
            parts_b.append(f"{row_data['g_b_mean']:.1f}±{row_data['g_b_std']:.1f}")
        if show_median and row_data["g_b_median"] is not None:
            parts_b.append(f"Me: {row_data['g_b_median']:.1f}")
        if show_quartiles and row_data["g_b_q1"] is not None:
            parts_b.append(f"[{row_data['g_b_q1']:.1f}-{row_data['g_b_q3']:.1f}]")
        row[3].text = "\n".join(parts_b) if parts_b else "—"
        
        # Delta %
        if row_data["delta_pct"] is not None:
            sign = "+" if row_data["delta_pct"] > 0 else ""
            row[4].text = f"{sign}{row_data['delta_pct']:.1f}%"
        else:
            row[4].text = "—"
        
        # Delta abs
        if row_data["delta_abs"] is not None:
            sign = "+" if row_data["delta_abs"] > 0 else ""
            row[5].text = f"{sign}{row_data['delta_abs']:.2f}"
        else:
            row[5].text = "—"
        
        # P-value
        if row_data["p_value"] is not None:
            row[6].text = format_p_value(row_data["p_value"])
        else:
            row[6].text = "—"
    
    doc.add_paragraph()
    
    # Significant findings summary
    significant_vars = [r for r in table_rows if r.get("significant")]
    if significant_vars:
        doc.add_heading("Статистически значимые различия", level=2)
        for r in significant_vars:
            p = doc.add_paragraph()
            p.add_run(f"• {r['variable']}: ").bold = True
            if r["conclusion"]:
                p.add_run(r["conclusion"])
            else:
                p.add_run(f"p={r['p_value']:.3f}")
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_batch_report(
    results: Dict[str, Any],
    descriptives: List[Dict[str, Any]],
    dataset_name: str = "Dataset",
    options: Dict[str, Any] = None
) -> bytes:
    """
    Generate a Word document from batch analysis results.
    """
    doc = Document()
    options = options or {}
    
    # Title
    title = doc.add_heading("Статистический отчёт", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Файл данных: {dataset_name}")
    doc.add_paragraph(f"Количество показателей: {len(results)}")
    doc.add_paragraph()
    
    # Get unique groups
    groups = sorted(set(str(d.get("group")) for d in descriptives if d.get("group")))
    
    # Summary Table
    doc.add_heading("Таблица 1. Сводные результаты", level=2)
    
    num_groups = len(groups)
    table = doc.add_table(rows=1, cols=3 + num_groups)
    table.style = 'Table Grid'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = "Показатель"
    for i, g in enumerate(groups):
        header_cells[1 + i].text = f"Группа {g}"
    header_cells[-2].text = "p-value"
    header_cells[-1].text = "Знач."
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for var_name, result in results.items():
        row = table.add_row().cells
        row[0].text = var_name
        
        var_descs = [d for d in descriptives if d.get("variable") == var_name]
        
        for i, g in enumerate(groups):
            g_desc = next((d for d in var_descs if str(d.get("group")) == g), None)
            if g_desc and g_desc.get("mean") is not None:
                mean = g_desc.get("mean", 0)
                sd = g_desc.get("sd", 0)
                row[1 + i].text = f"{mean:.1f}±{sd:.1f}"
            else:
                row[1 + i].text = "—"
        
        p_value = result.get("p_value")
        row[-2].text = format_p_value(p_value) if p_value else "—"
        row[-1].text = "✓" if result.get("significant") else ""
    
    doc.add_paragraph()
    
    # Significant findings
    significant = [(v, r) for v, r in results.items() if r.get("significant")]
    if significant:
        doc.add_heading("Статистически значимые различия", level=2)
        for var_name, result in significant:
            method = result.get("method", {}).get("name", "")
            p = result.get("p_value", 0)
            para = doc.add_paragraph()
            para.add_run(f"• {var_name}: ").bold = True
            para.add_run(f"{method}, {format_p_value(p)}")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# For testing

if __name__ == "__main__":
    test_data = {
        "results": {
            "descriptive": {
                "type": "descriptive_compare",
                "target": "Возраст",
                "stats": {
                    "A": {"count": 30, "mean": 45.2, "median": 44.0, "std": 12.3, "min": 22, "max": 68, "ci_lower": 40.5, "ci_upper": 49.9},
                    "B": {"count": 28, "mean": 52.1, "median": 51.0, "std": 10.8, "min": 31, "max": 72, "ci_lower": 47.9, "ci_upper": 56.3},
                    "overall": {"count": 58, "mean": 48.5}
                }
            },
            "hypothesis": {
                "type": "hypothesis_test",
                "target": "Возраст",
                "method": "t_test_ind",
                "stat_value": 2.341,
                "p_value": 0.023,
                "significant": True,
                "effect_size": 0.58,
                "assumptions": {
                    "normality": {"A": {"passed": True}, "B": {"passed": True}},
                    "homogeneity": {"passed": True}
                }
            }
        }
    }
    
    doc_bytes = generate_word_report(test_data, "test_data.csv")
    with open("/tmp/test_report.docx", "wb") as f:
        f.write(doc_bytes)
    print(f"Test report generated: {len(doc_bytes)} bytes")
