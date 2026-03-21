import os
import json
import pandas as pd
import numpy as np
import base64
import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from typing import Dict, Any, List, Optional
from pathlib import Path
from jinja2 import Environment, FileSystemLoader
from app.schemas.analysis import AnalysisResult
from app.core.logging import logger

from app.modules.plot_with_brackets import add_significance_bracket, normalize_comparisons
from app.modules.plot_config import apply_publication_config

from fpdf import FPDF

TEMPLATE_DIR = Path(__file__).parent.parent / "templates"

def _render_plot_png_bytes(res: Dict[str, Any]) -> bytes:
    try:
        apply_publication_config()
        plt.figure(figsize=(8, 5))

        plot_data = []
        plot_config = {}

        if isinstance(res, dict):
            roc = res.get("roc")
            if isinstance(roc, dict) and isinstance(roc.get("plot_data"), list) and roc.get("plot_data"):
                plot_data = roc.get("plot_data")
                plot_config = roc.get("plot_config") if isinstance(roc.get("plot_config"), dict) else {}
            else:
                plot_data = res.get("plot_data", [])
                plot_config = res.get("plot_config") if isinstance(res.get("plot_config"), dict) else {}

        if plot_data:
            df_plot = pd.DataFrame(plot_data)

            if "group" in df_plot.columns and "value" in df_plot.columns:
                sns.boxplot(x="group", y="value", data=df_plot, showfliers=False, color="lightblue", width=0.5)
                sns.stripplot(
                    x="group",
                    y="value",
                    data=df_plot,
                    size=4,
                    alpha=0.6,
                    color="#0f172a",
                )
                plt.title("Group Comparison")

                comparisons_raw = None
                if isinstance(res, dict):
                    comparisons_raw = res.get("comparisons") or res.get("plot_comparisons") or res.get("post_hoc")

                comparisons = normalize_comparisons(comparisons_raw)
                if comparisons:
                    group_order = [str(g) for g in (df_plot["group"].dropna().unique().tolist() or [])]
                    group_index = {g: i for i, g in enumerate(group_order)}

                    values = df_plot["value"].dropna().astype(float)
                    if len(values) > 0:
                        min_val = float(values.min())
                        max_val = float(values.max())
                        y_range = (max_val - min_val) or 1.0
                        base_pad = y_range * 0.08
                        step_pad = y_range * 0.08
                        y_base = max_val + base_pad

                        ranges = []
                        placed = []
                        for c in comparisons:
                            ia = group_index.get(c.a)
                            ib = group_index.get(c.b)
                            if ia is None or ib is None:
                                continue
                            start = min(ia, ib)
                            end = max(ia, ib)

                            level = 0
                            while True:
                                taken = ranges[level] if level < len(ranges) else []
                                overlaps = any(not (end < r[0] or start > r[1]) for r in taken)
                                if not overlaps:
                                    break
                                level += 1
                            while level >= len(ranges):
                                ranges.append([])
                            ranges[level].append((start, end))
                            placed.append((start, end, level, c.p_value))

                        ax = plt.gca()
                        max_level = max((lvl for _, _, lvl, _ in placed), default=-1)
                        try:
                            y0, y1_lim = ax.get_ylim()
                            extra = (max_level + 2) * step_pad
                            ax.set_ylim(y0, max(y1_lim, max_val + base_pad + extra))
                        except Exception:
                            pass
                        for start, end, level, p_value in placed:
                            add_significance_bracket(
                                ax,
                                float(start),
                                float(end),
                                y_base + level * step_pad,
                                p_value,
                                h=0.02,
                                lw=1.2,
                                color="#0f172a",
                            )

            elif "x" in df_plot.columns and "y" in df_plot.columns:
                if plot_config.get("type") == "line":
                    df_sorted = df_plot.sort_values("x")
                    plt.plot(df_sorted["x"], df_sorted["y"], color="#8b5cf6", linewidth=2)
                    plt.plot([0, 1], [0, 1], linestyle="--", color="#666", linewidth=1)
                    plt.xlim(0, 1)
                    plt.ylim(0, 1)
                    plt.title("ROC Curve")
                else:
                    sns.scatterplot(x="x", y="y", data=df_plot)
                    sns.regplot(x="x", y="y", data=df_plot, scatter=False, color="red")
                    plt.title("Correlation Analysis")

            elif "probability" in df_plot.columns and "time" in df_plot.columns and "group" in df_plot.columns:
                groups = df_plot["group"].unique()
                for g in groups:
                    sub = df_plot[df_plot["group"] == g]
                    plt.step(sub["time"], sub["probability"], where="post", label=f"Group {g}")
                plt.ylim(0, 1.05)
                plt.legend()
                plt.title("Kaplan-Meier Survival Curve")

        else:
            plot_stats = res.get("plot_stats", {}) if isinstance(res, dict) else {}
            if plot_stats:
                groups = list(plot_stats.keys())
                means = [s["mean"] for s in plot_stats.values()]
                sems = [s["sem"] for s in plot_stats.values()]
                plt.bar(groups, means, yerr=sems, capsize=5, color="skyblue", alpha=0.8)
                plt.title("Mean comparison (±SEM)")
            else:
                plt.text(0.5, 0.5, "No Visualization Available", ha="center", va="center", transform=plt.gca().transAxes)

        buf = io.BytesIO()
        plt.tight_layout()
        plt.savefig(buf, format="png")
        plt.close()
        return bytes(buf.getvalue())
    except Exception as e:
        logger.error(f"Plotting failed: {e}", exc_info=True)
        try:
            plt.close()
        except Exception:
            pass
        return b""

class ProtocolReport:
    """
    Generates a comprehensive HTML report from a Protocol Analysis Run.
    V2 Report Engine supporting multi-step protocols.
    """
    
    def __init__(self, run_data: Dict, dataset_name: str = "Dataset", style: str = "apa7",
                 dataset_id: Optional[str] = None):
        self.data = run_data # The full results.json
        self.dataset_name = dataset_name
        self.dataset_id = dataset_id
        self.style = style or "apa7"
        self.html_parts = []
        self.fig_counter = 0
        self.table_counter = 0

    def generate_html(self) -> str:
        self._add_header()

        # 0. Methods section from transform_log
        self._add_methods_section()

        results = self.data.get("results", {})

        # 1. Look for Table 1 (Descriptive)
        for step_id, res in results.items():
            rtype = res.get("type") or ""
            if rtype in ("table_1", "descriptive_compare"):
                self._add_table_one(res, step_id)

        # 2. Look for Hypothesis Tests / Analyses
        for step_id, res in results.items():
            if "error" in res:
                continue
            rtype = res.get("type") or ""
            if rtype in ("table_1", "descriptive_compare"):
                continue  # already handled above
            # Explicit type match
            if rtype in ("compare", "hypothesis_test", "correlation", "survival"):
                self._add_analysis_section(res, step_id)
            elif rtype == "regression":
                self._add_regression_section(res, step_id)
            elif rtype == "clustered_correlation":
                self._add_correlation_matrix_section(res, step_id)
            elif rtype in ("batch_compare_by_factor", "longitudinal_comparison"):
                self._add_longitudinal_section(res, step_id)
            # Fallback: if result has p_value and method, treat as analysis
            elif "p_value" in res and ("method" in res or "stats" in res):
                self._add_analysis_section(res, step_id)

        self._add_footer()
        return "\n".join(self.html_parts)

    def _add_methods_section(self):
        """Insert auto-generated Methods paragraph from transform_log.json."""
        if not self.dataset_id:
            return
        try:
            workspace_dir = os.getenv("STATWIZARD_WORKSPACE_DIR", "workspace")
            log_path = os.path.join(workspace_dir, "datasets", self.dataset_id, "processed", "transform_log.json")
            if not os.path.exists(log_path):
                return
            with open(log_path, "r", encoding="utf-8") as f:
                transform_log = json.load(f)
            if not transform_log:
                return
            methods_text = generate_methods_section(transform_log, dataset_name=self.dataset_name)
            style_key = str(self.style or "apa7").strip().lower()
            heading = "Методы" if style_key == "gost" else "Methods"
            self.html_parts.append(f"""
            <div class="card">
                <h2>{heading}</h2>
                <p style="text-align: justify;">{methods_text}</p>
            </div>
            """)
        except Exception as exc:
            logger.debug(f"_add_methods_section failed: {exc}")

    def _add_header(self):
        style_key = str(self.style or "apa7").strip().lower()
        if style_key == "nature":
            css = """
            <style>
                @import url('https://fonts.googleapis.com/css2?family=Source+Sans+Pro:wght@400;600;700&display=swap');
                body { font-family: 'Source Sans Pro', 'Helvetica Neue', Arial, sans-serif; font-size: 14px; line-height: 1.6; color: #1a1a1a; max-width: 860px; margin: 0 auto; padding: 48px 56px; background: #fff; }
                h1 { font-size: 22px; font-weight: 700; color: #003f5c; margin: 0 0 6px; letter-spacing: -0.3px; }
                h2 { font-size: 15px; font-weight: 700; color: #003f5c; margin-top: 32px; margin-bottom: 6px; padding-bottom: 4px; border-bottom: 2px solid #003f5c; text-transform: uppercase; letter-spacing: 0.5px; }
                h3 { font-size: 13px; font-weight: 600; color: #1a1a1a; margin-top: 14px; }
                .card { margin-bottom: 28px; padding: 0; }
                table { width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 12.5px; }
                thead tr { border-top: 2px solid #1a1a1a; border-bottom: 1px solid #1a1a1a; }
                tbody tr:last-child td { border-bottom: 2px solid #1a1a1a; }
                th, td { padding: 8px 10px; text-align: left; vertical-align: top; }
                th { font-weight: 700; background: transparent; }
                .fig-caption { font-size: 11.5px; color: #444; margin-top: 6px; font-style: italic; }
                .table-caption { font-size: 12px; font-weight: 700; margin-bottom: 4px; }
                .stat-val { font-weight: 700; }
                .sig-yes { font-weight: 700; }
                .sig-no { color: #666; }
                .plot-container { text-align: center; margin: 18px 0 4px; }
                img { max-width: 100%; height: auto; }
                .ai-box { background: #f7f7f7; border-left: 3px solid #003f5c; padding: 10px 14px; margin-top: 14px; font-size: 13px; }
                .meta-info { color: #555; font-size: 12px; margin-bottom: 28px; }
                @media print { body { padding: 0; max-width: 100%; } .card { break-inside: avoid; } }
            </style>
            """
        elif style_key == "lancet":
            css = """
            <style>
                body { font-family: 'Georgia', 'Times New Roman', serif; font-size: 14px; line-height: 1.65; color: #111; max-width: 840px; margin: 0 auto; padding: 48px 56px; background: #fff; }
                h1 { font-size: 22px; font-weight: 700; color: #8b0000; margin: 0 0 6px; }
                h2 { font-size: 16px; font-weight: 700; color: #8b0000; margin-top: 30px; margin-bottom: 6px; padding-bottom: 5px; border-bottom: 1.5px solid #8b0000; }
                h3 { font-size: 14px; font-weight: 700; font-style: italic; margin-top: 14px; }
                .card { margin-bottom: 28px; }
                table { width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 13px; }
                thead tr { border-top: 2px solid #111; border-bottom: 1px solid #111; }
                tbody tr:last-child td { border-bottom: 2px solid #111; }
                th, td { padding: 8px 12px; text-align: left; vertical-align: top; }
                th { font-weight: 700; background: transparent; }
                .fig-caption { font-size: 12px; color: #333; margin-top: 5px; font-style: italic; }
                .table-caption { font-size: 13px; font-weight: 700; margin-bottom: 4px; }
                .stat-val { font-weight: 700; font-family: 'Courier New', monospace; }
                .sig-yes { font-weight: 700; }
                .sig-no { color: #555; }
                .plot-container { text-align: center; margin: 18px 0 4px; }
                img { max-width: 100%; height: auto; }
                .ai-box { background: #fff8f8; border-left: 3px solid #8b0000; padding: 10px 14px; margin-top: 14px; font-size: 13px; }
                .meta-info { color: #555; font-size: 12px; margin-bottom: 24px; }
                @media print { body { padding: 0; max-width: 100%; } .card { break-inside: avoid; } }
            </style>
            """
        elif style_key == "nejm":
            css = """
            <style>
                body { font-family: 'Arial', 'Helvetica Neue', sans-serif; font-size: 13.5px; line-height: 1.6; color: #1a1a1a; max-width: 860px; margin: 0 auto; padding: 48px 56px; background: #fff; }
                h1 { font-size: 20px; font-weight: 700; color: #002b5c; margin: 0 0 6px; }
                h2 { font-size: 14px; font-weight: 700; color: #002b5c; text-transform: uppercase; letter-spacing: 0.6px; margin-top: 28px; margin-bottom: 6px; padding-bottom: 4px; border-bottom: 2px solid #002b5c; }
                h3 { font-size: 13px; font-weight: 700; margin-top: 14px; }
                .card { margin-bottom: 26px; }
                table { width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 12.5px; }
                thead tr { border-top: 2px solid #1a1a1a; border-bottom: 1px solid #1a1a1a; }
                tbody tr:last-child td { border-bottom: 2px solid #1a1a1a; }
                th, td { padding: 7px 10px; text-align: left; vertical-align: top; }
                th { font-weight: 700; background: transparent; }
                .fig-caption { font-size: 11.5px; color: #444; margin-top: 6px; }
                .table-caption { font-size: 12.5px; font-weight: 700; margin-bottom: 4px; }
                .stat-val { font-weight: 700; }
                .sig-yes { font-weight: 700; }
                .sig-no { color: #666; }
                .plot-container { text-align: center; margin: 18px 0 4px; }
                img { max-width: 100%; height: auto; }
                .ai-box { background: #f0f4f8; border-left: 3px solid #002b5c; padding: 10px 14px; margin-top: 14px; font-size: 12.5px; }
                .meta-info { color: #555; font-size: 12px; margin-bottom: 24px; }
                @media print { body { padding: 0; max-width: 100%; } .card { break-inside: avoid; } }
            </style>
            """
        elif style_key == "gost":
            css = """
            <style>
                body { font-family: 'Times New Roman', 'Times', serif; line-height: 1.5; color: #111; max-width: 820px; margin: 0 auto; padding: 48px; font-size: 14px; }
                h1 { font-size: 22px; font-weight: 700; margin: 0 0 18px; padding-bottom: 10px; border-bottom: 1px solid #111; }
                h2 { font-size: 18px; font-weight: 700; margin-top: 28px; padding-bottom: 6px; border-bottom: 1px solid #ddd; }
                h3 { font-size: 15px; font-weight: 700; margin-top: 16px; }
                .card { background: #fff; border: 1px solid #ddd; padding: 18px 20px; margin-bottom: 18px; }
                table { width: 100%; border-collapse: collapse; margin-top: 12px; font-size: 13px; }
                th, td { padding: 10px 12px; border-bottom: 1px solid #e6e6e6; text-align: left; vertical-align: top; }
                th { background-color: #f7f7f7; font-weight: 700; color: #111; }
                .fig-caption { font-size: 12px; color: #444; margin-top: 6px; font-style: italic; }
                .table-caption { font-size: 13px; font-weight: 700; margin-bottom: 4px; }
                .stat-val { font-family: 'Courier New', monospace; font-weight: 700; }
                .sig-yes { color: #0f5132; font-weight: 700; }
                .sig-no { color: #495057; }
                .plot-container { text-align: center; margin-top: 14px; }
                img { max-width: 100%; height: auto; border: 1px solid #e6e6e6; }
                .ai-box { background: #fafafa; border-left: 3px solid #111; padding: 12px 14px; margin-top: 14px; }
                .meta-info { color: #333; font-size: 13px; margin-bottom: 22px; }
                @media print { body { padding: 0; max-width: 100%; } .card { break-inside: avoid; border: none; padding: 0; margin-bottom: 26px; } }
            </style>
            """
        elif style_key == "simple":
            css = """
            <style>
                body { font-family: ui-sans-serif, system-ui, -apple-system, 'Segoe UI', sans-serif; line-height: 1.55; color: #111; max-width: 920px; margin: 0 auto; padding: 44px; }
                h1 { font-size: 20px; margin: 0 0 16px; padding-bottom: 10px; border-bottom: 1px solid #e5e7eb; }
                h2 { font-size: 16px; margin-top: 26px; padding-bottom: 6px; border-bottom: 1px solid #eef2f7; }
                h3 { font-size: 13px; margin-top: 14px; }
                .card { border: 1px solid #e5e7eb; padding: 16px 16px; margin-bottom: 14px; }
                table { width: 100%; border-collapse: collapse; margin-top: 10px; font-size: 13px; }
                th, td { padding: 9px 10px; border-bottom: 1px solid #f1f5f9; text-align: left; }
                th { font-weight: 700; color: #111; background: #fafafa; }
                .fig-caption { font-size: 11px; color: #64748b; margin-top: 5px; font-style: italic; }
                .table-caption { font-size: 12px; font-weight: 700; margin-bottom: 4px; }
                .stat-val { font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, monospace; font-weight: 600; }
                .sig-yes { color: #111; font-weight: 700; }
                .sig-no { color: #64748b; }
                .plot-container { text-align: center; margin-top: 12px; }
                img { max-width: 100%; height: auto; border: 1px solid #f1f5f9; }
                .ai-box { background: #fafafa; border-left: 2px solid #111; padding: 10px 12px; margin-top: 12px; }
                .meta-info { color: #475569; font-size: 12px; margin-bottom: 18px; }
                @media print { body { padding: 0; max-width: 100%; } .card { break-inside: avoid; border: none; padding: 0; margin-bottom: 22px; } }
            </style>
            """
        else:
            # apa7 (default)
            css = """
            <style>
                body { font-family: 'Helvetica Neue', 'Helvetica', 'Arial', sans-serif; line-height: 1.6; color: #333; max-width: 900px; margin: 0 auto; padding: 40px; }
                h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; margin-bottom: 20px; }
                h2 { color: #2980b9; margin-top: 40px; border-bottom: 1px solid #eee; padding-bottom: 5px; }
                h3 { color: #16a085; font-size: 1.1em; margin-top: 20px; }
                .card { background: #fff; border: 1px solid #e1e4e8; padding: 25px; border-radius: 8px; margin-bottom: 30px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
                table { width: 100%; border-collapse: collapse; margin-top: 15px; font-size: 0.95em; }
                th, td { padding: 12px 15px; border-bottom: 1px solid #e1e4e8; text-align: left; }
                th { background-color: #f8f9fa; font-weight: 600; color: #444; }
                tr:last-child td { border-bottom: none; }
                .fig-caption { font-size: 12px; color: #555; margin-top: 8px; font-style: italic; }
                .table-caption { font-size: 13px; font-weight: 700; margin-bottom: 6px; }
                .stat-val { font-family: 'SF Mono', 'Monaco', monospace; font-weight: 600; }
                .sig-yes { color: #27ae60; font-weight: bold; background: #eafaf1; padding: 2px 6px; border-radius: 4px; }
                .sig-no { color: #7f8c8d; }
                .plot-container { text-align: center; margin-top: 20px; background: #fff; padding: 10px; }
                img { max-width: 100%; height: auto; border-radius: 4px; border: 1px solid #eee; }
                .ai-box { background: #f0f7fb; border-left: 4px solid #3498db; padding: 15px; margin-top: 20px; border-radius: 0 4px 4px 0; }
                .meta-info { color: #666; font-size: 0.9em; margin-bottom: 30px; }
                @media print { body { padding: 0; max-width: 100%; } .card { break-inside: avoid; border: none; box-shadow: none; padding: 0; margin-bottom: 40px; } h1 { margin-top: 0; } }
            </style>
            """
        self.html_parts.append(f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Analysis Report - {self.dataset_name}</title>
            {css}
        </head>
        <body>
            <h1>Statistical Analysis Report</h1>
            <div class="meta-info">
                <p><strong>Protocol:</strong> {self.data.get('protocol_name', 'Custom Analysis')}</p>
                <p><strong>Dataset:</strong> {self.dataset_name}</p>
                <p><strong>Date Generated:</strong> {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}</p>
            </div>
        """)

    def _add_table_one(self, res: Dict, step_id: str):
        stats = res.get("data", {})
        if not stats: return

        self.table_counter += 1
        groups = [k for k in stats.keys() if k != 'overall']

        html = f"""
        <div class="card">
            <div class="table-caption">Table {self.table_counter}. Descriptive Statistics</div>
            <table>
                <thead>
                    <tr>
                        <th style="width: 30%">Characteristic</th>
                        {''.join([f'<th>{g} (n={stats[g]["count"]})</th>' for g in groups])}
                        <th>Overall (n={stats['overall']['count']})</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        metrics = [
            ("Mean (SD)", lambda s: f"{s['mean']:.2f} ({s['std']:.2f})"),
            ("95% CI (Mean)", lambda s: f"[{s.get('ci_95_low', 0):.2f}, {s.get('ci_95_high', 0):.2f}]"),
            ("Median [Q1, Q3]", lambda s: f"{s['median']:.2f} [{s['q1']:.2f}, {s['q3']:.2f}]"),
            ("IQR", lambda s: f"{s.get('iqr', 0):.2f}"),
            ("Range (Min-Max)", lambda s: f"{s['min']:.2f} - {s['max']:.2f}"),
            ("Normality (Shapiro p)", lambda s: f"{s['shapiro_p']:.3f} {'(!)' if s['shapiro_p'] < 0.05 else ''}")
        ]
        
        for name, formatter in metrics:
            row = f"<tr><td>{name}</td>"
            for g in groups:
                 row += f"<td>{formatter(stats[g])}</td>"
            row += f"<td>{formatter(stats['overall'])}</td></tr>"
            html += row
            
        html += """
                </tbody>
            </table>
        </div>
        """
        self.html_parts.append(html)

    def _add_analysis_section(self, res: Dict, step_id: str):
        sig_class = "sig-yes" if res.get("significant") else "sig-no"
        sig_text = "SIGNIFICANT" if res.get("significant") else "Not Significant"

        # method can be a dict, a string repr, or missing
        method_raw = res.get('method')
        if isinstance(method_raw, dict):
            method_name = method_raw.get('name', 'Statistical Test')
        elif isinstance(method_raw, str):
            # Parse "id='...' name='...'" format
            import re
            m = re.search(r"name='([^']+)'", method_raw)
            method_name = m.group(1) if m else method_raw
        else:
            method_name = 'Statistical Test'
        p_val = res.get('p_value', 1.0)
        p_display = "< 0.001" if p_val < 0.001 else f"{p_val:.4f}"
        
        html = f"""
        <div class="card">
            <h2>Analysis Step: {step_id}</h2>
            <div style="display: flex; justify-content: space-between; align-items: flex-start;">
                <div>
                    <h3>{method_name}</h3>
                    <table style="width: auto; margin-top: 10px;">
                        <tr>
                            <td><strong>P-Value:</strong></td>
                            <td><span class="stat-val {sig_class}">{p_display}</span></td>
                        </tr>
                        <tr>
                            <td><strong>Statistic:</strong></td>
                            <td>{float(res.get('stat_value', res.get('stats', 0)) or 0):.3f}</td>
                        </tr>
                        <tr>
                            <td><strong>Effect size:</strong></td>
                            <td>
                                {(
                                    f"{res.get('effect_size_name') or 'effect'} = {float(res.get('effect_size')):.2f}"
                                    if res.get('effect_size') is not None
                                    else "-"
                                )}
                            </td>
                        </tr>
                        <tr>
                            <td><strong>CI (effect):</strong></td>
                            <td>
                                {(
                                    f"[{float(res.get('effect_size_ci_lower')):.2f}, {float(res.get('effect_size_ci_upper')):.2f}]"
                                    if (res.get('effect_size_ci_lower') is not None and res.get('effect_size_ci_upper') is not None)
                                    else "-"
                                )}
                            </td>
                        </tr>
                        <tr>
                            <td><strong>Power:</strong></td>
                            <td>{(f"{float(res.get('power')):.2f}" if res.get('power') is not None else "-")}</td>
                        </tr>
                        <tr>
                            <td><strong>BF10:</strong></td>
                            <td>{(str(res.get('bf10')) if res.get('bf10') is not None else "-")}</td>
                        </tr>
                        <tr>
                            <td><strong>Result:</strong></td>
                            <td>{sig_text}</td>
                        </tr>
                    </table>
                </div>
            </div>
        """
        
        # Generate Plot
        img_b64 = self._generate_plot_image(res)
        if img_b64:
            self.fig_counter += 1
            p_caption = ""
            p_val = res.get("p_value")
            if p_val is not None:
                try:
                    p_fmt = "< 0.001" if float(p_val) < 0.001 else f"{float(p_val):.4f}"
                    p_caption = f" (p = {p_fmt})"
                except Exception:
                    pass
            html += (
                f'<div class="plot-container">'
                f'<img src="data:image/png;base64,{img_b64}" alt="Figure {self.fig_counter}" />'
                f'<div class="fig-caption">Figure {self.fig_counter}. {method_name}{p_caption}</div>'
                f'</div>'
            )

        if res.get("conclusion"):
            html += f'<div class="ai-box"><strong>AI Interpretation:</strong><br>{res["conclusion"]}</div>'
            
        html += "</div>"
        self.html_parts.append(html)

    def _add_longitudinal_section(self, res: Dict, step_id: str):
        self.table_counter += 1
        html = f"""
        <div class="card">
            <h2>Longitudinal Analysis: {step_id}</h2>
            <div class="table-caption">Table {self.table_counter}. Results split by: {res.get('split_by', '')}</div>
            <table>
                <thead>
                    <tr>
                        <th>Timepoint / Split</th>
                        <th>Method</th>
                        <th>P-Value</th>
                        <th>Result</th>
                    </tr>
                </thead>
                <tbody>
        """
        
        for slice_key, slice_res in res.get("slices", {}).items():
            is_sig = slice_res.get("significant", False)
            p_val = slice_res.get('p_value', 1.0)
            p_display = "< 0.001" if p_val is not None and p_val < 0.001 else (f"{p_val:.4f}" if p_val is not None else "-")

            # method can be dict or str
            method_raw = slice_res.get('method')
            if isinstance(method_raw, dict):
                method_name = method_raw.get('name', '-')
            elif isinstance(method_raw, str):
                import re
                m = re.search(r"name='([^']+)'", method_raw)
                method_name = m.group(1) if m else method_raw
            else:
                method_name = '-'

            html += f"""
                <tr>
                    <td><strong>{slice_key}</strong></td>
                    <td>{method_name}</td>
                    <td><span class="stat-val { 'sig-yes' if is_sig else 'sig-no' }">{p_display}</span></td>
                    <td>{ 'Difference Detected' if is_sig else 'No Difference' }</td>
                </tr>
            """
            
        html += "</tbody></table></div>"
        self.html_parts.append(html)

    def _add_regression_section(self, res: Dict, step_id: str):
        """Render regression results (linear or logistic)."""
        method_raw = res.get('method')
        if isinstance(method_raw, dict):
            method_name = method_raw.get('name', 'Regression')
        elif isinstance(method_raw, str):
            import re
            m = re.search(r"name='([^']+)'", method_raw)
            method_name = m.group(1) if m else 'Regression'
        else:
            method_name = 'Regression'

        p_val = res.get('p_value')
        p_display = "< 0.001" if p_val is not None and p_val < 0.001 else (f"{p_val:.4f}" if p_val is not None else "-")
        r_sq = res.get('r_squared')

        self.table_counter += 1
        html = f"""
        <div class="card">
            <h2>{method_name}: {step_id}</h2>
            <p>P-value: <span class="stat-val {'sig-yes' if res.get('significant') or (p_val is not None and p_val < 0.05) else 'sig-no'}">{p_display}</span></p>
        """
        if r_sq is not None:
            html += f"<p>R² = {float(r_sq):.4f}</p>"

        coefficients = res.get('coefficients')
        if coefficients and isinstance(coefficients, (list, dict)):
            html += f"""
            <div class="table-caption">Table {self.table_counter}. Regression Coefficients</div>
            <table>
                <thead><tr><th>Variable</th><th>Coefficient</th><th>Std Error</th><th>P-value</th></tr></thead>
                <tbody>
            """
            coefs = coefficients if isinstance(coefficients, list) else [coefficients]
            for c in coefs:
                if isinstance(c, dict):
                    name = c.get('variable', c.get('name', '?'))
                    coef = c.get('coef', c.get('coefficient', '-'))
                    se = c.get('std_err', c.get('se', '-'))
                    cp = c.get('p_value', '-')
                    html += f"<tr><td>{name}</td><td>{coef}</td><td>{se}</td><td>{cp}</td></tr>"
            html += "</tbody></table>"

        if res.get('conclusion'):
            html += f'<div class="ai-box"><strong>AI Interpretation:</strong><br>{res["conclusion"]}</div>'

        html += "</div>"
        self.html_parts.append(html)

    def _add_correlation_matrix_section(self, res: Dict, step_id: str):
        """Render clustered correlation matrix."""
        corr_matrix = res.get('correlation_matrix') or {}
        variables = corr_matrix.get('variables', [])
        if not variables:
            return  # nothing to render
        values = corr_matrix.get('values', [])
        n_clusters = res.get('optimal_n_clusters', res.get('n_clusters', '?'))

        self.table_counter += 1
        html = f"""
        <div class="card">
            <h2>Correlation Matrix: {step_id}</h2>
            <p>{len(variables)} variables, {n_clusters} clusters</p>
            <div class="table-caption">Table {self.table_counter}. Correlation Matrix (Spearman)</div>
            <table style="font-size: 11px;">
                <thead><tr><th></th>{''.join(f'<th>{v[:12]}</th>' for v in variables)}</tr></thead>
                <tbody>
        """
        for i, var in enumerate(variables):
            row_vals = values[i] if i < len(values) else []
            cells = ""
            for j, val in enumerate(row_vals):
                if val is None:
                    cells += "<td>-</td>"
                else:
                    color = f"rgba(59,130,246,{min(abs(val), 1) * 0.5})" if val >= 0 else f"rgba(239,68,68,{min(abs(val), 1) * 0.5})"
                    cells += f'<td style="background:{color}; text-align:center">{val:.2f}</td>'
            html += f"<tr><td><strong>{var[:12]}</strong></td>{cells}</tr>"

        html += "</tbody></table></div>"
        self.html_parts.append(html)

    def _add_footer(self):
        self.html_parts.append("""
        <div style="margin-top: 50px; color: #888; text-align: center; font-size: 0.8em; border-top: 1px solid #eee; padding-top: 20px;">
            Generated by AI Biostatistics Platform &bull; Antigravity Agent
        </div>
        </body></html>
        """)

    def _generate_plot_image(self, res: Dict) -> str:
        """
        Uses matplotlib/seaborn to render the plot stats into a base64 string.
        """
        try:
            png_bytes = _render_plot_png_bytes(res)
            if not png_bytes:
                return ""
            return base64.b64encode(png_bytes).decode("utf-8")
        except Exception as e:
            logger.error(f"Plotting failed: {e}", exc_info=True)
            return ""

def generate_protocol_docx_report(run_data: Dict[str, Any], dataset_name: str = "Dataset", style: Optional[str] = None) -> bytes:
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches

    def _fmt_p(value: Any) -> str:
        try:
            if value is None:
                return "-"
            p = float(value)
            if not np.isfinite(p):
                return "-"
            return "< 0.001" if p < 0.001 else f"{p:.4f}"
        except Exception:
            return "-"

    def _fmt_num(value: Any, digits: int = 3) -> str:
        try:
            if value is None:
                return "-"
            num = float(value)
            if not np.isfinite(num):
                return "-"
            return f"{num:.{digits}f}"
        except Exception:
            return "-"

    def _txt(value: Any) -> str:
        return "-" if value is None else str(value)

    style_key = str(style or "gost").strip().lower()
    is_ru = style_key in {"gost"}

    doc = Document()
    doc.add_heading("Результаты статистического анализа" if is_ru else "Statistical Analysis Results", level=0)
    doc.add_paragraph(("Набор данных" if is_ru else "Dataset") + f": {dataset_name}")
    protocol_name = run_data.get("protocol_name") if isinstance(run_data, dict) else None
    if protocol_name:
        doc.add_paragraph(("Протокол" if is_ru else "Protocol") + f": {protocol_name}")
    doc.add_paragraph(("Дата" if is_ru else "Date") + f": {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")

    results = run_data.get("results", {}) if isinstance(run_data, dict) else {}
    for step_id, res in (results or {}).items():
        doc.add_heading(("Шаг" if is_ru else "Step") + f": {step_id}", level=1)

        if not isinstance(res, dict):
            doc.add_paragraph("Нет структурированного результата" if is_ru else "No structured result")
            continue

        if res.get("type") == "table_1":
            stats_map = res.get("data", {})
            if isinstance(stats_map, dict) and stats_map:
                groups = [k for k in stats_map.keys() if k != "overall"]
                cols = 2 + len(groups)
                table = doc.add_table(rows=1, cols=cols)
                hdr = table.rows[0].cells
                hdr[0].text = "Показатель"
                for i, g in enumerate(groups):
                    n = _txt(stats_map.get(g, {}).get("count"))
                    hdr[i + 1].text = f"{g} (n={n})"
                overall_n = _txt(stats_map.get("overall", {}).get("count"))
                hdr[-1].text = f"Итого (n={overall_n})"

                def _cell_for(metric_key: str, s: Dict[str, Any]) -> str:
                    if metric_key == "mean_sd":
                        return f"{_fmt_num(s.get('mean'), 2)} ({_fmt_num(s.get('std'), 2)})"
                    if metric_key == "ci_95":
                        return f"[{_fmt_num(s.get('ci_95_low'), 2)}, {_fmt_num(s.get('ci_95_high'), 2)}]"
                    if metric_key == "median_q1_q3":
                        return f"{_fmt_num(s.get('median'), 2)} [{_fmt_num(s.get('q1'), 2)}, {_fmt_num(s.get('q3'), 2)}]"
                    if metric_key == "iqr":
                        return _fmt_num(s.get("iqr"), 2)
                    if metric_key == "min_max":
                        return f"{_fmt_num(s.get('min'), 2)} – {_fmt_num(s.get('max'), 2)}"
                    if metric_key == "shapiro":
                        return _fmt_p(s.get("shapiro_p"))
                    return "-"

                metrics = [
                    ("Mean (SD)", "mean_sd"),
                    ("95% CI (Mean)", "ci_95"),
                    ("Median [Q1, Q3]", "median_q1_q3"),
                    ("IQR", "iqr"),
                    ("Range (Min-Max)", "min_max"),
                    ("Normality (Shapiro p)", "shapiro"),
                ]

                for label, key in metrics:
                    row = table.add_row().cells
                    row[0].text = label
                    for i, g in enumerate(groups):
                        row[i + 1].text = _cell_for(key, stats_map.get(g, {}) or {})
                    row[-1].text = _cell_for(key, stats_map.get("overall", {}) or {})
            continue

        method = res.get("method")
        method_name = "Statistical Test"
        if isinstance(method, dict):
            method_name = method.get("name") or method.get("id") or method_name
        elif isinstance(method, str):
            method_name = method
        doc.add_paragraph(f"Метод: {method_name}")

        summary = doc.add_table(rows=0, cols=2)
        for k, v in [
            ("p-value", _fmt_p(res.get("p_value"))),
            ("stat", _fmt_num(res.get("stat_value", res.get("stats")), 3)),
            ("effect", f"{_txt(res.get('effect_size_name') or 'effect')} {_fmt_num(res.get('effect_size'), 2)}" if res.get("effect_size") is not None else "-"),
            ("power", _fmt_num(res.get("power"), 2)),
            ("BF10", _txt(res.get("bf10"))),
        ]:
            r = summary.add_row().cells
            r[0].text = str(k)
            r[1].text = str(v)

        warnings = res.get("warnings")
        if isinstance(warnings, list) and warnings:
            doc.add_paragraph("Предупреждения:")
            for w in warnings:
                doc.add_paragraph(str(w), style="List Bullet")

        roc = res.get("roc")
        if isinstance(roc, dict) and isinstance(roc.get("plot_data"), list) and roc.get("plot_data"):
            auc_val = roc.get("auc")
            if auc_val is not None:
                doc.add_paragraph(f"AUC: {_fmt_num(auc_val, 3)}")
            roc_png = _render_plot_png_bytes({"plot_data": roc.get("plot_data"), "plot_config": roc.get("plot_config")})
            if roc_png:
                bio = BytesIO(roc_png)
                doc.add_picture(bio, width=Inches(5.8))

        png_bytes = _render_plot_png_bytes(res)
        if png_bytes:
            bio = BytesIO(png_bytes)
            doc.add_picture(bio, width=Inches(5.8))

        conclusion = res.get("conclusion")
        if conclusion:
            doc.add_paragraph("Интерпретация:")
            doc.add_paragraph(str(conclusion))

    out = BytesIO()
    doc.save(out)
    return bytes(out.getvalue())

def generate_legacy_plot_image(plot_data: List[Dict[str, Any]], method_id: str) -> str:
    """
    Legacy: Generates a matplotlib plot based on plot_data and returns base64 string.
    """
    if not plot_data:
        return ""
    
    df = pd.DataFrame(plot_data)
    
    plt.figure(figsize=(8, 6))
    sns.set_theme(style="whitegrid")
    
    is_parametric = method_id in ["t_test_ind", "t_test_rel"]
    
    ax = sns.stripplot(
        data=df, 
        x="group", 
        y="value", 
        jitter=True, 
        alpha=0.6, 
        size=8,
        color="#0f172a"
    )
    
    sns.boxplot(
        data=df,
        x="group",
        y="value",
        showfliers=False,
        boxprops={'facecolor':'none', 'edgecolor':'grey'},
        width=0.4,
        ax=ax
    )

    plt.title(f"Distribution by Group ({method_id})")
    plt.xlabel("Group")
    plt.ylabel("Value")
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', dpi=100)
    plt.close()
    
    buf.seek(0)
    return base64.b64encode(buf.read()).decode('utf-8')

def render_report(
    analysis_result: AnalysisResult,
    target_col: str,
    group_col: str,
    dataset_name: str = "Dataset"
) -> str:
    """
    Legacy: Renders the HTML report using Jinja2 template.
    """
    
    plot_img = ""
    if analysis_result.plot_data:
        try:
            plot_img = generate_legacy_plot_image(analysis_result.plot_data, analysis_result.method.id)
        except Exception as e:
            logger.error(f"Error generating plot: {e}", exc_info=True)

    context = {
        "title": "Stat Analyzer Report",
        "dataset_name": dataset_name,
        "target_col": target_col,
        "group_col": group_col,
        "result": analysis_result,
        "image_base64": plot_img,
        "method_name": analysis_result.method.name,
        "method_desc": analysis_result.method.description, 
        "p_value_fmt": f"{analysis_result.p_value:.4f}" if analysis_result.p_value >= 0.001 else "< 0.001"
    }
    
    env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))
    template = env.get_template("report.html")
    return template.render(**context)

def render_protocol_report(run_data: Dict, dataset_name: str, style: Optional[str] = None,
                           dataset_id: Optional[str] = None) -> str:
    report = ProtocolReport(run_data, dataset_name, style=style or "apa7", dataset_id=dataset_id)
    return report.generate_html()

# ── Common PDF helpers ─────────────────────────────────────────────────────────

def _safe_text(value: Any) -> str:
    text = str(value) if value is not None else ""
    return text.encode("latin-1", errors="replace").decode("latin-1")

def _fmt_num(value: Any, digits: int = 3) -> str:
    try:
        if value is None:
            return "-"
        num = float(value)
        if not np.isfinite(num):
            return "-"
        return f"{num:.{digits}f}"
    except Exception:
        return "-"

def _fmt_p(value: Any) -> str:
    try:
        if value is None:
            return "-"
        p = float(value)
        if not np.isfinite(p):
            return "-"
        return "< 0.001" if p < 0.001 else f"{p:.4f}"
    except Exception:
        return "-"

def _pdf_bytes(pdf: FPDF) -> bytes:
    try:
        out = pdf.output()
    except TypeError:
        out = pdf.output(dest="S")
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    return str(out).encode("latin-1", errors="replace")

# ── Legacy single-analysis PDF ────────────────────────────────────────────────

def generate_pdf_report(results, variables, dataset_id):

    method = None
    if isinstance(results, dict):
        method = results.get("method")
    method_name = "Statistical Test"
    if isinstance(method, dict):
        method_name = method.get("name") or method.get("id") or method_name
    elif isinstance(method, str):
        method_name = method

    target = variables.get("target") if isinstance(variables, dict) else None
    group = variables.get("group") if isinstance(variables, dict) else None
    feature = variables.get("feature") if isinstance(variables, dict) else None

    pdf = FPDF(unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 9, _safe_text("Statistical Analysis Report"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, _safe_text(f"Dataset: {dataset_id}"), new_x="LMARGIN", new_y="NEXT")
    if target:
        pdf.cell(0, 6, _safe_text(f"Target: {target}"), new_x="LMARGIN", new_y="NEXT")
    if group:
        pdf.cell(0, 6, _safe_text(f"Group: {group}"), new_x="LMARGIN", new_y="NEXT")
    if feature and not group:
        pdf.cell(0, 6, _safe_text(f"Feature: {feature}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 7, _safe_text("Results"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, _safe_text(f"Method: {method_name}"), new_x="LMARGIN", new_y="NEXT")

    if isinstance(results, dict):
        pdf.cell(0, 6, _safe_text(f"P-value: {_fmt_p(results.get('p_value'))}"), new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 6, _safe_text(f"Statistic: {_fmt_num(results.get('stat_value'))}"), new_x="LMARGIN", new_y="NEXT")
        sig = results.get("significant")
        if isinstance(sig, bool):
            pdf.cell(0, 6, _safe_text(f"Significant: {'YES' if sig else 'NO'}"), new_x="LMARGIN", new_y="NEXT")

        effect_size = results.get("effect_size")
        effect_name = results.get("effect_size_name")
        if effect_size is not None:
            label = effect_name or "effect"
            pdf.cell(0, 6, _safe_text(f"Effect size: {label} {_fmt_num(effect_size, 2)}"), new_x="LMARGIN", new_y="NEXT")
        ci_lo = results.get("effect_size_ci_lower")
        ci_hi = results.get("effect_size_ci_upper")
        if ci_lo is not None and ci_hi is not None:
            pdf.cell(0, 6, _safe_text(f"Effect CI: [{_fmt_num(ci_lo, 2)}, {_fmt_num(ci_hi, 2)}]"), new_x="LMARGIN", new_y="NEXT")
        power = results.get("power")
        if power is not None:
            pdf.cell(0, 6, _safe_text(f"Power: {_fmt_num(power, 2)}"), new_x="LMARGIN", new_y="NEXT")
        bf10 = results.get("bf10")
        if bf10 is not None:
            pdf.cell(0, 6, _safe_text(f"BF10: {bf10}"), new_x="LMARGIN", new_y="NEXT")

        conclusion = results.get("conclusion")
        if conclusion:
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 7, _safe_text("Interpretation"), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, _safe_text(conclusion))

    return _pdf_bytes(pdf)


def generate_protocol_pdf_report(run_data: Dict[str, Any], dataset_name: str = "Dataset", style: Optional[str] = None) -> bytes:
    def _safe_text(value: Any) -> str:
        if value is None:
            return ""
        text = str(value)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    def _fmt_num(value: Any, digits: int = 3) -> str:
        try:
            if value is None:
                return "-"
            num = float(value)
            if not np.isfinite(num):
                return "-"
            return f"{num:.{digits}f}"
        except Exception:
            return "-"

    def _fmt_p(value: Any) -> str:
        try:
            if value is None:
                return "-"
            p = float(value)
            if not np.isfinite(p):
                return "-"
            return "< 0.001" if p < 0.001 else f"{p:.4f}"
        except Exception:
            return "-"

    def _pdf_bytes(pdf: FPDF) -> bytes:
        try:
            out = pdf.output()
        except TypeError:
            out = pdf.output(dest="S")
        if isinstance(out, (bytes, bytearray)):
            return bytes(out)
        return str(out).encode("latin-1", errors="replace")

    style_key = str(style or "apa7").strip().lower()
    is_ru = style_key in {"gost"}

    pdf = FPDF(unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 9, _safe_text("Отчёт по протоколу" if is_ru else "Protocol Analysis Report"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, _safe_text(("Набор данных" if is_ru else "Dataset") + f": {dataset_name}"), new_x="LMARGIN", new_y="NEXT")
    protocol_name = run_data.get("protocol_name") if isinstance(run_data, dict) else None
    if protocol_name:
        pdf.cell(0, 6, _safe_text(("Протокол" if is_ru else "Protocol") + f": {protocol_name}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    results = run_data.get("results", {}) if isinstance(run_data, dict) else {}
    for step_id, res in (results or {}).items():
        pdf.set_font("Helvetica", "B", 12)
        pdf.multi_cell(0, 6, _safe_text(("Шаг" if is_ru else "Step") + f": {step_id}"))

        if not isinstance(res, dict):
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, _safe_text("No structured result"))
            pdf.ln(2)
            continue

        pdf.set_font("Helvetica", "", 10)
        method = res.get("method")
        method_name = "Statistical Test"
        if isinstance(method, dict):
            method_name = method.get("name") or method.get("id") or method_name
        pdf.cell(0, 6, _safe_text(f"Method: {method_name}"), new_x="LMARGIN", new_y="NEXT")

        if "p_value" in res:
            pdf.cell(0, 6, _safe_text(f"P-value: {_fmt_p(res.get('p_value'))}"), new_x="LMARGIN", new_y="NEXT")
        if "stat_value" in res or "stats" in res:
            pdf.cell(0, 6, _safe_text(f"Statistic: {_fmt_num(res.get('stat_value', res.get('stats')))}"), new_x="LMARGIN", new_y="NEXT")

        effect_size = res.get("effect_size")
        if effect_size is not None:
            label = res.get("effect_size_name") or "effect"
            pdf.cell(0, 6, _safe_text(f"Effect size: {label} {_fmt_num(effect_size, 2)}"), new_x="LMARGIN", new_y="NEXT")
        ci_lo = res.get("effect_size_ci_lower")
        ci_hi = res.get("effect_size_ci_upper")
        if ci_lo is not None and ci_hi is not None:
            pdf.cell(0, 6, _safe_text(f"Effect CI: [{_fmt_num(ci_lo, 2)}, {_fmt_num(ci_hi, 2)}]"), new_x="LMARGIN", new_y="NEXT")
        if res.get("power") is not None:
            pdf.cell(0, 6, _safe_text(f"Power: {_fmt_num(res.get('power'), 2)}"), new_x="LMARGIN", new_y="NEXT")
        if res.get("bf10") is not None:
            pdf.cell(0, 6, _safe_text(f"BF10: {res.get('bf10')}"), new_x="LMARGIN", new_y="NEXT")

        conclusion = res.get("conclusion")
        if conclusion:
            pdf.ln(1)
            pdf.multi_cell(0, 5, _safe_text(f"Conclusion: {conclusion}"))
        pdf.ln(3)

    return _pdf_bytes(pdf)


# ── PDF с графиками (fpdf2 + PIL) ──────────────────────────────────────────────

def generate_protocol_pdf_with_plots(
    run_data: Dict[str, Any],
    dataset_name: str = "Dataset",
    style: Optional[str] = None,
    dataset_id: Optional[str] = None,
) -> bytes:
    """
    Генерирует PDF-отчёт с графиками.

    Стратегия:
    - Строим FPDF как обычно (текст + таблицы)
    - Дополнительно рисуем matplotlib-графики из plot_data каждого шага
      и вставляем их как PNG-изображения через PIL + FPDF.image()
    - Требует только fpdf2 + Pillow (matplotlib уже есть)
    """
    try:
        from PIL import Image as PILImage
        _pil_available = True
    except ImportError:
        _pil_available = False

    def _make_plot_png(res: Dict[str, Any]) -> Optional[bytes]:
        """Рисуем график из plot_data шага через matplotlib → PNG bytes."""
        try:
            plot_data = res.get("plot_data") or []
            if not plot_data:
                return None

            apply_publication_config()
            fig, ax = plt.subplots(figsize=(6.5, 3.8))

            groups: Dict[str, List] = {}
            for pt in plot_data:
                g = str(pt.get("group", ""))
                groups.setdefault(g, []).append(float(pt.get("value", 0)))

            colors = ["#2563eb", "#dc2626", "#16a34a", "#d97706", "#7c3aed", "#0891b2"]
            for i, (grp, vals) in enumerate(groups.items()):
                color = colors[i % len(colors)]
                ax.scatter([grp] * len(vals), vals, alpha=0.45, s=20, color=color, zorder=3)
                mean_val = np.mean(vals)
                ax.plot([grp], [mean_val], "D", color=color, ms=9, zorder=5)
                ci = 1.96 * np.std(vals) / max(np.sqrt(len(vals)), 1)
                ax.errorbar([grp], [mean_val], yerr=[[ci], [ci]], color=color, capsize=5, lw=1.5, zorder=4)

            # p-value annotation
            p_val = res.get("p_value")
            if p_val is not None:
                try:
                    p = float(p_val)
                    sig = "***" if p < 0.001 else ("**" if p < 0.01 else ("*" if p < 0.05 else "ns"))
                    ax.set_title(f"p = {_fmt_p(p_val)}  {sig}", fontsize=9, pad=4)
                except Exception:
                    pass

            ax.set_xlabel("")
            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
            plt.close(fig)
            buf.seek(0)
            return buf.read()
        except Exception as exc:
            logger.debug(f"_make_plot_png failed: {exc}")
            try:
                plt.close("all")
            except Exception:
                pass
            return None

    def _insert_png(pdf: FPDF, png_bytes: bytes) -> None:
        """Вставить PNG bytes в FPDF — fpdf2 принимает BytesIO напрямую."""
        try:
            buf = io.BytesIO(png_bytes)
            page_w = pdf.w - pdf.l_margin - pdf.r_margin
            img_w = min(page_w * 0.85, 160)
            pdf.image(buf, w=img_w)
            pdf.ln(4)
        except Exception as exc:
            logger.debug(f"_insert_png failed: {exc}")

    style_key = str(style or "apa7").strip().lower()
    is_ru = style_key in {"gost"}

    pdf = FPDF(unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Заголовок
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, _safe_text("Отчёт по протоколу" if is_ru else "Protocol Analysis Report"),
             new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 6, _safe_text(f"Dataset: {dataset_name}"), new_x="LMARGIN", new_y="NEXT")
    protocol_name = run_data.get("protocol_name") if isinstance(run_data, dict) else None
    if protocol_name:
        pdf.cell(0, 6, _safe_text(f"Protocol: {protocol_name}"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Methods section from transform_log
    if dataset_id:
        try:
            ws_dir = os.getenv("STATWIZARD_WORKSPACE_DIR", "workspace")
            log_path = os.path.join(ws_dir, "datasets", dataset_id, "processed", "transform_log.json")
            if os.path.exists(log_path):
                with open(log_path, "r", encoding="utf-8") as fh:
                    tlog = json.load(fh)
                if tlog:
                    methods_text = generate_methods_section(tlog, dataset_name=dataset_name)
                    pdf.ln(2)
                    pdf.set_font("Helvetica", "B", 13)
                    pdf.cell(0, 8, _safe_text("Methods"), new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", "", 10)
                    pdf.multi_cell(0, 5, _safe_text(methods_text))
                    pdf.ln(4)
        except Exception:
            pass

    results = run_data.get("results", {}) if isinstance(run_data, dict) else {}
    fig_counter = 0

    for step_id, res in (results or {}).items():
        if not isinstance(res, dict):
            continue
        if "error" in res:
            continue

        # Шапка шага
        pdf.set_font("Helvetica", "B", 12)
        pdf.multi_cell(0, 7, _safe_text(f"{'Шаг' if is_ru else 'Step'}: {step_id}"))
        pdf.set_font("Helvetica", "", 10)

        method = res.get("method")
        if isinstance(method, dict):
            method_name = method.get("name") or method.get("id") or "Statistical Test"
        elif isinstance(method, str):
            import re as _re
            _m = _re.search(r"name='([^']+)'", method)
            method_name = _m.group(1) if _m else method
        else:
            method_name = "Statistical Test"
        pdf.cell(0, 6, _safe_text(f"Method: {method_name}"), new_x="LMARGIN", new_y="NEXT")

        if "p_value" in res:
            pdf.cell(0, 6, _safe_text(f"P-value: {_fmt_p(res.get('p_value'))}"), new_x="LMARGIN", new_y="NEXT")
        if "stat_value" in res or "stats" in res:
            pdf.cell(0, 6, _safe_text(f"Statistic: {_fmt_num(res.get('stat_value', res.get('stats')))}"),
                     new_x="LMARGIN", new_y="NEXT")

        effect_size = res.get("effect_size")
        if effect_size is not None:
            label = res.get("effect_size_name") or "effect"
            pdf.cell(0, 6, _safe_text(f"Effect size ({label}): {_fmt_num(effect_size, 2)}"),
                     new_x="LMARGIN", new_y="NEXT")
        ci_lo, ci_hi = res.get("effect_size_ci_lower"), res.get("effect_size_ci_upper")
        if ci_lo is not None and ci_hi is not None:
            pdf.cell(0, 6, _safe_text(f"Effect CI: [{_fmt_num(ci_lo, 2)}, {_fmt_num(ci_hi, 2)}]"),
                     new_x="LMARGIN", new_y="NEXT")
        if res.get("power") is not None:
            pdf.cell(0, 6, _safe_text(f"Power: {_fmt_num(res['power'], 2)}"), new_x="LMARGIN", new_y="NEXT")
        if res.get("bf10") is not None:
            pdf.cell(0, 6, _safe_text(f"BF10: {res['bf10']}"), new_x="LMARGIN", new_y="NEXT")

        conclusion = res.get("conclusion")
        if conclusion:
            pdf.ln(1)
            pdf.multi_cell(0, 5, _safe_text(f"Conclusion: {conclusion}"))

        # График
        png_bytes = _make_plot_png(res)
        if png_bytes:
            fig_counter += 1
            pdf.ln(3)
            _insert_png(pdf, png_bytes)
            # Подпись рисунка
            pdf.set_font("Helvetica", "I", 8)
            pdf.multi_cell(0, 4, _safe_text(
                f"Figure {fig_counter}. {method_name}"
                + (f" (p = {_fmt_p(res.get('p_value'))})" if "p_value" in res else "")
            ))
            pdf.set_font("Helvetica", "", 10)

        pdf.ln(5)

    return _pdf_bytes(pdf)


# ── Phase 4.1: Auto-Methods from transform_log ────────────────────────────────

_ACTION_TEMPLATES: Dict[str, str] = {
    "split_column": (
        "Values in the column '{column}' were split by the delimiter '{sep}' "
        "into {mode} (trim={trim}). "
        "The dataset changed from {rows_before} rows / {cols_before} columns "
        "to {rows_after} rows / {cols_after} columns."
    ),
    "recode_values": (
        "Values in '{column}' were recoded according to a user-defined mapping table "
        "({n_mappings} substitution(s) defined). "
        "Missing mapping entries were left unchanged."
    ),
    "derive_column": (
        "A new variable '{new_col}' was derived using the formula: {formula}."
    ),
    "bin_variable": (
        "The continuous variable '{column}' was discretised into {n_bins} bins "
        "using the '{method}' method{labels_note}. "
        "The resulting variable was stored as '{new_col}'."
    ),
    "string_clean": (
        "Text values in '{column}' were cleaned: {ops}."
    ),
}


def _render_action_sentence(entry: Dict[str, Any]) -> str:
    """Convert a single transform_log entry into one English sentence."""
    action = entry.get("action", "")
    col = entry.get("column") or ""
    cfg = entry.get("config") or {}

    rows_before = entry.get("rows_before", "?")
    rows_after  = entry.get("rows_after",  "?")
    cols_before = entry.get("cols_before", "?")
    cols_after  = entry.get("cols_after",  "?")

    if action == "split_column":
        return _ACTION_TEMPLATES["split_column"].format(
            column=col,
            sep=cfg.get("separator", ","),
            mode=cfg.get("mode", "rows"),
            trim=cfg.get("trim", True),
            rows_before=rows_before, rows_after=rows_after,
            cols_before=cols_before, cols_after=cols_after,
        )
    elif action == "recode_values":
        mappings = cfg.get("mapping") or {}
        return _ACTION_TEMPLATES["recode_values"].format(
            column=col,
            n_mappings=len(mappings) if isinstance(mappings, dict) else "?",
        )
    elif action == "derive_column":
        return _ACTION_TEMPLATES["derive_column"].format(
            new_col=cfg.get("new_column") or col,
            formula=cfg.get("formula") or "—",
        )
    elif action == "bin_variable":
        labels = cfg.get("labels")
        labels_note = f" with labels {labels}" if labels else ""
        # For custom method, derive n_bins from the bins list
        bins_list = cfg.get("bins")
        if bins_list and isinstance(bins_list, list):
            n_bins = len(bins_list) - 1
        else:
            n_bins = cfg.get("n_bins", "?")
        return _ACTION_TEMPLATES["bin_variable"].format(
            column=col,
            n_bins=n_bins,
            method=cfg.get("method", "equal_width"),
            labels_note=labels_note,
            new_col=cfg.get("new_column") or f"{col}_bin",
        )
    elif action == "string_clean":
        operations = cfg.get("operations") or []
        ops_list = []
        if "trim" in operations:      ops_list.append("leading/trailing whitespace removed")
        if "lowercase" in operations:  ops_list.append("converted to lowercase")
        if "uppercase" in operations:  ops_list.append("converted to uppercase")
        if "replace" in operations:
            replace_from = cfg.get("replace_from", "")
            replace_to = cfg.get("replace_to", "")
            if replace_from:
                ops_list.append(f"'{replace_from}' replaced with '{replace_to}'")
            else:
                ops_list.append("substring replacement applied")
        ops = "; ".join(ops_list) if ops_list else "no specific operations recorded"
        return _ACTION_TEMPLATES["string_clean"].format(column=col, ops=ops)
    else:
        return f"Action '{action}' was applied to column '{col}'."


def generate_methods_section(
    transform_log: List[Dict[str, Any]],
    dataset_name: str = "the dataset",
) -> str:
    """
    Generate a publication-ready Methods paragraph from a transform_log list.

    Parameters
    ----------
    transform_log : list of dicts loaded from transform_log.json
    dataset_name  : display name for the dataset

    Returns
    -------
    str — plain-text paragraph suitable for embedding in a Methods section.
    """
    if not transform_log:
        return (
            f"Data wrangling: no preprocessing transformations were recorded "
            f"for {dataset_name}."
        )

    sentences = []
    for entry in transform_log:
        try:
            sentences.append(_render_action_sentence(entry))
        except Exception as exc:
            logger.debug(f"generate_methods_section: failed to render entry: {exc}")

    if not sentences:
        return (
            f"Data were loaded from {dataset_name}. "
            "No structured preprocessing steps were recorded."
        )

    preamble = (
        f"Data preparation was performed in the following sequence "
        f"({len(sentences)} step(s) recorded)."
    )
    body = " ".join(sentences)
    return f"{preamble} {body}"
