"""
reporting_docx.py — DOCX report generation extracted from reporting.py.

Contains generate_protocol_docx_report() which creates Word documents
from protocol analysis results.
"""
from __future__ import annotations

import io
import json
import logging
import os
import re
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd

from app.modules.analysis_result_v2 import normalize_run_data_results
from app.modules.reporting_contracts import (
    build_report_integrity_context,
    filter_step_pairs_for_report,
)
from app.modules.reporting import (
    _render_plot_png_bytes,
    _dedupe_step_payloads,
    _safe_float,
    _extract_protocol_findings,
    _build_discussion_conclusion,
    _method_label_from_type,
    _method_label_from_id,
    _normalize_correction_key,
    _format_correction_label,
    _format_boolean_label,
    _format_validation_status_label,
    _extract_bootstrap_policy,
    _format_bootstrap_policy_text,
    _extract_multiplicity_policy,
    _format_multiplicity_policy_text,
    _build_hypothesis_discovery_context,
    _protocol_validation_provenance_rows,
    _protocol_validation_section_context,
    _bootstrap_trace_lines,
    _is_placeholder_interpretation,
    _extract_step_context_value,
    _build_step_display,
    _generate_fallback_interpretation,
    _step_scope_summary,
    _batch_method_selection_rationale,
    _extract_report_methods,
    _build_report_limitations,
    _method_selection_rationale_ru,
    _normalize_report_density,
    _coerce_alpha,
    _batch_target_label,
    _collect_batch_inferential_rows,
    _build_pairwise_comparison_rows,
    _resolve_dataset_dir_path,
    _build_provenance_file_context,
    _parse_accent_rgb,
    _parse_accent_css,
    _interpret_bf10_ru,
)

logger = logging.getLogger(__name__)

def generate_protocol_docx_report(
    run_data: Dict[str, Any],
    dataset_name: str = "Dataset",
    style: Optional[str] = None,
    options: Optional[Dict[str, Any]] = None,
) -> bytes:
    from io import BytesIO
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    import re
    run_data = normalize_run_data_results(run_data if isinstance(run_data, dict) else {})

    def _fmt_p(value: Any) -> str:
        try:
            if value is None:
                return "-"
            p = float(value)
            if not np.isfinite(p):
                return "-"
            return "< 0.001" if p < 0.001 else f"{p:.4f}"
        except Exception:
            return "-"

    def _txt(value: Any) -> str:
        return "-" if value is None else str(value)

    def _bookmark_name(step_id: str) -> str:
        cleaned = re.sub(r"[^0-9A-Za-z_]+", "_", str(step_id or "").strip())
        if not cleaned:
            cleaned = "step"
        if not cleaned[0].isalpha():
            cleaned = f"s_{cleaned}"
        # MS Word bookmark names are limited and should avoid punctuation.
        return f"step_{cleaned[:32]}"

    bookmark_seq = 1

    def _add_bookmark(paragraph: Any, bookmark_name: str) -> None:
        nonlocal bookmark_seq
        if not bookmark_name:
            return
        try:
            p = paragraph._p
            start = OxmlElement("w:bookmarkStart")
            start.set(qn("w:id"), str(bookmark_seq))
            start.set(qn("w:name"), bookmark_name)
            end = OxmlElement("w:bookmarkEnd")
            end.set(qn("w:id"), str(bookmark_seq))
            p.insert(0, start)
            p.append(end)
            bookmark_seq += 1
        except Exception:
            return

    def _add_internal_hyperlink(paragraph: Any, anchor: str, text: str) -> None:
        if not anchor:
            paragraph.add_run(str(text or "-"))
            return
        try:
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("w:anchor"), str(anchor))

            run = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            rstyle = OxmlElement("w:rStyle")
            rstyle.set(qn("w:val"), "Hyperlink")
            rpr.append(rstyle)
            run.append(rpr)

            text_el = OxmlElement("w:t")
            text_el.text = str(text or "-")
            run.append(text_el)
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
        except Exception:
            paragraph.add_run(str(text or "-"))

    def _fmt_num(value: Any, digits: int = 3) -> str:
        try:
            if value is None:
                return "-"
            num = float(value)
            if not np.isfinite(num):
                return "-"
            return f"{num:.{digits}f}"
        except Exception:
            return "-"

    def _txt(value: Any) -> str:
        return "-" if value is None else str(value)

    style_key = str(style or "gost").strip().lower()
    is_ru = style_key in {"gost"}

    density = _normalize_report_density((options or {}).get("density"))

    font_name = "Calibri"
    base_pt = 11
    if style_key in {"gost", "apa7", "editorial"}:
        font_name = "Times New Roman"
        base_pt = 12
    if style_key == "gost":
        base_pt = 14
    if style_key == "brutal":
        font_name = "Courier New"
        base_pt = 10

    if density == "compact":
        base_pt = max(9, base_pt - 1)
    elif density == "spacious":
        base_pt = min(16, base_pt + 1)

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(base_pt)
    normal.paragraph_format.space_after = Pt(6 if density != "compact" else 3)

    for style_name, size in [("Title", base_pt + 8), ("Heading 1", base_pt + 4), ("Heading 2", base_pt + 2), ("Heading 3", base_pt + 1)]:
        try:
            st = doc.styles[style_name]
            st.font.name = font_name
            st.font.size = Pt(size)
        except Exception:
            pass

    doc.add_heading("Результаты статистического анализа" if is_ru else "Statistical Analysis Results", level=0)
    doc.add_paragraph(("Набор данных" if is_ru else "Dataset") + f": {dataset_name}")
    protocol_name = run_data.get("protocol_name") if isinstance(run_data, dict) else None
    if protocol_name:
        doc.add_paragraph(("Протокол" if is_ru else "Protocol") + f": {protocol_name}")
    doc.add_paragraph(("Дата" if is_ru else "Date") + f": {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")

    dataset_id = run_data.get("dataset_id") if isinstance(run_data, dict) else None
    run_id = run_data.get("run_id") if isinstance(run_data, dict) else None
    ds_dir = _resolve_dataset_dir_path(dataset_id) if dataset_id else None
    step_meta_map = run_data.get("step_meta") if isinstance(run_data.get("step_meta"), dict) else {}
    results_map = run_data.get("results") if isinstance(run_data.get("results"), dict) else {}
    all_step_ids: List[str] = []
    for sid in step_meta_map.keys():
        if isinstance(sid, str) and sid and sid not in all_step_ids:
            all_step_ids.append(sid)
    for sid in results_map.keys():
        if isinstance(sid, str) and sid and sid not in all_step_ids:
            all_step_ids.append(sid)
    step_bookmarks: Dict[str, str] = {sid: _bookmark_name(sid) for sid in all_step_ids}

    source_meta: Dict[str, Any] = {}
    source_files: List[str] = []
    if ds_dir and os.path.isdir(ds_dir):
        source_dir = os.path.join(ds_dir, "source")
        source_meta_path = os.path.join(source_dir, "meta.json")
        if os.path.exists(source_meta_path):
            try:
                with open(source_meta_path, "r", encoding="utf-8") as f:
                    loaded_meta = json.load(f)
                if isinstance(loaded_meta, dict):
                    source_meta = loaded_meta
            except Exception:
                source_meta = {}
        if os.path.isdir(source_dir):
            try:
                source_files = [
                    str(name)
                    for name in sorted(os.listdir(source_dir))
                    if str(name) not in {"meta.json", ".", ".."} and os.path.isfile(os.path.join(source_dir, str(name)))
                ]
            except Exception:
                source_files = []

    reproducibility = run_data.get("reproducibility") if isinstance(run_data.get("reproducibility"), dict) else {}
    analysis_dataset = run_data.get("analysis_dataset") if isinstance(run_data.get("analysis_dataset"), dict) else {}
    if not analysis_dataset and isinstance(reproducibility.get("analysis_dataset"), dict):
        analysis_dataset = reproducibility.get("analysis_dataset")
    analysis_set = run_data.get("analysis_set") if isinstance(run_data.get("analysis_set"), dict) else {}
    integrity_ctx_pdf = build_report_integrity_context(run_data)
    verification_ctx_pdf = (
        integrity_ctx_pdf.get("verification")
        if isinstance(integrity_ctx_pdf.get("verification"), dict)
        else {}
    )
    verification_status_pdf = str(verification_ctx_pdf.get("status") or "missing").strip().lower()
    if verification_status_pdf in {"passed", "ok"}:
        verification_label_pdf = "пройдено" if is_ru else "passed"
    elif verification_status_pdf in {"failed", "error", "blocked"}:
        verification_label_pdf = "ошибка" if is_ru else "failed"
    else:
        verification_label_pdf = "отсутствует" if is_ru else "missing"
    failed_steps_pdf = verification_ctx_pdf.get("failed_steps")
    excluded_steps_pdf = (
        len([str(x) for x in failed_steps_pdf if isinstance(x, str) and str(x).strip()])
        if isinstance(failed_steps_pdf, list)
        else 0
    )

    source_name = source_meta.get("original_filename") or source_meta.get("filename")
    source_rel = "-"
    if dataset_id and source_files:
        source_rel = os.path.join("workspace", "datasets", str(dataset_id), "source", source_files[0])
    elif source_name:
        source_rel = str(source_name)

    doc.add_heading("Воспроизводимость и provenance" if is_ru else "Reproducibility and provenance", level=1)
    doc.add_paragraph(("ID набора данных" if is_ru else "Dataset ID") + f": {_txt(dataset_id)}")
    doc.add_paragraph(("ID запуска" if is_ru else "Run ID") + f": {_txt(run_id)}")
    doc.add_paragraph(("Исходный файл" if is_ru else "Source file") + f": {_txt(source_name)}")
    doc.add_paragraph(("Путь источника" if is_ru else "Source path") + f": {_txt(source_rel)}")
    doc.add_paragraph(
        ("Параметры импорта" if is_ru else "Import settings")
        + f": sheet={_txt(source_meta.get('sheet_name'))}; header_row={_txt(source_meta.get('header_row'))}"
    )
    doc.add_paragraph(
        ("Рабочая выборка (analysis_dataset)" if is_ru else "Analysis dataset")
        + f": rows={_txt(analysis_dataset.get('rows'))}, cols={_txt(analysis_dataset.get('columns'))}, "
        + f"xlsx={_txt(analysis_dataset.get('xlsx'))}, parquet={_txt(analysis_dataset.get('parquet'))}"
    )
    doc.add_paragraph(
        ("Замороженная выборка (analysis_set)" if is_ru else "Frozen analysis set")
        + f": id={_txt(analysis_set.get('analysis_set_id'))}; n_selected={_txt(analysis_set.get('n_selected'))}; "
        + f"mode={_txt(analysis_set.get('mode'))}; enforce={_txt(analysis_set.get('enforce'))}"
    )
    doc.add_paragraph(
        ("Воспроизводимость" if is_ru else "Reproducibility")
        + f": ready={_txt(reproducibility.get('ready'))}; script={_txt(reproducibility.get('script'))}; "
        + f"payload={_txt(reproducibility.get('payload'))}; manifest={_txt(reproducibility.get('manifest'))}"
    )
    doc.add_paragraph(
        ("Артефакт bootstrap-трассировки" if is_ru else "Bootstrap trace artifact")
        + f": {_txt(reproducibility.get('bootstrap_trace'))}"
    )
    doc.add_paragraph(
        ("Артефакт гипотез" if is_ru else "Hypothesis discovery artifact")
        + f": {_txt(reproducibility.get('hypothesis_discovery'))}"
    )
    for key, value in _protocol_validation_provenance_rows(
        run_data if isinstance(run_data, dict) else {},
        is_ru=is_ru,
        dataset_id=dataset_id,
        run_id=run_id,
    ):
        doc.add_paragraph(f"{_txt(key)}: {_txt(value)}")
    if dataset_id and run_id:
        doc.add_paragraph(
            ("Путь артефактов" if is_ru else "Artifacts path")
            + ": "
            + os.path.join("workspace", "datasets", str(dataset_id), "analysis", str(run_id), "artifacts")
        )

    protocol_validation_ctx_docx = _protocol_validation_section_context(
        run_data if isinstance(run_data, dict) else {},
        is_ru=is_ru,
    )
    if protocol_validation_ctx_docx.get("present"):
        doc.add_heading("Валидация протокола" if is_ru else "Protocol Validation", level=1)
        for key, value in (protocol_validation_ctx_docx.get("summary_rows") or []):
            doc.add_paragraph(f"{_txt(key)}: {_txt(value)}")

        global_errors_docx = protocol_validation_ctx_docx.get("global_errors")
        if isinstance(global_errors_docx, list) and global_errors_docx:
            doc.add_paragraph("Глобальные ошибки валидации:" if is_ru else "Global validation errors:")
            for msg in global_errors_docx:
                doc.add_paragraph("• " + _txt(msg))

        issues_docx = protocol_validation_ctx_docx.get("issues")
        if isinstance(issues_docx, list) and issues_docx:
            doc.add_paragraph("Проблемные шаги:" if is_ru else "Validation findings by step:")
            for row in issues_docx:
                if not isinstance(row, dict):
                    continue
                step_id_ref = str(row.get("step_id") or "").strip()
                p = doc.add_paragraph("• ")
                if step_id_ref:
                    _add_internal_hyperlink(
                        p,
                        step_bookmarks.get(step_id_ref, ""),
                        step_id_ref,
                    )
                else:
                    p.add_run("-")
                p.add_run(
                    " | "
                    + _txt(row.get("method"))
                    + " | "
                    + _txt(row.get("status"))
                    + " | "
                    + _txt(row.get("issues"))
                )
        else:
            doc.add_paragraph(
                "Критичных замечаний по шагам не найдено."
                if is_ru
                else "No critical step findings were detected."
            )

    if step_meta_map or results_map:
        doc.add_paragraph(
            "Карта шагов протокола (выборка/подгруппы/сравнения):"
            if is_ru
            else "Protocol map (cohort/subgroups/comparisons):"
        )
        for sid in all_step_ids[:120]:
            meta = step_meta_map.get(sid) if isinstance(step_meta_map.get(sid), dict) else {}
            res = results_map.get(sid) if isinstance(results_map.get(sid), dict) else {}
            cfg = meta.get("config") if isinstance(meta.get("config"), dict) else {}
            method_hint = cfg.get("method_id") or meta.get("method") or res.get("method_id") or res.get("type")
            method_label = _method_label_from_id(method_hint, is_ru) or _method_label_from_type(res.get("type"), is_ru) or _txt(method_hint)
            scope = _step_scope_summary(meta, res, is_ru)
            rationale = None
            if str(res.get("type") or "").strip().lower() in {"batch_analysis", "timepoint_batch_analysis"}:
                rationale = _batch_method_selection_rationale(meta, res, is_ru)
            elif is_ru:
                rationale = _method_selection_rationale_ru(res)
            line = f"{sid}: {method_label}; {scope}"
            if isinstance(rationale, str) and rationale.strip():
                line += f"; {rationale}"
            p = doc.add_paragraph("• ")
            if sid:
                _add_internal_hyperlink(
                    p,
                    step_bookmarks.get(sid, ""),
                    sid,
                )
                p.add_run(": " + line.split(": ", 1)[1] if ": " in line else "")
            else:
                p.add_run(line)

    study_design = None
    design_warning = None
    if dataset_id:
        try:
            study_path = os.path.join(ds_dir, "processed", "study_design.json") if ds_dir else ""
            if study_path and os.path.exists(study_path):
                with open(study_path, "r", encoding="utf-8") as f:
                    study_design = json.load(f)
            else:
                design_warning = (
                    "Предупреждение: файл study_design.json не найден. Раздел Design неполный."
                    if is_ru
                    else "Warning: study_design.json not found. Design section is incomplete."
                )
        except Exception:
            study_design = None
            design_warning = (
                "Предупреждение: не удалось прочитать study_design.json. Раздел Design неполный."
                if is_ru
                else "Warning: failed to read study_design.json. Design section is incomplete."
            )
    else:
        design_warning = (
            "Предупреждение: dataset_id отсутствует. Раздел Design неполный."
            if is_ru
            else "Warning: dataset_id is missing. Design section is incomplete."
        )

    doc.add_heading("Дизайн исследования" if is_ru else "Study Design", level=1)
    design_rendered = False
    if isinstance(study_design, dict):
        design = study_design.get("design") if isinstance(study_design.get("design"), dict) else {}
        policy = study_design.get("analysis_policy") if isinstance(study_design.get("analysis_policy"), dict) else {}

        if design:
            design_rendered = True
            doc.add_paragraph(("Тип дизайна" if is_ru else "Design type") + f": {_txt(design.get('design_type'))}")
            if design.get("group_column"):
                doc.add_paragraph(("Группировка" if is_ru else "Group") + f": {_txt(design.get('group_column'))}")
            if design.get("time_column"):
                doc.add_paragraph(("Время/визит" if is_ru else "Time") + f": {_txt(design.get('time_column'))}")
            if design.get("subject_column"):
                doc.add_paragraph(("ID субъекта" if is_ru else "Subject ID") + f": {_txt(design.get('subject_column'))}")

            outcomes = design.get("outcomes")
            if isinstance(outcomes, list) and outcomes:
                doc.add_paragraph(("Числовые исходы" if is_ru else "Numeric outcomes") + f": {', '.join([_txt(o) for o in outcomes[:20]])}")

            cat_outcomes = design.get("categorical_outcomes")
            if isinstance(cat_outcomes, list) and cat_outcomes:
                doc.add_paragraph(("Категориальные исходы" if is_ru else "Categorical outcomes") + f": {', '.join([_txt(o) for o in cat_outcomes[:15]])}")

            endpoint_groups = design.get("endpoint_groups") if isinstance(design.get("endpoint_groups"), list) else []
            if endpoint_groups:
                doc.add_paragraph("Endpoint-группы по визитам:" if is_ru else "Endpoint groups by visit:")
                for item in endpoint_groups[:12]:
                    if not isinstance(item, dict):
                        continue
                    ep = _txt(item.get("endpoint") or "endpoint")
                    tps = item.get("timepoints") if isinstance(item.get("timepoints"), list) else []
                    tp_line = ", ".join([_txt(t) for t in tps]) if tps else "-"
                    doc.add_paragraph(f"• {ep}: {tp_line}")

        if policy:
            design_rendered = True
            alpha = policy.get("alpha")
            runtime_multiplicity_policy = _extract_multiplicity_policy(
                run_data if isinstance(run_data, dict) else {},
                (
                    run_data.get("protocol_validation")
                    if isinstance(run_data, dict) and isinstance(run_data.get("protocol_validation"), dict)
                    else None
                ),
            )
            multiplicity = (
                policy.get("multiplicity_correction")
                or runtime_multiplicity_policy.get("correction")
                or runtime_multiplicity_policy.get("multiplicity_correction")
            )
            post_hoc = policy.get("post_hoc")
            post_hoc_correction = (
                policy.get("post_hoc_correction")
                or runtime_multiplicity_policy.get("post_hoc_correction")
            )
            runtime_bootstrap_policy = _extract_bootstrap_policy(
                run_data if isinstance(run_data, dict) else {},
                (
                    run_data.get("protocol_validation")
                    if isinstance(run_data, dict) and isinstance(run_data.get("protocol_validation"), dict)
                    else None
                ),
            )
            bootstrap_enabled = policy.get("bootstrap_ci")
            if bootstrap_enabled is None:
                bootstrap_enabled = runtime_bootstrap_policy.get("enabled")
            bootstrap_samples = policy.get("bootstrap_samples")
            if bootstrap_samples is None:
                bootstrap_samples = runtime_bootstrap_policy.get("samples")
            doc.add_paragraph(("Политика анализа" if is_ru else "Analysis policy") + ":")
            if alpha is not None:
                doc.add_paragraph(f"• α={_txt(alpha)}")
            if multiplicity:
                corr_label = _format_correction_label(multiplicity, is_ru) or _txt(multiplicity)
                doc.add_paragraph(
                    ("• Поправка множественных сравнений: " if is_ru else "• Multiple-testing correction: ")
                    + f"{_txt(corr_label)}"
                )
            if post_hoc:
                doc.add_paragraph(f"• post-hoc={_txt(post_hoc)}")
            if post_hoc_correction:
                ph_corr_label = _format_correction_label(post_hoc_correction, is_ru) or _txt(post_hoc_correction)
                doc.add_paragraph(
                    ("• Post-hoc поправка: " if is_ru else "• Post-hoc correction: ")
                    + f"{_txt(ph_corr_label)}"
                )
            multiplicity_applied = runtime_multiplicity_policy.get("n_applied_steps")
            if multiplicity_applied is not None:
                doc.add_paragraph(
                    ("• Применено шагов (multiplicity): " if is_ru else "• Multiplicity applied steps: ")
                    + _txt(multiplicity_applied)
                )
            if bootstrap_enabled is not None:
                doc.add_paragraph(
                    ("• Bootstrap CI: " if is_ru else "• Bootstrap CI: ")
                    + _txt(_format_boolean_label(bool(bootstrap_enabled), is_ru))
                )
            if bootstrap_samples is not None:
                doc.add_paragraph(("• Bootstrap samples: " if is_ru else "• Bootstrap samples: ") + _txt(bootstrap_samples))

    if not design_rendered:
        warning_line = design_warning or (
            "Предупреждение: раздел Design не заполнен."
            if is_ru
            else "Warning: Design section is empty."
        )
        doc.add_paragraph(warning_line)

    def _add_data_quality_section(doc_obj: Any) -> None:
        """Add Data Quality section to DOCX report."""
        scan_report: Dict[str, Any] = {}
        cleaning_log: Dict[str, Any] = {}
        if ds_dir:
            scan_path = os.path.join(ds_dir, "processed", "scan_report.json")
            if os.path.exists(scan_path):
                try:
                    with open(scan_path, "r", encoding="utf-8") as f:
                        loaded = json.load(f)
                    if isinstance(loaded, dict):
                        scan_report = loaded
                except Exception:
                    scan_report = {}
            cl_path = os.path.join(ds_dir, "processed", "cleaning_log.json")
            if not os.path.exists(cl_path):
                cl_path = os.path.join(ds_dir, "processed", "dataset_cleaning_log.json")
            if os.path.exists(cl_path):
                try:
                    with open(cl_path, "r", encoding="utf-8") as f:
                        loaded = json.load(f)
                    if isinstance(loaded, dict):
                        cleaning_log = loaded
                except Exception:
                    cleaning_log = {}

        protocol_run = run_data.get("protocol_plan") if isinstance(run_data, dict) else None
        if not isinstance(protocol_run, dict):
            protocol_run = {}
        col_report = protocol_run.get("column_selection_report", {})
        if not isinstance(col_report, dict):
            col_report = run_data.get("column_selection_report", {}) if isinstance(run_data, dict) else {}
        if not isinstance(col_report, dict):
            col_report = {}

        doc_obj.add_heading("Качество данных и очистка" if is_ru else "Data Quality and Cleaning", level=1)

        rows_orig = cleaning_log.get("rows_original", "?")
        rows_final = cleaning_log.get("rows_final", "?")
        cols_orig = cleaning_log.get("cols_original", "?")
        cols_final = cleaning_log.get("cols_final", "?")
        analyzed_total = col_report.get("analyzed_total", "?")
        excluded_total = col_report.get("excluded_total", "?")
        summary = (
            f"{'Исходный датасет' if is_ru else 'Input dataset'}: {rows_orig} × {cols_orig}. "
            f"{'После очистки' if is_ru else 'After cleaning'}: {rows_final} × {cols_final}. "
            f"{'Включено в анализ' if is_ru else 'Included in analysis'}: {analyzed_total}, "
            f"{'исключено' if is_ru else 'excluded'}: {excluded_total}."
        )
        doc_obj.add_paragraph(summary)

        logic = col_report.get("selection_logic", "")
        if logic:
            p = doc_obj.add_paragraph()
            run = p.add_run(str(logic))
            run.italic = True

        cl_steps = cleaning_log.get("steps", [])
        if isinstance(cl_steps, list) and cl_steps:
            doc_obj.add_heading("Лог очистки" if is_ru else "Cleaning Log", level=2)
            table = doc_obj.add_table(rows=1, cols=2)
            table.style = "Light List Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Операция" if is_ru else "Operation"
            hdr[1].text = "Детали" if is_ru else "Details"
            for step in cl_steps:
                if not isinstance(step, dict):
                    continue
                action = str(step.get("action", "-"))
                details = step.get("details", {})
                if not isinstance(details, dict):
                    details = {}
                detail_str = ""
                if action == "remove_duplicates":
                    detail_str = f"Удалено дубликатов: {details.get('rows_removed', 0)}" if is_ru else f"Duplicates removed: {details.get('rows_removed', 0)}"
                elif action == "drop_high_missing_columns":
                    dropped = details.get("dropped", [])
                    threshold = details.get("threshold", 0.7)
                    try:
                        threshold = float(threshold)
                    except Exception:
                        threshold = 0.7
                    detail_str = (
                        f"Удалено столбцов (>{threshold:.0%} пропусков): {len(dropped)}"
                        if is_ru
                        else f"Dropped columns (>{threshold:.0%} missing): {len(dropped)}"
                    )
                elif action == "impute_missing_numeric":
                    n = len(details.get("columns", {}))
                    total = details.get("total_filled", 0)
                    strategy = details.get("strategy", "median")
                    detail_str = (
                        f"Импутация ({strategy}): {n} столбцов, {total} значений"
                        if is_ru
                        else f"Imputation ({strategy}): {n} columns, {total} values"
                    )
                elif action == "outlier_detection":
                    total = details.get("total_outliers", 0)
                    policy = details.get("policy", "flag")
                    n_cols = len(details.get("columns", {}))
                    detail_str = (
                        f"Выбросы (IQR, {policy}): {total} в {n_cols} столбцах"
                        if is_ru
                        else f"Outliers (IQR, {policy}): {total} in {n_cols} columns"
                    )
                else:
                    detail_str = action
                row = table.add_row().cells
                row[0].text = action
                row[1].text = detail_str

        excluded = col_report.get("excluded", {})
        if isinstance(excluded, dict):
            categories = [
                ("id_like", "ID-подобные" if is_ru else "ID-like"),
                ("group_time_subject", "Группировка / время / субъект" if is_ru else "Group / time / subject"),
                ("high_missing", "Высокая доля пропусков (>70%)" if is_ru else "High missing (>70%)"),
                ("constant", "Константные" if is_ru else "Constant"),
                ("mixed_types", "Смешанные типы" if is_ru else "Mixed types"),
                ("not_in_analysis", "Не включены (прочие)" if is_ru else "Not included (other)"),
            ]
            has_any = any(excluded.get(k) for k, _ in categories)
            if has_any:
                doc_obj.add_heading("Исключённые столбцы" if is_ru else "Excluded Columns", level=2)
                table = doc_obj.add_table(rows=1, cols=3)
                table.style = "Light List Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text = "Причина" if is_ru else "Reason"
                hdr[1].text = "Кол-во" if is_ru else "Count"
                hdr[2].text = "Столбцы" if is_ru else "Columns"
                for key, label in categories:
                    cols_list = excluded.get(key, [])
                    if not cols_list:
                        continue
                    if not isinstance(cols_list, list):
                        cols_list = [str(cols_list)]
                    preview = ", ".join(str(c) for c in cols_list[:8])
                    if len(cols_list) > 8:
                        preview += f" ... (+{len(cols_list) - 8})"
                    row = table.add_row().cells
                    row[0].text = label
                    row[1].text = str(len(cols_list))
                    row[2].text = preview

        recs = col_report.get("recommendations", [])
        if isinstance(recs, list) and recs:
            doc_obj.add_heading("Рекомендации по доработке первички" if is_ru else "Recommendations", level=2)
            for r in recs:
                doc_obj.add_paragraph(str(r), style="List Bullet")

    _add_data_quality_section(doc)

    methods_summary = _extract_report_methods(run_data if isinstance(run_data, dict) else {}, is_ru=is_ru)
    method_rows = methods_summary.get("rows") if isinstance(methods_summary.get("rows"), list) else []
    missing_method_steps = methods_summary.get("missing_inferential_steps") if isinstance(methods_summary.get("missing_inferential_steps"), list) else []

    doc.add_heading("Методы" if is_ru else "Methods", level=1)
    if method_rows:
        methods_table = doc.add_table(rows=1, cols=4)
        m_hdr = methods_table.rows[0].cells
        m_hdr[0].text = "Метод" if is_ru else "Method"
        m_hdr[1].text = "Шагов" if is_ru else "Steps"
        m_hdr[2].text = "ID шагов" if is_ru else "Step IDs"
        m_hdr[3].text = "Показатели" if is_ru else "Targets"
        for row in method_rows:
            rr = methods_table.add_row().cells
            rr[0].text = _txt(row.get("method"))
            rr[1].text = _txt(row.get("count"))
            rr[2].text = ", ".join([_txt(x) for x in (row.get("steps") or [])[:8]]) or "-"
            rr[3].text = ", ".join([_txt(x) for x in (row.get("targets") or [])[:10]]) or "-"
    else:
        doc.add_paragraph(
            "Методы не определены: в результатах нет исполнимых аналитических шагов."
            if is_ru
            else "Methods are unavailable: no analyzable execution steps were found."
        )

    if missing_method_steps:
        preview = ", ".join([str(x) for x in missing_method_steps[:6]])
        if len(missing_method_steps) > 6:
            preview += ", ..."
        doc.add_paragraph(
            ("Предупреждение: отсутствует metadata метода для шагов: " if is_ru else "Warning: method metadata is missing for steps: ")
            + preview
        )

    hypothesis_ctx_docx = _build_hypothesis_discovery_context(
        run_data if isinstance(run_data, dict) else {},
        is_ru=is_ru,
    )
    if hypothesis_ctx_docx.get("present"):
        hyp_rows_docx = hypothesis_ctx_docx.get("rows") if isinstance(hypothesis_ctx_docx.get("rows"), list) else []
        doc.add_heading(
            "Гипотезы и проверяемые утверждения" if is_ru else "Hypothesis discovery and traceability",
            level=1,
        )
        doc.add_paragraph(
            (
                ("Режим" if is_ru else "Mode")
                + f": {_txt(hypothesis_ctx_docx.get('analysis_mode'))}; "
                + ("Дизайн" if is_ru else "Design")
                + f": {_txt(hypothesis_ctx_docx.get('design_type'))}; "
                + ("Всего" if is_ru else "Total")
                + f": {_txt(hypothesis_ctx_docx.get('count'))}; "
                + ("Покрыто шагами" if is_ru else "Covered by steps")
                + f": {_txt(hypothesis_ctx_docx.get('covered'))}"
            )
        )
        doc.add_paragraph(
            (
                ("Подтверждено" if is_ru else "Supported")
                + f": {_txt(hypothesis_ctx_docx.get('supported'))}; "
                + ("Не подтверждено" if is_ru else "Not supported")
                + f": {_txt(hypothesis_ctx_docx.get('not_supported'))}; "
                + ("Не оценено" if is_ru else "Not evaluated")
                + f": {_txt(hypothesis_ctx_docx.get('not_evaluated'))}"
            )
        )
        for row in hyp_rows_docx[:10]:
            if not isinstance(row, dict):
                continue
            matched_steps = ", ".join([_txt(v) for v in (row.get("matched_steps") or [])]) or "-"
            doc.add_paragraph("• " + _txt(row.get("title")))
            doc.add_paragraph("  H0: " + _txt(row.get("h0")))
            doc.add_paragraph("  H1: " + _txt(row.get("h1")))
            doc.add_paragraph(
                ("  Метод: " if is_ru else "  Method: ")
                + _txt(row.get("suggested_method"))
                + ("; шаги: " if is_ru else "; steps: ")
                + matched_steps
            )
            doc.add_paragraph(
                ("  Вердикт: " if is_ru else "  Verdict: ")
                + _txt(row.get("verdict_label"))
                + ("; доказательства: " if is_ru else "; evidence: ")
                + _txt(row.get("evidence"))
            )

    try:
        from app.core.pipeline import PipelineManager

        result_ir = PipelineManager.build_result_ir(run_data)
    except Exception:
        result_ir = {}

    blocks = result_ir.get("blocks") if isinstance(result_ir, dict) else None
    if not isinstance(blocks, list):
        blocks = []

    results = run_data.get("results", {}) if isinstance(run_data, dict) else {}
    step_meta_map = run_data.get("step_meta") if isinstance(run_data, dict) else None
    if not isinstance(step_meta_map, dict):
        step_meta_map = {}
    protocol_goal = run_data.get("protocol_goal") if isinstance(run_data, dict) else None

    def _extract_target(meta: Dict[str, Any], res: Dict[str, Any]) -> Optional[str]:
        cfg = meta.get("config") if isinstance(meta, dict) else None
        if isinstance(cfg, dict):
            for k in ["target", "outcome", "endpoint", "y"]:
                v = cfg.get(k)
                if isinstance(v, str) and v.strip():
                    return v.strip()
        for k in ["target", "outcome", "endpoint", "y"]:
            v = meta.get(k) if isinstance(meta, dict) else None
            if isinstance(v, str) and v.strip():
                return v.strip()
        for k in ["target", "outcome", "endpoint", "y"]:
            v = res.get(k) if isinstance(res, dict) else None
            if isinstance(v, str) and v.strip():
                return v.strip()
        return None

    def _extract_visit(step_id: str, meta: Dict[str, Any], res: Dict[str, Any]) -> Optional[str]:
        for k in ["visit", "timepoint", "time", "visit_label", "time_label", "v"]:
            v = meta.get(k) if isinstance(meta, dict) else None
            if isinstance(v, (int, float)):
                try:
                    return f"V{int(v)}"
                except Exception:
                    pass
            if isinstance(v, str) and v.strip():
                s = v.strip()
                m = re.search(r"\bV\s*(\d+)\b", s, flags=re.IGNORECASE)
                if m:
                    return f"V{m.group(1)}"
                return s

        where = meta.get("filter") if isinstance(meta, dict) else None
        if isinstance(where, dict):
            col = str(where.get("col") or where.get("column") or "").strip().lower()
            val = where.get("value") if "value" in where else where.get("val")
            if col and any(x in col for x in ["visit", "визит", "time", "точк", "v"]):
                if isinstance(val, (int, float)):
                    try:
                        return f"V{int(val)}"
                    except Exception:
                        pass
                if isinstance(val, str) and val.strip():
                    s = val.strip()
                    m = re.search(r"\bV\s*(\d+)\b", s, flags=re.IGNORECASE)
                    if m:
                        return f"V{m.group(1)}"
                    return s

        m = re.search(r"(?:^|[_\-])v\s*(\d+)(?:$|[_\-])", step_id, flags=re.IGNORECASE)
        if m:
            return f"V{m.group(1)}"
        m = re.search(r"\bV\s*(\d+)\b", step_id, flags=re.IGNORECASE)
        if m:
            return f"V{m.group(1)}"
        return None

    def _extract_groups(res: Dict[str, Any]) -> List[str]:
        plot_stats = res.get("plot_stats")
        if isinstance(plot_stats, dict) and plot_stats:
            return [str(k) for k in plot_stats.keys()]
        groups = res.get("groups")
        if isinstance(groups, list) and groups:
            return [str(g) for g in groups]
        if isinstance(groups, dict) and groups:
            return [str(k) for k in groups.keys()]
        return []

    def _format_goal(goal: str) -> str:
        g = str(goal or "").strip()
        if not g:
            return ""
        if not is_ru:
            return g
        m = {
            "compare_groups": "Сравнение групп",
            "descriptive": "Описательная статистика",
            "correlation": "Корреляционный анализ",
            "regression": "Регрессионный анализ",
            "survival": "Анализ выживаемости",
            "longitudinal": "Динамика по визитам",
        }
        ru = m.get(g)
        return f"{ru} ({g})" if ru else g

    def _extract_task(meta: Dict[str, Any]) -> Optional[str]:
        for k in ["task", "analysis_task", "section", "domain", "objective", "goal"]:
            v = meta.get(k) if isinstance(meta, dict) else None
            if isinstance(v, str) and v.strip():
                return v.strip()
        if isinstance(protocol_goal, str) and protocol_goal.strip():
            formatted = _format_goal(protocol_goal)
            return formatted if formatted else protocol_goal.strip()
        return None

    def iter_steps():
        if blocks:
            for block in blocks:
                if not isinstance(block, dict):
                    continue
                step_id = block.get("id")
                payload = block.get("payload")
                if isinstance(step_id, str) and isinstance(payload, dict):
                    yield step_id, payload
            return
        if isinstance(results, dict):
            for step_id, res in (results or {}).items():
                if isinstance(step_id, str) and isinstance(res, dict):
                    yield step_id, res

    raw_steps = list(iter_steps())
    integrity_ctx = build_report_integrity_context(run_data)
    filtered_steps, filter_meta = filter_step_pairs_for_report(raw_steps, integrity_ctx)
    deduped = _dedupe_step_payloads(filtered_steps)
    total_steps = len(filtered_steps)
    source_total_steps = int(filter_meta.get("source_total_steps") or len(raw_steps))
    excluded_steps = len(filter_meta.get("excluded_step_ids") or [])
    unique_steps = len(deduped)
    removed = max(0, total_steps - unique_steps)

    table1_rows: List[Dict[str, Any]] = []
    table1_groups: List[str] = []
    table1_group_col: Optional[str] = None
    table1_pmap: Dict[str, Dict[str, Any]] = {}

    for e in deduped:
        step_id = e.get("step_id") if isinstance(e, dict) else None
        res = e.get("res") if isinstance(e, dict) else None
        if not isinstance(step_id, str) or not isinstance(res, dict):
            continue
        if res.get("type") == "table_1":
            stats = res.get("data")
            if not isinstance(stats, dict) or not stats:
                continue
            meta = step_meta_map.get(step_id) if isinstance(step_meta_map, dict) else None
            meta = meta if isinstance(meta, dict) else {}
            target = _extract_target(meta, res) or step_id
            if not target:
                continue
            if not table1_groups:
                table1_groups = [k for k in stats.keys() if k != "overall"]
            if table1_group_col is None:
                cfg = meta.get("config") if isinstance(meta.get("config"), dict) else {}
                table1_group_col = cfg.get("group")
            table1_rows.append({"target": target, "stats": stats})

    for e in deduped:
        res = e.get("res") if isinstance(e, dict) else None
        if not isinstance(res, dict):
            continue
        if res.get("type") != "batch_analysis":
            continue
        g = res.get("group") or res.get("group_column")
        if not g:
            continue
        items = res.get("items")
        if not isinstance(items, list):
            continue
        table1_pmap.setdefault(str(g), {})
        for it in items:
            if not isinstance(it, dict):
                continue
            target = it.get("target") or it.get("outcome")
            if not target:
                continue
            table1_pmap[str(g)][str(target)] = {
                "p_raw": it.get("p_value"),
                "p_adj": it.get("p_value_adj"),
            }

    skip_table1_steps = False
    if len(table1_rows) >= 2 and table1_groups:
        skip_table1_steps = True
        doc.add_heading("Таблица 1. Описательная статистика (сводная)" if is_ru else "Table 1. Descriptive statistics (summary)", level=1)

        any_adj = False
        for g in table1_pmap.values():
            for v in g.values():
                if isinstance(v, dict) and v.get("p_adj") is not None:
                    any_adj = True
                    break

        header_cols = 2 + len(table1_groups) + 1  # variable + groups + overall + p
        table = doc.add_table(rows=1, cols=header_cols)
        hdr = table.rows[0].cells
        hdr[0].text = "Показатель" if is_ru else "Variable"
        for idx, g in enumerate(table1_groups):
            hdr[idx + 1].text = str(g)
        hdr[len(table1_groups) + 1].text = "Итого" if is_ru else "Overall"
        hdr[len(table1_groups) + 2].text = "p(adj)" if any_adj else "p"

        def _summary(stats: Dict[str, Any], use_mean: bool) -> str:
            n = stats.get("count")
            n_s = f"n={int(n)}" if isinstance(n, (int, float)) else ""
            if use_mean:
                return f"{_fmt_num(stats.get('mean'), 2)} ± {_fmt_num(stats.get('std'), 2)} {n_s}".strip()
            return f"{_fmt_num(stats.get('median'), 2)} [{_fmt_num(stats.get('q1'), 2)}; {_fmt_num(stats.get('q3'), 2)}] {n_s}".strip()

        for row in table1_rows:
            stats = row.get("stats") if isinstance(row.get("stats"), dict) else {}
            all_normal = True
            for g in table1_groups:
                sp = stats.get(g, {}).get("shapiro_p")
                if sp is None:
                    all_normal = False
                    break
                try:
                    if float(sp) < 0.05:
                        all_normal = False
                        break
                except Exception:
                    all_normal = False
                    break

            r = table.add_row().cells
            r[0].text = _txt(row.get("target"))
            for idx, g in enumerate(table1_groups):
                r[idx + 1].text = _summary(stats.get(g, {}), all_normal)
            r[len(table1_groups) + 1].text = _summary(stats.get("overall", {}), all_normal)
            p_val = None
            if table1_group_col and table1_group_col in table1_pmap:
                p_info = table1_pmap.get(table1_group_col, {}).get(str(row.get("target")))
                if isinstance(p_info, dict):
                    p_val = p_info.get("p_adj") if p_info.get("p_adj") is not None else p_info.get("p_raw")
            r[len(table1_groups) + 2].text = _fmt_p(p_val)

        # Append categorical variables (counts and %)
        try:
            if dataset_id and table1_group_col:
                ds_dir = _resolve_dataset_dir_path(dataset_id)
                if not ds_dir:
                    raise FileNotFoundError("dataset_dir_not_found")
                study_path = os.path.join(ds_dir, "processed", "study_design.json")
                cat_cols = []
                if os.path.exists(study_path):
                    with open(study_path, "r", encoding="utf-8") as f:
                        sd = json.load(f)
                    if isinstance(sd, dict):
                        design = sd.get("design") if isinstance(sd.get("design"), dict) else {}
                        policy = sd.get("analysis_policy") if isinstance(sd.get("analysis_policy"), dict) else {}
                        cat_cols = design.get("categorical_outcomes") if isinstance(design.get("categorical_outcomes"), list) else []
                        max_cat = int(policy.get("max_table1_categorical") or 40)
                    else:
                        max_cat = 40
                else:
                    max_cat = 40

                parquet_path = os.path.join(ds_dir, "processed", f"{dataset_id}.parquet")
                use_cols = [table1_group_col] + [c for c in cat_cols if c != table1_group_col]
                use_cols = list(dict.fromkeys([c for c in use_cols if c]))
                if os.path.exists(parquet_path) and use_cols:
                    df_cat = pd.read_parquet(parquet_path, columns=use_cols)
                    df_cat = df_cat[df_cat[table1_group_col].notna()]
                    if not df_cat.empty:
                        group_series = df_cat[table1_group_col].astype(str)
                        group_vals = [str(g) for g in table1_groups]
                        group_totals = {str(g): int(df_cat[group_series == str(g)].shape[0]) for g in group_vals}
                        overall_total = int(df_cat.shape[0])

                        def _fmt_cnt_pct(count: int, total: int) -> str:
                            if total <= 0:
                                return "-"
                            pct = (float(count) / float(total)) * 100.0
                            return f"{int(count)} ({pct:.1f}%)"

                        def _bold_cell(cell):
                            try:
                                if cell.paragraphs and cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].bold = True
                            except Exception:
                                pass

                        for col in cat_cols[:max_cat]:
                            if col not in df_cat.columns:
                                continue
                            series = df_cat[col].fillna("Missing").astype(str)
                            overall_counts = series.value_counts(dropna=False)
                            categories = [str(c) for c in overall_counts.index.tolist()]
                            if not categories:
                                continue

                            p_val = None
                            if table1_group_col in table1_pmap:
                                p_info = table1_pmap.get(table1_group_col, {}).get(str(col))
                                if isinstance(p_info, dict):
                                    p_val = p_info.get("p_adj") if p_info.get("p_adj") is not None else p_info.get("p_raw")

                            # Variable header row
                            r = table.add_row().cells
                            r[0].text = _txt(col)
                            _bold_cell(r[0])
                            for idx in range(len(table1_groups) + 1):
                                r[idx + 1].text = ""
                            r[len(table1_groups) + 2].text = _fmt_p(p_val)

                            for cat in categories[:30]:
                                rr = table.add_row().cells
                                rr[0].text = f"  {cat}"
                                for idx, g in enumerate(group_vals):
                                    cnt = int(df_cat[(group_series == str(g)) & (df_cat[col].fillna('Missing').astype(str) == cat)].shape[0])
                                    rr[idx + 1].text = _fmt_cnt_pct(cnt, group_totals.get(str(g), 0))
                                cnt_overall = int(overall_counts.get(cat, 0))
                                rr[len(table1_groups) + 1].text = _fmt_cnt_pct(cnt_overall, overall_total)
                                rr[len(table1_groups) + 2].text = ""
        except Exception:
            pass

        doc.add_paragraph(
            "Примечание: если нормальность нарушена, используется медиана [Q1; Q3], иначе среднее ± SD."
            if is_ru
            else "Note: non-normal variables use median [Q1; Q3], otherwise mean ± SD."
        )

    doc.add_paragraph(("Шагов (всего)" if is_ru else "Steps (total)") + f": {total_steps}")
    doc.add_paragraph(("Шагов (до фильтра)" if is_ru else "Steps (source)") + f": {source_total_steps}")
    doc.add_paragraph(("Исключено верификатором" if is_ru else "Excluded by verifier") + f": {excluded_steps}")
    doc.add_paragraph(("Шагов (уникальных)" if is_ru else "Steps (unique)") + f": {unique_steps}")
    doc.add_paragraph(("Свернуто повторов" if is_ru else "Duplicates collapsed") + f": {removed}")

    by_type = Counter([str((e.get("res") or {}).get("type") or "result") for e in deduped if isinstance(e, dict)])
    if by_type:
        doc.add_paragraph(
            ("Состав по типам" if is_ru else "By type")
            + ": "
            + ", ".join(
                [
                    f"{k}={int(v)}"
                    for k, v in sorted(by_type.items(), key=lambda x: (-int(x[1]), str(x[0])))
                ]
            )
        )

    steps = deduped
    if steps:
        doc.add_heading("Навигация по шагам" if is_ru else "Step Navigation", level=1)
        for e in steps:
            sid = e.get("step_id") if isinstance(e, dict) else None
            res_for_nav = e.get("res") if isinstance(e, dict) else None
            if not isinstance(sid, str) or not isinstance(res_for_nav, dict):
                continue
            meta_for_nav = step_meta_map.get(sid) if isinstance(step_meta_map.get(sid), dict) else {}
            nav_title = _build_step_display(sid, res_for_nav, meta_for_nav, is_ru)
            p = doc.add_paragraph(style="List Bullet")
            _add_internal_hyperlink(p, step_bookmarks.get(sid, ""), f"{sid}: {nav_title}")

    for idx, e in enumerate(steps):
        step_id = e.get("step_id") if isinstance(e, dict) else None
        res = e.get("res") if isinstance(e, dict) else None
        dup_count = int(e.get("dup_count") or 1) if isinstance(e, dict) else 1
        if not isinstance(step_id, str) or not isinstance(res, dict):
            continue

        step_meta = step_meta_map.get(step_id) if isinstance(step_id, str) else None
        step_meta = step_meta if isinstance(step_meta, dict) else {}
        step_title = _build_step_display(step_id, res if isinstance(res, dict) else {}, step_meta, is_ru)
        suffix = f" (×{dup_count})" if dup_count > 1 else ""
        step_heading = doc.add_heading(("Шаг" if is_ru else "Step") + f": {step_title}{suffix}", level=1)
        _add_bookmark(step_heading, step_bookmarks.get(step_id, ""))
        doc.add_paragraph(f"ID: {step_id}")
        task = _extract_task(step_meta)
        visit = _extract_visit(step_id, step_meta, res if isinstance(res, dict) else {})
        group_levels = _extract_groups(res if isinstance(res, dict) else {})

        if task:
            doc.add_paragraph(("Задача" if is_ru else "Task") + f": {task}")
        if visit:
            doc.add_paragraph(("Точка" if is_ru else "Timepoint") + f": {visit}")
        if group_levels:
            if len(group_levels) == 2:
                grp_s = f"{group_levels[0]} vs {group_levels[1]}"
            else:
                grp_s = ", ".join(group_levels)
            doc.add_paragraph(("Сравниваемые группы" if is_ru else "Compared groups") + f": {grp_s}")

        if not isinstance(res, dict):
            doc.add_paragraph("Нет структурированного результата" if is_ru else "No structured result")
            continue

        if res.get("type") == "table_1":
            if skip_table1_steps:
                continue
            stats_map = res.get("data", {})
            if isinstance(stats_map, dict) and stats_map:
                groups = [k for k in stats_map.keys() if k != "overall"]
                cols = 2 + len(groups)
                table = doc.add_table(rows=1, cols=cols)
                hdr = table.rows[0].cells
                hdr[0].text = "Показатель"
                for i, g in enumerate(groups):
                    n = _txt(stats_map.get(g, {}).get("count"))
                    hdr[i + 1].text = f"{g} (n={n})"
                overall_n = _txt(stats_map.get("overall", {}).get("count"))
                hdr[-1].text = f"Итого (n={overall_n})"

                def _cell_for(metric_key: str, s: Dict[str, Any]) -> str:
                    if metric_key == "mean_sd":
                        return f"{_fmt_num(s.get('mean'), 2)} ({_fmt_num(s.get('std'), 2)})"
                    if metric_key == "ci_95":
                        return f"[{_fmt_num(s.get('ci_95_low'), 2)}, {_fmt_num(s.get('ci_95_high'), 2)}]"
                    if metric_key == "median_q1_q3":
                        return f"{_fmt_num(s.get('median'), 2)} [{_fmt_num(s.get('q1'), 2)}, {_fmt_num(s.get('q3'), 2)}]"
                    if metric_key == "iqr":
                        return _fmt_num(s.get("iqr"), 2)
                    if metric_key == "min_max":
                        return f"{_fmt_num(s.get('min'), 2)} – {_fmt_num(s.get('max'), 2)}"
                    if metric_key == "shapiro":
                        return _fmt_p(s.get("shapiro_p"))
                    return "-"

                metrics = [
                    ("Mean (SD)", "mean_sd"),
                    ("95% CI (Mean)", "ci_95"),
                    ("Median [Q1, Q3]", "median_q1_q3"),
                    ("IQR", "iqr"),
                    ("Range (Min-Max)", "min_max"),
                    ("Normality (Shapiro p)", "shapiro"),
                ]

                for label, key in metrics:
                    row = table.add_row().cells
                    row[0].text = label
                    for i, g in enumerate(groups):
                        row[i + 1].text = _cell_for(key, stats_map.get(g, {}) or {})
                    row[-1].text = _cell_for(key, stats_map.get("overall", {}) or {})
            continue

        if res.get("type") == "responders":
            outcome = res.get("outcome")
            baseline = res.get("baseline")
            baseline_time = baseline.get("time") if isinstance(baseline, dict) else None
            threshold = res.get("threshold")
            direction = res.get("direction")

            if outcome:
                doc.add_paragraph(("Показатель" if is_ru else "Outcome") + f": {_txt(outcome)}")
            head = []
            if baseline_time is not None:
                head.append(("база" if is_ru else "baseline") + f"={_txt(baseline_time)}")
            if threshold is not None:
                head.append(("порог" if is_ru else "threshold") + f"={_txt(threshold)}")
            if direction:
                head.append(("направление" if is_ru else "direction") + f"={_txt(direction)}")
            if head:
                doc.add_paragraph(" • ".join(head))

            continue

        def _render_batch_step(batch_res: Dict[str, Any], title: Optional[str] = None) -> None:
            alpha_val = _coerce_alpha(
                batch_res.get("alpha"),
                run_data.get("alpha") if isinstance(run_data, dict) else None,
            )

            group_col = batch_res.get("group") or batch_res.get("group_column") or ""
            multiplicity = batch_res.get("multiplicity_correction")
            post_hoc = batch_res.get("post_hoc")
            post_hoc_correction = batch_res.get("post_hoc_correction")
            corr_label = _format_correction_label(multiplicity, is_ru) if multiplicity else ""
            post_hoc_corr_label = _format_correction_label(post_hoc_correction, is_ru) if post_hoc_correction else ""

            rows_payload = _collect_batch_inferential_rows(batch_res.get("items"), alpha_val)
            rows = rows_payload.get("rows") if isinstance(rows_payload, dict) else []
            rows = rows if isinstance(rows, list) else []
            sig_count = int(rows_payload.get("sig_count") or 0) if isinstance(rows_payload, dict) else 0
            group_labels = rows_payload.get("group_labels") if isinstance(rows_payload, dict) else None

            if title:
                doc.add_heading(title, level=2)

            if group_col:
                doc.add_paragraph(("Группировка" if is_ru else "Group") + f": {_txt(group_col)}")
            doc.add_paragraph(("Альфа" if is_ru else "Alpha") + f": {_txt(alpha_val)}")
            if multiplicity:
                doc.add_paragraph(("Поправка" if is_ru else "Correction") + f": {_txt(corr_label or multiplicity)}")
            if post_hoc and post_hoc != "none":
                ph = _txt(post_hoc)
                if post_hoc_correction and post_hoc_correction != "none":
                    ph = f"{ph} ({_txt(post_hoc_corr_label or post_hoc_correction)})"
                doc.add_paragraph(("Post-hoc" if not is_ru else "Пост-хок") + f": {ph}")
            doc.add_paragraph(("Показателей" if is_ru else "Targets") + f": {_txt(len(rows))}")
            doc.add_paragraph(("Значимых" if is_ru else "Significant") + f": {_txt(sig_count)}")

            cols = 5
            has_group_stats = bool(group_labels and len(group_labels) == 2)
            if has_group_stats:
                cols = 7
            table = doc.add_table(rows=1, cols=cols)
            hdr = table.rows[0].cells
            hdr[0].text = "Показатель" if is_ru else "Target"
            if has_group_stats:
                hdr[1].text = f"{group_labels[0]}: M±SD; Me[Q1;Q3]"
                hdr[2].text = f"{group_labels[1]}: M±SD; Me[Q1;Q3]"
                hdr[3].text = "p"
                hdr[4].text = "q" if not is_ru else "p(adj)"
                hdr[5].text = "Значимо" if is_ru else "Sig"
                hdr[6].text = "Тест" if is_ru else "Test"
            else:
                hdr[1].text = "p"
                hdr[2].text = "q" if not is_ru else "p(adj)"
                hdr[3].text = "Значимо" if is_ru else "Sig"
                hdr[4].text = "Тест" if is_ru else "Test"

            def _fmt_group_stats(stats: Any) -> str:
                if not isinstance(stats, dict):
                    return "-"
                mean = _fmt_num(stats.get("mean"), 2)
                sd = _fmt_num(stats.get("sd"), 2)
                med = _fmt_num(stats.get("median"), 2)
                q1 = _fmt_num(stats.get("q1"), 2)
                q3 = _fmt_num(stats.get("q3"), 2)
                n = stats.get("count")
                n_s = f"n={int(n)}" if isinstance(n, (int, float)) else ""
                return f"{mean} ± {sd}; {med} [{q1}; {q3}] {n_s}".strip()

            for r in rows:
                row = table.add_row().cells
                row[0].text = _txt(r.get("target"))
                if has_group_stats:
                    group_stats = r.get("group_stats") if isinstance(r, dict) else None
                    row[1].text = _fmt_group_stats(group_stats.get(group_labels[0]) if isinstance(group_stats, dict) else None)
                    row[2].text = _fmt_group_stats(group_stats.get(group_labels[1]) if isinstance(group_stats, dict) else None)
                    row[3].text = _fmt_p(r.get("p_raw"))
                    row[4].text = _fmt_p(r.get("p_adj"))
                    row[5].text = ("да" if is_ru else "yes") if r.get("sig") else ("нет" if is_ru else "no")
                    row[6].text = _method_label_from_id(r.get("method"), is_ru) or _txt(r.get("method"))
                else:
                    row[1].text = _fmt_p(r.get("p_raw"))
                    row[2].text = _fmt_p(r.get("p_adj"))
                    row[3].text = ("да" if is_ru else "yes") if r.get("sig") else ("нет" if is_ru else "no")
                    row[4].text = _method_label_from_id(r.get("method"), is_ru) or _txt(r.get("method"))

        if res.get("type") == "batch_analysis":
            _render_batch_step(res)
            continue

        if res.get("type") == "timepoint_batch_analysis":
            split_by = res.get("split_by")
            group_col = res.get("group")
            if group_col:
                doc.add_paragraph(("Группировка" if is_ru else "Group") + f": {_txt(group_col)}")
            if split_by:
                doc.add_paragraph(("Разбиение" if is_ru else "Split by") + f": {_txt(split_by)}")

            slices = res.get("slices")
            slices = slices if isinstance(slices, dict) else {}
            for slice_key in sorted(slices.keys(), key=lambda x: str(x)):
                slice_res = slices.get(slice_key)
                if not isinstance(slice_res, dict):
                    continue
                _render_batch_step(slice_res, title=("Точка" if is_ru else "Slice") + f": {_txt(slice_key)}")
            continue

        step_type = res.get("type")
        if step_type and step_type not in {"compare", "hypothesis_test", "correlation", "regression", "survival", "mixed_effects", "clustered_correlation", "batch_compare_by_factor", "batch_analysis", "timepoint_batch_analysis", "delta_batch_analysis", "agreement", "assumption_test", "time_series", "responders"}:
            doc.add_paragraph(("Тип" if is_ru else "Type") + f": {_txt(step_type)}")
            err = res.get("error")
            if isinstance(err, str) and err.strip():
                doc.add_paragraph(("Ошибка" if is_ru else "Error") + f": {_txt(err)}")
            try:
                raw = json.dumps(res, ensure_ascii=False, indent=2, default=str)
            except Exception:
                raw = str(res)
            raw = raw[:8000]
            doc.add_paragraph(raw)
            continue

            by_visit = res.get("by_visit")
            if isinstance(by_visit, dict) and by_visit:
                def _sort_key(v: Any) -> Any:
                    try:
                        return float(v)
                    except Exception:
                        return str(v)

                for vk in sorted(by_visit.keys(), key=_sort_key):
                    v = by_visit.get(vk)
                    if not isinstance(v, dict):
                        continue
                    groups = v.get("groups")
                    if not isinstance(groups, dict) or not groups:
                        continue

                    test = v.get("test")
                    test_method = test.get("method") if isinstance(test, dict) else None
                    test_p = test.get("p_value") if isinstance(test, dict) else None
                    doc.add_heading(("Визит" if is_ru else "Visit") + f" {vk}", level=2)
                    doc.add_paragraph(("Тест" if is_ru else "Test") + f": {_txt(test_method or 'chi_square')}, p={_fmt_p(test_p)}")

                    table = doc.add_table(rows=1, cols=4)
                    hdr = table.rows[0].cells
                    hdr[0].text = "Группа" if is_ru else "Group"
                    hdr[1].text = "Респондеры" if is_ru else "Responders"
                    hdr[2].text = "Всего" if is_ru else "Total"
                    hdr[3].text = "Доля" if is_ru else "Rate"

                    for g in sorted(groups.keys(), key=_sort_key):
                        st = groups.get(g)
                        if not isinstance(st, dict):
                            continue
                        total = st.get("total")
                        responders = st.get("responders")
                        rate = st.get("rate")
                        try:
                            rate_s = f"{float(rate) * 100.0:.1f}%" if rate is not None and np.isfinite(float(rate)) else "-"
                        except Exception:
                            rate_s = "-"
                        r = table.add_row().cells
                        r[0].text = _txt(g)
                        r[1].text = _txt(responders)
                        r[2].text = _txt(total)
                        r[3].text = rate_s

            continue

        method = res.get("method")
        method_default = "Статистический тест" if is_ru else "Statistical Test"
        method_name = method_default
        if isinstance(method, dict):
            method_name = _method_label_from_id(method.get("id") or method.get("name"), is_ru) or method.get("name") or method.get("id") or method_name
        elif isinstance(method, str):
            method_name = _method_label_from_id(method, is_ru) or method
        doc.add_paragraph(("Метод" if is_ru else "Method") + f": {method_name}")

        if res.get("type") == "mixed_effects" and res.get("formula"):
            doc.add_paragraph(("Формула" if is_ru else "Formula") + f": {str(res.get('formula'))}")

        summary = doc.add_table(rows=0, cols=2)
        for k, v in [
            ("p-value", _fmt_p(res.get("p_value"))),
            (("статистика" if is_ru else "stat"), _fmt_num(res.get("stat_value", res.get("stats")), 3)),
            (("эффект" if is_ru else "effect"), f"{_txt(res.get('effect_size_name') or 'effect')} {_fmt_num(res.get('effect_size'), 2)}" if res.get("effect_size") is not None else "-"),
            (("мощность" if is_ru else "power"), _fmt_num(res.get("power"), 2)),
            ("BF10", _txt(res.get("bf10"))),
        ]:
            r = summary.add_row().cells
            r[0].text = str(k)
            r[1].text = str(v)
        printed_time_warnings = False

        if step_type == "time_series":
            trend = res.get("trend") if isinstance(res.get("trend"), dict) else {}
            diagnostics = res.get("diagnostics") if isinstance(res.get("diagnostics"), dict) else {}
            ljung = diagnostics.get("ljung_box") if isinstance(diagnostics.get("ljung_box"), dict) else {}
            forecast = res.get("forecast") if isinstance(res.get("forecast"), dict) else {}
            forecast_n = len(forecast.get("points") or []) if isinstance(forecast.get("points"), list) else 0
            time_quality = res.get("time_quality") if isinstance(res.get("time_quality"), dict) else {}
            if not time_quality and isinstance(diagnostics.get("time_quality"), dict):
                time_quality = diagnostics.get("time_quality")

            axis_kind_raw = str(res.get("time_axis_kind") or time_quality.get("time_axis_kind") or "").strip().lower()
            axis_kind_label = {
                "datetime": ("календарная дата/время" if is_ru else "calendar date/time"),
                "numeric": ("числовая последовательность" if is_ru else "numeric sequence"),
                "categorical": ("категориальная ось" if is_ru else "categorical axis"),
                "index": ("порядковый индекс" if is_ru else "row index"),
            }.get(axis_kind_raw, "-")

            quality_raw = str(time_quality.get("quality") or "").strip().lower()
            quality_label = {
                "ok": ("норма" if is_ru else "ok"),
                "caution": ("осторожно" if is_ru else "caution"),
                "warning": ("высокий риск" if is_ru else "high risk"),
            }.get(quality_raw, "-")

            parse_ratio_txt = "-"
            try:
                parse_ratio = float(time_quality.get("datetime_parse_ratio"))
                if np.isfinite(parse_ratio):
                    parse_ratio_txt = f"{parse_ratio * 100.0:.1f}%"
            except Exception:
                parse_ratio_txt = "-"

            year_range_txt = "-"
            try:
                min_year = int(time_quality.get("min_year"))
                max_year = int(time_quality.get("max_year"))
                year_range_txt = f"{min_year}-{max_year}"
            except Exception:
                year_range_txt = "-"

            flags_raw = time_quality.get("flags")
            if isinstance(flags_raw, list):
                flags = [str(v).strip() for v in flags_raw if str(v).strip()]
            else:
                flags = []
            flags_txt = ", ".join(flags) if flags else "-"
            inferred_freq = str(time_quality.get("inferred_frequency") or "").strip() or "-"

            series_diag = doc.add_table(rows=0, cols=2)
            for k, v in [
                (("Диагностика временного ряда" if is_ru else "Series diagnostics"), ""),
                (("Тип временной оси" if is_ru else "Time axis type"), axis_kind_label),
                (("Качество временной оси" if is_ru else "Time axis quality"), quality_label),
                (("Парсинг datetime" if is_ru else "Datetime parse ratio"), parse_ratio_txt),
                (("Диапазон лет" if is_ru else "Year range"), year_range_txt),
                (("Интервал (infer)" if is_ru else "Inferred frequency"), inferred_freq),
                (("Флаги качества" if is_ru else "Quality flags"), flags_txt),
                (("Ljung-Box (p)" if is_ru else "Ljung-Box (p)"), _fmt_p(ljung.get("p_value"))),
                (("Белый шум" if is_ru else "White-noise-like"), ("да" if is_ru else "yes") if isinstance(ljung.get("white_noise_like"), bool) and bool(ljung.get("white_noise_like")) else (("нет" if is_ru else "no") if isinstance(ljung.get("white_noise_like"), bool) else "-")),
                (("Прогноз (точек)" if is_ru else "Forecast (points)"), str(forecast_n) if forecast_n > 0 else "-"),
                (("Тренд (наклон)" if is_ru else "Trend (slope)"), _fmt_num(trend.get("slope"), 4)),
            ]:
                row = series_diag.add_row().cells
                row[0].text = str(k)
                row[1].text = str(v)

            warning_items = res.get("warnings")
            warning_lines = []
            if isinstance(warning_items, list):
                for item in warning_items:
                    txt = str(item).strip()
                    if txt:
                        warning_lines.append(txt)
            if warning_lines:
                doc.add_paragraph("Предупреждения по хронологии:" if is_ru else "Chronology warnings:")
                for msg in warning_lines[:6]:
                    doc.add_paragraph(msg, style="List Bullet")
                printed_time_warnings = True

        compare_rows = _build_pairwise_comparison_rows(res)
        if compare_rows:
            is_median = compare_rows[0].get("center_label") == "median"
            a_hdr = ("Me [Q1; Q3]" if is_median else "M ± SD")

            doc.add_paragraph("Сравнение групп (сводная таблица):" if is_ru else "Group Comparison (Summary Table):")
            table = doc.add_table(rows=1, cols=8)
            hdr = table.rows[0].cells
            hdr[0].text = "Сравнение" if is_ru else "Comparison"
            hdr[1].text = f"{a_hdr} A"
            hdr[2].text = f"{a_hdr} B"
            hdr[3].text = "Δ (A−B)"
            hdr[4].text = "Δ%"
            hdr[5].text = "p"
            hdr[6].text = "BF10"
            hdr[7].text = "Эффект" if is_ru else "Effect"

            def _fmt(v: Any, d: int = 2) -> str:
                try:
                    if v is None:
                        return "-"
                    f = float(v)
                    return f"{f:.{d}f}" if np.isfinite(f) else "-"
                except Exception:
                    return "-"

            for r in compare_rows[:80]:
                row = table.add_row().cells
                a = str(r.get("a") or "-")
                b = str(r.get("b") or "-")
                row[0].text = f"{a} vs {b}"

                a_n = r.get("a_n")
                b_n = r.get("b_n")
                a_n_s = f"n={int(a_n)}" if isinstance(a_n, (int, float)) else ""
                b_n_s = f"n={int(b_n)}" if isinstance(b_n, (int, float)) else ""

                if is_median:
                    q1a, q3a = r.get("a_spread") if isinstance(r.get("a_spread"), tuple) else (None, None)
                    q1b, q3b = r.get("b_spread") if isinstance(r.get("b_spread"), tuple) else (None, None)
                    row[1].text = f"{_fmt(r.get('a_center'))} [{_fmt(q1a)}; {_fmt(q3a)}] {a_n_s}".strip()
                    row[2].text = f"{_fmt(r.get('b_center'))} [{_fmt(q1b)}; {_fmt(q3b)}] {b_n_s}".strip()
                else:
                    row[1].text = f"{_fmt(r.get('a_center'))} ± {_fmt(r.get('a_spread'))} {a_n_s}".strip()
                    row[2].text = f"{_fmt(r.get('b_center'))} ± {_fmt(r.get('b_spread'))} {b_n_s}".strip()

                row[3].text = _fmt(r.get("diff"), 2)
                diff_pct = r.get("diff_pct")
                row[4].text = (f"{_fmt(diff_pct, 1)}%" if diff_pct is not None else "-")
                row[5].text = _fmt_p(r.get("p_value"))
                row[6].text = _fmt(r.get("bf10"), 3)
                eff = r.get("effect_size")
                eff_name = r.get("effect_size_name")
                row[7].text = (f"{_txt(eff_name or 'effect')}={_fmt(eff)}" if eff is not None else "-")

            doc.add_paragraph(
                "Пояснения: Δ — абсолютная разница (A−B); Δ% — (A−B)/B; p — уровень значимости; BF10 — сила свидетельства в пользу H1 (значения <1 поддерживают H0); эффект — размер эффекта." if is_ru else "Notes: Δ is (A−B); Δ% is (A−B)/B; p is the p-value; BF10 quantifies evidence for H1 (values <1 support H0); effect is effect size."
            )
        else:
            group_levels = _extract_groups(res if isinstance(res, dict) else {})
            if group_levels and len(group_levels) >= 3:
                doc.add_paragraph(
                    "Есть 3+ группы: попарные сравнения (post-hoc) не выполнены или отсутствуют в результате." if is_ru else "3+ groups: pairwise post-hoc comparisons are not available for this step."
                )

        if is_ru:
            rationale_ru = _method_selection_rationale_ru(res)
            if isinstance(rationale_ru, str) and rationale_ru.strip():
                doc.add_paragraph("Обоснование выбора теста: " + rationale_ru)

            bf10_text = _interpret_bf10_ru(res.get("bf10"))
            if isinstance(bf10_text, str) and bf10_text.strip():
                doc.add_paragraph("Интерпретация BF10: " + bf10_text)

        try:
            plot_png = _render_plot_png_bytes(res, is_ru=is_ru)
            if plot_png:
                bio = BytesIO(plot_png)
                doc.add_picture(bio, width=Inches(5.8))
        except Exception:
            pass

        warnings = res.get("warnings")
        if isinstance(warnings, list) and warnings and not printed_time_warnings:
            doc.add_paragraph("Предупреждения:" if is_ru else "Warnings:")
            for w in warnings:
                doc.add_paragraph(str(w), style="List Bullet")

        roc = res.get("roc")
        if isinstance(roc, dict) and isinstance(roc.get("plot_data"), list) and roc.get("plot_data"):
            auc_val = roc.get("auc")
            if auc_val is not None:
                doc.add_paragraph(f"AUC: {_fmt_num(auc_val, 3)}")
            roc_png = _render_plot_png_bytes(
                {"plot_data": roc.get("plot_data"), "plot_config": roc.get("plot_config")},
                is_ru=is_ru,
            )
            if roc_png:
                bio = BytesIO(roc_png)
                doc.add_picture(bio, width=Inches(5.8))

        if res.get("type") == "mixed_effects":
            em = res.get("estimated_means")
            if isinstance(em, dict) and em:
                doc.add_heading("Оценённые средние" if is_ru else "Estimated Means", level=2)
                table = doc.add_table(rows=1, cols=5)
                hdr = table.rows[0].cells
                hdr[0].text = "Группа" if is_ru else "Group"
                hdr[1].text = "Визит" if is_ru else "Time"
                hdr[2].text = "Оценка" if is_ru else "Estimate"
                hdr[3].text = "95% ДИ" if is_ru else "95% CI"
                hdr[4].text = "n"

                def _sort_key(v: Any) -> Any:
                    try:
                        return float(v)
                    except Exception:
                        return str(v)

                for g in sorted(em.keys(), key=_sort_key):
                    times = em.get(g)
                    if not isinstance(times, dict):
                        continue
                    for t in sorted(times.keys(), key=_sort_key):
                        stats = times.get(t)
                        if not isinstance(stats, dict):
                            continue
                        r = table.add_row().cells
                        r[0].text = str(g)
                        r[1].text = str(t)
                        r[2].text = _fmt_num(stats.get("estimate"), 2)
                        r[3].text = f"[{_fmt_num(stats.get('ci_lower'), 2)}, {_fmt_num(stats.get('ci_upper'), 2)}]"
                        r[4].text = _txt(stats.get("n"))

        if res.get("type") == "regression":
            coefs = res.get("coefficients")
            if isinstance(coefs, list) and coefs:
                doc.add_heading("Коэффициенты" if is_ru else "Coefficients", level=2)
                has_or = any(isinstance(c, dict) and c.get("odds_ratio") is not None for c in coefs)
                cols = 5 + (1 if has_or else 0)
                table = doc.add_table(rows=1, cols=cols)
                hdr = table.rows[0].cells
                hdr[0].text = "Переменная" if is_ru else "Term"
                hdr[1].text = "Коэф." if is_ru else "Coef"
                hdr[2].text = "SE"
                hdr[3].text = "p"
                hdr[4].text = "95% ДИ" if is_ru else "95% CI"
                if has_or:
                    hdr[5].text = "OR"

                for c in coefs[:40]:
                    if not isinstance(c, dict):
                        continue
                    r = table.add_row().cells
                    r[0].text = _txt(c.get("variable"))
                    r[1].text = _fmt_num(c.get("coefficient"), 3)
                    r[2].text = _fmt_num(c.get("std_err"), 3)
                    r[3].text = _fmt_p(c.get("p_value"))
                    r[4].text = f"[{_fmt_num(c.get('ci_lower'), 3)}, {_fmt_num(c.get('ci_upper'), 3)}]"
                    if has_or:
                        r[5].text = _fmt_num(c.get("odds_ratio"), 3) if c.get("odds_ratio") is not None else "-"

        bootstrap_lines_docx = _bootstrap_trace_lines(res.get("bootstrap"), is_ru=is_ru)
        if bootstrap_lines_docx:
            doc.add_paragraph("Bootstrap-трассировка:" if is_ru else "Bootstrap trace:")
            for line in bootstrap_lines_docx:
                doc.add_paragraph(str(line), style="List Bullet")

        png_bytes = _render_plot_png_bytes(res, is_ru=is_ru)
        if png_bytes:
            try:
                section = doc.sections[-1]
                available_emu = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
                available_in = max(1.0, float(available_emu) / 914400.0)
                available_h_emu = int(section.page_height) - int(section.top_margin) - int(section.bottom_margin)
                available_h_in = max(1.0, float(available_h_emu) / 914400.0)
            except Exception:
                available_in = 5.8
                available_h_in = 7.5

            doc.add_paragraph(("График" if is_ru else "Plot") + (": на следующей странице" if is_ru else ": next page"))
            doc.add_page_break()
            bio = BytesIO(png_bytes)

            try:
                from PIL import Image
                from io import BytesIO as _Bio

                with Image.open(_Bio(png_bytes)) as im:
                    img_w_px, img_h_px = im.size
                ratio = float(img_h_px) / float(img_w_px) if img_w_px else None
            except Exception:
                ratio = None

            if ratio and ratio > 0:
                if (available_in * ratio) <= available_h_in:
                    doc.add_picture(bio, width=Inches(available_in))
                else:
                    doc.add_picture(bio, height=Inches(available_h_in))
            else:
                doc.add_picture(bio, width=Inches(available_in))
            if idx < len(steps) - 1:
                doc.add_page_break()

        interpretation = (res.get("ai_interpretation") or res.get("conclusion")) if is_ru else res.get("conclusion")
        if _is_placeholder_interpretation(interpretation):
            interpretation = _generate_fallback_interpretation(res if isinstance(res, dict) else {}, step_meta, is_ru)
        if interpretation:
            doc.add_paragraph(("Интерпретация" if is_ru else "Interpretation") + ":")
            doc.add_paragraph(str(interpretation))

    summary = _extract_protocol_findings(run_data if isinstance(run_data, dict) else {})
    text = _build_discussion_conclusion(summary, is_ru=is_ru)
    discussion = text.get("discussion") if isinstance(text.get("discussion"), list) else []
    conclusion = text.get("conclusion") if isinstance(text.get("conclusion"), list) else []

    if discussion:
        doc.add_heading("Обсуждение" if is_ru else "Discussion", level=1)
        for line in discussion:
            doc.add_paragraph(str(line))

    if conclusion:
        doc.add_heading("Выводы" if is_ru else "Conclusions", level=1)
        for line in conclusion:
            doc.add_paragraph(str(line), style="List Bullet")

    limitations = _build_report_limitations(summary, methods_summary, is_ru=is_ru)
    if limitations:
        doc.add_heading("Ограничения" if is_ru else "Limitations", level=1)
        for line in limitations:
            doc.add_paragraph(str(line), style="List Bullet")

    out = BytesIO()
    doc.save(out)
    return bytes(out.getvalue())
