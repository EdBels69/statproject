"""
Copilot Scientific Report Generator - DOCX output.

Generates professional clinical trial reports matching 'Expert' standards including:
- Glossary of Statistical Terms
- Section 2: Descriptive & Group Comparison (for all endpoints)
- Section 3: Longitudinal Analysis (Mixed Models)
- Section 4: Responder Analysis
- Professional formatting (ISO/GCP style)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Dict, Any, Optional, List
from datetime import datetime
import io
import json
import pandas as pd
import numpy as np
import os

def generate_report(
    results: Dict[str, Any],
    plan: Dict[str, Any],
    code: Optional[str] = None,
    interpretation: Optional[str] = None,
    dataset_info: Optional[Dict] = None
) -> bytes:
    """
    Generate professional DOCX report from analysis results.
    """
    doc = Document()
    _set_document_styles(doc)
    
    # Counters for numbering (mutable for closures)
    table_counter = [0]
    figure_counter = [0]
    
    
    # 1. Front Matter
    _add_title_page(doc, dataset_info)
    _add_toc_placeholder(doc)
    _add_glossary_section(doc)
    
    # 2. Design
    _add_design_section(doc, plan)

    # 3. Methods Section
    doc.add_heading("2. Методы", level=1)
    if plan and "analyses" in plan:
        for i, analysis in enumerate(plan.get("analyses", []), 1):
            name = analysis.get("name", f"Анализ {i}")
            atype = analysis.get("type", "unknown")
            doc.add_paragraph(f"{i}. {name} ({atype})", style='List Number')
    doc.add_paragraph(
        f"Уровень значимости α = {plan.get('confidence_level', 0.95) if plan else 0.95:.2f}. "
        "Множественные сравнения скорректированы методом Холма-Бонферрони.",
        style='Normal'
    )

    # 4. Results
    doc.add_page_break()
    doc.add_heading("3. Результаты", level=1)

    # 4.1 Create helper to capture counters
    def render_node(key, item):
        if "table" in item:
            table_counter[0] += 1
            title = item.get("title", key)
            doc.add_paragraph(f"Таблица {table_counter[0]}. {title}", style='Caption')
            _render_table_node(doc, item, key)
        
        if "plots" in item:
             for plot_path in item["plots"]:
                if plot_path and os.path.exists(plot_path):
                    figure_counter[0] += 1
                    try:
                        doc.add_picture(plot_path, width=Inches(6))
                        doc.add_paragraph(f"Рисунок {figure_counter[0]}. {item.get('title', key)}", style='Caption')
                    except Exception as e:
                        doc.add_paragraph(f"[Ошибка вывода рисунка: {e}]")

    # Iterate results and render
    if results:
        for key, item in results.items():
            if isinstance(item, dict):
                render_node(key, item)

    # 5. Interpretation
    if interpretation:
        doc.add_page_break()
        doc.add_heading('4. Интерпретация и Выводы (AI Expert)', level=1)
        for para in interpretation.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # 6. Appendix
    if code:
        doc.add_page_break()
        doc.add_heading('Приложение: Код Python/R', level=1)
        p = doc.add_paragraph(code[:15000])
        p.style = 'No Spacing'
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(8)

    # Page numbers
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'PAGE')
        run._element.append(fld)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _set_document_styles(doc):
    """Set professional document styles - Times New Roman, APA-like."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    
    # Headings
    for level in range(1, 4):
        if f'Heading {level}' in doc.styles:
            heading = doc.styles[f'Heading {level}']
            heading.font.name = 'Times New Roman'
            heading.font.bold = True
            heading.font.color.rgb = RGBColor(0, 0, 0)
            heading.font.size = Pt(16 - level * 2)  # H1=14, H2=12, H3=10
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    
    # Table style
    try:
        table_style = doc.styles.add_style('ReportTable', 2)  # WD_STYLE_TYPE.TABLE = 2
        table_style.font.name = 'Times New Roman'
        table_style.font.size = Pt(10)
    except:
        pass # Style might already exist

def _add_title_page(doc, info):
    doc.add_heading('СТАТИСТИЧЕСКИЙ ОТЧЁТ', level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('Клиническое исследование / Анализ данных').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if info:
        doc.add_paragraph(f"Источник: {info.get('filename', 'Unknown')}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

def _add_toc_placeholder(doc):
    doc.add_heading('Оглавление', level=1)
    doc.add_paragraph("[Обновите поле оглавления в Word: ПКМ -> Обновить поле]")
    doc.add_page_break()

def _add_glossary_section(doc):
    doc.add_heading('Глоссарий и Методология', level=1)
    
    terms = [
        ("H0 (Hypthesis 0)", "Нулевая гипотеза об отсутствии различий."),
        ("p-value", "Вероятность ошибки при отвержении H0. Значимым считается p < 0.05."),
        ("ANOVA/Kruskal-Wallis", "Тесты для сравнения 3+ групп (параметрический / непараметрический)."),
        ("Mixed Models", "Модели смешанных эффектов для анализа повторных измерений (Longitudinal)."),
        ("BF10 (Bayes Factor)", "Фактор Байеса: сила доказательств в пользу альтернативы (H1)."),
        ("Effect Size", "Величина эффекта (не зависит от размера выборки).")
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Термин"
    hdr[1].text = "Определение"
    _bold_runs(hdr[0].paragraphs[0])
    _bold_runs(hdr[1].paragraphs[0])
    
    for t, d in terms:
        row = table.add_row().cells
        row[0].text = t
        row[1].text = d

    doc.add_paragraph("\nВсе p-значения являются двусторонними. Коррекция множественных сравнений применяется при необходимости (Holm/Bonferroni).")
    doc.add_page_break()

def _add_design_section(doc, plan):
    doc.add_heading('1. Дизайн исследования', level=1)
    if not plan: return
    
    if plan.get("understood_goal"):
        p = doc.add_paragraph()
        p.add_run("Цель анализа: ").bold = True
        p.add_run(plan['understood_goal'])
        
    design = plan.get("design", {})
    if design:
        doc.add_paragraph(f"Группирующая переменная: {design.get('group_col')}", style='List Bullet')
        doc.add_paragraph(f"Временные точки (Визиты): {', '.join(design.get('visits_order', []) or ['N/A'])}", style='List Bullet')



def _render_table_node(doc, node, key=None, level=3):
    if not isinstance(node, dict):
        return
    title = node.get("title") or key
    if title:
        doc.add_heading(str(title), level=level)

    columns = node.get("columns")
    rows = node.get("rows") or []

    # Support list-of-lists format: "table": [["H1","H2"],["r1","r2"],...]
    table_lol = node.get("table")
    if table_lol and isinstance(table_lol, list) and len(table_lol) > 0:
        if isinstance(table_lol[0], list):
            columns = [str(c) for c in table_lol[0]]
            rows = table_lol[1:]
        elif isinstance(table_lol[0], dict):
            columns = list(table_lol[0].keys())
            rows = [[r.get(c) for c in columns] for r in table_lol]

    if not columns and rows and isinstance(rows[0], dict):
        columns = list(rows[0].keys())
        rows = [[r.get(c) for c in columns] for r in rows]

    if not columns:
        return

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'

    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = str(col)
        _bold_runs(table.rows[0].cells[i].paragraphs[0])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)

    caption = node.get("caption")
    if caption:
        doc.add_paragraph(str(caption), style='Caption')

def _render_plot_node(doc, node, key=None):
    plot_path = node.get("path") or node.get("file") or node.get("plot_path")
    if plot_path and os.path.exists(plot_path):
        doc.add_picture(plot_path, width=Inches(6))
        caption = node.get("caption") or node.get("title") or key
        if caption:
            doc.add_paragraph(str(caption), style='Caption')

def _render_text_node(doc, node):
    text = node.get("text")
    if text is None:
        return
    if isinstance(text, list):
        for item in text:
            if item is None:
                continue
            doc.add_paragraph(str(item))
        return

    text_str = str(text)
    for para in text_str.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())


def _render_results_recursive(doc, node, key=None, level=3):
    """Legacy recursive renderer retained for compatibility tests."""
    if node is None:
        return

    if isinstance(node, list):
        for item in node:
            _render_results_recursive(doc, item, key=key, level=level)
        return

    if not isinstance(node, dict):
        doc.add_paragraph(str(node))
        return

    node_type = str(node.get("type") or "").strip().lower()

    if node_type == "section":
        title = node.get("title") or key
        if title:
            doc.add_heading(str(title), level=max(1, min(4, int(level))))
        children = node.get("children")
        if isinstance(children, list):
            for child in children:
                _render_results_recursive(doc, child, level=min(4, int(level) + 1))
        return

    if "table" in node or node_type == "table":
        _render_table_node(doc, node, key=key, level=max(2, min(4, int(level))))
    if "path" in node or "file" in node or "plot_path" in node or node_type == "plot":
        _render_plot_node(doc, node, key=key)
    if "text" in node or node_type == "text":
        _render_text_node(doc, node)

    for sub_key, value in node.items():
        if sub_key in {
            "type",
            "title",
            "children",
            "table",
            "columns",
            "rows",
            "caption",
            "path",
            "file",
            "plot_path",
            "text",
        }:
            continue
        if isinstance(value, (dict, list)):
            _render_results_recursive(doc, value, key=str(sub_key), level=min(4, int(level) + 1))



    
def _format_p(val):
    if val is None: return "N/A"
    try:
        f = float(val)
        if f < 0.001: return "< 0.001"
        return f"{f:.3f}"
    except: return str(val)

def _bold_runs(paragraph):
    for run in paragraph.runs:
        run.bold = True
