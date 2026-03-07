import io
import json
import os
import sys
import tempfile

from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from app.modules.reporting import generate_protocol_docx_report, generate_protocol_pdf_report, render_protocol_report


def test_render_protocol_report_adds_design_warning_when_study_missing(monkeypatch):
    with tempfile.TemporaryDirectory() as tmpdir:
        monkeypatch.setenv("CLINIMETRIA_WORKSPACE_DIR", tmpdir)

        run_data = {
            "dataset_id": "dataset_no_design",
            "results": {},
        }
        html = render_protocol_report(run_data, dataset_name="Demo", style="gost")

        assert 'id="design"' in html
        assert "Предупреждение" in html
        assert "Секция дизайна" in html


def test_generate_protocol_docx_report_adds_design_warning_when_study_missing(monkeypatch):
    with tempfile.TemporaryDirectory() as tmpdir:
        monkeypatch.setenv("CLINIMETRIA_WORKSPACE_DIR", tmpdir)

        run_data = {
            "dataset_id": "dataset_no_design",
            "results": {},
            "protocol_name": "Protocol",
        }
        payload = generate_protocol_docx_report(run_data, dataset_name="Demo", style="gost")
        doc = Document(io.BytesIO(payload))
        text = "\n".join([p.text for p in doc.paragraphs if p.text])

        assert "Дизайн исследования" in text
        assert "Предупреждение" in text
        assert "Раздел Design неполный" in text


def test_design_policy_renders_human_readable_multiplicity_label(monkeypatch):
    with tempfile.TemporaryDirectory() as tmpdir:
        monkeypatch.setenv("CLINIMETRIA_WORKSPACE_DIR", tmpdir)
        dataset_id = "dataset_design_policy"
        processed_dir = os.path.join(tmpdir, "datasets", dataset_id, "processed")
        os.makedirs(processed_dir, exist_ok=True)
        with open(os.path.join(processed_dir, "study_design.json"), "w", encoding="utf-8") as f:
            json.dump(
                {
                    "design": {
                        "design_type": "parallel_groups",
                        "group_column": "group",
                        "outcomes": ["y"],
                    },
                    "analysis_policy": {
                        "alpha": 0.05,
                        "multiplicity_correction": "fdr_bh",
                        "post_hoc": "tukey",
                        "bootstrap_ci": True,
                        "bootstrap_samples": 1500,
                    },
                },
                f,
                ensure_ascii=False,
                indent=2,
            )

        run_data = {
            "dataset_id": dataset_id,
            "protocol_name": "Policy Protocol",
            "results": {},
        }
        html = render_protocol_report(run_data, dataset_name="Demo", style="gost")
        assert "FDR (Benjamini-Hochberg)" in html
        assert "[fdr_bh]" not in html
        assert "Bootstrap samples" in html
        assert "1500" in html

        payload = generate_protocol_docx_report(run_data, dataset_name="Demo", style="gost")
        doc = Document(io.BytesIO(payload))
        text = "\n".join([p.text for p in doc.paragraphs if p.text])
        assert "Поправка множественных сравнений" in text
        assert "FDR (Benjamini-Hochberg)" in text
        assert "[fdr_bh]" not in text
        assert "Bootstrap samples" in text
        assert "1500" in text


def test_generate_protocol_docx_report_includes_protocol_validation_provenance(monkeypatch):
    with tempfile.TemporaryDirectory() as tmpdir:
        monkeypatch.setenv("CLINIMETRIA_WORKSPACE_DIR", tmpdir)

        run_data = {
            "dataset_id": "dataset_validation_docx",
            "run_id": "run_validation_docx_1",
            "results": {},
            "protocol_validation": {
                "schema": "clinimetria.protocol_validation",
                "status": "failed",
                "enabled": True,
                "strict": True,
                "policy_profile": "focused",
                "policy": {
                    "profile": "focused",
                    "multiplicity_correction": "fdr_bh",
                    "post_hoc_correction": "holm",
                    "repair_correction": "holm",
                    "reflection_enabled": False,
                },
                "multiplicity_policy": {
                    "enabled": True,
                    "correction": "fdr_bh",
                    "post_hoc_correction": "holm",
                    "n_applied_steps": 1,
                },
                "summary": {
                    "steps_total": 2,
                    "steps_checked": 2,
                    "steps_failed": 1,
                    "global_errors": 1,
                },
                "steps": [
                    {"step_id": "s1", "status": "failed", "warnings": ["w1"], "errors": ["e1"]},
                    {"step_id": "s2", "status": "passed", "warnings": [], "errors": []},
                ],
            },
        }

        payload = generate_protocol_docx_report(run_data, dataset_name="Demo", style="gost")
        doc = Document(io.BytesIO(payload))
        text = "\n".join([p.text for p in doc.paragraphs if p.text])

        assert "Валидация протокола" in text
        assert "status=ошибка" in text
        assert "Сводка валидации протокола" in text
        assert "warnings=1" in text
        assert "Политика валидации" in text
        assert "Multiplicity-политика" in text
        assert "FDR (Benjamini-Hochberg)" in text
        assert "Проблемные шаги" in text


def test_generate_protocol_docx_report_includes_bootstrap_trace():
    run_data = {
        "dataset_id": "dataset_bootstrap_docx",
        "results": {
            "reg_step": {
                "type": "regression",
                "method": {"id": "linear_regression", "name": "linear_regression"},
                "p_value": 0.02,
                "coefficients": [
                    {"variable": "x1", "coefficient": 0.5, "std_err": 0.1, "p_value": 0.004, "ci_lower": 0.3, "ci_upper": 0.7}
                ],
                "bootstrap": {
                    "enabled": True,
                    "method": "bootstrap_percentile",
                    "samples": 250,
                    "ci_level": 0.95,
                    "metrics": {
                        "coefficients": [
                            {"variable": "x1", "estimate": 0.5, "ci_lower": 0.35, "ci_upper": 0.65, "n_valid": 240}
                        ]
                    },
                },
            }
        },
    }

    payload = generate_protocol_docx_report(run_data, dataset_name="Demo", style="gost")
    doc = Document(io.BytesIO(payload))
    text = "\n".join([p.text for p in doc.paragraphs if p.text])

    assert "Bootstrap-трассировка" in text
    assert "samples=250" in text


def test_generate_protocol_docx_report_includes_hypothesis_discovery_trace():
    run_data = {
        "dataset_id": "dataset_hypothesis_docx",
        "hypotheses": {
            "schema": "clinimetria.hypothesis_discovery",
            "analysis_mode": "publication",
            "design_type": "parallel_groups",
            "count": 1,
            "items": [
                {
                    "id": "h_group_numeric",
                    "title": "Сравнить outcome между group",
                    "h0": "H0: различий нет",
                    "h1": "H1: различия есть",
                    "suggested_method": "t_test_ind / mann_whitney / anova",
                    "priority": "high",
                }
            ],
        },
        "step_meta": {
            "s1": {"id": "s1", "method": "t_test_ind", "config": {"outcome": "outcome", "group": "group"}},
        },
        "results": {
            "s1": {
                "type": "hypothesis_test",
                "method": {"id": "t_test_ind", "name": "t_test_ind"},
                "p_value": 0.02,
            }
        },
    }

    payload = generate_protocol_docx_report(run_data, dataset_name="Demo", style="gost")
    doc = Document(io.BytesIO(payload))
    text = "\n".join([p.text for p in doc.paragraphs if p.text])

    assert "Гипотезы и проверяемые утверждения" in text
    assert "Сравнить outcome между group" in text
    assert "H0: различий нет" in text
    assert "H1: различия есть" in text
    assert "Вердикт:" in text
    assert "подтверждена" in text


def test_protocol_reports_include_time_series_quality_trace_docx_pdf():
    run_data = {
        "dataset_id": "dataset_ts_quality",
        "results": {
            "ts_step": {
                "type": "time_series",
                "method": {"id": "time_series_analysis", "name": "Time Series"},
                "time_axis_kind": "datetime",
                "p_value": 0.02,
                "trend": {"slope": 0.18},
                "diagnostics": {"ljung_box": {"p_value": 0.03, "white_noise_like": False}},
                "forecast": {"points": [{"x": "2026-01-04", "y": 12.4}]},
                "time_quality": {
                    "quality": "warning",
                    "datetime_parse_ratio": 1.0,
                    "min_year": 1970,
                    "max_year": 1970,
                    "inferred_frequency": "D",
                    "flags": ["epoch_artifact_risk"],
                },
                "warnings": [
                    "Calendar years are concentrated in 1970-1985; verify date parsing to avoid Unix-epoch artifacts."
                ],
                "plot_data": [
                    {"x": "2026-01-01", "y": 10.0, "trend": 9.9},
                    {"x": "2026-01-02", "y": 10.7, "trend": 10.3},
                    {"x": "2026-01-03", "y": 11.4, "trend": 10.8},
                ],
                "plot_config": {"type": "line", "x_label": "date", "y_label": "value"},
            }
        },
    }

    docx_payload = generate_protocol_docx_report(run_data, dataset_name="Demo", style="gost")
    doc = Document(io.BytesIO(docx_payload))
    docx_parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    docx_parts.append(cell.text)
    docx_text = "\n".join(docx_parts)

    assert "Диагностика временного ряда" in docx_text
    assert "Качество временной оси" in docx_text
    assert "epoch_artifact_risk" in docx_text
    assert "Предупреждения по хронологии" in docx_text

    pdf_payload = generate_protocol_pdf_report(run_data, dataset_name="Demo", style="apa7")
    assert pdf_payload[:4] == b"%PDF"
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(pdf_payload))
    pdf_text = "\n".join((page.extract_text() or "") for page in reader.pages)

    assert "Series diagnostics" in pdf_text
    assert "Time axis quality" in pdf_text
    assert "Chronology warnings" in pdf_text
    assert "epoch_artifact_risk" in pdf_text
