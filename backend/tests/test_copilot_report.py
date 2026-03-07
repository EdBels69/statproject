"""
Tests for Copilot Report Generator - DOCX generation.
"""

import pytest
import os
import sys
from io import BytesIO

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from app.copilot.report import (
    generate_report,
    _format_p,
    _render_results_recursive,
)


class TestReportFormatters:
    """Tests for report formatting utilities."""

    def test_format_p_none(self):
        """Test p-value formatting with None."""
        assert _format_p(None) == "N/A"

    def test_format_p_very_small(self):
        """Test p-value formatting for p < 0.001."""
        assert _format_p(0.0001) == "< 0.001"
        assert _format_p(0.00001) == "< 0.001"

    def test_format_p_normal(self):
        """Test p-value formatting for normal values."""
        assert _format_p(0.05) == "0.050"
        assert _format_p(0.123) == "0.123"

    def test_format_p_edge_cases(self):
        """Test p-value formatting edge cases."""
        assert _format_p(1.0) == "1.000"
        assert _format_p(0.001) == "0.001"


class TestReportGeneration:
    """Tests for DOCX report generation."""

    @pytest.fixture
    def sample_plan(self):
        """Sample analysis plan."""
        return {
            "understood_goal": "Analyze treatment effects on patient outcomes",
            "design": {
                "group_col": "Treatment",
                "visits_order": ["V1", "V2", "V3"]
            }
        }

    @pytest.fixture
    def sample_results(self):
        """Sample analysis results."""
        return {
            "demographics": {
                "descriptive": {
                    "Control": {"n": 50, "mean": 45.2, "sd": 12.3, "median": 44.0, "iqr": 15.0, "shapiro_p": 0.12, "normal": True},
                    "Treatment": {"n": 48, "mean": 46.8, "sd": 11.9, "median": 45.5, "iqr": 14.2, "shapiro_p": 0.08, "normal": True}
                },
                "omnibus": {
                    "method": "ANOVA",
                    "p_value": 0.42,
                    "statistic": 0.65,
                    "effect_size": 0.007,
                    "significant": False
                },
                "plots": []
            }
        }

    @pytest.fixture
    def sample_dataset_info(self):
        """Sample dataset info."""
        return {
            "filename": "clinical_trial_data.xlsx",
            "n_rows": 98,
            "n_cols": 25
        }

    def test_generate_report_returns_bytes(self, sample_results, sample_plan, sample_dataset_info):
        """Test that generate_report returns bytes (DOCX content)."""
        result = generate_report(
            results=sample_results,
            plan=sample_plan,
            dataset_info=sample_dataset_info
        )
        
        assert isinstance(result, bytes)
        assert len(result) > 0
        # DOCX files start with PK (zip magic bytes)
        assert result[:2] == b'PK'

    def test_generate_report_with_code(self, sample_results, sample_plan, sample_dataset_info):
        """Test report generation with Python code included."""
        test_code = "import pandas as pd\ndf = pd.read_parquet('data.parquet')"
        
        result = generate_report(
            results=sample_results,
            plan=sample_plan,
            code=test_code,
            dataset_info=sample_dataset_info
        )
        
        assert isinstance(result, bytes)
        assert len(result) > 1000  # Should be larger with code

    def test_generate_report_with_interpretation(self, sample_results, sample_plan, sample_dataset_info):
        """Test report generation with AI interpretation."""
        interpretation = """
        Результаты анализа показывают отсутствие статистически значимых различий 
        между группами лечения (p = 0.42). Размер эффекта минимален (η² = 0.007).
        """
        
        result = generate_report(
            results=sample_results,
            plan=sample_plan,
            interpretation=interpretation,
            dataset_info=sample_dataset_info
        )
        
        assert isinstance(result, bytes)

    def test_generate_report_empty_results(self, sample_plan, sample_dataset_info):
        """Test report generation with empty results."""
        result = generate_report(
            results={},
            plan=sample_plan,
            dataset_info=sample_dataset_info
        )
        
        assert isinstance(result, bytes)
        # Should still produce valid DOCX
        assert result[:2] == b'PK'

    def test_generate_report_with_mixed_model_results(self, sample_plan, sample_dataset_info):
        """Test report generation with mixed model results."""
        results = {
            "longitudinal": {
                "mixed_effects": {
                    "Score": {
                        "method": "Linear Mixed Effects (RI)",
                        "formula": "Value ~ Group * Visit + (1|Subject)",
                        "n_observations": 294,
                        "p_interaction": 0.023,
                        "significant_interaction": True,
                        "converged": True
                    }
                }
            }
        }
        
        result = generate_report(
            results=results,
            plan=sample_plan,
            dataset_info=sample_dataset_info
        )
        
        assert isinstance(result, bytes)

    def test_generate_report_with_responder_analysis(self, sample_plan, sample_dataset_info):
        """Test report generation with responder analysis."""
        results = {
            "responders": {
                "type": "responder_analysis",
                "threshold_pct": 20,
                "data": [
                    {"group": "Control", "n_responders": 15, "n_total": 50, "pct_responders": 30.0},
                    {"group": "Treatment", "n_responders": 28, "n_total": 48, "pct_responders": 58.3}
                ],
                "p_value": 0.008,
                "plots": []
            }
        }
        
        result = generate_report(
            results=results,
            plan=sample_plan,
            dataset_info=sample_dataset_info
        )
        
        assert isinstance(result, bytes)


class TestRenderResultsRecursive:
    """Tests for recursive results rendering."""

    def test_render_handles_none(self):
        """Test that None results don't crash."""
        from docx import Document
        doc = Document()
        # Should not raise
        _render_results_recursive(doc, None)

    def test_render_handles_list(self):
        """Test rendering of list results."""
        from docx import Document
        doc = Document()
        results = [
            {"type": "text", "text": "Item 1"},
            {"type": "text", "text": "Item 2"}
        ]
        # Should not raise
        _render_results_recursive(doc, results)

    def test_render_handles_nested_sections(self):
        """Test rendering of nested section structure."""
        from docx import Document
        doc = Document()
        results = {
            "type": "section",
            "title": "Main Section",
            "children": [
                {"type": "text", "text": "Some content"},
                {
                    "type": "section",
                    "title": "Subsection",
                    "children": [
                        {"type": "text", "text": "Nested content"}
                    ]
                }
            ]
        }
        # Should not raise
        _render_results_recursive(doc, results)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
