#!/usr/bin/env python3
"""
Structured COVID report focused on outcomes and glycemia.

Goals:
- deterministic analysis pipeline (no LLM dependency),
- explicit coverage of treatment and comorbidity blocks,
- figures with per-figure interpretation,
- coherent report with linked sections.
"""

from __future__ import annotations

import argparse
import json
import math
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
import statsmodels.api as sm
from docx import Document
from docx.shared import Inches
from scipy import stats
from sklearn.metrics import brier_score_loss, roc_auc_score, roc_curve


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_INPUT = PROJECT_ROOT / "docs" / "Общая таблица Ковид19.xlsx"
DEFAULT_OUTPUT_ROOT = PROJECT_ROOT / "docs" / "exports"


NUMERIC_SPECS: List[Tuple[str, str, str]] = [
    ("age", "Возраст, лет", "возраст"),
    ("hosp_days", "Длительность госпитализации, дней", "Длительность госпитализации"),
    ("icu_days", "Дней в ОРИТ", "Дней в ОРИТ"),
    ("spo2", "SpO2, %", "SpO2 %"),
    ("news2", "NEWS2", "NEWS2"),
    ("qsofa", "qSOFA", "qSOFA"),
    ("crp1", "СРБ1", "СРБ1"),
    ("fibrinogen1", "Фибриноген1", "фибриноген1"),
    ("creatinine1", "Креатинин1", "креатинин1"),
    ("glucose_adm", "Глюкоза при поступлении", "Глюкоза при поступлении"),
    ("glucose_last", "Глюкоза последний результат", "Глюкоза последний результат"),
]

GLYCEMIA_BINARY_SPECS: List[Tuple[str, str, str]] = [
    ("hypergly_twice", "Гипергликемия >11.1 дважды", "Гипергликемия >11,1 дважды по результатам бх крови"),
    ("dm2_pre", "СД2 до госпитализации", "Сахарный диабет 2 типа перед госпитализацией (да/нет)"),
    ("insulin_tx", "Инсулинотерапия", "Инсулинотерапия (да, нет)"),
]

COMORBIDITY_SPECS: List[Tuple[str, str, str]] = [
    ("dm2_pre", "СД2 до госпитализации", "Сахарный диабет 2 типа перед госпитализацией (да/нет)"),
    ("hypertension", "Артериальная гипертензия", "Гипертоническая болезнь (да/нет)"),
    ("ibs", "ИБС", "ИБС (да/нет)"),
    ("obesity", "Ожирение", "Ожирение"),
    ("oncology", "Онкология", "Онкология (да/нет)"),
    ("anemia", "Анемия", "Анемия (да/нет)"),
    ("chronic_pyelonephritis", "Хронический пиелонефрит", "Хронический пиелонефрит"),
    ("icu", "Поступление/лечение в ОРИТ", "В ОРИТ"),
]

TREATMENT_SPECS: List[Tuple[str, str, str]] = [
    ("abx_pre", "Антибиотики до госпитализации", "антибиотики перед госпитализацией"),
    ("abx_hosp", "Антибиотики в госпитализации", "антибиотики во время госпитализации"),
    ("anticoag", "Антикоагулянты в госпитализации", "антикоагулянты во время госпитализаци (да/нет)"),
    ("gcs_pre", "ГКС до госпитализации", "ГКС перед госпитализацией (да/нет)"),
    ("gcs_hosp", "ГКС в госпитализации", "ГКС во время госпитализации (да/нет)"),
    ("insulin_tx", "Инсулинотерапия", "Инсулинотерапия (да, нет)"),
    ("anti_cytokine", "Антицитокиновая терапия", "Антицитокиновая терапия (да/нет)"),
    ("eufillin", "Эуфиллин", "Эуфиллин"),
    ("pentoxy", "Пентоксифиллин", "Пентоксифиллин"),
    ("statins", "Статины", "Статины"),
    ("beta_blockers", "Бетаблокаторы", "Бетаблокаторы"),
    ("vitamin_c", "Витамин C", "Витамин С"),
]

ADDITIONAL_BINARY_SPECS: List[Tuple[str, str, str]] = [
    (
        "dm_discharge",
        "СД при выписке (по диагнозу ИБ)",
        "Сахарный диабет при выписке (да/нет) указанный так в диагнозе ИБ",
    ),
]

VARIABLE_LABELS: Dict[str, str] = {}
for _specs in (
    NUMERIC_SPECS,
    GLYCEMIA_BINARY_SPECS,
    COMORBIDITY_SPECS,
    TREATMENT_SPECS,
    ADDITIONAL_BINARY_SPECS,
):
    for _key, _label, _source in _specs:
        if _key not in VARIABLE_LABELS:
            VARIABLE_LABELS[_key] = _label
VARIABLE_LABELS.setdefault("sex_male", "Пол (мужской)")
VARIABLE_LABELS.setdefault("glucose_dm2_interaction", "Глюкоза при поступлении × СД2")
VARIABLE_LABELS.setdefault("death", "Летальный исход")

NUMERIC_KEYS = {k for k, _, _ in NUMERIC_SPECS}
BINARY_KEYS = {k for k, _, _ in (GLYCEMIA_BINARY_SPECS + COMORBIDITY_SPECS + TREATMENT_SPECS + ADDITIONAL_BINARY_SPECS)}
BINARY_KEYS.update({"sex_male"})


def _norm_text(value: Any) -> str:
    if value is None:
        return ""
    s = str(value).strip().lower()
    s = s.replace("\xa0", " ")
    s = re.sub(r"\s+", " ", s)
    return s


def _parse_float(value: Any) -> float:
    if value is None:
        return np.nan
    if isinstance(value, (int, float, np.integer, np.floating)):
        f = float(value)
        if math.isnan(f) or math.isinf(f):
            return np.nan
        return f
    s = _norm_text(value)
    if not s or s in {"nan", "none", "na", "n/a", "-", "--"}:
        return np.nan
    s = s.replace(",", ".")
    match = re.search(r"[-+]?\d*\.?\d+", s)
    if not match:
        return np.nan
    try:
        return float(match.group(0))
    except Exception:
        return np.nan


def _to_numeric(series: pd.Series) -> pd.Series:
    return series.map(_parse_float).astype(float)


def _to_binary(series: pd.Series) -> pd.Series:
    positive = {
        "да",
        "yes",
        "true",
        "1",
        "y",
        "имеется",
        "положительный",
        "полож",
    }
    negative = {
        "нет",
        "no",
        "false",
        "0",
        "n",
        "отрицательный",
        "отриц",
        "не было",
        "отсутствует",
    }
    out: List[float] = []
    for value in series.tolist():
        t = _norm_text(value)
        if not t or t in {"nan", "none", "na", "n/a", "-", "--"}:
            out.append(np.nan)
            continue

        if t in positive or t.startswith("да"):
            out.append(1.0)
            continue
        if t in negative or t.startswith("нет"):
            out.append(0.0)
            continue

        numeric = _parse_float(value)
        if not np.isnan(numeric):
            out.append(1.0 if numeric > 0 else 0.0)
            continue

        # Fallback for text-coded treatments/comorbidities:
        # any non-empty non-negative token is considered "present".
        out.append(1.0)
    return pd.Series(out, index=series.index, dtype="float64")


def _sex_to_male(series: pd.Series) -> pd.Series:
    out: List[float] = []
    for value in series.tolist():
        t = _norm_text(value)
        if t in {"м", "male", "man", "муж", "мужской", "1"}:
            out.append(1.0)
        elif t in {"ж", "female", "woman", "жен", "женский", "0"}:
            out.append(0.0)
        else:
            numeric = _parse_float(value)
            if np.isnan(numeric):
                out.append(np.nan)
            else:
                out.append(1.0 if numeric == 1 else 0.0)
    return pd.Series(out, index=series.index, dtype="float64")


def _build_death_outcome(df: pd.DataFrame) -> pd.Series:
    death = pd.Series(np.nan, index=df.index, dtype="float64")

    primary = df["Исход"] if "Исход" in df.columns else pd.Series(np.nan, index=df.index)
    for idx, value in primary.items():
        t = _norm_text(value)
        if not t:
            continue
        if "мертв" in t or "умер" in t or "неблагоприят" in t:
            death.at[idx] = 1.0
        elif "выписан" in t or "благоприят" in t:
            death.at[idx] = 0.0

    # fallback using secondary endpoint coding
    if "Исход.1" in df.columns:
        sec = df["Исход.1"]
        for idx, value in sec.items():
            if not np.isnan(death.at[idx]):
                continue
            t = _norm_text(value)
            if "неблагоприят" in t:
                death.at[idx] = 1.0
            elif "благоприят" in t:
                death.at[idx] = 0.0

    if "Исход.2" in df.columns:
        sec2 = df["Исход.2"]
        for idx, value in sec2.items():
            if not np.isnan(death.at[idx]):
                continue
            t = _norm_text(value)
            if "мертв" in t or "умер" in t:
                death.at[idx] = 1.0
            elif "выписан" in t:
                death.at[idx] = 0.0

    return death


def _format_p(p: Any) -> str:
    if p is None:
        return "-"
    try:
        p = float(p)
    except Exception:
        return "-"
    if np.isnan(p):
        return "-"
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


def _median_iqr(series: pd.Series) -> str:
    clean = series.dropna().astype(float)
    if clean.empty:
        return "NA"
    return f"{clean.median():.2f} [{clean.quantile(0.25):.2f}; {clean.quantile(0.75):.2f}]"


def _pct(numerator: int, denominator: int) -> str:
    if denominator <= 0:
        return "NA"
    return f"{100.0 * numerator / denominator:.1f}%"


def _safe_or(a: int, b: int, c: int, d: int) -> float:
    # Haldane-Anscombe correction for zero cells
    a2, b2, c2, d2 = a + 0.5, b + 0.5, c + 0.5, d + 0.5
    return (a2 * d2) / (b2 * c2)


def _fisher_p(a: int, b: int, c: int, d: int) -> float:
    try:
        _, p = stats.fisher_exact([[a, b], [c, d]])
        return float(p)
    except Exception:
        return np.nan


def _prepare_analysis_frame(raw: pd.DataFrame) -> Tuple[pd.DataFrame, Dict[str, str]]:
    raw = raw.copy()
    raw.columns = [str(c).strip() for c in raw.columns]

    frame = pd.DataFrame(index=raw.index)
    source_map: Dict[str, str] = {}

    frame["death"] = _build_death_outcome(raw)
    source_map["death"] = "Исход / Исход.1 / Исход.2"

    if "пол" in raw.columns:
        frame["sex_male"] = _sex_to_male(raw["пол"])
        source_map["sex_male"] = "пол"

    for key, _, source in NUMERIC_SPECS:
        if source in raw.columns:
            frame[key] = _to_numeric(raw[source])
            source_map[key] = source

    for key, _, source in GLYCEMIA_BINARY_SPECS + COMORBIDITY_SPECS + TREATMENT_SPECS + ADDITIONAL_BINARY_SPECS:
        if source in raw.columns and key not in frame.columns:
            frame[key] = _to_binary(raw[source])
            source_map[key] = source

    frame["glucose_adm_cat"] = pd.cut(
        frame.get("glucose_adm"),
        bins=[-np.inf, 7.8, 11.1, np.inf],
        labels=["<7.8", "7.8-11.1", ">=11.1"],
        right=False,
    )
    frame["glucose_last_cat"] = pd.cut(
        frame.get("glucose_last"),
        bins=[-np.inf, 7.8, 11.1, np.inf],
        labels=["<7.8", "7.8-11.1", ">=11.1"],
        right=False,
    )

    frame["glucose_dm2_interaction"] = frame.get("glucose_adm") * frame.get("dm2_pre")
    return frame, source_map


def _numeric_table(frame: pd.DataFrame, specs: Sequence[Tuple[str, str, str]]) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    alive_total = int((frame["death"] == 0).sum())
    dead_total = int((frame["death"] == 1).sum())

    for key, label, _ in specs:
        if key not in frame.columns:
            continue
        alive = frame.loc[frame["death"] == 0, key]
        dead = frame.loc[frame["death"] == 1, key]
        alive_n = int(alive.notna().sum())
        dead_n = int(dead.notna().sum())

        p_value = np.nan
        test_name = "-"
        if alive_n >= 8 and dead_n >= 8:
            try:
                _, p_value = stats.mannwhitneyu(alive.dropna(), dead.dropna(), alternative="two-sided")
                test_name = "Mann-Whitney U"
            except Exception:
                p_value = np.nan
                test_name = "Mann-Whitney U"

        rows.append(
            {
                "Показатель": label,
                f"Выписан (n={alive_total})": f"{_median_iqr(alive)}; n={alive_n}",
                f"Летальный исход (n={dead_total})": f"{_median_iqr(dead)}; n={dead_n}",
                "p": _format_p(p_value),
                "Тест": test_name,
                "_p_raw": p_value,
                "_alive_median": float(alive.dropna().median()) if alive_n else np.nan,
                "_dead_median": float(dead.dropna().median()) if dead_n else np.nan,
            }
        )
    out = pd.DataFrame(rows)
    return out


def _binary_or_table(
    frame: pd.DataFrame, specs: Sequence[Tuple[str, str, str]], table_name: str
) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []

    for key, label, _ in specs:
        if key not in frame.columns:
            continue
        sub = frame[["death", key]].dropna()
        if sub.empty:
            continue
        alive_total = int((sub["death"] == 0).sum())
        dead_total = int((sub["death"] == 1).sum())
        if alive_total <= 0 or dead_total <= 0:
            continue

        exposed_dead = int(((sub["death"] == 1) & (sub[key] == 1)).sum())
        exposed_alive = int(((sub["death"] == 0) & (sub[key] == 1)).sum())
        unexp_dead = int(((sub["death"] == 1) & (sub[key] == 0)).sum())
        unexp_alive = int(((sub["death"] == 0) & (sub[key] == 0)).sum())

        if (exposed_dead + exposed_alive + unexp_dead + unexp_alive) < 30:
            continue

        odds_ratio = _safe_or(exposed_dead, exposed_alive, unexp_dead, unexp_alive)
        p_value = _fisher_p(exposed_dead, exposed_alive, unexp_dead, unexp_alive)

        missing_alive = int(((frame["death"] == 0) & (frame[key].isna())).sum())
        missing_dead = int(((frame["death"] == 1) & (frame[key].isna())).sum())
        zero_cell = int(exposed_dead == 0 or exposed_alive == 0 or unexp_dead == 0 or unexp_alive == 0)

        rows.append(
            {
                "Показатель": label,
                "N (available: alive/dead)": f"{alive_total}/{dead_total}",
                "Missing (alive/dead)": f"{missing_alive}/{missing_dead}",
                "Экспозиция среди выписанных": f"{exposed_alive}/{alive_total} ({_pct(exposed_alive, alive_total)})",
                "Экспозиция при летальном исходе": f"{exposed_dead}/{dead_total} ({_pct(exposed_dead, dead_total)})",
                "OR (crude)": round(float(odds_ratio), 3),
                "p (Fisher)": _format_p(p_value),
                "_p_raw": p_value,
                "_or_raw": odds_ratio,
                "_zero_cell": zero_cell,
                "_missing_total": missing_alive + missing_dead,
                "_available_total": alive_total + dead_total,
                "_domain": table_name,
            }
        )
    return pd.DataFrame(rows)


def _fit_univariate_logit(frame: pd.DataFrame, predictors: Sequence[Tuple[str, str]]) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    for key, label in predictors:
        if key not in frame.columns:
            continue
        sub = frame[["death", key]].dropna().copy()
        if len(sub) < 80 or sub["death"].nunique() < 2 or sub[key].nunique() < 2:
            continue
        sub[key] = sub[key].astype(float)
        try:
            X = sm.add_constant(sub[[key]], has_constant="add")
            model = sm.Logit(sub["death"].astype(float), X).fit(disp=False, maxiter=200)
            coef = float(model.params[key])
            ci_low, ci_high = model.conf_int().loc[key].tolist()
            p_val = float(model.pvalues[key])
            rows.append(
                {
                    "Предиктор": label,
                    "N": int(len(sub)),
                    "OR": round(float(np.exp(coef)), 3),
                    "95% CI": f"{np.exp(ci_low):.3f}–{np.exp(ci_high):.3f}",
                    "p": _format_p(p_val),
                    "_p_raw": p_val,
                    "_or_raw": float(np.exp(coef)),
                    "_coef_raw": coef,
                }
            )
        except Exception:
            continue
    return pd.DataFrame(rows)


def _paired_glucose_delta_table(frame: pd.DataFrame) -> Tuple[pd.DataFrame, float]:
    sub = frame[["death", "glucose_adm", "glucose_last"]].dropna().copy()
    if sub.empty or sub["death"].nunique() < 2:
        return pd.DataFrame(), np.nan

    sub["delta_glucose"] = sub["glucose_last"].astype(float) - sub["glucose_adm"].astype(float)

    rows: List[Dict[str, Any]] = []
    for outcome_value, label in [(0.0, "Выписан"), (1.0, "Летальный исход")]:
        g = sub[sub["death"] == outcome_value].copy()
        if g.empty:
            continue
        delta = g["delta_glucose"].astype(float)
        n = int(len(g))

        p_within = np.nan
        if n >= 10 and delta.nunique(dropna=True) > 1:
            try:
                _, p_within = stats.wilcoxon(delta, alternative="two-sided", zero_method="wilcox")
            except Exception:
                p_within = np.nan

        up = int((delta > 0).sum())
        rows.append(
            {
                "Группа": label,
                "N (pairs)": n,
                "Глюкоза при поступлении, median [IQR]": _median_iqr(g["glucose_adm"]),
                "Глюкоза последний результат, median [IQR]": _median_iqr(g["glucose_last"]),
                "Δ (last - admission), median [IQR]": _median_iqr(delta),
                "Δ>0, n (%)": f"{up} ({_pct(up, n)})",
                "p (within, Wilcoxon Δ vs 0)": _format_p(p_within),
                "_p_within_raw": p_within,
            }
        )

    p_between = np.nan
    d0 = sub[sub["death"] == 0.0]["delta_glucose"].astype(float)
    d1 = sub[sub["death"] == 1.0]["delta_glucose"].astype(float)
    if len(d0) >= 10 and len(d1) >= 10 and (d0.nunique(dropna=True) > 1 or d1.nunique(dropna=True) > 1):
        try:
            _, p_between = stats.mannwhitneyu(d0, d1, alternative="two-sided")
        except Exception:
            p_between = np.nan

    out = pd.DataFrame(rows)
    return out, float(p_between) if not np.isnan(p_between) else np.nan


@dataclass
class ModelResult:
    outcome_key: str
    outcome_label: str
    model_id: str
    title: str
    predictors: List[str]
    n: int
    events: int
    aic: float
    auc: float
    brier: float
    coef_table: pd.DataFrame
    y_true: np.ndarray
    y_pred: np.ndarray


def _fit_multivariable_model(
    frame: pd.DataFrame,
    outcome_key: str,
    outcome_label: str,
    model_id: str,
    title: str,
    predictors: Sequence[str],
) -> Optional[ModelResult]:
    cols = [outcome_key] + [p for p in predictors if p in frame.columns]
    if len(cols) <= 2:
        return None
    sub = frame[cols].dropna().copy()
    if len(sub) < 100:
        return None
    if sub[outcome_key].nunique() < 2:
        return None

    valid_predictors: List[str] = []
    for p in cols[1:]:
        if sub[p].nunique() >= 2:
            valid_predictors.append(p)
    if len(valid_predictors) < 2:
        return None

    try:
        X = sm.add_constant(sub[valid_predictors].astype(float), has_constant="add")
        y = sub[outcome_key].astype(float)
        model = sm.Logit(y, X).fit(disp=False, maxiter=300)
    except Exception:
        return None

    pred = model.predict(X)
    try:
        auc = float(roc_auc_score(y, pred))
    except Exception:
        auc = np.nan
    try:
        brier = float(brier_score_loss(y, pred))
    except Exception:
        brier = np.nan

    rows: List[Dict[str, Any]] = []
    conf = model.conf_int()
    for var in valid_predictors:
        coef = float(model.params[var])
        ci_low = float(conf.loc[var, 0])
        ci_high = float(conf.loc[var, 1])
        p = float(model.pvalues[var])
        rows.append(
            {
                "_outcome": outcome_label,
                "Model": model_id,
                "_term": var,
                "Предиктор": var,
                "OR": float(np.exp(coef)),
                "CI_low": float(np.exp(ci_low)),
                "CI_high": float(np.exp(ci_high)),
                "p_raw": p,
                "p": _format_p(p),
            }
        )
    coef_table = pd.DataFrame(rows)
    if not coef_table.empty:
        coef_table["Предиктор"] = coef_table["_term"].map(lambda t: VARIABLE_LABELS.get(str(t), str(t)))

    return ModelResult(
        outcome_key=outcome_key,
        outcome_label=outcome_label,
        model_id=model_id,
        title=title,
        predictors=valid_predictors,
        n=int(len(sub)),
        events=int(y.sum()),
        aic=float(model.aic),
        auc=auc,
        brier=brier,
        coef_table=coef_table,
        y_true=np.asarray(y),
        y_pred=np.asarray(pred),
    )


def _fit_multivariable_model_imputed(
    frame: pd.DataFrame,
    outcome_key: str,
    outcome_label: str,
    model_id: str,
    title: str,
    predictors: Sequence[str],
) -> Optional[ModelResult]:
    cols = [outcome_key] + [p for p in predictors if p in frame.columns]
    if len(cols) <= 2:
        return None
    sub = frame[cols].copy()
    sub = sub.dropna(subset=[outcome_key]).copy()
    if len(sub) < 100:
        return None
    if sub[outcome_key].nunique() < 2:
        return None

    for p in cols[1:]:
        if p in NUMERIC_KEYS:
            med = float(sub[p].dropna().astype(float).median()) if sub[p].notna().any() else 0.0
            if np.isnan(med) or np.isinf(med):
                med = 0.0
            sub[p] = sub[p].astype(float).fillna(med)
        else:
            vals = sub[p].dropna().astype(float)
            if vals.empty:
                fill = 0.0
            else:
                mode = vals.mode()
                fill = float(mode.iloc[0]) if not mode.empty else float(round(float(vals.mean())))
            sub[p] = sub[p].astype(float).fillna(fill)

    valid_predictors: List[str] = []
    for p in cols[1:]:
        if sub[p].nunique() >= 2:
            valid_predictors.append(p)
    if len(valid_predictors) < 2:
        return None

    try:
        X = sm.add_constant(sub[valid_predictors].astype(float), has_constant="add")
        y = sub[outcome_key].astype(float)
        model = sm.Logit(y, X).fit(disp=False, maxiter=300)
    except Exception:
        return None

    pred = model.predict(X)
    try:
        auc = float(roc_auc_score(y, pred))
    except Exception:
        auc = np.nan
    try:
        brier = float(brier_score_loss(y, pred))
    except Exception:
        brier = np.nan

    rows: List[Dict[str, Any]] = []
    conf = model.conf_int()
    for var in valid_predictors:
        coef = float(model.params[var])
        ci_low = float(conf.loc[var, 0])
        ci_high = float(conf.loc[var, 1])
        p_val = float(model.pvalues[var])
        rows.append(
            {
                "_outcome": outcome_label,
                "Model": model_id,
                "_term": var,
                "Предиктор": VARIABLE_LABELS.get(var, var),
                "OR": float(np.exp(coef)),
                "CI_low": float(np.exp(ci_low)),
                "CI_high": float(np.exp(ci_high)),
                "p_raw": p_val,
                "p": _format_p(p_val),
            }
        )
    coef_table = pd.DataFrame(rows)

    return ModelResult(
        outcome_key=outcome_key,
        outcome_label=outcome_label,
        model_id=model_id,
        title=title,
        predictors=valid_predictors,
        n=int(len(sub)),
        events=int(y.sum()),
        aic=float(model.aic),
        auc=auc,
        brier=brier,
        coef_table=coef_table,
        y_true=np.asarray(y),
        y_pred=np.asarray(pred),
    )


def _plot_missingness(frame: pd.DataFrame, vars_for_plot: List[str], out_path: Path) -> pd.DataFrame:
    valid = [c for c in vars_for_plot if c in frame.columns]
    miss = frame[valid].isna().mean().mul(100).sort_values(ascending=False)
    miss_df = miss.reset_index()
    miss_df.columns = ["variable", "missing_pct"]

    plt.figure(figsize=(9, 6))
    sns.barplot(data=miss_df, x="missing_pct", y="variable", color="#457b9d")
    plt.xlabel("Missing, %")
    plt.ylabel("Variable")
    plt.title("Доля пропусков по ключевым переменным")
    plt.tight_layout()
    plt.savefig(out_path, dpi=170)
    plt.close()
    return miss_df


def _plot_glucose_category_mortality(frame: pd.DataFrame, category_col: str, title: str, out_path: Path) -> pd.DataFrame:
    sub = frame[[category_col, "death"]].dropna()
    tab = (
        sub.groupby(category_col, observed=False)["death"]
        .agg(n="count", mortality="mean")
        .reset_index()
        .rename(columns={category_col: "category"})
    )
    tab["mortality_pct"] = tab["mortality"] * 100.0

    plt.figure(figsize=(7, 4.8))
    sns.barplot(data=tab, x="category", y="mortality_pct", color="#e76f51")
    for idx, row in tab.iterrows():
        plt.text(idx, row["mortality_pct"] + 1.2, f"n={int(row['n'])}", ha="center", fontsize=9)
    plt.ylim(0, max(5.0, float(tab["mortality_pct"].max()) + 8.0))
    plt.ylabel("Летальность, %")
    plt.xlabel("Категория глюкозы (ммоль/л)")
    plt.title(title)
    plt.tight_layout()
    plt.savefig(out_path, dpi=170)
    plt.close()
    return tab


def _plot_glucose_delta_by_outcome(frame: pd.DataFrame, out_path: Path) -> pd.DataFrame:
    sub = frame[["death", "glucose_adm", "glucose_last"]].dropna().copy()
    if sub.empty or sub["death"].nunique() < 2:
        return pd.DataFrame()
    sub["delta_glucose"] = sub["glucose_last"].astype(float) - sub["glucose_adm"].astype(float)
    sub["outcome"] = sub["death"].map({0.0: "Выписан", 1.0: "Летальный исход"}).astype(str)

    plt.figure(figsize=(7.2, 4.9))
    sns.boxplot(data=sub, x="outcome", y="delta_glucose", color="#f4a261", showfliers=False)
    sns.stripplot(data=sub, x="outcome", y="delta_glucose", color="#264653", alpha=0.35, jitter=0.25, size=3)
    plt.axhline(0.0, color="gray", linestyle="--", linewidth=1)
    plt.xlabel("")
    plt.ylabel("Δ глюкоза (последняя - при поступлении), ммоль/л")
    plt.title("Динамика глюкозы: paired сравнение admission -> last")
    plt.tight_layout()
    plt.savefig(out_path, dpi=170)
    plt.close()
    return sub[["outcome", "delta_glucose"]]


def _plot_spearman_corr_heatmap(
    frame: pd.DataFrame,
    keys: List[str],
    label_map: Dict[str, str],
    out_path: Path,
) -> Tuple[pd.DataFrame, pd.DataFrame]:
    use_keys = [k for k in keys if k in frame.columns]
    if len(use_keys) < 3:
        return pd.DataFrame(), pd.DataFrame()

    data = frame[use_keys].astype(float)
    rho = data.corr(method="spearman")
    mask = data.notna().astype(int)
    n_mat = mask.T.dot(mask)

    rename = {k: label_map.get(k, k) for k in use_keys}
    rho = rho.rename(index=rename, columns=rename)
    n_mat = n_mat.rename(index=rename, columns=rename)

    annot = rho.copy().astype(object)
    for i in rho.index:
        for j in rho.columns:
            r = float(rho.loc[i, j]) if not pd.isna(rho.loc[i, j]) else np.nan
            n = int(n_mat.loc[i, j]) if not pd.isna(n_mat.loc[i, j]) else 0
            if np.isnan(r) or n < 10:
                annot.loc[i, j] = ""
            else:
                annot.loc[i, j] = f"{r:.2f}\n(n={n})"

    plt.figure(figsize=(10.2, 7.8))
    sns.heatmap(
        rho,
        vmin=-1.0,
        vmax=1.0,
        center=0.0,
        cmap="vlag",
        square=True,
        linewidths=0.5,
        linecolor="white",
        annot=annot,
        fmt="",
        cbar_kws={"label": "Spearman rho"},
    )
    plt.title("Корреляционная матрица ключевых количественных переменных (Spearman)")
    plt.xticks(rotation=45, ha="right")
    plt.yticks(rotation=0)
    plt.tight_layout()
    plt.savefig(out_path, dpi=170)
    plt.close()
    return rho, n_mat


def _plot_forest(coef_table: pd.DataFrame, title: str, out_path: Path) -> pd.DataFrame:
    if coef_table.empty:
        return coef_table
    plot_df = coef_table.copy()
    plot_df = plot_df.sort_values("OR", ascending=True)

    plt.figure(figsize=(8.5, max(3.5, 0.6 * len(plot_df))))
    y_pos = np.arange(len(plot_df))
    x = plot_df["OR"].astype(float).values
    xerr_low = x - plot_df["CI_low"].astype(float).values
    xerr_high = plot_df["CI_high"].astype(float).values - x

    plt.errorbar(
        x,
        y_pos,
        xerr=[xerr_low, xerr_high],
        fmt="o",
        color="#1d3557",
        ecolor="#457b9d",
        capsize=3,
    )
    plt.axvline(1.0, color="gray", linestyle="--", linewidth=1)
    plt.yticks(y_pos, plot_df["Предиктор"])
    plt.xscale("log")
    plt.xlabel("OR (log scale)")
    plt.title(title)
    plt.tight_layout()
    plt.savefig(out_path, dpi=170)
    plt.close()
    return plot_df


def _plot_roc(y_true: np.ndarray, y_pred: np.ndarray, title: str, out_path: Path) -> float:
    fpr, tpr, _ = roc_curve(y_true, y_pred)
    auc = float(roc_auc_score(y_true, y_pred))

    plt.figure(figsize=(6.5, 5.2))
    plt.plot(fpr, tpr, color="#2a9d8f", linewidth=2, label=f"AUC = {auc:.3f}")
    plt.plot([0, 1], [0, 1], color="gray", linestyle="--")
    plt.xlabel("1 - Specificity")
    plt.ylabel("Sensitivity")
    plt.title(title)
    plt.legend(loc="lower right")
    plt.tight_layout()
    plt.savefig(out_path, dpi=170)
    plt.close()
    return auc


def _spline_nonlinearity_glucose_last(
    frame: pd.DataFrame,
    outcome_key: str,
    covariates: Sequence[str],
    out_path: Path,
    spline_df: int = 4,
) -> pd.DataFrame:
    try:
        from patsy import build_design_matrices, dmatrix
    except Exception:
        return pd.DataFrame()

    cols = [outcome_key, "glucose_last"] + [c for c in covariates if c in frame.columns]
    if len(cols) < 4:
        return pd.DataFrame()

    sub = frame[cols].dropna().copy()
    if len(sub) < 120 or sub[outcome_key].nunique() < 2:
        return pd.DataFrame()

    y = sub[outcome_key].astype(float)
    cov = [c for c in covariates if c in sub.columns]

    try:
        X_lin = sm.add_constant(sub[["glucose_last"] + cov].astype(float), has_constant="add")
        m_lin = sm.Logit(y, X_lin).fit(disp=False, maxiter=300)
    except Exception:
        return pd.DataFrame()

    df_candidates: List[int] = []
    for cand in [int(spline_df), 5, 4, 3]:
        if cand >= 3 and cand not in df_candidates:
            df_candidates.append(cand)

    m_spl = None
    basis = None
    X_spl = None
    used_df = None
    for df_try in df_candidates:
        try:
            basis = dmatrix(
                f"bs(glucose_last, df={int(df_try)}, include_intercept=False)",
                sub,
                return_type="dataframe",
            )
            X_spl = pd.concat([basis, sub[cov].astype(float)], axis=1)
            X_spl = sm.add_constant(X_spl, has_constant="add")
            m_spl = sm.Logit(y, X_spl).fit(disp=False, maxiter=300)
            used_df = int(df_try)
            break
        except Exception:
            m_spl = None
            basis = None
            X_spl = None
            used_df = None
            continue
    if m_spl is None or basis is None or X_spl is None or used_df is None:
        return pd.DataFrame()

    # LRT for non-linearity (spline vs linear)
    lr_stat = float(2.0 * (m_spl.llf - m_lin.llf))
    df_diff = int(len(m_spl.params) - len(m_lin.params))
    p_lr = float(stats.chi2.sf(lr_stat, df_diff)) if df_diff > 0 else np.nan

    # Plot: predicted probability across glucose_last grid with covariates fixed at typical values
    g_low = float(sub["glucose_last"].quantile(0.05))
    g_high = float(sub["glucose_last"].quantile(0.95))
    grid = np.linspace(g_low, g_high, 120)
    grid_df = pd.DataFrame({"glucose_last": grid})

    basis_grid = build_design_matrices([basis.design_info], grid_df)[0]
    basis_grid = pd.DataFrame(np.asarray(basis_grid), columns=basis.columns)

    for c in cov:
        if c in NUMERIC_KEYS:
            fill = float(sub[c].astype(float).median())
        else:
            mode = sub[c].dropna().astype(float).mode()
            fill = float(mode.iloc[0]) if not mode.empty else float(round(float(sub[c].astype(float).mean())))
        basis_grid[c] = fill

    Xg = sm.add_constant(basis_grid, has_constant="add")
    pred = m_spl.predict(Xg)

    plt.figure(figsize=(7.2, 4.9))
    plt.plot(grid, pred, color="#1d3557", linewidth=2)
    plt.ylim(0, 1)
    plt.xlabel("Глюкоза последний результат, ммоль/л")
    plt.ylabel("Предсказанная вероятность исхода")
    plt.title("Нелинейность: spline эффект последней глюкозы (с поправкой на ковариаты)")
    plt.tight_layout()
    plt.savefig(out_path, dpi=170)
    plt.close()

    try:
        auc_lin = float(roc_auc_score(y, m_lin.predict(X_lin)))
    except Exception:
        auc_lin = np.nan
    try:
        auc_spl = float(roc_auc_score(y, m_spl.predict(X_spl)))
    except Exception:
        auc_spl = np.nan

    out = pd.DataFrame(
        [
            {
                "N": int(len(sub)),
                "Spline df": int(used_df),
                "AIC (linear)": float(m_lin.aic),
                "AIC (spline)": float(m_spl.aic),
                "ΔAIC (spline-linear)": float(m_spl.aic - m_lin.aic),
                "AUC (linear)": auc_lin,
                "AUC (spline)": auc_spl,
                "LRT df": df_diff,
                "LRT p (non-linearity)": p_lr,
            }
        ]
    )
    return out


def _interpret_numeric_table(table: pd.DataFrame) -> str:
    if table.empty:
        return "Интерпретация: недостаточно данных для сравнения ключевых количественных показателей."
    sig = table[table["_p_raw"].astype(float) < 0.05].copy()
    if sig.empty:
        return "Интерпретация: по ключевым количественным показателям статистически значимые различия между исходами не выявлены."
    sig = sig.sort_values("_p_raw")
    chunks: List[str] = []
    for _, row in sig.head(4).iterrows():
        direction = "выше" if row["_dead_median"] > row["_alive_median"] else "ниже"
        chunks.append(f"{row['Показатель']} ({direction} при летальном исходе, p={_format_p(row['_p_raw'])})")
    return "Интерпретация: значимые различия выявлены для " + "; ".join(chunks) + "."


def _interpret_binary_table(table: pd.DataFrame, domain_name: str) -> str:
    if table.empty:
        return f"Интерпретация: блок '{domain_name}' имеет недостаточное покрытие для устойчивого вывода."
    sig = table[table["_p_raw"].astype(float) < 0.05].copy()
    if sig.empty:
        return f"Интерпретация: в блоке '{domain_name}' статистически значимых ассоциаций с летальным исходом не найдено."
    sig = sig.sort_values("_p_raw")
    parts: List[str] = []
    for _, row in sig.head(4).iterrows():
        trend = "более высокой" if float(row["_or_raw"]) > 1 else "более низкой"
        parts.append(
            f"{row['Показатель']} (OR={row['OR (crude)']}, p={row['p (Fisher)']}, "
            f"ассоциировано с {trend} вероятностью летального исхода)"
        )
    note_bits: List[str] = []
    if "_zero_cell" in table.columns and bool(table["_zero_cell"].astype(int).sum()):
        note_bits.append("в части 2×2 таблиц есть нулевые ячейки (квазисепарация)")
    if "_missing_total" in table.columns and "_available_total" in table.columns:
        high_missing = (table["_missing_total"].astype(int) > table["_available_total"].astype(int)).sum()
        if high_missing:
            note_bits.append("по части факторов пропусков больше, чем наблюдений с данными")
    note = ""
    if note_bits:
        note = " Примечание: " + "; ".join(note_bits) + ", поэтому crude OR может быть неустойчивым."
    return f"Интерпретация: в блоке '{domain_name}' значимы: " + "; ".join(parts) + "." + note


def _interpret_univariate_glycemia(table: pd.DataFrame) -> str:
    if table.empty:
        return "Интерпретация: однофакторная логистическая оценка гликемических предикторов не выполнена из-за пропусков/нестабильности модели."
    sig = table[table["_p_raw"].astype(float) < 0.05].copy()
    if sig.empty:
        return "Интерпретация: в однофакторных моделях гликемические предикторы не показали статистически значимой связи с летальным исходом."
    sig = sig.sort_values("_p_raw")
    parts = [f"{r['Предиктор']} (OR={r['OR']}, p={r['p']})" for _, r in sig.head(4).iterrows()]
    return "Интерпретация: однофакторно значимы: " + "; ".join(parts) + "."


def _interpret_model_metrics(metrics_df: pd.DataFrame) -> str:
    if metrics_df.empty:
        return "Интерпретация: многофакторные модели не стабилизировались на текущем покрытии данных."
    best = metrics_df.sort_values("AIC").iloc[0]
    return (
        "Интерпретация: лучшая из собранных моделей — "
        f"{best['Model']} ({best['Описание']}), AUC={best['AUC']:.3f}, "
        f"AIC={best['AIC']:.1f}, N={int(best['N'])}."
    )


def _models_to_metrics_df(model_results: Sequence[ModelResult]) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    for mr in model_results:
        rows.append(
            {
                "Model": mr.model_id,
                "Описание": mr.title,
                "N": mr.n,
                "Events": mr.events,
                "AIC": round(mr.aic, 2),
                "AUC": round(mr.auc, 3) if not np.isnan(mr.auc) else np.nan,
                "Brier": round(mr.brier, 4) if not np.isnan(mr.brier) else np.nan,
            }
        )
    return pd.DataFrame(rows)


def _interpret_paired_glucose(frame: pd.DataFrame, p_between: float) -> str:
    sub = frame[["death", "glucose_adm", "glucose_last"]].dropna().copy()
    if sub.empty or sub["death"].nunique() < 2:
        return "Интерпретация: парное сравнение глюкозы (admission vs last) недоступно из-за пропусков."

    sub["delta"] = sub["glucose_last"].astype(float) - sub["glucose_adm"].astype(float)
    d0 = sub[sub["death"] == 0.0]["delta"].astype(float)
    d1 = sub[sub["death"] == 1.0]["delta"].astype(float)

    def _summ(series: pd.Series) -> str:
        if series.empty:
            return "NA"
        return _median_iqr(series)

    text = (
        "Интерпретация: парное сравнение (внутри пациента) показывает, как менялась глюкоза от поступления к последнему измерению. "
        f"Для выписанных Δ={_summ(d0)}; для летальных исходов Δ={_summ(d1)}."
    )
    if p_between is not None and not np.isnan(p_between):
        text += f" Различие Δ между исходами: p={_format_p(p_between)} (Mann-Whitney)."
    text += " Важно: это описывает динамику и тяжесть/терапию, а не причинный эффект вмешательств."
    return text


def _interpret_corr_matrix(rho: pd.DataFrame, n_mat: pd.DataFrame) -> str:
    if rho.empty or n_mat.empty:
        return "Интерпретация: корреляционная матрица не построена (недостаточно переменных/наблюдений)."
    pairs: List[Tuple[float, float, int, str, str]] = []
    cols = list(rho.columns)
    for i in range(len(cols)):
        for j in range(i + 1, len(cols)):
            a, b = cols[i], cols[j]
            r = rho.loc[a, b]
            n = int(n_mat.loc[a, b])
            if pd.isna(r) or n < 80:
                continue
            pairs.append((abs(float(r)), float(r), n, a, b))
    pairs.sort(reverse=True, key=lambda x: x[0])
    if not pairs:
        return (
            "Интерпретация: корреляции рассчитаны попарно (pairwise complete); "
            "на выбранном пороге N устойчивых сильных связей не выявлено."
        )
    top = pairs[:4]
    bits = [f"{a}—{b} (rho={r:.2f}, N={n})" for _, r, n, a, b in top]
    return (
        "Интерпретация: корреляционная структура (Spearman) помогает оценить коллинеарность и возможную избыточность ковариат. "
        "Наиболее выраженные связи: " + "; ".join(bits) + ". "
        "Примечание: из-за пропусков N в разных ячейках различается, поэтому матрица носит разведочный характер."
    )


def _discussion_text(
    outcome_n: int,
    death_n: int,
    glycemia_uni: pd.DataFrame,
    comorb_table: pd.DataFrame,
    treatment_table: pd.DataFrame,
    metrics_df: pd.DataFrame,
) -> List[str]:
    lines: List[str] = []
    lines.append(
        f"В анализ включено {outcome_n} пациентов с определенным исходом, из них {death_n} с летальным исходом. "
        "Основной фокус был сделан на гликемии и ее связи с прогнозом."
    )

    if not glycemia_uni.empty:
        sig_g = glycemia_uni[glycemia_uni["_p_raw"].astype(float) < 0.05]
        if sig_g.empty:
            lines.append(
                "В однофакторной постановке устойчивой статистически значимой связи гликемических маркеров с летальностью не получено, "
                "что указывает на вероятный вклад конфаундинга и неполноты данных."
            )
        else:
            top = sig_g.sort_values("_p_raw").iloc[0]
            lines.append(
                f"Наиболее выраженный однофакторный сигнал наблюдался для '{top['Предиктор']}' "
                f"(OR={top['OR']}, p={top['p']})."
            )

    if not comorb_table.empty:
        sig_c = comorb_table[comorb_table["_p_raw"].astype(float) < 0.05]
        if sig_c.empty:
            lines.append("По блоку коморбидности сильных независимых сигналов в crude-сравнениях не найдено.")
        else:
            picks = "; ".join(sig_c.sort_values("_p_raw").head(3)["Показатель"].astype(str).tolist())
            lines.append(f"В коморбидном профиле статистически выделяются: {picks}.")

    if not treatment_table.empty:
        sig_t = treatment_table[treatment_table["_p_raw"].astype(float) < 0.05]
        if sig_t.empty:
            lines.append(
                "По блоку терапии значимых различий в crude-оценке немного; это ожидаемо для ретроспективной когорты с лечением по показаниям."
            )
        else:
            picks = "; ".join(sig_t.sort_values("_p_raw").head(3)["Показатель"].astype(str).tolist())
            lines.append(f"В блоке терапии статистически выделяются: {picks}.")
        lines.append(
            "Для терапии интерпретация принципиально ассоциативная: выраженные OR могут отражать confounding-by-indication "
            "(более интенсивное лечение у более тяжелых пациентов)."
        )

    if not metrics_df.empty:
        best = metrics_df.sort_values("AIC").iloc[0]
        lines.append(
            f"Многофакторная модель {best['Model']} показывает дискриминацию уровня AUC={best['AUC']:.3f}. "
            "Это указывает на умеренную прогностическую ценность комбинации признаков, но не доказывает причинность."
        )

    lines.append(
        "Патофизиологически полученный паттерн согласуется с концепцией стресс-гипергликемии как маркера тяжести: "
        "гликемия может отражать воспалительный и нейроэндокринный ответ, а не только диабетический статус."
    )
    return lines


def _add_dataframe_table(doc: Document, df: pd.DataFrame, drop_internal: bool = True) -> None:
    view = df.copy()
    if drop_internal:
        internal = [c for c in view.columns if c.startswith("_")]
        view = view.drop(columns=internal, errors="ignore")
    if view.empty:
        doc.add_paragraph("Нет данных для таблицы.")
        return
    table = doc.add_table(rows=1, cols=len(view.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(view.columns):
        hdr_cells[i].text = str(col)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(view.columns):
            val = row[col]
            if isinstance(val, float):
                if np.isnan(val):
                    cells[i].text = ""
                else:
                    cells[i].text = f"{val:.3f}"
            else:
                cells[i].text = str(val)


def _write_report_docx(
    out_docx: Path,
    cohort_meta: Dict[str, Any],
    coverage_df: pd.DataFrame,
    numeric_table: pd.DataFrame,
    comorb_table: pd.DataFrame,
    treatment_table: pd.DataFrame,
    paired_table: pd.DataFrame,
    glycemia_uni: pd.DataFrame,
    model_metrics: pd.DataFrame,
    coef_table_best: pd.DataFrame,
    sensitivity_table: pd.DataFrame,
    spline_table: pd.DataFrame,
    icu_model_metrics: pd.DataFrame,
    icu_coef_best: pd.DataFrame,
    newdm_model_metrics: pd.DataFrame,
    newdm_coef_best: pd.DataFrame,
    interpretations: Dict[str, str],
    figure_paths: Dict[str, Path],
    discussion_lines: List[str],
    conclusions: List[str],
) -> None:
    doc = Document()
    doc.add_heading("COVID-19: исходы и влияние гликемии (структурированный отчет)", level=1)
    doc.add_paragraph(
        f"Дата генерации: {datetime.now().strftime('%Y-%m-%d %H:%M')}. "
        f"N={cohort_meta['n_total']}, с известным исходом N={cohort_meta['n_outcome']}, летальных исходов N={cohort_meta['n_death']}."
    )

    doc.add_heading("1. Дизайн и переменные", level=2)
    doc.add_paragraph(
        "Отчет включает блоки: когортный профиль, гликемия, коморбидность, лечение, многофакторное моделирование и графическая диагностика."
    )
    doc.add_paragraph(
        "Аналитическая схема: (1) нормализация исхода и ключевых переменных; "
        "(2) baseline-сравнение по исходу; (3) crude OR для коморбидности и терапии; "
        "(4) однофакторные модели для гликемии и тяжести; (5) многофакторные модели с оценкой AUC/Brier; "
        "(6) интерпретация каждого артефакта и сводное обсуждение."
    )
    _add_dataframe_table(doc, coverage_df, drop_internal=False)
    doc.add_paragraph(interpretations["coverage"])

    doc.add_heading("2. Количественные показатели по исходу", level=2)
    _add_dataframe_table(doc, numeric_table)
    doc.add_paragraph(interpretations["numeric"])

    doc.add_heading("3. Коморбидность и исход", level=2)
    _add_dataframe_table(doc, comorb_table)
    doc.add_paragraph(interpretations["comorb"])

    doc.add_heading("4. Терапия и исход", level=2)
    _add_dataframe_table(doc, treatment_table)
    doc.add_paragraph(interpretations["treatment"])

    doc.add_heading("5. Динамика гликемии (парные сравнения)", level=2)
    _add_dataframe_table(doc, paired_table)
    doc.add_paragraph(interpretations["paired"])

    doc.add_heading("6. Однофакторные модели (гликемия/связанные предикторы)", level=2)
    _add_dataframe_table(doc, glycemia_uni)
    doc.add_paragraph(interpretations["glycemia_uni"])

    doc.add_heading("7. Многофакторные модели", level=2)
    _add_dataframe_table(doc, model_metrics, drop_internal=False)
    doc.add_paragraph(interpretations["models"])
    doc.add_paragraph("Коэффициенты лучшей модели:")
    _add_dataframe_table(doc, coef_table_best)

    if not sensitivity_table.empty:
        doc.add_heading("7.1. Чувствительность к пропускам", level=3)
        _add_dataframe_table(doc, sensitivity_table, drop_internal=False)
        doc.add_paragraph(interpretations["sensitivity"])

    if not spline_table.empty or (figure_paths.get("spline") and figure_paths["spline"].exists()):
        doc.add_heading("7.2. Нелинейность последней глюкозы (spline)", level=3)
        if not spline_table.empty:
            _add_dataframe_table(doc, spline_table, drop_internal=False)
        doc.add_paragraph(interpretations["spline"])
        spline_path = figure_paths.get("spline")
        if spline_path and spline_path.exists():
            doc.add_paragraph("Spline-эффект последней глюкозы (визуализация):")
            doc.add_picture(str(spline_path), width=Inches(6.5))
            doc.add_paragraph(interpretations.get("fig_spline", ""))

    if not icu_model_metrics.empty:
        doc.add_heading("7.3. Прогнозирование исхода: ОРИТ (exploratory)", level=3)
        _add_dataframe_table(doc, icu_model_metrics, drop_internal=False)
        doc.add_paragraph(interpretations.get("icu_models", ""))
        if not icu_coef_best.empty:
            doc.add_paragraph("Коэффициенты лучшей модели (ОРИТ):")
            _add_dataframe_table(doc, icu_coef_best)
        forest_icu = figure_paths.get("forest_icu")
        if forest_icu and forest_icu.exists():
            doc.add_paragraph("Forest plot (ОРИТ):")
            doc.add_picture(str(forest_icu), width=Inches(6.5))
            doc.add_paragraph(interpretations.get("fig_forest_icu", ""))
        roc_icu = figure_paths.get("roc_icu")
        if roc_icu and roc_icu.exists():
            doc.add_paragraph("ROC (ОРИТ):")
            doc.add_picture(str(roc_icu), width=Inches(6.5))
            doc.add_paragraph(interpretations.get("fig_roc_icu", ""))

    if not newdm_model_metrics.empty:
        doc.add_heading("7.4. Прогнозирование исхода: СД при выписке (без СД2 до госпитализации)", level=3)
        _add_dataframe_table(doc, newdm_model_metrics, drop_internal=False)
        doc.add_paragraph(interpretations.get("newdm_models", ""))
        if not newdm_coef_best.empty:
            doc.add_paragraph("Коэффициенты лучшей модели (СД при выписке, без СД2):")
            _add_dataframe_table(doc, newdm_coef_best)
        forest_newdm = figure_paths.get("forest_newdm")
        if forest_newdm and forest_newdm.exists():
            doc.add_paragraph("Forest plot (СД при выписке, без СД2):")
            doc.add_picture(str(forest_newdm), width=Inches(6.5))
            doc.add_paragraph(interpretations.get("fig_forest_newdm", ""))
        roc_newdm = figure_paths.get("roc_newdm")
        if roc_newdm and roc_newdm.exists():
            doc.add_paragraph("ROC (СД при выписке, без СД2):")
            doc.add_picture(str(roc_newdm), width=Inches(6.5))
            doc.add_paragraph(interpretations.get("fig_roc_newdm", ""))

    figure_order = [
        ("missingness", "Рисунок 1. Пропуски в ключевых переменных", "fig_missingness"),
        ("delta_glucose", "Рисунок 2. Динамика глюкозы: Δ(последняя - при поступлении) по исходу", "fig_delta_glucose"),
        ("corr", "Рисунок 3. Корреляционная матрица ключевых количественных переменных (Spearman)", "fig_corr"),
        ("glucose_adm_cat", "Рисунок 4. Летальность по категориям глюкозы при поступлении", "fig_glucose_adm"),
        ("glucose_last_cat", "Рисунок 5. Летальность по категориям последней глюкозы", "fig_glucose_last"),
        ("forest", "Рисунок 6. Forest plot OR лучшей модели", "fig_forest"),
        ("roc", "Рисунок 7. ROC-кривая лучшей модели", "fig_roc"),
    ]
    doc.add_heading("8. Графики и интерпретация", level=2)
    for key, caption, interp_key in figure_order:
        path = figure_paths.get(key)
        if not path or not path.exists():
            continue
        doc.add_paragraph(caption)
        doc.add_picture(str(path), width=Inches(6.5))
        doc.add_paragraph(interpretations.get(interp_key, ""))

    doc.add_heading("9. Обсуждение", level=2)
    for line in discussion_lines:
        doc.add_paragraph(line)

    doc.add_heading("10. Практические выводы", level=2)
    for item in conclusions:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(out_docx)


def _safe_markdown(df: pd.DataFrame, drop_internal: bool = True) -> str:
    view = df.copy()
    if drop_internal:
        view = view[[c for c in view.columns if not c.startswith("_")]]
    if view.empty:
        return "_Нет данных_"
    try:
        return view.to_markdown(index=False)
    except Exception:
        return "```\n" + view.to_string(index=False) + "\n```"


def run_report(input_path: Path, output_root: Path, output_dir: Optional[Path] = None) -> Dict[str, Any]:
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    if output_dir is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        run_dir = output_root / f"covid_structured_{timestamp}"
    else:
        run_dir = output_dir
    figs_dir = run_dir / "figures"
    tables_dir = run_dir / "tables"
    run_dir.mkdir(parents=True, exist_ok=True)
    figs_dir.mkdir(parents=True, exist_ok=True)
    tables_dir.mkdir(parents=True, exist_ok=True)

    raw = pd.read_excel(input_path, sheet_name="Лист1")
    frame, source_map = _prepare_analysis_frame(raw)

    n_total = int(len(frame))
    n_outcome = int(frame["death"].notna().sum())
    n_death = int(frame["death"].sum(skipna=True))

    coverage_rows: List[Dict[str, Any]] = []
    registry = (
        [("death", "Исход (бинарный)")]
        + [(k, l) for k, l, _ in NUMERIC_SPECS]
        + [(k, l) for k, l, _ in GLYCEMIA_BINARY_SPECS]
        + [(k, l) for k, l, _ in COMORBIDITY_SPECS]
        + [(k, l) for k, l, _ in TREATMENT_SPECS]
        + [(k, l) for k, l, _ in ADDITIONAL_BINARY_SPECS]
    )
    seen = set()
    for key, label in registry:
        if key in seen:
            continue
        seen.add(key)
        if key not in frame.columns:
            continue
        non_null = int(frame[key].notna().sum())
        coverage_rows.append(
            {
                "key": key,
                "label": label,
                "source_column": source_map.get(key, key),
                "non_null_n": non_null,
                "coverage_pct": round(100.0 * non_null / n_total, 1),
            }
        )
    coverage_df = pd.DataFrame(coverage_rows).sort_values(["coverage_pct", "label"], ascending=[False, True])

    numeric_table = _numeric_table(frame, NUMERIC_SPECS)
    comorb_table = _binary_or_table(frame, COMORBIDITY_SPECS, "коморбидность")
    treatment_table = _binary_or_table(frame, TREATMENT_SPECS, "лечение")

    uni_predictors = [
        ("glucose_adm", "Глюкоза при поступлении"),
        ("glucose_last", "Глюкоза последний результат"),
        ("hypergly_twice", "Гипергликемия >11.1 дважды"),
        ("dm2_pre", "СД2 до госпитализации"),
        ("age", "Возраст"),
        ("spo2", "SpO2"),
        ("news2", "NEWS2"),
        ("qsofa", "qSOFA"),
    ]
    glycemia_uni = _fit_univariate_logit(frame, uni_predictors)

    model_specs = [
        ("M1", "Глюкоза при поступлении + демография", ["glucose_adm", "age", "sex_male"]),
        ("M2", "Глюкоза при поступлении + тяжесть", ["glucose_adm", "age", "sex_male", "spo2", "news2"]),
        (
            "M3",
            "Последняя глюкоза + тяжесть + диабет/гипергликемия",
            ["glucose_last", "age", "sex_male", "spo2", "news2", "dm2_pre", "hypergly_twice"],
        ),
        (
            "M4",
            "Interaction: глюкоза при поступлении * СД2",
            ["glucose_adm", "dm2_pre", "glucose_dm2_interaction", "age", "sex_male", "spo2", "news2"],
        ),
    ]

    model_results: List[ModelResult] = []
    for model_id, title, predictors in model_specs:
        mr = _fit_multivariable_model(frame, "death", "Летальный исход", model_id, title, predictors)
        if mr is not None:
            model_results.append(mr)

    metrics_rows: List[Dict[str, Any]] = []
    for mr in model_results:
        metrics_rows.append(
            {
                "Model": mr.model_id,
                "Описание": mr.title,
                "N": mr.n,
                "Events": mr.events,
                "AIC": round(mr.aic, 2),
                "AUC": round(mr.auc, 3) if not np.isnan(mr.auc) else np.nan,
                "Brier": round(mr.brier, 4) if not np.isnan(mr.brier) else np.nan,
            }
        )
    model_metrics_df = pd.DataFrame(metrics_rows)

    best_model: Optional[ModelResult] = None
    if model_results:
        best_model = sorted(model_results, key=lambda m: (m.aic, -m.auc if not np.isnan(m.auc) else 999.0))[0]
    coef_table_best = best_model.coef_table.copy() if best_model is not None else pd.DataFrame()

    paired_table, paired_p_between = _paired_glucose_delta_table(frame)

    imputed_best_model: Optional[ModelResult] = None
    sensitivity_df = pd.DataFrame()
    if best_model is not None:
        imputed_best_model = _fit_multivariable_model_imputed(
            frame,
            "death",
            "Летальный исход",
            f"{best_model.model_id}_imp",
            f"{best_model.title} (simple imputation)",
            best_model.predictors,
        )
        if imputed_best_model is not None:
            sensitivity_df = pd.DataFrame(
                [
                    {
                        "Approach": "Complete-case",
                        "N": best_model.n,
                        "Events": best_model.events,
                        "AIC": round(best_model.aic, 2),
                        "AUC": round(best_model.auc, 3) if not np.isnan(best_model.auc) else np.nan,
                        "Brier": round(best_model.brier, 4) if not np.isnan(best_model.brier) else np.nan,
                    },
                    {
                        "Approach": "Simple imputation",
                        "N": imputed_best_model.n,
                        "Events": imputed_best_model.events,
                        "AIC": round(imputed_best_model.aic, 2),
                        "AUC": round(imputed_best_model.auc, 3) if not np.isnan(imputed_best_model.auc) else np.nan,
                        "Brier": round(imputed_best_model.brier, 4) if not np.isnan(imputed_best_model.brier) else np.nan,
                    },
                ]
            )

    icu_models: List[ModelResult] = []
    icu_model_metrics_df = pd.DataFrame()
    icu_best_model: Optional[ModelResult] = None
    icu_coef_best = pd.DataFrame()
    if "icu" in frame.columns:
        icu_model_specs = [
            ("ICU1", "ОРИТ: глюкоза при поступлении + демография", ["glucose_adm", "age", "sex_male"]),
            ("ICU2", "ОРИТ: глюкоза при поступлении + тяжесть", ["glucose_adm", "age", "sex_male", "spo2", "news2"]),
        ]
        for model_id, title, predictors in icu_model_specs:
            mr = _fit_multivariable_model(frame, "icu", "ОРИТ", model_id, title, predictors)
            if mr is not None:
                icu_models.append(mr)
        icu_model_metrics_df = _models_to_metrics_df(icu_models)
        if icu_models:
            icu_best_model = sorted(icu_models, key=lambda m: (m.aic, -m.auc if not np.isnan(m.auc) else 999.0))[0]
            icu_coef_best = icu_best_model.coef_table.copy()

    newdm_models: List[ModelResult] = []
    newdm_model_metrics_df = pd.DataFrame()
    newdm_best_model: Optional[ModelResult] = None
    newdm_coef_best = pd.DataFrame()
    if "dm_discharge" in frame.columns and "dm2_pre" in frame.columns:
        newdm_frame = frame[frame["dm2_pre"] == 0].copy()
        newdm_model_specs = [
            ("NDM1", "СД при выписке (без СД2): глюкоза при поступлении + демография", ["glucose_adm", "age", "sex_male"]),
            ("NDM2", "СД при выписке (без СД2): глюкоза при поступлении + тяжесть", ["glucose_adm", "age", "sex_male", "spo2", "news2"]),
            ("NDM3", "СД при выписке (без СД2): последняя глюкоза + тяжесть", ["glucose_last", "age", "sex_male", "spo2", "news2"]),
        ]
        for model_id, title, predictors in newdm_model_specs:
            mr = _fit_multivariable_model(newdm_frame, "dm_discharge", "СД при выписке (без СД2)", model_id, title, predictors)
            if mr is not None:
                newdm_models.append(mr)
        newdm_model_metrics_df = _models_to_metrics_df(newdm_models)
        if newdm_models:
            newdm_best_model = sorted(newdm_models, key=lambda m: (m.aic, -m.auc if not np.isnan(m.auc) else 999.0))[0]
            newdm_coef_best = newdm_best_model.coef_table.copy()

    fig_spline = figs_dir / "spline_glucose_last.png"
    spline_table = _spline_nonlinearity_glucose_last(
        frame,
        outcome_key="death",
        covariates=["age", "sex_male", "spo2", "news2", "dm2_pre"],
        out_path=fig_spline,
        spline_df=4,
    )

    numeric_label_map: Dict[str, str] = {k: l for k, l, _ in NUMERIC_SPECS}
    corr_keys_pref = [
        "age",
        "spo2",
        "news2",
        "qsofa",
        "crp1",
        "creatinine1",
        "glucose_adm",
        "glucose_last",
        "icu_days",
        "hosp_days",
    ]

    vars_for_missing = [r[0] for r in NUMERIC_SPECS] + [r[0] for r in GLYCEMIA_BINARY_SPECS]
    vars_for_missing += [r[0] for r in COMORBIDITY_SPECS] + [r[0] for r in TREATMENT_SPECS]
    vars_for_missing += [r[0] for r in ADDITIONAL_BINARY_SPECS]
    vars_for_missing = sorted(set([v for v in vars_for_missing if v in frame.columns]))

    fig_missingness = figs_dir / "missingness.png"
    miss_df = _plot_missingness(frame, vars_for_missing, fig_missingness)

    fig_glucose_adm = figs_dir / "mortality_by_glucose_adm.png"
    adm_cat_df = _plot_glucose_category_mortality(
        frame, "glucose_adm_cat", "Летальность по категориям глюкозы при поступлении", fig_glucose_adm
    )

    fig_glucose_last = figs_dir / "mortality_by_glucose_last.png"
    last_cat_df = _plot_glucose_category_mortality(
        frame, "glucose_last_cat", "Летальность по категориям последней глюкозы", fig_glucose_last
    )

    fig_delta_glucose = figs_dir / "delta_glucose_by_outcome.png"
    _plot_glucose_delta_by_outcome(frame, fig_delta_glucose)

    fig_corr = figs_dir / "corr_spearman_key.png"
    corr_rho, corr_n = _plot_spearman_corr_heatmap(frame, corr_keys_pref, numeric_label_map, fig_corr)

    fig_forest = figs_dir / "forest_best_model.png"
    if best_model is not None and not best_model.coef_table.empty:
        _plot_forest(best_model.coef_table, f"Forest plot: {best_model.model_id}", fig_forest)

    fig_roc = figs_dir / "roc_best_model.png"
    if best_model is not None:
        _plot_roc(best_model.y_true, best_model.y_pred, f"ROC: {best_model.model_id}", fig_roc)

    fig_forest_icu = figs_dir / "forest_icu_best.png"
    if icu_best_model is not None and not icu_best_model.coef_table.empty:
        _plot_forest(icu_best_model.coef_table, f"Forest plot: ОРИТ ({icu_best_model.model_id})", fig_forest_icu)

    fig_roc_icu = figs_dir / "roc_icu_best.png"
    if icu_best_model is not None:
        _plot_roc(icu_best_model.y_true, icu_best_model.y_pred, f"ROC: ОРИТ ({icu_best_model.model_id})", fig_roc_icu)

    fig_forest_newdm = figs_dir / "forest_newdm_best.png"
    if newdm_best_model is not None and not newdm_best_model.coef_table.empty:
        _plot_forest(newdm_best_model.coef_table, f"Forest plot: СД при выписке ({newdm_best_model.model_id})", fig_forest_newdm)

    fig_roc_newdm = figs_dir / "roc_newdm_best.png"
    if newdm_best_model is not None:
        _plot_roc(
            newdm_best_model.y_true,
            newdm_best_model.y_pred,
            f"ROC: СД при выписке ({newdm_best_model.model_id})",
            fig_roc_newdm,
        )

    # Save tables to CSV
    coverage_df.to_csv(tables_dir / "coverage.csv", index=False)
    numeric_table.drop(columns=[c for c in numeric_table.columns if c.startswith("_")], errors="ignore").to_csv(
        tables_dir / "numeric_by_outcome.csv", index=False
    )
    comorb_table.drop(columns=[c for c in comorb_table.columns if c.startswith("_")], errors="ignore").to_csv(
        tables_dir / "comorbidity_or.csv", index=False
    )
    treatment_table.drop(columns=[c for c in treatment_table.columns if c.startswith("_")], errors="ignore").to_csv(
        tables_dir / "treatment_or.csv", index=False
    )
    paired_table.drop(columns=[c for c in paired_table.columns if c.startswith("_")], errors="ignore").to_csv(
        tables_dir / "glucose_paired_delta_by_outcome.csv", index=False
    )
    glycemia_uni.drop(columns=[c for c in glycemia_uni.columns if c.startswith("_")], errors="ignore").to_csv(
        tables_dir / "glycemia_univariate_logit.csv", index=False
    )
    model_metrics_df.to_csv(tables_dir / "model_metrics.csv", index=False)
    coef_table_best.to_csv(tables_dir / "best_model_coefficients.csv", index=False)
    if not sensitivity_df.empty:
        sensitivity_df.to_csv(tables_dir / "sensitivity_missingness.csv", index=False)
    if not spline_table.empty:
        spline_table.to_csv(tables_dir / "spline_glucose_last_nonlinearity.csv", index=False)
    if not icu_model_metrics_df.empty:
        icu_model_metrics_df.to_csv(tables_dir / "icu_model_metrics.csv", index=False)
    if not icu_coef_best.empty:
        icu_coef_best.to_csv(tables_dir / "icu_best_model_coefficients.csv", index=False)
    if not newdm_model_metrics_df.empty:
        newdm_model_metrics_df.to_csv(tables_dir / "newdm_model_metrics.csv", index=False)
    if not newdm_coef_best.empty:
        newdm_coef_best.to_csv(tables_dir / "newdm_best_model_coefficients.csv", index=False)
    miss_df.to_csv(tables_dir / "missingness.csv", index=False)
    adm_cat_df.to_csv(tables_dir / "glucose_adm_categories.csv", index=False)
    last_cat_df.to_csv(tables_dir / "glucose_last_categories.csv", index=False)
    if not corr_rho.empty:
        corr_rho.to_csv(tables_dir / "correlation_spearman_rho.csv")
    if not corr_n.empty:
        corr_n.to_csv(tables_dir / "correlation_pairwise_n.csv")

    interpretations: Dict[str, str] = {}
    interpretations["coverage"] = (
        "Интерпретация: покрытие переменных неравномерное; поэтому все модели представлены в complete-case постановке, "
        "а стабильность выводов оценивается через согласованность между таблицами и графиками."
    )
    interpretations["numeric"] = _interpret_numeric_table(numeric_table)
    interpretations["comorb"] = _interpret_binary_table(comorb_table, "коморбидность")
    interpretations["treatment"] = _interpret_binary_table(treatment_table, "лечение")
    interpretations["paired"] = _interpret_paired_glucose(frame, paired_p_between)
    interpretations["glycemia_uni"] = _interpret_univariate_glycemia(glycemia_uni)
    interpretations["models"] = _interpret_model_metrics(model_metrics_df)
    if sensitivity_df.empty:
        interpretations["sensitivity"] = (
            "Интерпретация: чувствительность к пропускам не оценена (недостаточно данных или модель не стабилизировалась)."
        )
    else:
        interpretations["sensitivity"] = (
            "Интерпретация: сравнение complete-case и простой импутации показывает, насколько выводы зависят от пропусков. "
            "Если направление и порядок величин эффектов сохраняются, результат устойчивее; если сильно меняются — пропуски критичны."
        )
    if spline_table.empty:
        interpretations["spline"] = "Интерпретация: spline-анализ нелинейности последней глюкозы не выполнен (пропуски/нестабильность)."
    else:
        row = spline_table.iloc[0].to_dict()
        delta_aic = float(row.get("ΔAIC (spline-linear)", np.nan))
        p_lr = row.get("LRT p (non-linearity)", np.nan)
        interpretations["spline"] = (
            "Интерпретация: spline-модель проверяет, есть ли у последней глюкозы нелинейная связь с исходом после поправки на ковариаты. "
            f"Тест нелинейности: p={_format_p(p_lr)}; ΔAIC={delta_aic:.1f} (отрицательное значение в пользу spline)."
        )
    interpretations["icu_models"] = _interpret_model_metrics(icu_model_metrics_df) if not icu_model_metrics_df.empty else (
        "Интерпретация: модели для ОРИТ не построены (недостаточно данных/вариабельности)."
    )
    interpretations["newdm_models"] = (
        _interpret_model_metrics(newdm_model_metrics_df)
        if not newdm_model_metrics_df.empty
        else "Интерпретация: модели для СД при выписке (без СД2 до госпитализации) не построены (недостаточно данных)."
    )
    interpretations["fig_missingness"] = (
        "Интерпретация: на графике видна выраженная гетерогенность пропусков; "
        "это ограничивает глубину ковариатной корректировки и требует аккуратного чтения OR."
    )
    interpretations["fig_glucose_adm"] = (
        "Интерпретация: график показывает распределение летальности по категориям стартовой глюкозы; "
        "оценивается наличие порогового паттерна >=11.1 ммоль/л."
    )
    interpretations["fig_glucose_last"] = (
        "Интерпретация: категория последней глюкозы отражает динамический метаболический контроль; "
        "градиент летальности здесь клинически важен как возможный сигнал risk-stratification."
    )
    interpretations["fig_delta_glucose"] = (
        "Интерпретация: распределение Δ(последняя - при поступлении) показывает различия в динамике гликемии между исходами "
        "и может отражать стресс-гипергликемию/тяжесть и влияние терапии; интерпретация строго ассоциативная."
    )
    interpretations["fig_corr"] = _interpret_corr_matrix(corr_rho, corr_n)
    interpretations["fig_spline"] = (
        "Интерпретация: кривая показывает, как меняется предсказанная вероятность исхода при росте последней глюкозы "
        "при фиксированных ковариатах; это иллюстрация возможной нелинейности, а не причинное утверждение."
    )
    interpretations["fig_forest"] = (
        "Интерпретация: forest plot показывает направление и неопределенность OR в лучшей многофакторной модели."
    )
    interpretations["fig_roc"] = (
        "Интерпретация: ROC-кривая отражает дискриминационную способность лучшей модели, но не решает вопрос причинности."
    )
    interpretations["fig_forest_icu"] = "Интерпретация: forest plot показывает OR предикторов модели исхода 'ОРИТ'."
    interpretations["fig_roc_icu"] = "Интерпретация: ROC-кривая отражает дискриминационную способность модели исхода 'ОРИТ'."
    interpretations["fig_forest_newdm"] = "Интерпретация: forest plot показывает OR предикторов модели исхода 'СД при выписке (без СД2)'."
    interpretations["fig_roc_newdm"] = "Интерпретация: ROC-кривая отражает дискриминационную способность модели исхода 'СД при выписке (без СД2)'."

    discussion_lines = _discussion_text(
        outcome_n=n_outcome,
        death_n=n_death,
        glycemia_uni=glycemia_uni,
        comorb_table=comorb_table,
        treatment_table=treatment_table,
        metrics_df=model_metrics_df,
    )
    conclusions = [
        "Для клинической практики ключевой вывод должен опираться на согласованный сигнал из однофакторных и многофакторных оценок.",
        "При триаже пациентов целесообразно учитывать гликемию в связке с тяжестью (SpO2/NEWS2), а не изолированно.",
        "Отдельный мониторинг пациентов без известного СД2 с дисгликемией оправдан как гипотеза risk-based маршрутизации.",
        "Необходима валидация на независимой когорте и/или проспективном дизайне с лучшим контролем пропусков.",
    ]

    figures = {
        "missingness": fig_missingness,
        "glucose_adm_cat": fig_glucose_adm,
        "glucose_last_cat": fig_glucose_last,
        "delta_glucose": fig_delta_glucose,
        "corr": fig_corr,
        "spline": fig_spline,
        "forest": fig_forest,
        "roc": fig_roc,
        "forest_icu": fig_forest_icu,
        "roc_icu": fig_roc_icu,
        "forest_newdm": fig_forest_newdm,
        "roc_newdm": fig_roc_newdm,
    }

    out_docx = run_dir / "covid_glycemia_structured_report.docx"
    _write_report_docx(
        out_docx=out_docx,
        cohort_meta={"n_total": n_total, "n_outcome": n_outcome, "n_death": n_death},
        coverage_df=coverage_df,
        numeric_table=numeric_table,
        comorb_table=comorb_table,
        treatment_table=treatment_table,
        paired_table=paired_table,
        glycemia_uni=glycemia_uni,
        model_metrics=model_metrics_df,
        coef_table_best=coef_table_best,
        sensitivity_table=sensitivity_df,
        spline_table=spline_table,
        icu_model_metrics=icu_model_metrics_df,
        icu_coef_best=icu_coef_best,
        newdm_model_metrics=newdm_model_metrics_df,
        newdm_coef_best=newdm_coef_best,
        interpretations=interpretations,
        figure_paths=figures,
        discussion_lines=discussion_lines,
        conclusions=conclusions,
    )

    # Markdown report
    md_lines: List[str] = []
    md_lines.append("# COVID-19: исходы и влияние гликемии (структурированный отчет)")
    md_lines.append("")
    md_lines.append(f"- Дата: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    md_lines.append(f"- Файл: `{input_path}`")
    md_lines.append(f"- N всего: {n_total}; N с известным исходом: {n_outcome}; летальные исходы: {n_death}")
    md_lines.append("")
    md_lines.append("## 1. Карта переменных и покрытие")
    md_lines.append(
        "Аналитическая схема: (1) нормализация исхода и ключевых переменных; "
        "(2) baseline-сравнение по исходу; (3) crude OR для коморбидности и терапии; "
        "(4) однофакторные модели для гликемии и тяжести; (5) многофакторные модели с оценкой AUC/Brier; "
        "(6) интерпретация каждого артефакта и сводное обсуждение."
    )
    md_lines.append("")
    md_lines.append(_safe_markdown(coverage_df, drop_internal=False))
    md_lines.append("")
    md_lines.append(interpretations["coverage"])
    md_lines.append("")

    md_lines.append("## 2. Количественные показатели по исходу")
    md_lines.append(_safe_markdown(numeric_table))
    md_lines.append("")
    md_lines.append(interpretations["numeric"])
    md_lines.append("")

    md_lines.append("## 3. Коморбидность и исход")
    md_lines.append(_safe_markdown(comorb_table))
    md_lines.append("")
    md_lines.append(interpretations["comorb"])
    md_lines.append("")

    md_lines.append("## 4. Лечение и исход")
    md_lines.append(_safe_markdown(treatment_table))
    md_lines.append("")
    md_lines.append(interpretations["treatment"])
    md_lines.append("")

    md_lines.append("## 5. Динамика гликемии (парные сравнения)")
    md_lines.append(_safe_markdown(paired_table))
    md_lines.append("")
    md_lines.append(interpretations["paired"])
    md_lines.append("")

    md_lines.append("## 6. Однофакторные модели (гликемия)")
    md_lines.append(_safe_markdown(glycemia_uni))
    md_lines.append("")
    md_lines.append(interpretations["glycemia_uni"])
    md_lines.append("")

    md_lines.append("## 7. Многофакторные модели")
    md_lines.append(_safe_markdown(model_metrics_df, drop_internal=False))
    md_lines.append("")
    md_lines.append(interpretations["models"])
    md_lines.append("")
    md_lines.append("Коэффициенты лучшей модели:")
    md_lines.append(_safe_markdown(coef_table_best))
    md_lines.append("")

    if not sensitivity_df.empty:
        md_lines.append("### 7.1. Чувствительность к пропускам")
        md_lines.append(_safe_markdown(sensitivity_df, drop_internal=False))
        md_lines.append("")
        md_lines.append(interpretations["sensitivity"])
        md_lines.append("")

    if not spline_table.empty or (figures.get("spline") and figures["spline"].exists()):
        md_lines.append("### 7.2. Нелинейность последней глюкозы (spline)")
        if not spline_table.empty:
            md_lines.append(_safe_markdown(spline_table, drop_internal=False))
            md_lines.append("")
        md_lines.append(interpretations["spline"])
        md_lines.append("")
        fig_path = figures.get("spline")
        if fig_path is not None and fig_path.exists():
            title = "Spline-эффект последней глюкозы"
            md_lines.append(f"![{title}]({fig_path})")
            md_lines.append("")
            md_lines.append(interpretations.get("fig_spline", ""))
            md_lines.append("")

    if not icu_model_metrics_df.empty:
        md_lines.append("### 7.3. Прогнозирование исхода: ОРИТ (exploratory)")
        md_lines.append(_safe_markdown(icu_model_metrics_df, drop_internal=False))
        md_lines.append("")
        md_lines.append(interpretations.get("icu_models", ""))
        md_lines.append("")
        if not icu_coef_best.empty:
            md_lines.append("Коэффициенты лучшей модели (ОРИТ):")
            md_lines.append(_safe_markdown(icu_coef_best))
            md_lines.append("")
        for key, title, interp_key in [
            ("forest_icu", "Forest plot (ОРИТ)", "fig_forest_icu"),
            ("roc_icu", "ROC (ОРИТ)", "fig_roc_icu"),
        ]:
            fig_path = figures.get(key)
            if fig_path is None or not fig_path.exists():
                continue
            md_lines.append(f"#### {title}")
            md_lines.append(f"![{title}]({fig_path})")
            md_lines.append("")
            md_lines.append(interpretations.get(interp_key, ""))
            md_lines.append("")

    if not newdm_model_metrics_df.empty:
        md_lines.append("### 7.4. Прогнозирование исхода: СД при выписке (без СД2 до госпитализации)")
        md_lines.append(_safe_markdown(newdm_model_metrics_df, drop_internal=False))
        md_lines.append("")
        md_lines.append(interpretations.get("newdm_models", ""))
        md_lines.append("")
        if not newdm_coef_best.empty:
            md_lines.append("Коэффициенты лучшей модели (СД при выписке, без СД2):")
            md_lines.append(_safe_markdown(newdm_coef_best))
            md_lines.append("")
        for key, title, interp_key in [
            ("forest_newdm", "Forest plot (СД при выписке, без СД2)", "fig_forest_newdm"),
            ("roc_newdm", "ROC (СД при выписке, без СД2)", "fig_roc_newdm"),
        ]:
            fig_path = figures.get(key)
            if fig_path is None or not fig_path.exists():
                continue
            md_lines.append(f"#### {title}")
            md_lines.append(f"![{title}]({fig_path})")
            md_lines.append("")
            md_lines.append(interpretations.get(interp_key, ""))
            md_lines.append("")

    md_lines.append("## 8. Графики и интерпретация")
    figure_desc = [
        ("missingness", "Пропуски в ключевых переменных", "fig_missingness"),
        ("delta_glucose", "Динамика глюкозы: Δ(последняя - при поступлении) по исходу", "fig_delta_glucose"),
        ("corr", "Корреляционная матрица ключевых количественных переменных (Spearman)", "fig_corr"),
        ("glucose_adm_cat", "Летальность по категориям глюкозы при поступлении", "fig_glucose_adm"),
        ("glucose_last_cat", "Летальность по категориям последней глюкозы", "fig_glucose_last"),
        ("forest", "Forest plot OR лучшей модели", "fig_forest"),
        ("roc", "ROC-кривая лучшей модели", "fig_roc"),
    ]
    for key, title, interp_key in figure_desc:
        fig_path = figures.get(key)
        if fig_path is None or not fig_path.exists():
            continue
        md_lines.append(f"### {title}")
        md_lines.append(f"![{title}]({fig_path})")
        md_lines.append("")
        md_lines.append(interpretations.get(interp_key, ""))
        md_lines.append("")

    md_lines.append("## 9. Обсуждение")
    for line in discussion_lines:
        md_lines.append(f"- {line}")
    md_lines.append("")

    md_lines.append("## 10. Практические выводы")
    for i, c in enumerate(conclusions, start=1):
        md_lines.append(f"{i}. {c}")
    md_lines.append("")

    out_md = run_dir / "covid_glycemia_structured_report.md"
    out_md.write_text("\n".join(md_lines), encoding="utf-8")

    summary = {
        "run_dir": str(run_dir),
        "input": str(input_path),
        "docx_report": str(out_docx),
        "markdown_report": str(out_md),
        "tables_dir": str(tables_dir),
        "figures_dir": str(figs_dir),
        "n_total": n_total,
        "n_outcome": n_outcome,
        "n_death": n_death,
        "models_fitted": [m.model_id for m in model_results],
        "best_model": best_model.model_id if best_model is not None else None,
    }
    (run_dir / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate structured COVID glycemia report.")
    parser.add_argument("--input", type=Path, default=DEFAULT_INPUT, help="Path to source Excel file.")
    parser.add_argument("--output-root", type=Path, default=DEFAULT_OUTPUT_ROOT, help="Output root directory.")
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=None,
        help="Optional fixed output directory (overwrites artifacts in-place).",
    )
    args = parser.parse_args()

    output_dir = args.output_dir
    if output_dir is not None and not output_dir.is_absolute():
        candidate = (PROJECT_ROOT / output_dir).resolve()
        if str(candidate).startswith(str(args.output_root.resolve())):
            output_dir = candidate
        else:
            output_dir = args.output_root / output_dir
    summary = run_report(args.input, args.output_root, output_dir=output_dir)
    print(json.dumps(summary, ensure_ascii=False))


if __name__ == "__main__":
    main()
