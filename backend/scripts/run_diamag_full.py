#!/usr/bin/env python3
"""
DIAMAG Clinical Trial - COMPREHENSIVE Analysis Script

Полный статистический анализ с:
- Анализом КАЖДОЙ временной точки (V2, V3, V4, V5, V6)
- Попарными сравнениями между группами
- Изменениями внутри групп (парные тесты)
- Полной интерпретацией H0/H1
- Всеми возможными статистиками
"""

import os
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional
import warnings

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import pandas as pd
import numpy as np
from scipy import stats
import pingouin as pg
import json
import httpx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns

warnings.filterwarnings('ignore')

try:
    from app.stats.mixed_effects import MixedEffectsEngine
except Exception:
    MixedEffectsEngine = None

try:
    from app.core.config import settings as app_settings
except Exception:
    app_settings = None

# ============================================================
# CONFIGURATION
# ============================================================

EXCEL_PATH = PROJECT_ROOT.parent / "docs" / "Первичка для анализа работа.xlsx"
OUTPUT_DIR = PROJECT_ROOT / "output"

ENDPOINTS = {
    "updrs_part3": {
        "name": "UPDRS часть 3 (Двигательные функции)",
        "short": "UPDRS III",
        "cols": {
            "V2": "УШОБП часть 3 «Двигательные функции» баллы V2",
            "V3": "УШОБП часть 3 «Двигательные функции» баллы V3",
            "V4": "УШОБП часть 3 «Двигательные функции» баллы V4",
            "V5": "УШОБП часть 3 «Двигательные функции» баллы V5",
            "V6": "УШОБП часть 3 «Двигательные функции» баллы V6",
        },
        "primary": True,
        "direction": "lower_is_better",
    },
    "updrs_part2": {
        "name": "UPDRS часть 2 (Повседневная активность)",
        "short": "UPDRS II",
        "cols": {
            "V2": "УШОБП часть 2 «Повседневная активность» баллы V2",
            "V3": "УШОБП часть 2 «Повседневная активность» баллы V3",
            "V4": "УШОБП часть 2 «Повседневная активность» баллы V4",
            "V5": "УШОБП часть 2 «Повседневная активность» баллы V5",
            "V6": "УШОБП часть 2 «Повседневная активность» баллы V6",
        },
        "primary": True,
        "direction": "lower_is_better",
    },
    "dass21": {
        "name": "DASS-21 (Депрессия, тревога, стресс)",
        "short": "DASS-21",
        "cols": {
            "V2": "Шкала депрессии тревоги и стресса DASS-21 баллы V2",
            "V3": "Шкала депрессии тревоги и стресса DASS-21 баллы V3",
            "V4": "Шкала депрессии тревоги и стресса DASS-21 баллы V4",
            "V5": "Шкала депрессии тревоги и стресса DASS-21 баллы V5",
            "V6": "Шкала депрессии, тревоги и стресса DASS-21баллы V6",
        },
        "primary": False,
        "direction": "lower_is_better",
    },
    "epworth": {
        "name": "Шкала сонливости Эпуорта",
        "short": "Epworth",
        "cols": {
            "V2": "Шкала оценки сонливости Эпуорта баллы V2",
            "V3": "Шкала оценки сонливости Эпуорта баллы V3",
            "V4": "Шкала оценки сонливости Эпуорта баллы V4",
            "V5": "Шкала оценки сонливости Эпуорта баллы V5",
            "V6": "Шкала оценки сонливости Эпуорта баллы V6",
        },
        "primary": False,
        "direction": "lower_is_better",
    },
    "apathy": {
        "name": "Шкала апатии Старкстейна",
        "short": "Апатия",
        "cols": {
            "V2": "Шкала апатии Старкстейна баллы V2",
            "V3": "Шкала апатии Старкстейна баллы V3",
            "V4": "Шкала апатии Старкстейна баллы V4",
            "V5": "Шкала апатии Старкстейна баллы V5",
            "V6": "Шкала апатии Старкстейна баллы V6",
        },
        "primary": False,
        "direction": "lower_is_better",
    },
    "fatigue": {
        "name": "Шкала утомляемости при БП",
        "short": "Утомляемость",
        "cols": {
            "V2": "Шкала оценки утомляемости при БП баллы V2",
            "V3": "Шкала оценки утомляемости при БП баллы V3",
            "V4": "Шкала оценки утомляемости при БП баллы V4",
            "V5": "Шкала оценки утомляемости при БП баллы V5",
            "V6": "Шкала оценки утомляемости при БП баллы V6",
        },
        "primary": False,
        "direction": "lower_is_better",
    },
    "pdq39": {
        "name": "PDQ-39 (Качество жизни)",
        "short": "PDQ-39",
        "cols": {
            "V2": "Шкала оценки качества жизни при БП баллы V2",
            "V3": "Шкала оценки качества жизни при БП баллы V3",
            "V4": "Шкала оценки качества жизни при БП баллы",
            "V5": "Шкала оценки качества жизни при БП баллы V5",
            "V6": "Шкала оценки качества жизни при БП баллы V6",
        },
        "primary": False,
        "direction": "lower_is_better",
    },
}

GROUP_COL = "Группа"
ID_COL = "ID № участника исследования"
VISITS = ["V2", "V3", "V4", "V5", "V6"]

# ============================================================
# STATISTICS FUNCTIONS
# ============================================================

def descriptive(values):
    """Full descriptive statistics."""
    clean = pd.Series(values).dropna()
    n = len(clean)
    if n == 0:
        return {"n": 0}
    return {
        "n": n,
        "mean": float(clean.mean()),
        "sd": float(clean.std()),
        "se": float(clean.std() / np.sqrt(n)) if n > 1 else np.nan,
        "median": float(clean.median()),
        "q1": float(clean.quantile(0.25)),
        "q3": float(clean.quantile(0.75)),
        "min": float(clean.min()),
        "max": float(clean.max()),
        "skew": float(clean.skew()) if n > 2 else np.nan,
        "kurtosis": float(clean.kurtosis()) if n > 3 else np.nan,
    }


def normality_test(values):
    """Shapiro-Wilk normality test."""
    clean = pd.Series(values).dropna()
    n = len(clean)
    if n < 3 or n > 5000:
        return {"normal": None, "p": np.nan}
    try:
        stat, p = stats.shapiro(clean)
        return {"normal": p > 0.05, "p": float(p), "W": float(stat)}
    except:
        return {"normal": None, "p": np.nan}


def kruskal_wallis(groups_data):
    """Kruskal-Wallis test with effect size."""
    valid = [g.dropna() for g in groups_data if len(g.dropna()) > 0]
    if len(valid) < 2:
        return {"error": "Not enough groups"}
    
    stat, p = stats.kruskal(*valid)
    n_total = sum(len(g) for g in valid)
    k = len(valid)
    epsilon_sq = (stat - k + 1) / (n_total - k) if n_total > k else np.nan
    
    return {
        "H": float(stat),
        "p": float(p),
        "epsilon_sq": float(epsilon_sq) if np.isfinite(epsilon_sq) else None,
        "significant": p < 0.05,
    }


def mann_whitney(group1, group2):
    """Mann-Whitney U test for two groups."""
    g1 = pd.Series(group1).dropna()
    g2 = pd.Series(group2).dropna()
    
    if len(g1) < 2 or len(g2) < 2:
        return {"error": "Not enough data"}
    
    stat, p = stats.mannwhitneyu(g1, g2, alternative='two-sided')
    
    # Effect size: rank-biserial correlation
    n1, n2 = len(g1), len(g2)
    r = 1 - (2 * stat) / (n1 * n2)
    
    return {
        "U": float(stat),
        "p": float(p),
        "r": float(r),  # rank-biserial
        "significant": p < 0.05,
        "n1": n1,
        "n2": n2,
    }


def wilcoxon_signed_rank(before, after):
    """Wilcoxon signed-rank test for paired data."""
    b = pd.Series(before)
    a = pd.Series(after)
    
    # Align by index
    valid = pd.DataFrame({"before": b, "after": a}).dropna()
    if len(valid) < 5:
        return {"error": "Not enough paired data"}
    
    diff = valid["after"] - valid["before"]
    
    try:
        stat, p = stats.wilcoxon(valid["before"], valid["after"])
        
        # Effect size: r = Z / sqrt(N)
        n = len(valid)
        z = stats.norm.ppf(p / 2)
        r = abs(z) / np.sqrt(n)
        
        return {
            "W": float(stat),
            "p": float(p),
            "r": float(r),
            "significant": p < 0.05,
            "n_pairs": n,
            "mean_diff": float(diff.mean()),
            "median_diff": float(diff.median()),
        }
    except:
        return {"error": "Test failed"}


def bayes_factor(p_value):
    """BF10 from p-value (Sellke bound)."""
    if p_value is None or not np.isfinite(p_value) or p_value <= 0 or p_value >= 1:
        return np.nan
    try:
        return min(-1 / (np.e * p_value * np.log(p_value)), 1000)
    except:
        return np.nan


def holm_adjust(p_values: List[float]) -> List[float]:
    ps = np.array([float(p) if p is not None and np.isfinite(p) else np.nan for p in p_values], dtype=float)
    m = len(ps)
    if m == 0:
        return []
    order = np.argsort(ps)
    adj = np.full(m, np.nan, dtype=float)
    prev = 0.0
    for rank, idx in enumerate(order):
        p = ps[idx]
        if not np.isfinite(p):
            continue
        a = (m - rank) * p
        if a < prev:
            a = prev
        prev = a
        adj[idx] = min(a, 1.0)
    return adj.tolist()


def add_toc(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_sep = OxmlElement('w:fldChar')
    fld_char_sep.set(qn('w:fldCharType'), 'separate')

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def enable_word_update_fields_on_open(doc: Document) -> None:
    try:
        settings_elm = doc.settings.element
        for child in settings_elm.iterchildren():
            if child.tag == qn('w:updateFields'):
                child.set(qn('w:val'), 'true')
                return
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        settings_elm.append(update_fields)
    except Exception:
        return


def _is_openrouter_model(model: str) -> bool:
    return "/" in str(model or "")


def _resolve_llm_target(model: str) -> tuple[str, Optional[str]]:
    if app_settings is None:
        return "", None
    if _is_openrouter_model(model):
        url = getattr(app_settings, "OPENROUTER_API_URL", None) or getattr(app_settings, "GLM_API_URL", "")
        api_key = getattr(app_settings, "OPENROUTER_API_KEY", None) or getattr(app_settings, "GLM_API_KEY", None)
        return url, api_key
    return getattr(app_settings, "GLM_API_URL", ""), getattr(app_settings, "GLM_API_KEY", None)


def llm_chat_completion_sync(
    *,
    model: str,
    prompt: str,
    temperature: float = 0.3,
    max_tokens: int = 1600,
    timeout_s: float = 60.0,
) -> Optional[str]:
    if app_settings is None:
        return None
    if not getattr(app_settings, "GLM_ENABLED", True):
        return None
    url, api_key = _resolve_llm_target(model)
    if not url or not api_key:
        return None
    if "/chat/completions" not in str(url):
        u = str(url).rstrip("/")
        if u.endswith("/paas/v4") or u.endswith("/coding/paas/v4"):
            url = u + "/chat/completions"

    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": float(temperature),
        "max_tokens": int(max_tokens),
    }
    if "api.z.ai" in str(url):
        payload["thinking"] = {"type": "disabled"}
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "Accept-Language": "ru-RU,ru",
    }
    try:
        with httpx.Client(timeout=timeout_s) as client:
            resp = client.post(url, json=payload, headers=headers)
            resp.raise_for_status()
            data = resp.json()
            msg = (data.get("choices") or [{}])[0].get("message") or {}
            content = str(msg.get("content") or "").strip()
            if not content:
                content = str(msg.get("reasoning_content") or "").strip()
            return content or None
    except Exception:
        return None


def _doc_add_heading(doc: Document, text: str, level: int = 1):
    t = str(text or "")
    is_numbered_level1 = (level == 1) and bool(__import__("re").match(r"^\s*\d+\.", t))
    if is_numbered_level1:
        def _has_trailing_page_break() -> bool:
            try:
                if not getattr(doc, "paragraphs", None):
                    return False
                p = doc.paragraphs[-1]
                for br in p._p.xpath(".//w:br"):
                    if br.get(qn("w:type")) == "page":
                        return True
            except Exception:
                return False
            return False

        started = bool(getattr(doc, "_statproject_numbered_sections_started", False))
        if started:
            if not _has_trailing_page_break():
                doc.add_page_break()
        else:
            setattr(doc, "_statproject_numbered_sections_started", True)
    return doc.add_heading(t, level=level)


def _try_parse_json(text: str):
    if not text:
        return None
    s = str(text).strip()
    try:
        return json.loads(s)
    except Exception:
        pass
    try:
        start = s.find("[")
        end = s.rfind("]")
        if start != -1 and end != -1 and end > start:
            return json.loads(s[start : end + 1])
    except Exception:
        pass
    try:
        start = s.find("{")
        end = s.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(s[start : end + 1])
    except Exception:
        return None
    return None


def generate_ai_artifact_explanations(
    *,
    artifacts: List[Dict[str, Any]],
    preferred_model: str = "glm-4.7",
    chunk_size: int = 6,
) -> Dict[str, Dict[str, Any]]:
    if not artifacts:
        return {}

    model_candidates = [preferred_model]
    if app_settings is not None:
        model_candidates.append(getattr(app_settings, "GLM_MODEL", ""))
    model_candidates = [m for m in model_candidates if m]
    if not model_candidates:
        return {}

    out: Dict[str, Dict[str, Any]] = {}
    chunks = _chunk_list(artifacts, chunk_size=chunk_size)
    for part_idx, chunk in enumerate(chunks, start=1):
        prompt = (
            "Ты — научный редактор статистического отчёта. Сформулируй интерпретацию каждого артефакта в официально-академическом стиле.\n"
            "Строго используй только предоставленные данные (не придумывай чисел).\n"
            "Нельзя: списки с маркерами, многоточия, разговорные формулировки, упоминания AI/LLM/модели.\n\n"
            "Верни ТОЛЬКО валидный JSON-массив, без пояснений и без Markdown.\n"
            "Формат каждого элемента массива:\n"
            "{\"id\": string, \"title\": string, \"plain\": string, \"takeaways\": [string, ...], \"caveats\": [string, ...]}\n\n"
            "Правила текста:\n"
            "- title: коротко и официально (как подпись к блоку интерпретации).\n"
            "- plain: 3–6 предложений; опиши, что показано, где различия, и как это подтверждается (p_adj/ES/Δ/BF₁₀).\n"
            "- takeaways: 2–4 коротких факта-утверждения, без маркеров.\n"
            "- caveats: 0–2 формулировки об ограничениях (коррекция, мощность, вариабельность).\n\n"
            f"Данные (chunk {part_idx}/{len(chunks)}): {json.dumps(chunk, ensure_ascii=False, default=str)}"
        )

        content = None
        for m in model_candidates:
            content = llm_chat_completion_sync(model=m, prompt=prompt, temperature=0.2, max_tokens=1800, timeout_s=90.0)
            if content:
                break
        parsed = _try_parse_json(content or "")
        if not isinstance(parsed, list):
            continue
        for item in parsed:
            if not isinstance(item, dict):
                continue
            item_id = str(item.get("id") or "").strip()
            if not item_id:
                continue
            out[item_id] = item
    return out


def generate_ai_rewrites(
    *,
    items: List[Dict[str, Any]],
    preferred_model: str = "glm-4.7",
    chunk_size: int = 8,
) -> Dict[str, str]:
    if not items:
        return {}
    model_candidates = [preferred_model]
    if app_settings is not None:
        model_candidates.append(getattr(app_settings, "GLM_MODEL", ""))
    model_candidates = [m for m in model_candidates if m]
    if not model_candidates:
        return {}

    out: Dict[str, str] = {}
    chunks = _chunk_list(items, chunk_size=chunk_size)
    for part_idx, chunk in enumerate(chunks, start=1):
        prompt = (
            "Ты — научный редактор официального отчёта. Сформулируй краткое резюме к каждому абзацу.\n"
            "Стиль: нейтральный, академический, без разговорных формулировок.\n"
            "Нельзя: списки, маркеры, многоточия, обрыв слов, упоминания AI/LLM/модели.\n\n"
            "Верни ТОЛЬКО валидный JSON-массив, без пояснений и без Markdown.\n"
            "Формат элемента: {\"id\": string, \"rewritten\": string}\n"
            "rewritten: 1–2 предложения, до 240 символов, завершать точкой.\n\n"
            f"Тексты (chunk {part_idx}/{len(chunks)}): {json.dumps(chunk, ensure_ascii=False, default=str)}"
        )
        content = None
        for m in model_candidates:
            content = llm_chat_completion_sync(model=m, prompt=prompt, temperature=0.2, max_tokens=1600, timeout_s=90.0)
            if content:
                break
        parsed = _try_parse_json(content or "")
        if not isinstance(parsed, list):
            continue
        for item in parsed:
            if not isinstance(item, dict):
                continue
            item_id = str(item.get("id") or "").strip()
            rewritten = str(item.get("rewritten") or "").strip()
            if item_id and rewritten:
                out[item_id] = rewritten
    return out


def _doc_add_ai_block(doc: Document, payload: Optional[Dict[str, Any]]):
    if not isinstance(payload, dict):
        return
    title = str(payload.get("title") or "").strip()
    plain = str(payload.get("plain") or "").strip()
    takeaways = payload.get("takeaways")
    caveats = payload.get("caveats")

    def _clean_text(v: Any) -> str:
        s = str(v or "").strip()
        if not s:
            return ""
        s = s.replace("\u2022", "").replace("•", " ")
        s = " ".join(s.split())
        return s.strip()

    head = doc.add_paragraph()
    head_run = head.add_run("Интерпретация результатов")
    head_run.bold = True

    if title:
        doc.add_paragraph(_clean_text(title))
    if plain:
        doc.add_paragraph(_clean_text(plain))

    if isinstance(takeaways, list) and takeaways:
        items = [_clean_text(x) for x in takeaways[:6]]
        items = [x for x in items if x]
        if items:
            doc.add_paragraph("Ключевые выводы: " + "; ".join(items) + ".")

    if isinstance(caveats, list) and caveats:
        items = [_clean_text(x) for x in caveats[:4]]
        items = [x for x in items if x]
        if items:
            doc.add_paragraph("Ограничения интерпретации: " + "; ".join(items) + ".")

    doc.add_paragraph()


def _chunk_list(items: List[Any], chunk_size: int) -> List[List[Any]]:
    if chunk_size <= 0:
        return [items]
    return [items[i:i + chunk_size] for i in range(0, len(items), chunk_size)]


def _safe_json(obj: Any) -> str:
    return json.dumps(obj, ensure_ascii=False, default=str)


def generate_ai_discussion_chunked(
    *,
    exec_summary: Dict[str, Any],
    all_results: Dict[str, Any],
    active_vs_placebo: Dict[str, Any],
    responders: Dict[str, Any],
    preferred_model: str = "glm-4.7",
    chunk_size: int = 2,
) -> Optional[Dict[str, List[str]]]:
    endpoints_payload: List[Dict[str, Any]] = []
    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)

        kw_by_visit = []
        for v in VISITS:
            vdata = result.get("by_visit", {}).get(v, {})
            kw = vdata.get("kruskal", {})
            kw_by_visit.append({
                "visit": v,
                "p": kw.get("p"),
                "epsilon_sq": kw.get("epsilon_sq"),
                "bf10": vdata.get("bf10"),
            })

        sig_pairs = []
        for v in VISITS:
            pv = result.get("pairwise", {}).get(v, {})
            if not isinstance(pv, dict):
                continue
            for pair_key, st in pv.items():
                if not isinstance(st, dict) or st.get("error"):
                    continue
                p_adj = st.get("p_adj")
                if p_adj is None or not np.isfinite(p_adj) or float(p_adj) >= 0.05:
                    continue
                sig_pairs.append({
                    "visit": v,
                    "pair": str(pair_key).replace("_vs_", " vs "),
                    "p_adj": p_adj,
                    "bf10": st.get("bf10"),
                    "delta_abs": st.get("diff_median"),
                    "delta_pct": st.get("diff_pct"),
                    "es_r": st.get("r"),
                })

        pooled_visits = (active_vs_placebo.get(ep_key, {}) or {}).get("visits", {})
        pooled_rows = []
        if isinstance(pooled_visits, dict):
            for v, d in pooled_visits.items():
                if not isinstance(d, dict):
                    continue
                pooled_rows.append({
                    "visit": v,
                    "p_adj": d.get("p_adj", d.get("p_value")),
                    "bf10": d.get("bf10"),
                    "delta_abs": d.get("diff_median"),
                    "delta_pct": d.get("diff_pct"),
                    "es": d.get("effect_size"),
                    "significant": d.get("significant"),
                })

        mm4 = result.get("mixed_effects_4g", {})
        mmp = result.get("mixed_effects_pooled", {})

        resp_ep = responders.get(ep_key, {}) if isinstance(responders, dict) else {}
        resp_visits = []
        if isinstance(resp_ep, dict):
            for v in VISITS:
                vres = (resp_ep.get("visits", {}) or {}).get(v, {})
                if not isinstance(vres, dict):
                    continue
                resp_visits.append({
                    "visit": v,
                    "groups": vres.get("groups", []),
                    "pooled": vres.get("pooled", {}),
                    "test": vres.get("test", {}),
                })

        endpoints_payload.append({
            "endpoint": short,
            "primary": bool(result.get("primary")),
            "direction": result.get("direction"),
            "kw_by_visit": kw_by_visit,
            "significant_posthoc": sig_pairs,
            "pooled_active_vs_placebo": pooled_rows,
            "mixed_effects_4g": {
                "interaction_p": (mm4 or {}).get("interaction_p_value"),
                "interaction_significant": (mm4 or {}).get("interaction", {}).get("significant"),
                "interpretation": (mm4 or {}).get("interaction", {}).get("interpretation"),
            },
            "mixed_effects_pooled": {
                "interaction_p": (mmp or {}).get("interaction_p_value"),
                "interaction_significant": (mmp or {}).get("interaction", {}).get("significant"),
                "interpretation": (mmp or {}).get("interaction", {}).get("interpretation"),
            },
            "responders": resp_visits,
        })

    base_context = {
        "goals": [
            "Оценить эффективность аппарата ДИАМАГ по сравнению с плацебо",
            "Проверить различия между 4 рандомизированными группами по визитам V2–V6",
            "Оценить динамику (повторные измерения) и клиническую значимость (респондеры ≥20%)",
        ],
        "hypotheses": [
            "H0: различий между группами нет; H1: различия есть",
            "Для укрупнения (Active vs Placebo): H0 Active=Placebo; H1 Active≠Placebo",
            "Для mixed effects: H0 нет взаимодействия Визит×Группа; H1 взаимодействие есть",
        ],
        "notes": [
            "Используй p_adj (Холм) как основной критерий; BF₁₀ — сила доказательств; ES/Δ — величина и клиническая направленность",
            "Не упоминай 'AI/GLM/модель'; пиши как раздел отчёта",
        ],
        "n_patients": exec_summary.get("n_patients"),
    }

    model_candidates = [preferred_model]
    if app_settings is not None:
        model_candidates.append(getattr(app_settings, "GLM_MODEL", ""))
    model_candidates = [m for m in model_candidates if m]
    if not model_candidates:
        return None

    chunks = _chunk_list(endpoints_payload, chunk_size=chunk_size)
    discussion_parts: List[str] = []
    for part_idx, chunk in enumerate(chunks, start=1):
        prompt = (
            "Ты — клинический биостатистик. Напиши часть раздела 'Обсуждение' на русском, "
            "строго опираясь на данные ниже. Тон: научный, без воды.\n\n"
            "Требования:\n"
            "- Привяжи интерпретацию к целям, задачам и гипотезам.\n"
            "- Для каждого показателя в этом чанке: кратко опиши 4-групповое сравнение по визитам, укрупнение Active vs Placebo, post-hoc, mixed effects и респондеров (если есть).\n"
            "- Если значимых post-hoc нет для показателя, скажи, что после коррекции Холма значимых различий не выявлено, но укажи наиболее близкие сигналы (самые низкие p_adj из укрупнения или KW).\n"
            "- 5–10 абзацев, без списков.\n\n"
            f"Контекст: {_safe_json(base_context)}\n\n"
            f"Данные (chunk {part_idx}/{len(chunks)}): {_safe_json(chunk)}\n"
        )

        content = None
        for m in model_candidates:
            content = llm_chat_completion_sync(model=m, prompt=prompt, temperature=0.2, max_tokens=1800, timeout_s=90.0)
            if content:
                break
        if content:
            discussion_parts.append(content.strip())

    if not discussion_parts:
        return None

    conclusion_prompt = (
        "Ты — клинический биостатистик. На основе данных ниже составь раздел 'Выводы' на русском.\n\n"
        "Требования:\n"
        "- 8–12 коротких пунктов (каждый с '• ').\n"
        "- Каждый пункт должен быть проверяемым по данным: что показали 4 группы, укрупнение, post-hoc, mixed effects, респондеры.\n"
        "- Обязательно: 2–3 пункта про ограничения и интерпретацию BF₁₀/ES.\n"
        "- Не упоминай 'AI/GLM/модель'.\n\n"
        f"Контекст: {_safe_json(base_context)}\n\n"
        f"Данные (все показатели): {_safe_json(endpoints_payload)}\n"
    )

    conclusion_text = None
    for m in model_candidates:
        conclusion_text = llm_chat_completion_sync(model=m, prompt=conclusion_prompt, temperature=0.2, max_tokens=1400, timeout_s=90.0)
        if conclusion_text:
            break

    conclusions: List[str] = []
    if conclusion_text:
        for line in conclusion_text.splitlines():
            t = line.strip()
            if not t:
                continue
            if t.startswith("•"):
                conclusions.append(t)
            elif t.startswith("-"):
                conclusions.append("• " + t.lstrip("- ").strip())
    if not conclusions:
        return None

    discussion_paragraphs: List[str] = []
    for block in discussion_parts:
        for para in block.split("\n\n"):
            t = para.strip()
            if t:
                discussion_paragraphs.append(t)

    return {"discussion": discussion_paragraphs, "conclusions": conclusions}


def fit_mixed_effects(df: pd.DataFrame, endpoint_key: str, pooled: bool) -> Dict[str, Any]:
    if MixedEffectsEngine is None:
        return {"error": "MixedEffectsEngine unavailable"}

    cfg = ENDPOINTS[endpoint_key]
    long_rows = []
    for visit, col in cfg.get("cols", {}).items():
        if not col or col not in df.columns:
            continue
        tmp = df[[ID_COL, GROUP_COL, col]].copy()
        tmp.columns = ["subject", "group", "value"]
        tmp["visit"] = visit
        long_rows.append(tmp)

    if not long_rows:
        return {"error": "No data"}

    long_df = pd.concat(long_rows, ignore_index=True)
    long_df["subject"] = long_df["subject"].astype(str)
    long_df["group"] = long_df["group"].astype(str)
    long_df["visit"] = long_df["visit"].astype(str)

    if pooled:
        long_df["group"] = long_df["group"].apply(lambda g: "Active" if g in ["1", "3"] else "Placebo")

    engine = MixedEffectsEngine()
    return engine.fit(
        long_df,
        outcome="value",
        time_col="visit",
        group_col="group",
        subject_col="subject",
        covariates=None,
        random_slope=False,
        alpha=0.05,
    )


def interpret_bf(bf):
    """Full Bayes Factor interpretation."""
    if not np.isfinite(bf):
        return "не определён"
    
    if bf > 100:
        return "экстремально сильное свидетельство в пользу H₁ (различия есть)"
    elif bf > 30:
        return "очень сильное свидетельство в пользу H₁"
    elif bf > 10:
        return "сильное свидетельство в пользу H₁"
    elif bf > 3:
        return "умеренное свидетельство в пользу H₁"
    elif bf > 1:
        return "слабое свидетельство в пользу H₁"
    elif bf > 1/3:
        return "неопределённое (данные нейтральны)"
    elif bf > 1/10:
        return "умеренное свидетельство в пользу H₀ (различий нет)"
    else:
        return "сильное свидетельство в пользу H₀"


def effect_size_interpret(es, es_type="r"):
    """Interpret effect size."""
    if es is None or not np.isfinite(es):
        return "—"
    
    abs_es = abs(es)
    
    if es_type in ("r", "rbc"):
        if abs_es < 0.1:
            return "незначительный"
        elif abs_es < 0.3:
            return "малый"
        elif abs_es < 0.5:
            return "средний"
        else:
            return "большой"
    elif es_type == "epsilon_sq":
        if abs_es < 0.01:
            return "незначительный"
        elif abs_es < 0.06:
            return "малый"
        elif abs_es < 0.14:
            return "средний"
        else:
            return "большой"
    else:  # Cohen's d
        if abs_es < 0.2:
            return "незначительный"
        elif abs_es < 0.5:
            return "малый"
        elif abs_es < 0.8:
            return "средний"
        else:
            return "большой"


# ============================================================
# PLOTTING FUNCTIONS
# ============================================================

def set_publication_plot_style():
    plt.rcParams.update(
        {
            "figure.dpi": 300,
            "savefig.dpi": 300,
            "font.size": 16,
            "axes.titlesize": 20,
            "axes.labelsize": 18,
            "xtick.labelsize": 16,
            "ytick.labelsize": 16,
            "legend.fontsize": 16,
            "axes.linewidth": 1.2,
        }
    )

def draw_significance_bracket(ax, x1, x2, y, h, text, color='black'):
    """Draw a statistical significance bracket."""
    ax.plot([x1, x1, x2, x2], [y, y+h, y+h, y], lw=1.5, c=color)
    ax.text((x1+x2)*.5, y+h, text, ha='center', va='bottom', color=color, fontsize=16, fontweight='bold')


def create_spaghetti_plot(df, endpoint_key, output_path, pairwise_by_visit=None):
    """Create mean ± SE spaghetti plot with significant p= annotations."""
    cfg = ENDPOINTS[endpoint_key]

    set_publication_plot_style()
    plt.figure(figsize=(10, 6))
    plt.style.use('seaborn-v0_8-whitegrid')
    
    # G1=Active 30d, G2=Placebo 30d, G3=Active 20d, G4=Placebo 20d
    colors = {'1': '#2ecc71', '2': '#e74c3c', '3': '#3498db', '4': '#9b59b6'}
    markers = {'1': 'o', '2': 's', '3': '^', '4': 'D'}
    labels = {
        '1': 'Группа 1 (ДИАМАГ 30д)', 
        '2': 'Группа 2 (Плацебо 30д)', 
        '3': 'Группа 3 (ДИАМАГ 20д)', 
        '4': 'Группа 4 (Плацебо 20д)'
    }
    
    groups = sorted(df[GROUP_COL].dropna().unique())
    
    # Plot lines
    max_y = -np.inf
    
    for g in groups:
        means = []
        ses = []
        visit_labels = []
        
        for visit in VISITS:
            col = cfg["cols"].get(visit)
            if col and col in df.columns:
                values = df[df[GROUP_COL] == g][col].dropna()
                if len(values) > 0:
                    means.append(values.mean())
                    ses.append(values.std() / np.sqrt(len(values)))
                    visit_labels.append(visit)
        
        if means:
            x = range(len(means))
            plt.errorbar(x, means, yerr=ses, 
                        marker=markers.get(str(g), 'o'),
                        color=colors.get(str(g), '#333'),
                        linewidth=2, markersize=8, capsize=4,
                        label=labels.get(str(g), f'Группа {g}'))
            
            # Track max y for annotations
            current_max = max([m + s for m, s in zip(means, ses)])
            if current_max > max_y:
                max_y = current_max

    # Add p-values (only significant)
    y_offset = max_y * 0.05
    for i, visit in enumerate(VISITS):
        col = cfg["cols"].get(visit)
        if not col or col not in df.columns:
            continue
            
        # Kruskal-Wallis
        group_arrays = [df[df[GROUP_COL] == g][col].dropna() for g in groups]
        kw = kruskal_wallis(group_arrays)
        
        label = None
        if kw.get("significant"):
            p_text = "p < 0.001" if kw["p"] < 0.001 else f"p = {kw['p']:.3f}"
            label = p_text
        else:
            pv = (pairwise_by_visit or {}).get(visit)
            if isinstance(pv, dict) and pv:
                best = None
                for pair_key, st in pv.items():
                    if not isinstance(st, dict) or "error" in st:
                        continue
                    p_adj = st.get("p_adj")
                    if p_adj is None or not np.isfinite(p_adj) or p_adj >= 0.05:
                        continue
                    if best is None or float(p_adj) < float(best["p_adj"]):
                        best = {"pair": pair_key, "p_adj": float(p_adj), "es": float(st.get("r", 0.0))}

                if best is not None:
                    try:
                        g1, g2 = str(best["pair"]).split("_vs_")
                        pair_txt = f"Г{g1}–Г{g2}"
                    except Exception:
                        pair_txt = "пара"
                    p_text = "p < 0.001" if best["p_adj"] < 0.001 else f"p = {best['p_adj']:.3f}"
                    label = f"{pair_txt}\n{p_text}\nES={best['es']:.2f}"

        if label:
            plt.text(i, max_y + y_offset, label, ha='center', va='bottom', fontsize=16, fontweight='bold', color='red')

    plt.xticks(range(len(VISITS)), VISITS)
    plt.xlabel('Визит', fontsize=18)
    plt.ylabel(f'{cfg["short"]} (Mean ± SE)', fontsize=18)
    plt.title(f'Динамика {cfg["short"]} по группам', fontsize=20, fontweight='bold')
    plt.legend(loc='best', frameon=True)
    plt.ylim(top=max_y * 1.2)  # Add space for annotations
    plt.tight_layout()
    
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


def create_visit_boxplot(df, endpoint_key, visit, output_path, pairwise_stats=None):
    """Create boxplot with swarmplot and significance brackets for a specific visit."""
    cfg = ENDPOINTS[endpoint_key]
    col = cfg["cols"].get(visit)
    
    if not col or col not in df.columns:
        return None
    
    set_publication_plot_style()
    plt.figure(figsize=(10, 7))
    plt.style.use('seaborn-v0_8-whitegrid')
    
    # Prepare data
    plot_data = []
    groups = sorted(df[GROUP_COL].dropna().unique())
    
    # Custom labels for x-axis
    group_labels = {
        '1': 'Гр.1\n(Актив 30)', 
        '2': 'Гр.2\n(Плацебо 30)', 
        '3': 'Гр.3\n(Актив 20)', 
        '4': 'Гр.4\n(Плацебо 20)'
    }
    
    for g in groups:
        values = df[df[GROUP_COL] == g][col].dropna()
        for v in values:
            plot_data.append({
                'GroupCode': str(g),
                'Группа': group_labels.get(str(g), f"Гр.{g}"), 
                'Значение': v
            })
    
    if not plot_data:
        plt.close()
        return None
        
    plot_df = pd.DataFrame(plot_data)
    
    # Colors
    colors = {'1': '#2ecc71', '2': '#e74c3c', '3': '#3498db', '4': '#9b59b6'}
    palette = [colors.get(g, '#333') for g in sorted(plot_df['GroupCode'].unique())]
    
    # Plot
    ax = sns.boxplot(x='Группа', y='Значение', data=plot_df, palette=palette, showfliers=False, width=0.5)
    sns.stripplot(x='Группа', y='Значение', data=plot_df, color='black', alpha=0.55, size=5, jitter=0.2)
    
    # Significance brackets (GraphPad-like: show all significant pairwise comparisons)
    y_max = plot_df['Значение'].max()
    y_range = y_max - plot_df['Значение'].min()
    bracket_h = y_range * 0.05
    step_h = y_range * 0.1
    current_y = y_max + y_range * 0.1
    
    # Map group labels to x-coordinates (0, 1, 2, 3)
    x_map = {g: i for i, g in enumerate(sorted(plot_df['Группа'].unique()))}
    
    # Map group code to label to find x
    code_to_label = {str(g): group_labels.get(str(g), f"Гр.{g}") for g in groups}
    
    significant_found = False

    sig_pairs = []
    if isinstance(pairwise_stats, dict) and pairwise_stats:
        for pair_key, st in pairwise_stats.items():
            if not isinstance(st, dict) or "error" in st:
                continue
            p_adj = st.get("p_adj")
            if p_adj is None or not np.isfinite(p_adj) or p_adj >= 0.05:
                continue
            try:
                g1, g2 = pair_key.split("_vs_")
            except Exception:
                continue
            sig_pairs.append((str(g1), str(g2), float(p_adj), float(st.get("r", 0.0))))
        sig_pairs.sort(key=lambda x: x[2])
    else:
        groups_codes = sorted(plot_df['GroupCode'].unique())
        for i, g1_code in enumerate(groups_codes):
            for g2_code in groups_codes[i + 1:]:
                group1_data = df[df[GROUP_COL].astype(str) == str(g1_code)][col].dropna()
                group2_data = df[df[GROUP_COL].astype(str) == str(g2_code)][col].dropna()
                if len(group1_data) < 2 or len(group2_data) < 2:
                    continue
                u, p = stats.mannwhitneyu(group1_data, group2_data, alternative='two-sided')
                if p < 0.05:
                    n1, n2 = len(group1_data), len(group2_data)
                    r = 1 - (2 * float(u)) / (n1 * n2)
                    sig_pairs.append((str(g1_code), str(g2_code), float(p), float(r)))
        sig_pairs.sort(key=lambda x: x[2])

    for g1_code, g2_code, p_used, r_used in sig_pairs:
        label1 = code_to_label.get(g1_code)
        label2 = code_to_label.get(g2_code)
        if label1 not in x_map or label2 not in x_map:
            continue
        x1 = x_map[label1]
        x2 = x_map[label2]
        p_text = "p < 0.001" if p_used < 0.001 else f"p = {p_used:.3f}"
        bracket_text = f"{p_text}\nES={r_used:.2f}"
        draw_significance_bracket(plt.gca(), x1, x2, current_y, bracket_h, bracket_text)
        current_y += step_h
        significant_found = True
    
    plt.xlabel('', fontsize=18)
    plt.ylabel(cfg["short"], fontsize=18)
    plt.title(f'{cfg["short"]} ({visit})', fontsize=20, fontweight='bold')
    
    # Adjust ylim to fit brackets
    if significant_found:
        plt.ylim(top=current_y + step_h)
        
    plt.tight_layout()
    
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


def create_pooled_spaghetti(pooled_results, endpoint_key, output_path):
    """Create longitudinal mean ± SE plot for Active vs Placebo."""
    res = pooled_results.get(endpoint_key, {})
    visits_data = res.get("visits", {})
    if not visits_data:
        return None
        
    cfg = ENDPOINTS[endpoint_key]
    visit_list = sorted(list(visits_data.keys()), key=lambda v: VISITS.index(v) if v in VISITS else 99)
    
    means_active = []
    ses_active = []
    means_placebo = []
    ses_placebo = []
    v_labels = []
    
    for v in visit_list:
        d = visits_data[v]
        means_active.append(d["active_mean"])
        ses_active.append(d["active_se"])
        means_placebo.append(d["placebo_mean"])
        ses_placebo.append(d["placebo_se"])
        v_labels.append(v)
        
    set_publication_plot_style()
    plt.figure(figsize=(10, 6))
    plt.style.use('seaborn-v0_8-whitegrid')
    
    x = range(len(v_labels))
    plt.errorbar(x, means_active, yerr=ses_active, marker='o', label='Active (укрупнение)', 
                 color='#2ecc71', linewidth=2.5, capsize=4)
    plt.errorbar(x, means_placebo, yerr=ses_placebo, marker='s', label='Placebo (укрупнение)', 
                 color='#95a5a6', linewidth=2.5, capsize=4, linestyle='--')
    
    # Add p-values (only significant; corrected)
    max_y = max(max(m+s for m,s in zip(means_active, ses_active)), max(m+s for m,s in zip(means_placebo, ses_placebo)))
    y_offset = max_y * 0.05
    
    for i, v in enumerate(v_labels):
        d = visits_data[v]
        if d["significant"]:
            p_used = float(d.get("p_adj", d.get("p_value", np.nan)))
            p_text = "p < 0.001" if np.isfinite(p_used) and p_used < 0.001 else (f"p = {p_used:.3f}" if np.isfinite(p_used) else "p = —")
            es = float(d.get("effect_size", 0.0))
            plt.text(i, max_y + y_offset, f"{p_text}\nES={es:.2f}", ha='center', va='bottom', fontsize=16, color='red', fontweight='bold')
            
    plt.xticks(x, v_labels)
    plt.ylabel(f'{cfg["short"]} (Mean ± SE)', fontsize=18)
    plt.title(f'Active vs Placebo: {cfg["short"]}', fontsize=20, fontweight='bold')
    plt.legend()
    plt.ylim(top=max_y * 1.25)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


def create_pooled_boxplot(df, endpoint_key, visit, output_path, visit_stats=None):
    """Create boxplot for Active vs Placebo at a specific visit."""
    cfg = ENDPOINTS[endpoint_key]
    col = cfg["cols"].get(visit)
    if not col or col not in df.columns:
        return None
        
    set_publication_plot_style()
    plt.figure(figsize=(8, 6))
    plt.style.use('seaborn-v0_8-whitegrid')
    
    df_plot = df.copy()
    df_plot["Treatment"] = df_plot[GROUP_COL].astype(str).apply(lambda g: "Active" if g in ["1", "3"] else "Placebo")
    plot_data = df_plot[[col, "Treatment"]].dropna()
    plot_data.columns = ["Value", "Treatment"]
    
    colors = {"Active": "#2ecc71", "Placebo": "#95a5a6"}
    
    sns.boxplot(x="Treatment", y="Value", data=plot_data, palette=colors, width=0.5, showfliers=False)
    sns.stripplot(x="Treatment", y="Value", data=plot_data, color='black', alpha=0.55, size=5, jitter=0.2)
    
    if isinstance(visit_stats, dict) and visit_stats.get("significant"):
        p_used = float(visit_stats.get("p_adj", visit_stats.get("p_value", np.nan)))
        es = float(visit_stats.get("effect_size", 0.0))

        y_max = plot_data["Value"].max()
        y_range = y_max - plot_data["Value"].min()
        h = y_range * 0.06
        y = y_max + h

        plt.plot([0, 0, 1, 1], [y, y + h, y + h, y], lw=1.5, c='black')
        p_text = "p < 0.001" if np.isfinite(p_used) and p_used < 0.001 else (f"p = {p_used:.3f}" if np.isfinite(p_used) else "p = —")
        plt.text(0.5, y + h, f"{p_text} | ES={es:.2f}", ha='center', va='bottom', fontsize=16, fontweight='bold')
        plt.ylim(top=y + h * 4)
        
    plt.title(f"Active vs Placebo: {cfg['short']} ({visit})", fontsize=20, fontweight='bold')
    plt.ylabel(cfg["short"])
    plt.xlabel("")
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


def create_duration_spaghetti(duration_results, endpoint_key, duration_key, output_path):
    res = duration_results.get(endpoint_key, {})
    comp = res.get("comparisons", {}).get(duration_key, {})
    visits_data = comp.get("visits", {})
    if not visits_data:
        return None

    cfg = ENDPOINTS[endpoint_key]
    v_labels = sorted(list(visits_data.keys()), key=lambda v: VISITS.index(v) if v in VISITS else 99)

    means_active = [visits_data[v]["active_mean"] for v in v_labels]
    ses_active = [visits_data[v]["active_se"] for v in v_labels]
    means_placebo = [visits_data[v]["placebo_mean"] for v in v_labels]
    ses_placebo = [visits_data[v]["placebo_se"] for v in v_labels]

    set_publication_plot_style()
    plt.figure(figsize=(10, 6))
    plt.style.use('seaborn-v0_8-whitegrid')

    x = range(len(v_labels))
    plt.errorbar(
        x,
        means_active,
        yerr=ses_active,
        marker='o',
        label=comp.get("active_label", "Active"),
        color='#2ecc71',
        linewidth=2.5,
        capsize=4,
    )
    plt.errorbar(
        x,
        means_placebo,
        yerr=ses_placebo,
        marker='s',
        label=comp.get("placebo_label", "Placebo"),
        color='#95a5a6',
        linewidth=2.5,
        capsize=4,
        linestyle='--',
    )

    max_y = max(
        max(m + s for m, s in zip(means_active, ses_active)),
        max(m + s for m, s in zip(means_placebo, ses_placebo)),
    )
    y_offset = max_y * 0.06

    for i, v in enumerate(v_labels):
        d = visits_data[v]
        if d.get("significant"):
            p_val = float(d.get("p_adj", d.get("p_value", np.nan)))
            p_text = "p < 0.001" if np.isfinite(p_val) and p_val < 0.001 else (f"p = {p_val:.3f}" if np.isfinite(p_val) else "p = —")
            es = float(d.get("effect_size", 0.0))
            plt.text(
                i,
                max_y + y_offset,
                f"{p_text}\nES={es:.2f}",
                ha='center',
                va='bottom',
                fontsize=16,
                color='red',
                fontweight='bold',
            )

    plt.xticks(x, v_labels)
    plt.ylabel(f'{cfg["short"]} (Mean ± SE)', fontsize=18)
    title = f"{comp.get('active_label','Active')} vs {comp.get('placebo_label','Placebo')}: {cfg['short']}"
    plt.title(title, fontsize=20, fontweight='bold')
    plt.legend()
    plt.ylim(top=max_y * 1.30)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


def create_duration_boxplot(df, endpoint_key, visit, duration_spec, output_path, visit_stats=None):
    cfg = ENDPOINTS[endpoint_key]
    col = cfg["cols"].get(visit)
    if not col or col not in df.columns:
        return None

    g_active = str(duration_spec["active_group"])
    g_placebo = str(duration_spec["placebo_group"])
    active_label = duration_spec.get("active_label", "Active")
    placebo_label = duration_spec.get("placebo_label", "Placebo")

    df_plot = df[df[GROUP_COL].astype(str).isin([g_active, g_placebo])].copy()
    if df_plot.empty:
        return None

    df_plot["Treatment"] = df_plot[GROUP_COL].astype(str).apply(lambda g: active_label if g == g_active else placebo_label)
    plot_data = df_plot[[col, "Treatment"]].dropna()
    plot_data.columns = ["Value", "Treatment"]
    if plot_data.empty:
        return None

    set_publication_plot_style()
    plt.figure(figsize=(8.5, 6.5))
    plt.style.use('seaborn-v0_8-whitegrid')

    colors = {active_label: "#2ecc71", placebo_label: "#95a5a6"}
    sns.boxplot(x="Treatment", y="Value", data=plot_data, palette=colors, width=0.5, showfliers=False)
    sns.stripplot(x="Treatment", y="Value", data=plot_data, color='black', alpha=0.55, size=5, jitter=0.2)

    if isinstance(visit_stats, dict) and visit_stats.get("significant"):
        p_used = float(visit_stats.get("p_adj", visit_stats.get("p_value", np.nan)))
        es = float(visit_stats.get("effect_size", 0.0))

        y_max = plot_data["Value"].max()
        y_range = y_max - plot_data["Value"].min()
        h = y_range * 0.06
        y = y_max + h
        plt.plot([0, 0, 1, 1], [y, y + h, y + h, y], lw=1.5, c='black')
        p_text = "p < 0.001" if np.isfinite(p_used) and p_used < 0.001 else (f"p = {p_used:.3f}" if np.isfinite(p_used) else "p = —")
        plt.text(0.5, y + h, f"{p_text} | ES={es:.2f}", ha='center', va='bottom', fontsize=16, fontweight='bold')
        plt.ylim(top=y + h * 4)

    plt.title(f"{active_label} vs {placebo_label}: {cfg['short']} ({visit})", fontsize=20, fontweight='bold')
    plt.ylabel(cfg["short"], fontsize=18)
    plt.xlabel("")
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


def create_responder_visit_plot(responder_data, endpoint_key, visit, output_path):
    """Create barplot for responders at a specific visit."""
    res = responder_data.get(endpoint_key, {})
    visit_data = res.get("visits", {}).get(visit, {})
    groups_data = visit_data.get("groups", [])
    
    if not groups_data:
        return None
        
    cfg = ENDPOINTS[endpoint_key]
    threshold = res.get("threshold", 20)
    
    set_publication_plot_style()
    plt.figure(figsize=(9.5, 6.5))
    plt.style.use('seaborn-v0_8-whitegrid')
    
    groups = [g['group'] for g in groups_data]
    pcts = [g['pct'] for g in groups_data]
    ci_lows = [g['ci_low'] for g in groups_data]
    ci_highs = [g['ci_high'] for g in groups_data]
    
    colors = ['#2ecc71', '#e74c3c', '#3498db', '#9b59b6'] 
    
    x = range(len(groups))
    bars = plt.bar(x, pcts, color=[colors[int(g)-1] if int(g)<=4 else '#333' for g in groups], 
                   edgecolor='black', width=0.6)
    
    # Error bars per bar
    for i, bar in enumerate(bars):
        y_err_lower = max(0, pcts[i] - ci_lows[i])
        y_err_upper = max(0, ci_highs[i] - pcts[i])
        
        plt.errorbar(bar.get_x() + bar.get_width()/2, pcts[i], 
                     yerr=[[y_err_lower], [y_err_upper]], 
                     fmt='none', color='black', capsize=4)
        
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 2,
                 f"{pcts[i]:.0f}%", ha='center', va='bottom', fontweight='bold', fontsize=16)
    
    # Add Chi-square p-value
    test = visit_data.get("test", {})
    if test.get("p") is not None:
        p_text = f"p = {float(test['p']):.4f}"
        verdict = "значимо" if test.get("significant") else "незначимо"
        plt.title(
            f"Респондеры (≥{threshold}%) на визите {visit}\n{cfg['short']} | {p_text} ({verdict})",
            fontsize=20,
            fontweight='bold',
        )
    else:
        plt.title(f"Респондеры (≥{threshold}%) на визите {visit}\n{cfg['short']}", fontsize=20, fontweight='bold')
    
    plt.xticks(x, [f'Группа {g}' for g in groups])
    plt.ylabel('% респондеров (95% ДИ)', fontsize=18)
    plt.ylim(0, 100)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


def create_responder_barplot(responder_data, endpoint_key, visit, output_path):
    """Create pooled Active vs Placebo responder barplot (95% CI) for one endpoint & visit."""
    res = responder_data.get(endpoint_key, {})
    cfg = ENDPOINTS.get(endpoint_key, {"short": endpoint_key})
    visit_data = res.get("visits", {}).get(visit, {})
    groups_data = visit_data.get("groups", [])
    if not groups_data:
        return None

    by_g = {str(r.get("group")): r for r in groups_data if isinstance(r, dict)}
    active_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["1", "3"])
    active_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["1", "3"])
    placebo_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["2", "4"])
    placebo_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["2", "4"])

    def wilson_ci(k: int, n: int, z: float = 1.96):
        if n <= 0:
            return 0.0, 0.0, 0.0
        p_hat = k / n
        denom = 1 + z**2 / n
        center = (p_hat + z**2 / (2*n)) / denom
        margin = z * np.sqrt((p_hat*(1-p_hat) + z**2/(4*n)) / n) / denom
        low = max(0.0, center - margin)
        high = min(1.0, center + margin)
        return p_hat * 100.0, low * 100.0, high * 100.0

    active_pct, active_low, active_high = wilson_ci(active_resp, active_n)
    placebo_pct, placebo_low, placebo_high = wilson_ci(placebo_resp, placebo_n)

    set_publication_plot_style()
    plt.figure(figsize=(7.5, 6.5))
    plt.style.use('seaborn-v0_8-whitegrid')

    labels = ["Active (G1+G3)", "Placebo (G2+G4)"]
    pcts = [active_pct, placebo_pct]
    ci_lows = [active_low, placebo_low]
    ci_highs = [active_high, placebo_high]

    x = range(2)
    bars = plt.bar(x, pcts, color=['#2ecc71', '#95a5a6'], edgecolor='black', width=0.6)
    for i, bar in enumerate(bars):
        y_err_lower = max(0, pcts[i] - ci_lows[i])
        y_err_upper = max(0, ci_highs[i] - pcts[i])
        plt.errorbar(bar.get_x() + bar.get_width()/2, pcts[i], yerr=[[y_err_lower], [y_err_upper]], fmt='none', color='black', capsize=4)
        plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 2, f"{pcts[i]:.0f}%", ha='center', va='bottom', fontsize=16, fontweight='bold')

    plt.xticks(list(x), labels)
    plt.ylabel('% респондеров (95% ДИ)', fontsize=18)
    plt.title(f"Респондеры (≥20%) | {cfg.get('short', endpoint_key)} | {visit}", fontsize=20, fontweight='bold')
    plt.ylim(0, 100)
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


def create_forest_plot(all_results, output_path):
    """Create forest plot of effect sizes at V6."""
    set_publication_plot_style()
    plt.figure(figsize=(10, 6))
    plt.style.use('seaborn-v0_8-whitegrid')
    
    endpoints = []
    effect_sizes = []
    
    for key, result in all_results.items():
        v6 = result.get("by_visit", {}).get("V6", {})
        kw = v6.get("kruskal", {})
        es = kw.get("epsilon_sq")
        
        if es is not None:
            endpoints.append(result["short"])
            effect_sizes.append(es)
    
    # Sort by effect size
    sorted_data = sorted(zip(endpoints, effect_sizes), key=lambda x: x[1], reverse=True)
    endpoints, effect_sizes = zip(*sorted_data) if sorted_data else ([], [])
    
    y = range(len(endpoints))
    colors = ['#2ecc71' if es > 0.06 else '#3498db' if es > 0.01 else '#95a5a6' for es in effect_sizes]
    
    plt.barh(y, effect_sizes, color=colors, edgecolor='black', height=0.6)
    plt.axvline(x=0.01, color='gray', linestyle='--', alpha=0.7, label='Малый эффект')
    plt.axvline(x=0.06, color='orange', linestyle='--', alpha=0.7, label='Средний эффект')
    plt.axvline(x=0.14, color='red', linestyle='--', alpha=0.7, label='Большой эффект')
    
    plt.yticks(y, endpoints)
    plt.xlabel('Размер эффекта (ε²)', fontsize=18)
    plt.title('Сравнение размеров эффекта на визите V6', fontsize=20, fontweight='bold')
    plt.legend(loc='lower right')
    plt.tight_layout()
    
    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


def create_forest_plot_by_visit(all_results, visit, output_path):
    set_publication_plot_style()
    plt.figure(figsize=(10, 6))
    plt.style.use('seaborn-v0_8-whitegrid')

    endpoints = []
    effect_sizes = []

    for _, result in all_results.items():
        vdata = result.get("by_visit", {}).get(str(visit), {})
        kw = vdata.get("kruskal", {})
        es = kw.get("epsilon_sq")
        if es is None or not np.isfinite(es):
            continue
        endpoints.append(result.get("short", "—"))
        effect_sizes.append(float(es))

    sorted_data = sorted(zip(endpoints, effect_sizes), key=lambda x: x[1], reverse=True)
    endpoints, effect_sizes = zip(*sorted_data) if sorted_data else ([], [])

    y = range(len(endpoints))
    colors = ['#2ecc71' if es > 0.06 else '#3498db' if es > 0.01 else '#95a5a6' for es in effect_sizes]

    plt.barh(y, effect_sizes, color=colors, edgecolor='black', height=0.6)
    plt.axvline(x=0.01, color='gray', linestyle='--', alpha=0.7, label='Малый эффект')
    plt.axvline(x=0.06, color='orange', linestyle='--', alpha=0.7, label='Средний эффект')
    plt.axvline(x=0.14, color='red', linestyle='--', alpha=0.7, label='Большой эффект')

    plt.yticks(y, endpoints)
    plt.xlabel('Размер эффекта (ε²)', fontsize=18)
    plt.title(f'Сравнение размеров эффекта на визите {visit}', fontsize=20, fontweight='bold')
    plt.legend(loc='lower right')
    plt.tight_layout()

    plt.savefig(output_path, dpi=300, bbox_inches='tight', facecolor='white')
    plt.close()
    return str(output_path)


# ============================================================
# COMPREHENSIVE ANALYSIS
# ============================================================

def analyze_endpoint_full(df, endpoint_key):
    """Full analysis for one endpoint at ALL timepoints."""
    cfg = ENDPOINTS[endpoint_key]
    print(f"\n📊 {cfg['name']}")
    
    results = {
        "name": cfg["name"],
        "short": cfg["short"],
        "primary": cfg["primary"],
        "direction": cfg["direction"],
        "by_visit": {},
        "by_group": {},
        "within_group_changes": {},
        "pairwise": {},
        "mixed_effects_4g": {},
        "mixed_effects_pooled": {},
    }
    
    groups = sorted(df[GROUP_COL].dropna().unique())
    
    # === 1. DESCRIPTIVE STATS BY GROUP AND VISIT ===
    for visit in VISITS:
        col = cfg["cols"].get(visit)
        if not col or col not in df.columns:
            continue
        
        visit_data = {"groups": {}}
        
        for g in groups:
            values = df[df[GROUP_COL] == g][col].dropna()
            visit_data["groups"][str(g)] = descriptive(values)
            visit_data["groups"][str(g)]["normality"] = normality_test(values)
        
        # Kruskal-Wallis for this visit
        group_arrays = [df[df[GROUP_COL] == g][col].dropna() for g in groups]
        visit_data["kruskal"] = kruskal_wallis(group_arrays)
        visit_data["bf10"] = bayes_factor(visit_data["kruskal"].get("p"))
        
        results["by_visit"][visit] = visit_data
    
    # === 2. WITHIN-GROUP CHANGES (V2 → each visit) ===
    baseline_col = cfg["cols"].get("V2")
    if baseline_col and baseline_col in df.columns:
        for g in groups:
            results["within_group_changes"][str(g)] = {}
            
            baseline = df[df[GROUP_COL] == g][baseline_col]
            baseline_median = float(pd.Series(baseline).dropna().median()) if len(pd.Series(baseline).dropna()) else np.nan
            
            visit_keys: List[str] = []
            visit_ps: List[float] = []
            for visit in ["V3", "V4", "V5", "V6"]:
                col = cfg["cols"].get(visit)
                if not col or col not in df.columns:
                    continue
                
                follow = df[df[GROUP_COL] == g][col]
                
                # Paired test
                wilcox = wilcoxon_signed_rank(baseline, follow)
                if isinstance(wilcox, dict) and "error" not in wilcox:
                    follow_median = float(pd.Series(follow).dropna().median()) if len(pd.Series(follow).dropna()) else np.nan
                    wilcox["baseline_median"] = baseline_median if np.isfinite(baseline_median) else None
                    wilcox["follow_median"] = follow_median if np.isfinite(follow_median) else None

                    if np.isfinite(baseline_median) and baseline_median != 0:
                        wilcox["delta_pct"] = float((wilcox.get("median_diff", np.nan) / baseline_median) * 100.0)
                    else:
                        wilcox["delta_pct"] = np.nan

                    p_val = wilcox.get("p")
                    wilcox["bf10"] = float(bayes_factor(p_val)) if p_val is not None and np.isfinite(p_val) else np.nan
                results["within_group_changes"][str(g)][visit] = wilcox

                if isinstance(wilcox, dict) and "error" not in wilcox:
                    visit_keys.append(visit)
                    visit_ps.append(float(wilcox.get("p", np.nan)))

            p_adj = holm_adjust(visit_ps)
            for v, padj in zip(visit_keys, p_adj):
                w = results["within_group_changes"].get(str(g), {}).get(v)
                if not isinstance(w, dict) or "error" in w:
                    continue
                w["p_adj"] = float(padj) if np.isfinite(padj) else np.nan
                w["significant"] = bool(np.isfinite(padj) and padj < 0.05)
    
    # === 3. PAIRWISE GROUP COMPARISONS AT EACH VISIT ===
    for visit in VISITS:
        col = cfg["cols"].get(visit)
        if not col or col not in df.columns:
            continue
        
        results["pairwise"][visit] = {}

        keys: List[str] = []
        ps: List[float] = []
        for i, g1 in enumerate(groups):
            for g2 in groups[i + 1:]:
                d1 = df[df[GROUP_COL] == g1][col].dropna()
                d2 = df[df[GROUP_COL] == g2][col].dropna()

                mw = mann_whitney(d1, d2)
                if isinstance(mw, dict) and "error" not in mw:
                    med1 = float(pd.Series(d1).dropna().median()) if len(pd.Series(d1).dropna()) else np.nan
                    med2 = float(pd.Series(d2).dropna().median()) if len(pd.Series(d2).dropna()) else np.nan
                    mw["g1_median"] = med1 if np.isfinite(med1) else None
                    mw["g2_median"] = med2 if np.isfinite(med2) else None

                    diff = med1 - med2 if (np.isfinite(med1) and np.isfinite(med2)) else np.nan
                    mw["diff_median"] = float(diff) if np.isfinite(diff) else np.nan
                    if np.isfinite(med2) and med2 != 0 and np.isfinite(diff):
                        mw["diff_pct"] = float((diff / med2) * 100.0)
                    else:
                        mw["diff_pct"] = np.nan

                k = f"{g1}_vs_{g2}"
                keys.append(k)
                ps.append(mw.get("p", np.nan))
                if "error" not in mw and mw.get("p") is not None:
                    mw["bf10"] = bayes_factor(mw.get("p"))
                results["pairwise"][visit][k] = mw

        p_adj = holm_adjust(ps)
        for k, padj in zip(keys, p_adj):
            s = results["pairwise"][visit].get(k)
            if not isinstance(s, dict) or "error" in s:
                continue
            s["p_adj"] = float(padj) if np.isfinite(padj) else np.nan
            s["significant"] = bool(np.isfinite(padj) and padj < 0.05)

    results["mixed_effects_4g"] = fit_mixed_effects(df, endpoint_key, pooled=False)
    results["mixed_effects_pooled"] = fit_mixed_effects(df, endpoint_key, pooled=True)
    
    return results


def analyze_responders(df, threshold=20):
    """
    Analyze responders (improvement >= threshold %) for EACH visit relative to baseline (V2).
    Returns nested dict: {endpoint: {visit: {groups: [], test: {}}}}
    """
    print(f"\n📊 Анализ респондеров (≥{threshold}% улучшения) по ВСЕМ визитам...")
    
    all_responders = {}
    groups = sorted(df[GROUP_COL].dropna().unique())
    target_visits = VISITS
    
    for endpoint_key, cfg in ENDPOINTS.items():
        v2_col = cfg["cols"].get("V2")
        if not v2_col or v2_col not in df.columns:
            continue
            
        endpoint_res = {"short": cfg["short"], "visits": {}}
        
        for visit in target_visits:
            v_col = cfg["cols"].get(visit)
            if not v_col or v_col not in df.columns:
                continue
                
            visit_res = {"groups": [], "test": {}}
            
            for g in groups:
                gdf = df[df[GROUP_COL] == g]
                if visit == "V2":
                    valid = gdf[[v2_col]].dropna()
                else:
                    valid = gdf[[v2_col, v_col]].dropna()
                n = len(valid)
                
                if n == 0:
                    continue

                if visit == "V2":
                    k = 0
                    pct = 0.0
                else:
                    denom_base = valid[v2_col].replace(0, np.nan)
                    if cfg.get("direction") == "higher_is_better":
                        improvement = (valid[v_col] - valid[v2_col]) / denom_base * 100
                    else:
                        improvement = (valid[v2_col] - valid[v_col]) / denom_base * 100
                    improvement = improvement.fillna(0)
                    k = int((improvement >= threshold).sum())
                    pct = k / n * 100 if n > 0 else 0
                
                # Wilson CI
                z = 1.96
                p_hat = k / n if n > 0 else 0
                wilson_denom = 1 + z**2 / n if n > 0 else 1
                center = (p_hat + z**2 / (2*n)) / wilson_denom if n > 0 else 0
                margin = z * np.sqrt((p_hat*(1-p_hat) + z**2/(4*n)) / n) / wilson_denom if n > 0 else 0
                ci_low = max(0, center - margin) * 100
                ci_high = min(1, center + margin) * 100
                
                visit_res["groups"].append({
                    "group": str(g),
                    "n": n,
                    "responders": int(k),
                    "pct": pct,
                    "ci_low": ci_low,
                    "ci_high": ci_high,
                })
            
            # Chi-square per visit (skip baseline)
            if visit != "V2" and len(visit_res["groups"]) >= 2:
                contingency = np.array([
                    [r["responders"], r["n"] - r["responders"]]
                    for r in visit_res["groups"]
                ])
                try:
                    # Remove zero rows
                    contingency = contingency[~np.all(contingency == 0, axis=1)]
                    if contingency.shape[0] >= 2:
                        chi2, p, dof, _ = stats.chi2_contingency(contingency)
                        visit_res["test"] = {
                            "name": "Chi-square",
                            "chi2": float(chi2),
                            "p": float(p),
                            "significant": p < 0.05,
                        }
                except:
                    pass
            
            endpoint_res["visits"][visit] = visit_res
            
        all_responders[endpoint_key] = endpoint_res
        
    return all_responders


def calculate_nnt(active_responder_pct, placebo_responder_pct):
    """Calculate Number Needed to Treat (NNT)."""
    arr = (active_responder_pct - placebo_responder_pct) / 100  # Absolute Risk Reduction
    if arr <= 0:
        return None, None  # No benefit
    nnt = 1 / arr
    return round(nnt, 1), round(arr * 100, 1)


def analyze_active_vs_placebo(df, all_results):
    """
    Longitudinal Pooled Analysis: Active (G1+G3) vs Placebo (G2+G4) for ALL visits.
    Returns nested dict: {endpoint: {visit: {stats...}}}
    """
    print("\n📊 Сравнение Active vs Placebo (укрупнение)...")
    
    results = {}
    groups = df[GROUP_COL].astype(str)
    df_copy = df.copy()
    df_copy["treatment"] = groups.apply(lambda g: "Active" if g in ["1", "3"] else "Placebo")
    
    for endpoint_key, cfg in ENDPOINTS.items():
        results[endpoint_key] = {"short": cfg["short"], "visits": {}}
        
        # Baseline V2
        v2_col = cfg["cols"].get("V2")
        
        p_keys: List[str] = []
        p_vals: List[float] = []
        for visit in VISITS:
            v_col = cfg["cols"].get(visit)
            if not v_col or v_col not in df.columns:
                continue
            
            active_data = df_copy[df_copy["treatment"] == "Active"][v_col].dropna()
            placebo_data = df_copy[df_copy["treatment"] == "Placebo"][v_col].dropna()
            
            # Mann-Whitney U
            if len(active_data) >= 3 and len(placebo_data) >= 3:
                u_stat, p_val = stats.mannwhitneyu(active_data, placebo_data, alternative='two-sided')
                
                # Effect size (rank-biserial)
                n1, n2 = len(active_data), len(placebo_data)
                r = 1 - (2 * u_stat) / (n1 * n2)
                
                # Descriptive
                active_median = active_data.median()
                placebo_median = placebo_data.median()
                active_mean = active_data.mean()
                placebo_mean = placebo_data.mean()
                active_se = active_data.std() / np.sqrt(n1)
                placebo_se = placebo_data.std() / np.sqrt(n2)
                
                visit_res = {
                    "active_n": n1,
                    "placebo_n": n2,
                    "active_median": float(active_median),
                    "placebo_median": float(placebo_median),
                    "active_mean": float(active_mean),
                    "placebo_mean": float(placebo_mean),
                    "active_se": float(active_se),
                    "placebo_se": float(placebo_se),
                    "u_stat": float(u_stat),
                    "p_value": float(p_val),
                    "effect_size": float(r),
                    "significant": bool(p_val < 0.05),
                    "bf10": bayes_factor(p_val),
                }

                diff = float(active_median - placebo_median)
                visit_res["diff_median"] = diff
                if float(placebo_median) != 0:
                    visit_res["diff_pct"] = float((diff / float(placebo_median)) * 100.0)
                else:
                    visit_res["diff_pct"] = np.nan
                
                results[endpoint_key]["visits"][visit] = visit_res

                p_keys.append(visit)
                p_vals.append(float(p_val))

        p_adj = holm_adjust(p_vals)
        for v, padj in zip(p_keys, p_adj):
            st = results[endpoint_key]["visits"].get(v)
            if not isinstance(st, dict):
                continue
            st["p_adj"] = float(padj) if np.isfinite(padj) else np.nan
            st["significant"] = bool(np.isfinite(padj) and padj < 0.05)
    
    return results


def analyze_active_vs_placebo_by_duration(df):
    results = {}
    duration_specs = {
        "30d": {
            "active_group": "1",
            "placebo_group": "2",
            "active_label": "ДИАМАГ 30 дней",
            "placebo_label": "Плацебо 30 дней",
        },
        "20d": {
            "active_group": "3",
            "placebo_group": "4",
            "active_label": "ДИАМАГ 20 дней",
            "placebo_label": "Плацебо 20 дней",
        },
    }

    for endpoint_key, cfg in ENDPOINTS.items():
        endpoint_out = {"short": cfg["short"], "comparisons": {}}
        for duration_key, spec in duration_specs.items():
            endpoint_out["comparisons"][duration_key] = {
                "active_group": spec["active_group"],
                "placebo_group": spec["placebo_group"],
                "active_label": spec["active_label"],
                "placebo_label": spec["placebo_label"],
                "visits": {},
            }

        for visit in VISITS:
            v_col = cfg["cols"].get(visit)
            if not v_col or v_col not in df.columns:
                continue

            for duration_key, spec in duration_specs.items():
                active_vals = df[df[GROUP_COL].astype(str) == spec["active_group"]][v_col].dropna()
                placebo_vals = df[df[GROUP_COL].astype(str) == spec["placebo_group"]][v_col].dropna()

                if len(active_vals) < 3 or len(placebo_vals) < 3:
                    continue

                u_stat, p_val = stats.mannwhitneyu(active_vals, placebo_vals, alternative="two-sided")
                n1, n2 = len(active_vals), len(placebo_vals)
                r = 1 - (2 * float(u_stat)) / (n1 * n2)

                active_median = float(active_vals.median())
                placebo_median = float(placebo_vals.median())
                active_mean = float(active_vals.mean())
                placebo_mean = float(placebo_vals.mean())
                active_se = float(active_vals.std() / np.sqrt(n1)) if n1 > 1 else float("nan")
                placebo_se = float(placebo_vals.std() / np.sqrt(n2)) if n2 > 1 else float("nan")

                diff = active_median - placebo_median
                diff_pct = (diff / placebo_median * 100.0) if placebo_median not in (0, 0.0) else np.nan

                endpoint_out["comparisons"][duration_key]["visits"][visit] = {
                    "active_n": int(n1),
                    "placebo_n": int(n2),
                    "active_median": active_median,
                    "placebo_median": placebo_median,
                    "active_mean": active_mean,
                    "placebo_mean": placebo_mean,
                    "active_se": active_se,
                    "placebo_se": placebo_se,
                    "u_stat": float(u_stat),
                    "p_value": float(p_val),
                    "effect_size": float(r),
                    "bf10": bayes_factor(p_val),
                    "diff_median": float(diff),
                    "diff_pct": float(diff_pct) if np.isfinite(diff_pct) else np.nan,
                    "significant": bool(p_val < 0.05),
                }

        for duration_key, comp in endpoint_out.get("comparisons", {}).items():
            visits = comp.get("visits", {})
            keys = []
            ps = []
            for v in sorted(list(visits.keys()), key=lambda vv: VISITS.index(vv) if vv in VISITS else 99):
                p_val = visits[v].get("p_value")
                keys.append(v)
                ps.append(float(p_val) if p_val is not None and np.isfinite(p_val) else np.nan)
            p_adj = holm_adjust(ps)
            for v, padj in zip(keys, p_adj):
                st = visits.get(v)
                if not isinstance(st, dict):
                    continue
                st["p_adj"] = float(padj) if np.isfinite(padj) else np.nan
                st["significant"] = bool(np.isfinite(padj) and padj < 0.05)

        results[endpoint_key] = endpoint_out

    return results


def generate_executive_summary(all_results, responders, active_vs_placebo, df):
    """Generate data for Executive Summary."""
    summary = {
        "n_patients": len(df),
        "n_groups": 4,
        "primary_findings": [],
        "secondary_findings": [],
        "responder_findings": [],
        "recommendation": "",
    }
    
    # Primary endpoints
    for key in ["updrs_part3", "updrs_part2"]:
        if key in all_results:
            result = all_results[key]
            v6 = result.get("by_visit", {}).get("V6", {})
            kw = v6.get("kruskal", {})

            avp_v6 = (
                active_vs_placebo
                .get(key, {})
                .get("visits", {})
                .get("V6", {})
            )

            finding = {
                "name": result["short"],
                "kw_p": kw.get("p", np.nan),
                "kw_significant": kw.get("significant", False),
                "active_better": bool(avp_v6.get("active_median", np.nan) < avp_v6.get("placebo_median", np.nan)),
                "avp_p": avp_v6.get("p_value", np.nan),
                "avp_p_adj": avp_v6.get("p_adj", np.nan),
                "avp_effect": avp_v6.get("effect_size", np.nan),
                "diff_median": avp_v6.get("diff_median", np.nan),
                "diff_pct": avp_v6.get("diff_pct", np.nan),
            }
            summary["primary_findings"].append(finding)
    
    # Responders NNT (по визиту V6)
    for key in ["updrs_part3", "updrs_part2"]:
        if key in responders:
            resp = responders[key]
            v6_groups = resp.get("visits", {}).get("V6", {}).get("groups", [])

            by_g = {str(r.get("group")): r for r in v6_groups if isinstance(r, dict)}
            active_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["1", "3"])
            active_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["1", "3"])
            placebo_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["2", "4"])
            placebo_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["2", "4"])

            active_pct = (active_resp / active_n * 100.0) if active_n > 0 else 0.0
            placebo_pct = (placebo_resp / placebo_n * 100.0) if placebo_n > 0 else 0.0

            nnt, arr = calculate_nnt(active_pct, placebo_pct)
            
            summary["responder_findings"].append({
                "name": resp["short"],
                "active_pct": active_pct,
                "placebo_pct": placebo_pct,
                "active_n": int(active_n),
                "placebo_n": int(placebo_n),
                "nnt": nnt,
                "arr": arr,
            })
    
    # Recommendation logic
    primary_sig = sum(1 for f in summary["primary_findings"] if f["kw_significant"])
    primary_avp_sig = sum(1 for f in summary["primary_findings"] if f.get("avp_p_adj", 1) < 0.05)
    
    if primary_sig >= 1 or primary_avp_sig >= 1:
        summary["recommendation"] = "Получены статистические сигналы эффективности; целесообразно продолжение исследований с увеличением выборки и заранее заданными первичными конечными точками."
    elif primary_sig == 0 and any(f.get("avp_p_adj", 1) < 0.1 for f in summary["primary_findings"]):
        summary["recommendation"] = "Убедительных доказательств эффективности не получено; наблюдаются тенденции, требующие подтверждения на большей выборке и с контролем множественных сравнений."
    else:
        summary["recommendation"] = "Убедительных статистических и клинически значимых доказательств эффективности по ключевым конечным точкам не получено."
    
    return summary


def generate_discussion_text(exec_summary, all_results, active_vs_placebo, responders):
    """Generate SCIENTIFIC discussion text for the report."""
    discussion: List[str] = []

    n_patients = exec_summary.get("n_patients")
    n_txt = str(n_patients) if n_patients is not None else "—"

    discussion.append(
        "В настоящем рандомизированном плацебо-контролируемом исследовании оценивалась эффективность применения аппарата ДИАМАГ "
        f"в комплексной терапии болезни Паркинсона. В анализ включено {n_txt} пациентов, рандомизированных на 4 группы "
        "с различной длительностью воздействия (30 и 20 дней) и типом терапии (активная/плацебо)."
    )

    discussion.append(
        "Для оценки различий между группами на каждом визите применялся глобальный тест Kruskal–Wallis с оценкой размера эффекта (ε²) и силы доказательств (BF₁₀). "
        "При наличии статистических сигналов выполнялись post-hoc сравнения с коррекцией Холма для контроля семейной ошибки I рода. "
        "Дополнительно анализ включал укрупнение Active vs Placebo, оценку внутригрупповой динамики и модели смешанных эффектов для повторных измерений."
    )

    lines: List[str] = []
    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)

        kw_best = None
        for v in VISITS:
            vdata = result.get("by_visit", {}).get(v, {})
            kw = vdata.get("kruskal", {})
            p = kw.get("p")
            if p is None or not np.isfinite(p):
                continue
            if kw_best is None or float(p) < float(kw_best["p"]):
                kw_best = {
                    "visit": v,
                    "p": float(p),
                    "epsilon_sq": kw.get("epsilon_sq"),
                    "bf10": vdata.get("bf10", np.nan),
                }

        avp_visits = active_vs_placebo.get(ep_key, {}).get("visits", {}) if isinstance(active_vs_placebo, dict) else {}
        avp_best = None
        for v, d in avp_visits.items():
            if not isinstance(d, dict):
                continue
            p = d.get("p_adj", d.get("p_value"))
            if p is None or not np.isfinite(p):
                continue
            if avp_best is None or float(p) < float(avp_best["p"]):
                avp_best = {
                    "visit": v,
                    "p": float(p),
                    "bf10": d.get("bf10", np.nan),
                    "es": d.get("effect_size", np.nan),
                    "diff": d.get("diff_median", np.nan),
                    "diff_pct": d.get("diff_pct", np.nan),
                }

        fragments: List[str] = []
        if kw_best is not None:
            p_txt = "p < 0.001" if kw_best["p"] < 0.001 else f"p={kw_best['p']:.3f}"
            es = kw_best.get("epsilon_sq")
            es_txt = f"ε²={float(es):.3f}" if es is not None and np.isfinite(es) else "ε²=—"
            bf = kw_best.get("bf10", np.nan)
            bf_txt = f"BF₁₀={float(bf):.2f} ({interpret_bf(float(bf))})" if bf is not None and np.isfinite(bf) else "BF₁₀=—"
            fragments.append(f"4 группы: {kw_best['visit']}, {p_txt}, {es_txt}, {bf_txt}")

        if avp_best is not None:
            p_txt = "p < 0.001" if avp_best["p"] < 0.001 else f"p_adj={avp_best['p']:.3f}"
            es = avp_best.get("es")
            es_txt = f"ES={float(es):.2f}" if es is not None and np.isfinite(es) else "ES=—"
            diff = avp_best.get("diff")
            diff_pct = avp_best.get("diff_pct")
            diff_txt = f"Δ={float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "Δ=—"
            diffp_txt = f" ({float(diff_pct):+.1f}%)" if diff_pct is not None and np.isfinite(diff_pct) else ""
            bf = avp_best.get("bf10", np.nan)
            bf_txt = f"BF₁₀={float(bf):.2f} ({interpret_bf(float(bf))})" if bf is not None and np.isfinite(bf) else "BF₁₀=—"
            fragments.append(f"укрупнение Active vs Placebo: {avp_best['visit']}, {p_txt}, {es_txt}, {diff_txt}{diffp_txt}, {bf_txt}")

        if fragments:
            lines.append(f"{short}: " + "; ".join(fragments) + ".")

    if lines:
        discussion.append("Сводка по показателям (статистическая значимость, сила доказательств и клиническая величина различий):")
        discussion.extend(lines)

    responder_lines: List[str] = []
    for r in exec_summary.get("responder_findings", []):
        if not isinstance(r, dict):
            continue
        nnt = r.get("nnt")
        if nnt is None:
            continue
        responder_lines.append(
            f"{r.get('name','—')}: NNT≈{float(nnt):.1f} при долях респондеров {float(r.get('active_pct',0)):.0f}% (Active) vs {float(r.get('placebo_pct',0)):.0f}% (Placebo)."
        )
    if responder_lines:
        discussion.append("Клиническая интерпретация по респондерам (улучшение ≥20% от baseline):")
        discussion.extend(responder_lines)

    mixed_lines: List[str] = []
    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)
        mm4 = result.get("mixed_effects_4g", {})
        mmp = result.get("mixed_effects_pooled", {})
        sig4 = isinstance(mm4, dict) and (mm4.get("interaction", {}) or {}).get("significant") is True
        sigp = isinstance(mmp, dict) and (mmp.get("interaction", {}) or {}).get("significant") is True
        if not (sig4 or sigp):
            continue
        parts = []
        if sig4:
            p = mm4.get("interaction_p_value")
            p_txt = f"p={float(p):.4f}" if p is not None and np.isfinite(p) else "p=—"
            parts.append(f"4 группы ({p_txt})")
        if sigp:
            p = mmp.get("interaction_p_value")
            p_txt = f"p={float(p):.4f}" if p is not None and np.isfinite(p) else "p=—"
            parts.append(f"укрупнение ({p_txt})")
        direction = str((ENDPOINTS.get(ep_key, {}) or {}).get("direction") or "").strip()
        if direction == "lower_is_better":
            dir_txt = " (ниже — лучше)"
        elif direction == "higher_is_better":
            dir_txt = " (выше — лучше)"
        else:
            dir_txt = ""
        mixed_lines.append(
            f"{short}: значимое взаимодействие Визит×Группа{dir_txt} в модели {', '.join(parts)}; это означает различие траекторий во времени, поэтому практический вывод формируется по follow-up оценкам по визитам (p_adj, Δ, ES)."
        )
    if mixed_lines:
        discussion.append("Подтверждение результатов моделями повторных измерений (Mixed Effects):")
        discussion.extend(mixed_lines)

    discussion.append(
        "Ограничения исследования включают ограниченный размер выборки при разделении на 4 группы, что снижает статистическую мощность и повышает риск ложнонегативных результатов, "
        "а также возможную клиническую гетерогенность пациентов. В этих условиях интерпретация должна учитывать согласованность направлений эффектов, величину ES и Δ, а также p_adj и BF₁₀, а не только номинальные p-value."
    )

    discussion.append(
        "Сопоставление результатов с опубликованными исследованиями не выполнено в рамках автоматической генерации, поскольку список источников не включён в входные данные; данный раздел следует дополнить при наличии заранее отобранной библиографии."
    )

    rec = exec_summary.get("recommendation", "—")
    discussion.append(f"Итоговое заключение: {rec}.")

    return discussion


def generate_conclusions_text(exec_summary, all_results, active_vs_placebo, responders) -> List[str]:
    out: List[str] = []

    rec = str((exec_summary or {}).get("recommendation") or "—").strip()
    out.append("Итоговая рекомендация: " + rec)

    sig_kw = 0
    sig_avp = 0
    any_data = 0
    for ep_key, r in (all_results or {}).items():
        if not isinstance(r, dict):
            continue
        any_data += 1
        kw_any = False
        for v in VISITS:
            kw = ((r.get("by_visit", {}) or {}).get(v, {}) or {}).get("kruskal", {})
            p = kw.get("p")
            if p is not None and np.isfinite(p) and float(p) < 0.05:
                kw_any = True
                break
        if kw_any:
            sig_kw += 1

        avp_visits = ((active_vs_placebo or {}).get(ep_key, {}) or {}).get("visits", {})
        avp_any = False
        if isinstance(avp_visits, dict):
            for v in VISITS:
                d = avp_visits.get(v, {})
                if not isinstance(d, dict):
                    continue
                p = d.get("p_adj", d.get("p_value"))
                if p is not None and np.isfinite(p) and float(p) < 0.05:
                    avp_any = True
                    break
        if avp_any:
            sig_avp += 1

    out.append(
        f"По совокупности анализируемых показателей (n={any_data}) статистические сигналы различий между 4 группами (Kruskal–Wallis, p<0.05 на отдельных визитах) зарегистрированы для {sig_kw} показателей, а в укрупнённом сравнении Active vs Placebo (p_adj<0.05 на отдельных визитах) — для {sig_avp} показателей."
    )

    resp_findings = (exec_summary or {}).get("responder_findings", [])
    resp_rows = [r for r in resp_findings if isinstance(r, dict) and r.get("nnt") is not None]
    resp_rows_sorted = sorted(resp_rows, key=lambda r: float(r.get("nnt")) if r.get("nnt") is not None else 1e9)
    if resp_rows_sorted:
        top = resp_rows_sorted[:2]
        parts = []
        for r in top:
            parts.append(
                f"{r.get('name','—')}: NNT≈{float(r.get('nnt')):.1f} при долях респондеров {float(r.get('active_pct',0)):.0f}% (Active) и {float(r.get('placebo_pct',0)):.0f}% (Placebo)"
            )
        out.append("Клиническая интерпретация по респондерам (улучшение ≥20% от baseline): " + "; ".join(parts) + ".")
    else:
        out.append("Клиническая интерпретация по респондерам (улучшение ≥20% от baseline): данных недостаточно для устойчивых оценок клинического выигрыша (NNT).")

    out.append(
        "При интерпретации следует учитывать множественные сравнения (ориентир — p_adj), согласованность направления эффектов, величину ES и Δ, а также ограничения мощности при разделении выборки на 4 группы."
    )

    out.append(
        "Результаты настоящего отчёта не заменяют заранее зарегистрированный протокол статистического анализа и должны рассматриваться в контексте дизайна исследования и качества исходных данных."
    )

    return out


# ============================================================
# WORD REPORT - COMPREHENSIVE
# ============================================================

def generate_comprehensive_report_legacy(all_results, responders, active_vs_placebo, duration_active_vs_placebo, exec_summary, figures, df, output_path):
    """Generate Word report with EVERYTHING including graphs and business insights."""
    print(f"\n📄 Generating COMPREHENSIVE report: {output_path}")
    
    doc = Document()
    groups = sorted(df[GROUP_COL].dropna().unique())
    
    # === TITLE ===
    title = doc.add_paragraph()
    title.add_run("ПОЛНЫЙ СТАТИСТИЧЕСКИЙ ОТЧЁТ").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    
    doc.add_paragraph()
    doc.add_paragraph("Клиническое исследование ДИАМАГ при болезни Паркинсона").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    doc.add_heading("ОГЛАВЛЕНИЕ", level=1)
    add_toc(doc)
    doc.add_page_break()
    
    # === EXECUTIVE SUMMARY ===
    doc.add_heading("EXECUTIVE SUMMARY", level=1)
    
    doc.add_heading("Ключевые выводы", level=2)
    
    # Primary endpoints summary
    p = doc.add_paragraph()
    p.add_run("Первичные конечные точки:").bold = True
    for f in exec_summary["primary_findings"]:
        status = "✅ ДОСТИГНУТО" if f["kw_significant"] else "⚠️ НЕ ДОСТИГНУТО"
        if not f["kw_significant"] and f.get("avp_p", 1) < 0.05:
            status = "⚠️ ТРЕНД (Active vs Placebo)"
        
        doc.add_paragraph(f"• {f['name']}: {status}")
        if f.get("avp_p", 1) < 0.05:
             doc.add_paragraph(f"  (Active vs Placebo: p={f['avp_p']:.4f}, эффект значимый)")

    # Responders summary
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Клиническая значимость (NNT):").bold = True
    for r in exec_summary["responder_findings"]:
        if r.get("nnt"):
             doc.add_paragraph(f"• {r['name']}: NNT = {r['nnt']:.1f} (нужно пролечить {r['nnt']:.0f} чел. для 1 респондера)")
             doc.add_paragraph(f"  Ответ на терапию: Активная {r['active_pct']:.0f}% vs Плацебо {r['placebo_pct']:.0f}%")
    
    # Recommendation
    doc.add_paragraph()
    panel = doc.add_paragraph()
    run = panel.add_run(f"РЕКОМЕНДАЦИЯ: {exec_summary['recommendation']}")
    run.bold = True
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # === GLOSSARY ===
    doc.add_heading("ГЛОССАРИЙ СТАТИСТИЧЕСКИХ ТЕРМИНОВ", level=1)
    
    terms = [
        ("H₀ (нулевая гипотеза)", "Гипотеза об отсутствии различий между группами. Если H₀ верна — терапия не работает."),
        ("H₁ (альтернативная гипотеза)", "Гипотеза о наличии различий между группами. Если H₁ верна — терапия эффективна."),
        ("p-value", "Вероятность получить такие или более экстремальные результаты, если H₀ верна. p < 0.05 → отвергаем H₀."),
        ("Bayes Factor (BF₁₀)", "Отношение вероятностей H₁ к H₀. BF₁₀ > 3 → данные поддерживают H₁. BF₁₀ < 1/3 → поддерживают H₀."),
        ("Effect size (размер эффекта)", "Величина различий. r = 0.1 малый, 0.3 средний, 0.5 большой."),
        ("Kruskal-Wallis", "Непараметрический тест сравнения 3+ групп (альтернатива ANOVA)."),
        ("Mann-Whitney U", "Непараметрический тест сравнения 2 групп (альтернатива t-test)."),
        ("Wilcoxon", "Непараметрический парный тест изменений внутри группы."),
        ("Mixed Model", "Модель для анализа повторных измерений с учётом индивидуальных различий."),
    ]
    
    for term, desc in terms:
        p = doc.add_paragraph()
        p.add_run(f"{term}: ").bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # === STUDY DESIGN ===
    doc.add_heading("1. ДИЗАЙН ИССЛЕДОВАНИЯ", level=1)
    
    doc.add_heading("1.1. Группы пациентов", level=2)
    doc.add_paragraph(f"Всего пациентов: {len(df)}")
    
    # Correct group assignments
    group_info = {
        '1': ('ДИАМАГ 30 дней', 'Активная терапия'),
        '2': ('Плацебо 30 дней', 'Контроль'),
        '3': ('ДИАМАГ 20 дней', 'Активная терапия'),
        '4': ('Плацебо 20 дней', 'Контроль'),
    }
    
    table = doc.add_table(rows=1, cols=4, style="Table Grid")
    hdr = table.rows[0].cells
    for i, h in enumerate(["Группа", "N", "Терапия", "Тип"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    
    for g in groups:
        n = len(df[df[GROUP_COL] == g])
        info = group_info.get(str(g), ('Неизвестно', 'Неизвестно'))
        row = table.add_row().cells
        row[0].text = f"Группа {g}"
        row[1].text = str(n)
        row[2].text = info[0]
        row[3].text = info[1]
    
    doc.add_paragraph()
    doc.add_paragraph("Основные сравнения:")
    doc.add_paragraph("• Активная vs Плацебо (30 дней): Группа 1 vs Группа 2")
    doc.add_paragraph("• Активная vs Плацебо (20 дней): Группа 3 vs Группа 4")
    doc.add_paragraph("• Длительность терапии: 30 дней vs 20 дней (Группы 1+2 vs Группы 3+4)")
    
    doc.add_heading("1.2. Временные точки", level=2)
    
    table = doc.add_table(rows=1, cols=3, style="Table Grid")
    hdr = table.rows[0].cells
    for i, h in enumerate(["Визит", "День", "Описание"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    
    visits_info = [
        ("V2", "0", "Baseline — исходный уровень"),
        ("V3", "+10", "После окончания курса терапии"),
        ("V4", "+20", "Оценка стабильности эффекта"),
        ("V5", "+30", "Оценка стабильности эффекта"),
        ("V6", "+30±1", "Endpoint — финальная оценка"),
    ]
    for v, day, desc in visits_info:
        row = table.add_row().cells
        row[0].text = v
        row[1].text = day
        row[2].text = desc
    
    doc.add_heading("1.3. Схема статистических сравнений", level=2)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("На каждом визите (V2–V6):").bold = True
    doc.add_paragraph("• Kruskal-Wallis — сравнение всех 4 групп одновременно (H₀: медианы равны)")
    doc.add_paragraph("• Mann-Whitney U — попарные сравнения (G1 vs G2, G1 vs G3, и т.д.)")
    doc.add_paragraph("• Bayes Factor — оценка силы доказательств (BF₁₀ > 3 → в пользу H₁)")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Продольный анализ (V2 → каждый визит):").bold = True
    doc.add_paragraph("• Wilcoxon signed-rank — изменения внутри каждой группы")
    doc.add_paragraph("• Δ абсолютное и % изменение от baseline")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Анализ респондеров:").bold = True
    doc.add_paragraph("• Критерий: улучшение UPDRS III ≥ 20% от baseline")
    doc.add_paragraph("• Chi-square — сравнение долей респондеров между группами")
    
    doc.add_heading("1.4. Конечные точки", level=2)
    
    p = doc.add_paragraph()
    p.add_run("Первичные:").bold = True
    doc.add_paragraph("• UPDRS часть 3 — двигательные функции")
    doc.add_paragraph("• UPDRS часть 2 — повседневная активность")
    
    p = doc.add_paragraph()
    p.add_run("Вторичные:").bold = True
    for key, cfg in ENDPOINTS.items():
        if not cfg["primary"]:
            doc.add_paragraph(f"• {cfg['short']} — {cfg['name']}")
    
    doc.add_page_break()
    
    # === EACH ENDPOINT ===
    endpoint_num = 2
    for key, result in all_results.items():
        cfg = ENDPOINTS[key]
        
        doc.add_heading(f"{endpoint_num}. {result['name'].upper()}", level=1)
        
        is_primary = "ПЕРВИЧНАЯ" if result["primary"] else "ВТОРИЧНАЯ"
        direction = "снижение = улучшение" if result["direction"] == "lower_is_better" else "повышение = улучшение"
        doc.add_paragraph(f"Тип: {is_primary} конечная точка. Интерпретация: {direction}.")
        
        # --- 2.1 DESCRIPTIVE BY VISIT ---
        doc.add_heading(f"{endpoint_num}.1. Описательные статистики по визитам", level=2)
        
        for visit in VISITS:
            if visit not in result["by_visit"]:
                continue
            
            vdata = result["by_visit"][visit]
            
            doc.add_heading(f"Визит {visit}", level=3)
            
            # Table: Group stats
            table = doc.add_table(rows=1, cols=6, style="Table Grid")
            hdr = table.rows[0].cells
            for i, h in enumerate(["Группа", "N", "Median", "[Q1–Q3]", "Mean±SD", "Норм.?"]):
                hdr[i].text = h
                hdr[i].paragraphs[0].runs[0].bold = True
            
            for g, gstats in vdata["groups"].items():
                row = table.add_row().cells
                row[0].text = f"Группа {g}"
                row[1].text = str(gstats.get("n", 0))
                row[2].text = f"{gstats.get('median', 0):.1f}"
                row[3].text = f"[{gstats.get('q1', 0):.1f}–{gstats.get('q3', 0):.1f}]"
                row[4].text = f"{gstats.get('mean', 0):.1f} ± {gstats.get('sd', 0):.1f}"
                norm = gstats.get("normality", {})
                row[5].text = "Да" if norm.get("normal") else "Нет" if norm.get("normal") is False else "—"
            
            doc.add_paragraph()
            
            # Kruskal-Wallis
            kw = vdata.get("kruskal", {})
            if "H" in kw:
                p_val = kw["p"]
                bf = vdata.get("bf10", np.nan)
                
                p = doc.add_paragraph()
                p.add_run("Межгрупповое сравнение (Kruskal-Wallis): ").bold = True
                
                if kw["significant"]:
                    conclusion = "H₀ ОТВЕРГАЕТСЯ — есть значимые различия между группами"
                else:
                    conclusion = "H₀ НЕ отвергается — значимых различий не выявлено"
                
                doc.add_paragraph(f"H = {kw['H']:.2f}, p = {p_val:.4f}")
                doc.add_paragraph(f"Вывод: {conclusion}")
                
                if kw.get("epsilon_sq") is not None:
                    es = kw["epsilon_sq"]
                    doc.add_paragraph(f"Размер эффекта: ε² = {es:.3f} ({effect_size_interpret(es, 'epsilon_sq')})")
                
                if np.isfinite(bf):
                    doc.add_paragraph(f"Bayes Factor: BF₁₀ = {bf:.2f} — {interpret_bf(bf)}")
            
            # Pairwise comparisons
            if visit in result["pairwise"] and result["pairwise"][visit]:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run("Попарные сравнения (Mann-Whitney U):").bold = True
                
                for pair, mw in result["pairwise"][visit].items():
                    if "error" in mw:
                        continue
                    
                    g1, g2 = pair.split("_vs_")
                    sig = "✓ значимо" if mw["significant"] else "✗ не значимо"
                    r = mw.get("r", 0)
                    
                    doc.add_paragraph(
                        f"  Группа {g1} vs {g2}: U = {mw['U']:.0f}, p = {mw['p']:.4f} ({sig}), "
                        f"r = {r:.2f} ({effect_size_interpret(r, 'r')})"
                    )
            
            doc.add_paragraph()
        
        # --- 2.2 WITHIN-GROUP CHANGES ---
        doc.add_heading(f"{endpoint_num}.2. Изменения внутри групп (относительно baseline V2)", level=2)
        
        for g, changes in result.get("within_group_changes", {}).items():
            doc.add_heading(f"Группа {g}", level=3)
            
            table = doc.add_table(rows=1, cols=5, style="Table Grid")
            hdr = table.rows[0].cells
            for i, h in enumerate(["V2 →", "Δ Median", "p (Wilcoxon)", "r", "Вывод"]):
                hdr[i].text = h
                hdr[i].paragraphs[0].runs[0].bold = True
            
            for visit, wtest in changes.items():
                if "error" in wtest:
                    continue
                
                row = table.add_row().cells
                row[0].text = visit
                row[1].text = f"{wtest.get('median_diff', 0):+.1f}"
                row[2].text = f"{wtest.get('p', 1):.4f}"
                row[3].text = f"{wtest.get('r', 0):.2f}"
                
                if wtest["significant"]:
                    direction = "улучшение ↓" if wtest.get("median_diff", 0) < 0 else "ухудшение ↑"
                    row[4].text = f"Значимо: {direction}"
                else:
                    row[4].text = "Не значимо"
            
            doc.add_paragraph()
        
        # --- 2.3 SUMMARY FOR THIS ENDPOINT ---
        doc.add_heading(f"{endpoint_num}.3. Резюме по {result['short']}", level=2)
        
        # Find best and worst groups at V6
        v6_data = result["by_visit"].get("V6", {}).get("groups", {})
        if v6_data:
            sorted_groups = sorted(v6_data.items(), key=lambda x: x[1].get("median", 999))
            best = sorted_groups[0]
            worst = sorted_groups[-1]
            
            doc.add_paragraph(
                f"• На финальном визите V6 лучшие показатели в группе {best[0]} "
                f"(медиана = {best[1].get('median', 0):.1f}), худшие — в группе {worst[0]} "
                f"(медиана = {worst[1].get('median', 0):.1f})."
            )
        
        # Significant within-group changes
        sig_improvements = []
        for g, changes in result.get("within_group_changes", {}).items():
            v6_change = changes.get("V6", {})
            if v6_change.get("significant") and v6_change.get("median_diff", 0) < 0:
                sig_improvements.append(f"группа {g} (Δ = {v6_change['median_diff']:+.1f})")
        
        if sig_improvements:
            doc.add_paragraph(f"• Значимое улучшение V2→V6 наблюдается в: {', '.join(sig_improvements)}.")
        else:
            doc.add_paragraph("• Значимых улучшений V2→V6 не выявлено ни в одной группе.")
        
        # --- GRAPHS ---
        if figures:
            doc.add_heading(f"{endpoint_num}.4. Графики (GraphPad Style)", level=2)
            
            # Spaghetti plot
            fig_key = f"{key}_spaghetti"
            if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
                doc.add_picture(figures[fig_key], width=Inches(6))
                doc.add_paragraph(f"Рисунок: Динамика {result['short']} по группам (Mean ± SE) + Kruskal-Wallis p-values")
            
            # Detailed Boxplots for each visit
            doc.add_heading("Детальный анализ по визитам:", level=3)
            doc.add_paragraph("Графики включают индивидуальные значения (точки), медиану (линия) и 25-75 квартили (ящик). "
                              "Скобы показывают значимые различия (Mann-Whitney U).")
            
            for visit in VISITS:
                fig_key = f"{key}_boxplot_{visit}"
                if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
                     doc.add_picture(figures[fig_key], width=Inches(6))
                     doc.add_paragraph(f"Рисунок: {result['short']} на визите {visit}")
        
        doc.add_page_break()
        endpoint_num += 1
    
    # === RESPONDERS (ALL ENDPOINTS) ===
    doc.add_heading(f"{endpoint_num}. АНАЛИЗ РЕСПОНДЕРОВ (ВСЕ ПОКАЗАТЕЛИ)", level=1)
    
    doc.add_paragraph(
        "Респондер — пациент с улучшением ≥ 20% от baseline (V2 → V6). "
        "Для шкал, где меньше = лучше, улучшение = (V2 − V6) / V2 × 100%."
    )
    doc.add_paragraph()
    
    # Summary table across all endpoints
    p = doc.add_paragraph()
    p.add_run("Сводная таблица респондеров по всем показателям:").bold = True
    
    table = doc.add_table(rows=1, cols=6, style="Table Grid")
    hdr = table.rows[0].cells
    for i, h in enumerate(["Показатель", "Гр.1 (%)", "Гр.2 (%)", "Гр.3 (%)", "Гр.4 (%)", "p (χ²)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    
    for ep_key, resp_data in responders.items():
        row = table.add_row().cells
        row[0].text = resp_data.get("short", ep_key)
        
        # Get percentages for each group
        for i, g in enumerate(['1', '2', '3', '4']):
            group_data = next((r for r in resp_data.get("groups", []) if r["group"] == g), None)
            if group_data:
                row[i + 1].text = f"{group_data['pct']:.0f}%"
            else:
                row[i + 1].text = "—"
        
        # p-value
        test = resp_data.get("test", {})
        if test.get("p") is not None:
            row[5].text = f"{test['p']:.4f}" + (" *" if test["significant"] else "")
        else:
            row[5].text = "—"
    
    doc.add_paragraph()
    doc.add_paragraph("* p < 0.05")
    
    # Detailed for each endpoint
    for ep_key, resp_data in responders.items():
        doc.add_heading(f"Респондеры: {resp_data.get('short', ep_key)}", level=2)
        
        table = doc.add_table(rows=1, cols=5, style="Table Grid")
        hdr = table.rows[0].cells
        for i, h in enumerate(["Группа", "N", "Респондеры", "%", "95% ДИ"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        
        for r in resp_data.get("groups", []):
            row = table.add_row().cells
            row[0].text = f"Группа {r['group']}"
            row[1].text = str(r["n"])
            row[2].text = str(r["responders"])
            row[3].text = f"{r['pct']:.0f}%"
            row[4].text = f"[{r['ci_low']:.0f}–{r['ci_high']:.0f}%]"
        
        doc.add_paragraph()
        
        test = resp_data.get("test", {})
        if test.get("p") is not None:
            sig = "значима" if test["significant"] else "не значима"
            doc.add_paragraph(f"Chi-square: χ² = {test['chi2']:.2f}, p = {test['p']:.4f}. Разница {sig}.")
        
        # Interpretation
        if resp_data.get("groups"):
            best = max(resp_data["groups"], key=lambda x: x["pct"])
            worst = min(resp_data["groups"], key=lambda x: x["pct"])
            doc.add_paragraph(
                f"Наибольшая доля респондеров в группе {best['group']} ({best['pct']:.0f}%), "
                f"наименьшая — в группе {worst['group']} ({worst['pct']:.0f}%)."
            )
    
    endpoint_num += 1
    doc.add_page_break()
    
    endpoint_num += 1
    doc.add_page_break()
    
    # === ACTIVE VS PLACEBO (COMBINED) ===
    doc.add_heading(f"{endpoint_num}. СРАВНЕНИЕ ACTIVE VS PLACEBO (ОБЪЕДИНЁННЫЕ)", level=1)
    doc.add_paragraph(
        "В данном разделе сравниваются объединённые группы активной терапии (Группа 1 + Группа 3) "
        "против объединённых групп плацебо (Группа 2 + Группа 4). Это повышает статистическую мощность "
        "для обнаружения общих эффектов терапии."
    )
    
    table = doc.add_table(rows=1, cols=6, style="Table Grid")
    hdr = table.rows[0].cells
    for i, h in enumerate(["Показатель", "Activ Med", "Plac Med", "Diff", "p-value", "Effect Size"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    
    for key, res in active_vs_placebo.items():
        row = table.add_row().cells
        row[0].text = res["short"]
        row[1].text = f"{res['active_median']:.1f}"
        row[2].text = f"{res['placebo_median']:.1f}"
        
        diff = res["active_median"] - res["placebo_median"]
        row[3].text = f"{diff:.1f}"
        
        p_val = res["p_value"]
        sig = " *" if res["significant"] else ""
        row[4].text = f"{p_val:.4f}{sig}"
        
        row[5].text = f"{res['effect_size']:.2f}"
    
    doc.add_paragraph()
    doc.add_paragraph("* p < 0.05 (Mann-Whitney U тест)")
    
    endpoint_num += 1
    doc.add_page_break()
    
    # === GRAND SUMMARY ===
    doc.add_heading(f"{endpoint_num}. ИТОГОВАЯ СВОДКА", level=1)
    
    # Forest plot
    if figures and "forest" in figures and figures["forest"] and os.path.exists(figures["forest"]):
        doc.add_picture(figures["forest"], width=Inches(6))
        doc.add_paragraph("Рисунок: Сравнение размеров эффекта по показателям")
        doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=7, style="Table Grid")
    hdr = table.rows[0].cells
    headers = ["Показатель", "Тип", "K-W V6 p", "BF₁₀", "ε²", "Лучшая гр.", "Вывод"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    
    for key, result in all_results.items():
        v6 = result.get("by_visit", {}).get("V6", {})
        kw = v6.get("kruskal", {})
        bf = v6.get("bf10", np.nan)
        
        # Best group
        v6_groups = v6.get("groups", {})
        best_g = "—"
        if v6_groups:
            best_g = min(v6_groups, key=lambda g: v6_groups[g].get("median", 999))
        
        # Conclusion
        if kw.get("significant"):
            conclusion = "Значимо"
        elif kw.get("p") and kw["p"] < 0.1:
            conclusion = "Тренд"
        else:
            conclusion = "Не знач."
        
        row = table.add_row().cells
        row[0].text = result["short"]
        row[1].text = "I" if result["primary"] else "II"
        row[2].text = f"{kw.get('p', 1):.3f}" if kw.get("p") else "—"
        row[3].text = f"{bf:.1f}" if np.isfinite(bf) else "—"
        row[4].text = f"{kw.get('epsilon_sq', 0):.3f}" if kw.get("epsilon_sq") else "—"
        row[5].text = str(best_g)
        row[6].text = conclusion
    
    doc.add_page_break()
    
# ============================================================
# CAPTION GENERATOR & NUMBERING
# ============================================================

class Numbering:
    def __init__(self):
        self.figure = 1
        self.table = 1
    
    def fig(self):
        s = f"Рисунок {self.figure}"
        self.figure += 1
        return s
        
    def tab(self):
        s = f"Таблица {self.table}"
        self.table += 1
        return s

def generate_figure_caption_text(all_results, endpoint_key, visit_key=None, active_vs_placebo=None):
    """Generate scientific caption with statistical interpretation."""
    res = all_results.get(endpoint_key, {})
    short_name = res.get("short", endpoint_key)
    
    if visit_key:
        # Boxplot interpretation
        caption = f"Распределение значений {short_name} на визите {visit_key}. "
        
        # Check pairwise significance
        pairwise = res.get("pairwise", {}).get(visit_key, {})
        sig_pairs = []
        for pair, stats in pairwise.items():
            if not isinstance(stats, dict) or "error" in stats:
                continue
            p_used = stats.get("p_adj", stats.get("p", 1))
            if p_used is None or not np.isfinite(p_used) or float(p_used) >= 0.05:
                continue
            try:
                g1, g2 = pair.split("_vs_")
            except Exception:
                continue
            diff = stats.get("diff_median", np.nan)
            diff_pct = stats.get("diff_pct", np.nan)
            es = stats.get("r", np.nan)
            p_text = "p < 0.001" if float(p_used) < 0.001 else f"p = {float(p_used):.3f}"
            diff_text = f"Δ={float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "Δ=—"
            diffp_text = f" ({float(diff_pct):+.1f}%)" if diff_pct is not None and np.isfinite(diff_pct) else ""
            es_text = f"ES={float(es):.2f}" if es is not None and np.isfinite(es) else "ES=—"
            sig_pairs.append(f"Гр.{g1} vs Гр.{g2}: {p_text}, {diff_text}{diffp_text}, {es_text}")
        
        if sig_pairs:
            caption += f"Выявлены статистически значимые различия: {', '.join(sig_pairs)}. "
            caption += "Центральная линия — медиана, ящик — межквартильный размах, точки — индивидуальные значения."
        else:
            caption += "Статистически значимых различий между группами не выявлено (p_adj ≥ 0.05). "
            caption += "Центральная линия — медиана, ящик — межквартильный размах."
            
    else:
        # Spaghetti plot interpretation
        caption = f"Динамика показателя {short_name} в течение периода наблюдения (V2–V6). "
        
        # Check overall V6 significance
        v6 = res.get("by_visit", {}).get("V6", {})
        kw = v6.get("kruskal", {})
        
        if kw.get("significant"):
            caption += f"На финальном визите V6 отмечены значимые различия между группами (Kruskal-Wallis p={kw['p']:.4f}). "
        else:
            caption += f"На финальном визите V6 значимых различий не выявлено (p={kw.get('p', 1):.3f}). "
            
        # Mention active vs placebo if significant
        if active_vs_placebo:
            avp = active_vs_placebo.get(endpoint_key, {})
            avp_v6 = avp.get("visits", {}).get("V6", {})
            if avp_v6.get("significant"):
                p_used = avp_v6.get("p_adj", avp_v6.get("p_value", np.nan))
                caption += f"Объединенный анализ Active vs Placebo (V6) показывает значимый эффект (p_adj={float(p_used):.4f}, ES={float(avp_v6.get('effect_size', 0.0)):.2f})."
    
    return caption


def generate_comprehensive_report(all_results, responders, active_vs_placebo, duration_active_vs_placebo, exec_summary, figures, df, output_path):
    """Generate Word report with EVERYTHING including graphs and business insights."""
    print(f"\n📄 Generating COMPREHENSIVE report: {output_path}")
    
    doc = Document()
    num = Numbering()
    groups = sorted(df[GROUP_COL].dropna().unique())
    
    # === TITLE ===
    title = doc.add_paragraph()
    title.add_run("ПОЛНЫЙ АНАЛИТИЧЕСКИЙ ОТЧЁТ ПО ИССЛЕДОВАНИЮ ДИАМАГ").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(22)
    
    doc.add_paragraph()
    doc.add_paragraph("Фокус: эффективность терапии, динамика ответа, сравнение с плацебо").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    doc.add_heading("ОГЛАВЛЕНИЕ", level=1)
    add_toc(doc)
    doc.add_page_break()
    
    # === 1. EXECUTIVE SUMMARY ===
    doc.add_heading("1. EXECUTIVE SUMMARY (РЕЗЮМЕ)", level=1)
    doc.add_paragraph("Данное резюме предназначено для быстрого ознакомления с ключевыми результатами исследования.")
    
    # Recommendation
    panel = doc.add_paragraph()
    run = panel.add_run(f"ИТОГОВОЕ ЗАКЛЮЧЕНИЕ: {exec_summary['recommendation']}")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 150)
    
    doc.add_heading("1.1 Ключевые находки", level=2)
    for f in exec_summary["primary_findings"]:
        # Status icon
        if f["kw_significant"]:
            icon = "✅ ДОКАЗАНО"
        elif f.get("avp_p_adj", 1) < 0.05:
            icon = "⚠️ ТРЕНД (Active vs Placebo)"
        else:
            icon = "❌ НЕ ДОКАЗАНО"
            
        p = doc.add_paragraph()
        p.add_run(f"{f['name']}: {icon}").bold = True
        
        if f.get("avp_p_adj", 1) < 0.05:
            diff = f.get("diff_median", np.nan)
            diff_pct = f.get("diff_pct", np.nan)
            diff_txt = f"Δ={float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "Δ=—"
            diffp_txt = f" ({float(diff_pct):+.1f}%)" if diff_pct is not None and np.isfinite(diff_pct) else ""
            doc.add_paragraph(
                f"  • Active vs Placebo (V6): p_adj={float(f['avp_p_adj']):.4f}, ES={float(f.get('avp_effect', 0.0)):.2f}, {diff_txt}{diffp_txt} (значимо)"
            )
    
    doc.add_heading("1.2 Клиническая значимость (NNT)", level=2)
    doc.add_paragraph("NNT (Number Needed to Treat) — сколько пациентов нужно пролечить, чтобы получить у одного улучшение ≥20%.")
    for r in exec_summary["responder_findings"]:
        if r.get("nnt"):
             doc.add_paragraph(f"• {r['name']}: NNT = {r['nnt']:.1f}. (Эффективность {r['active_pct']:.0f}% vs Плацебо {r['placebo_pct']:.0f}%)")
    
    doc.add_page_break()
    
    # === 2. EDUCATIONAL SECTION ===
    doc.add_heading("2. МЕТОДОЛОГИЯ: КАК ЧИТАТЬ ЭТОТ ОТЧЁТ (DATA STORYTELLING)", level=1)
    doc.add_paragraph(
        "Этот раздел поможет интерпретировать представленные данные и статистические тесты."
    )
    
    stories = [
        ("Kruskal-Wallis Test", "Мы используем его первым. Это 'широкий взгляд'. Он отвечает на вопрос: 'Есть ли вообще хоть какая-то разница между этими 4 группами?'. Если p < 0.05 — значит, группы точно разные, и мы можем искать дальше (кто именно лучше). Если p > 0.05 — различия не доказаны."),
        ("Mann-Whitney U Test", "Это 'снайперский прицел'. Мы берем две конкретные группы (например, ДИАМАГ vs Плацебо) и сравниваем их лоб в лоб. Это позволяет найти эффекты, которые 'Крускал-Уоллис' мог пропустить из-за большого количества групп."),
        ("Bayes Factor (BF₁₀)", "Коэффициент уверенности. Традиционное p-value говорит 'это не случайность', а Байес говорит 'насколько гипотеза об эффективности вероятнее гипотезы о пустышке'. BF > 3 — хорошие доказательства эффективности."),
        ("Spaghetti Plot", "График 'спагетти' показывает среднюю динамику. Если зеленая линия (Active) уходит вниз (улучшение) круче, чем серая (Placebo) — это визуальное подтверждение эффективности."),
        ("Boxplot (Ящик с усами)", "Показывает не только среднее, но и разброс. 'Ящик' — это где находятся 50% пациентов. 'Усы' — остальные. Точки — отдельные пациенты. Если ящики двух групп не перекрываются по вертикали — это признак сильного различия.")
    ]
    
    for title, desc in stories:
        p = doc.add_paragraph()
        p.add_run(f"🔹 {title}: ").bold = True
        p.add_run(desc)
        
    doc.add_page_break()
    
    # === 3. POOLED ANALYSIS ===
    doc.add_heading("3. СРАВНЕНИЕ ЭФФЕКТИВНОСТИ: ACTIVE VS PLACEBO (POOLED ANALYSIS)", level=1)
    doc.add_paragraph(
        "В этом разделе мы объединили данные всех пациентов, получавших ДИАМАГ (Группы 1 и 3), "
        "и сравнили их с теми, кто получал Плацебо (Группы 2 и 4). "
        "Это наиболее мощный статистический подход для ответа на главный вопрос: «Работает ли аппарат?»."
    )
    
    pooled_idx = 1
    for key, res in active_vs_placebo.items():
        doc.add_heading(f"3.{pooled_idx} Показатель: {res['short']}", level=2)
        
        # Spaghetti Plot
        fig_key = f"{key}_pooled_spaghetti"
        if fig_key in figures:
            doc.add_picture(figures[fig_key], width=Inches(6))
            caption_text = f"Динамика {res['short']} (Active vs Placebo). Подписи p= отмечают визиты со статистически значимыми различиями."
            p = doc.add_paragraph(num.fig() + ". " + caption_text)
            p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None
        
        # Table of longitudinal stats
        doc.add_paragraph(num.tab() + f". Статистика Active vs Placebo по визитам ({res['short']})")
        table = doc.add_table(rows=1, cols=9, style="Table Grid")
        hdr = table.rows[0].cells
        for i, h in enumerate(["Визит", "Active (Med)", "Placebo (Med)", "Δ (Abs)", "Δ (%)", "p_adj", "BF₁₀", "ES", "Вывод"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            
        visits_sorted = sorted(res["visits"].keys(), key=lambda v: VISITS.index(v) if v in VISITS else 99)
        for v in visits_sorted:
            d = res["visits"][v]
            row = table.add_row().cells
            row[0].text = v
            row[1].text = f"{d['active_median']:.1f}"
            row[2].text = f"{d['placebo_median']:.1f}"
            row[3].text = f"{d.get('diff_median', 0):+.1f}"
            dp = d.get('diff_pct', np.nan)
            row[4].text = f"{float(dp):+.1f}%" if dp is not None and np.isfinite(dp) else "—"
            p_val = float(d.get('p_adj', d.get('p_value', np.nan)))
            row[5].text = f"{p_val:.4f}" if np.isfinite(p_val) else "—"
            bf = d.get('bf10', np.nan)
            row[6].text = f"{bf:.2f}" if np.isfinite(bf) else "—"
            row[7].text = f"{float(d.get('effect_size', 0.0)):.2f}"
            row[8].text = "Значимо" if d.get('significant') else "Не значимо"
            
        doc.add_paragraph()

        doc.add_heading("Boxplot по визитам (Active vs Placebo)", level=3)
        for v in VISITS:
            fig_key = f"{key}_pooled_boxplot_{v}"
            if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
                doc.add_picture(figures[fig_key], width=Inches(6))

                vstats = res.get("visits", {}).get(v, {})
                p_used = vstats.get("p_adj", vstats.get("p_value", np.nan))
                p_txt = "p < 0.001" if p_used is not None and np.isfinite(p_used) and float(p_used) < 0.001 else (f"p = {float(p_used):.3f}" if p_used is not None and np.isfinite(p_used) else "p = —")
                diff = vstats.get("diff_median", np.nan)
                diff_pct = vstats.get("diff_pct", np.nan)
                es = vstats.get("effect_size", np.nan)
                diff_txt = f"Δ={float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "Δ=—"
                diffp_txt = f" ({float(diff_pct):+.1f}%)" if diff_pct is not None and np.isfinite(diff_pct) else ""
                es_txt = f"ES={float(es):.2f}" if es is not None and np.isfinite(es) else "ES=—"
                verdict = "значимо" if vstats.get("significant") else "незначимо"

                p = doc.add_paragraph(num.fig() + f". Распределение {res['short']} на визите {v} (Active vs Placebo): {p_txt} ({verdict}), {diff_txt}{diffp_txt}, {es_txt}.")
                p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None

        pooled_idx += 1
        
    doc.add_heading("3.2 ДОПОЛНИТЕЛЬНОЕ ИССЛЕДОВАНИЕ: ПОПАРНОЕ ACTIVE VS PLACEBO ВНУТРИ ДЛИТЕЛЬНОСТИ", level=1)
    doc.add_paragraph(
        "В этом дополнительном разделе сравнение проводится отдельно для двух длительностей курса: "
        "30 дней (Группа 1 vs Группа 2) и 20 дней (Группа 3 vs Группа 4). "
        "Данный подход увеличивает объём аналитики и позволяет оценить эффект терапии внутри фиксированной длительности, "
        "хотя формально он не заменяет глобальный 4-групповой тест."
    )

    duration_order = ["30d", "20d"]
    duration_titles = {
        "30d": "30 дней: ДИАМАГ vs Плацебо (Г1 vs Г2)",
        "20d": "20 дней: ДИАМАГ vs Плацебо (Г3 vs Г4)",
    }

    for key, res in duration_active_vs_placebo.items():
        doc.add_heading(f"3.2.{key} Показатель: {res.get('short', key)}", level=2)

        for dkey in duration_order:
            comp = res.get("comparisons", {}).get(dkey, {})
            if not comp:
                continue

            doc.add_heading(duration_titles.get(dkey, dkey), level=3)

            fig_key = f"{key}_{dkey}_spaghetti"
            if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
                doc.add_picture(figures[fig_key], width=Inches(6))
                p = doc.add_paragraph(num.fig() + f". Динамика {res.get('short', key)}: {duration_titles.get(dkey, dkey)}.")
                p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None

            doc.add_paragraph(num.tab() + f". Статистика по визитам: {res.get('short', key)} ({duration_titles.get(dkey, dkey)})")
            table = doc.add_table(rows=1, cols=9, style="Table Grid")
            hdr = table.rows[0].cells
            for i, h in enumerate(["Визит", "Active (Med)", "Placebo (Med)", "Δ (Abs)", "Δ (%)", "p-value", "BF₁₀", "Effect Size", "Вывод"]):
                hdr[i].text = h
                hdr[i].paragraphs[0].runs[0].bold = True

            visits = comp.get("visits", {})
            visits_sorted = sorted(list(visits.keys()), key=lambda v: VISITS.index(v) if v in VISITS else 99)
            for v in visits_sorted:
                d = visits[v]
                row = table.add_row().cells
                row[0].text = v
                row[1].text = f"{d['active_median']:.1f}"
                row[2].text = f"{d['placebo_median']:.1f}"
                row[3].text = f"{d.get('diff_median', 0):+.1f}"
                dp = d.get('diff_pct', np.nan)
                row[4].text = f"{dp:+.1f}%" if np.isfinite(dp) else "—"
                p_val = float(d.get('p_adj', d.get('p_value', np.nan)))
                row[5].text = (f"{p_val:.4f}" if np.isfinite(p_val) else "—")
                bf = d.get('bf10', np.nan)
                row[6].text = f"{bf:.2f}" if np.isfinite(bf) else "—"
                row[7].text = f"{d.get('effect_size', 0):.2f}"
                row[8].text = "Значимо" if d.get('significant') else "Не значимо"

            doc.add_paragraph(
                "Визит-специфические различия интерпретируются одновременно по p-value и размеру эффекта, "
                "а также по BF₁₀ как мере силы доказательств в пользу гипотезы об эффективности терапии."
            )

            for visit in VISITS:
                fig_key = f"{key}_{dkey}_boxplot_{visit}"
                if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
                    doc.add_picture(figures[fig_key], width=Inches(6))
                    p = doc.add_paragraph(num.fig() + f". {res.get('short', key)} на визите {visit}: {duration_titles.get(dkey, dkey)}.")
                    p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None

        doc.add_page_break()

    doc.add_page_break()
    
    # === 4. DETAILED ENDPOINT ANALYSIS ===
    doc.add_heading("4. ДЕТАЛЬНЫЙ АНАЛИЗ ПО ПОКАЗАТЕЛЯМ (4 ГРУППЫ)", level=1)
    
    endpoint_num = 1
    for key, result in all_results.items():
        doc.add_heading(f"4.{endpoint_num} {result['name']}", level=2)
        
        # --- Description ---
        doc.add_paragraph(f"Анализ показателя '{result['short']}' для 4-х рандомизированных групп.")
        
        # --- 4.X.1 Kruskal-Wallis ---
        v6 = result.get("by_visit", {}).get("V6", {})
        kw = v6.get("kruskal", {})
        if kw.get("significant"):
            doc.add_paragraph(f"✅ Тест Краскела-Уоллиса на V6 показал значимые различия (p={kw['p']:.4f}). Необходимо смотреть попарные сравнения.")
        else:
            doc.add_paragraph(f"ℹ️ Тест Краскела-Уоллиса на V6 не выявил глобальных различий (p={kw.get('p',1):.4f}).")
            
        # --- 4.X.2 Pairwise Comparisons (Detailed Matrix) ---
        doc.add_heading("Попарные сравнения (Mann-Whitney U)", level=3)
        doc.add_paragraph("Ниже представлена детализация различий между конкретными парами групп на каждом визите.")

        doc.add_paragraph(num.tab() + f". Матрица попарных сравнений (p_adj) по визитам: {result['short']}.")
        
        table = doc.add_table(rows=1, cols=len(VISITS)+1, style="Table Grid")
        hdr = table.rows[0].cells
        hdr[0].text = "Сравнение"
        for i, v in enumerate(VISITS):
            hdr[i+1].text = v
            hdr[i+1].paragraphs[0].runs[0].bold = True
            
        pairwise_keys = sorted(list(result.get("pairwise", {}).get("V6", {}).keys()))
        for pair_key in pairwise_keys:
            row = table.add_row().cells
            row[0].text = pair_key.replace("_vs_", " vs ")
            
            for i, v in enumerate(VISITS):
                res_mw = result["pairwise"].get(v, {}).get(pair_key)
                if res_mw and res_mw.get("p_adj") is not None:
                    p_used = float(res_mw.get("p_adj"))
                    row[i+1].text = f"{p_used:.4f}" if np.isfinite(p_used) else "—"
                else:
                     row[i+1].text = "—"
        
        doc.add_paragraph("Критерий значимости: p_adj < 0.05 (Холм). Сравнения: Группа 1 (Активная 30д), Группа 2 (Плацебо 30д), Группа 3 (Активная 20д), Группа 4 (Плацебо 20д).")
        
        # --- 4.X.3 Within-Group Changes (Wilcoxon) ---
        doc.add_heading("Внутригрупповая динамика (Wilcoxon)", level=3)
        doc.add_paragraph("Значимость изменений внутри каждой группы относительно baseline (V2).")

        doc.add_paragraph(num.tab() + f". Внутригрупповая динамика (V2 → визит): {result['short']}.")
        
        table = doc.add_table(rows=1, cols=len(VISITS)+1, style="Table Grid")
        hdr = table.rows[0].cells
        hdr[0].text = "Группа / Визит (vs V2)"
        for i, v in enumerate(VISITS): # V2 vs V2 is trivial, but let's keep visits structure or skip V2
             if v == "V2":
                 hdr[i+1].text = "V2 (Base)"
             else:
                 hdr[i+1].text = f"Δ {v}"
             hdr[i+1].paragraphs[0].runs[0].bold = True
             
        for g in groups:
            row = table.add_row().cells
            row[0].text = f"Группа {g}"
            
            row[1].text = "—" # V2
            
            # V3..V6
            for i, v in enumerate(VISITS[1:], start=2):
                 res_w = result.get("within_group_changes", {}).get(g, {}).get(v)
                 if isinstance(res_w, dict) and "error" not in res_w:
                     p_used = res_w.get("p_adj", res_w.get("p", np.nan))
                     delta = res_w.get("median_diff", np.nan)
                     delta_pct = res_w.get("delta_pct", np.nan)
                     es = res_w.get("r", np.nan)

                     p_txt = f"p_adj={float(p_used):.4f}" if p_used is not None and np.isfinite(p_used) else "p_adj=—"
                     dp_txt = f" ({float(delta_pct):+.1f}%)" if delta_pct is not None and np.isfinite(delta_pct) else ""
                     d_txt = f"Δ={float(delta):+.1f}{dp_txt}" if delta is not None and np.isfinite(delta) else "Δ=—"

                     if res_w.get("significant"):
                         es_txt = f"ES={float(es):.2f}" if es is not None and np.isfinite(es) else "ES=—"
                         row[i].text = f"{p_txt}\n{d_txt}\n{es_txt}\nзначимо"
                     else:
                         row[i].text = f"{p_txt}\n{d_txt}\nнезначимо"
                 else:
                     row[i].text = "—"

        doc.add_heading("Mixed Effects (повторные измерения)", level=3)
        doc.add_paragraph("Модель оценивает взаимодействие Визит×Группа с учётом повторных измерений у пациентов.")

        mm4 = result.get("mixed_effects_4g", {})
        mmp = result.get("mixed_effects_pooled", {})

        doc.add_paragraph(num.tab() + f". Mixed effects: {result['short']} (4 группы и pooled).")
        table = doc.add_table(rows=1, cols=6, style="Table Grid")
        hdr = table.rows[0].cells
        for i, h in enumerate(["Модель", "N (obs)", "N (subjects)", "p (interaction)", "Значимость", "Интерпретация"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        def _add_mm_row(label: str, mm: Dict[str, Any]):
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(mm.get("n_observations", "—")) if isinstance(mm, dict) and "error" not in mm else "—"
            row[2].text = str(mm.get("n_subjects", "—")) if isinstance(mm, dict) and "error" not in mm else "—"
            p_int = mm.get("interaction_p_value") if isinstance(mm, dict) else None
            row[3].text = f"{float(p_int):.4f}" if p_int is not None and np.isfinite(p_int) else "—"
            sig = mm.get("interaction", {}).get("significant") if isinstance(mm, dict) else None
            if sig is True:
                row[4].text = "значимо"
            elif sig is False:
                row[4].text = "незначимо"
            else:
                row[4].text = "—"
            interp = mm.get("interaction", {}).get("interpretation") if isinstance(mm, dict) else None
            row[5].text = str(interp) if interp else (str(mm.get("error")) if isinstance(mm, dict) and mm.get("error") else "—")

        _add_mm_row("4 группы", mm4)
        _add_mm_row("Укрупнение (Active vs Placebo)", mmp)
        
        # --- 4.X.4 Detailed Graphs (Standard) ---
        fig_key = f"{key}_spaghetti"
        if fig_key in figures:
            doc.add_picture(figures[fig_key], width=Inches(6))
            doc.add_paragraph(num.fig() + f". Динамика показателя для 4 групп.")

        doc.add_heading("Boxplot по визитам (4 группы)", level=3)
        for v in VISITS:
            fig_key = f"{key}_boxplot_{v}"
            if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
                doc.add_picture(figures[fig_key], width=Inches(6))
                caption = generate_figure_caption_text(all_results, key, visit_key=v, active_vs_placebo=None)
                p = doc.add_paragraph(num.fig() + ". " + caption)
                p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None
        
        # --- 4.X.3 Responder Analysis (Longitudinal) ---
        doc.add_heading("Динамика терапевтического ответа (Респондеры)", level=3)
        doc.add_paragraph("Ниже представлена доля пациентов, достигших улучшения ≥20% на каждом визите.")
        
        for v in ["V3", "V4", "V5", "V6"]:
             chart_key = f"{key}_resp_barplot_{v}"
             if chart_key in figures and os.path.exists(figures[chart_key]):
                 doc.add_picture(figures[chart_key], width=Inches(5))
                 doc.add_paragraph(num.fig() + f". Доля респондеров на визите {v}.")

        resp = responders.get(key, {})
        if isinstance(resp, dict) and resp.get("visits"):
            doc.add_paragraph(num.tab() + f". Респондеры по визитам: {result['short']} (4 группы + укрупнение)")
            table = doc.add_table(rows=1, cols=8, style="Table Grid")
            hdr = table.rows[0].cells
            for i, h in enumerate(["Визит", "Г1 %", "Г2 %", "Г3 %", "Г4 %", "p (χ²)", "Active %", "Placebo %"]):
                hdr[i].text = h
                hdr[i].paragraphs[0].runs[0].bold = True

            for v in ["V3", "V4", "V5", "V6"]:
                vres = resp.get("visits", {}).get(v, {})
                groups_list = vres.get("groups", []) if isinstance(vres, dict) else []
                by_g = {str(r.get("group")): r for r in groups_list if isinstance(r, dict)}
                row = table.add_row().cells
                row[0].text = v
                for col_i, g in enumerate(["1", "2", "3", "4"], start=1):
                    pct = by_g.get(g, {}).get("pct")
                    row[col_i].text = f"{float(pct):.0f}%" if pct is not None and np.isfinite(pct) else "—"
                p_val = vres.get("test", {}).get("p") if isinstance(vres, dict) else None
                row[5].text = f"{float(p_val):.4f}" if p_val is not None and np.isfinite(p_val) else "—"

                active_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["1", "3"])
                active_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["1", "3"])
                placebo_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["2", "4"])
                placebo_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["2", "4"])
                active_pct = (active_resp / active_n * 100.0) if active_n > 0 else np.nan
                placebo_pct = (placebo_resp / placebo_n * 100.0) if placebo_n > 0 else np.nan
                row[6].text = f"{active_pct:.0f}%" if np.isfinite(active_pct) else "—"
                row[7].text = f"{placebo_pct:.0f}%" if np.isfinite(placebo_pct) else "—"

            best_visit = None
            best_arr = None
            for v in ["V3", "V4", "V5", "V6"]:
                vres = resp.get("visits", {}).get(v, {})
                groups_list = vres.get("groups", []) if isinstance(vres, dict) else []
                by_g = {str(r.get("group")): r for r in groups_list if isinstance(r, dict)}
                active_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["1", "3"])
                active_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["1", "3"])
                placebo_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["2", "4"])
                placebo_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["2", "4"])
                if active_n <= 0 or placebo_n <= 0:
                    continue
                active_pct = active_resp / active_n * 100.0
                placebo_pct = placebo_resp / placebo_n * 100.0
                arr = active_pct - placebo_pct
                if best_arr is None or arr > best_arr:
                    best_arr = arr
                    best_visit = v

            if best_visit is not None and best_arr is not None:
                vres = resp.get("visits", {}).get(best_visit, {})
                groups_list = vres.get("groups", []) if isinstance(vres, dict) else []
                by_g = {str(r.get("group")): r for r in groups_list if isinstance(r, dict)}

                active_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["1", "3"])
                active_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["1", "3"])
                placebo_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["2", "4"])
                placebo_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["2", "4"])

                active_pct = (active_resp / active_n * 100.0) if active_n > 0 else np.nan
                placebo_pct = (placebo_resp / placebo_n * 100.0) if placebo_n > 0 else np.nan

                nnt, _ = calculate_nnt(float(active_pct), float(placebo_pct)) if np.isfinite(active_pct) and np.isfinite(placebo_pct) else (None, None)
                nnt_txt = f"NNT≈{nnt:.1f}" if nnt is not None else "NNT=—"
                doc.add_paragraph(
                    f"Интерпретация: максимальный разрыв долей респондеров Active vs Placebo наблюдается на {best_visit} "
                    f"(Δ={best_arr:+.1f} п.п.), что соответствует {nnt_txt} при трактовке как ARR."
                )
        
        doc.add_page_break()
        endpoint_num += 1

    # === 5. FINAL SUMMARY (BY VISIT) ===
    doc.add_heading("5. ИТОГОВАЯ СВОДКА ПО ВИЗИТАМ", level=1)
    doc.add_paragraph(
        "Раздел агрегирует результаты по каждой временной точке (V2–V6): глобальный тест 4 групп (Kruskal-Wallis), "
        "лучшие пары по post-hoc (p_adj по Холму) и объединённый анализ Active vs Placebo (p_adj, ES, Δ)."
    )

    for idx, v in enumerate(VISITS, start=1):
        doc.add_heading(f"5.{idx} Визит {v}", level=2)
        doc.add_paragraph(num.tab() + f". Сводка по визиту {v} (все показатели)")

        table = doc.add_table(rows=1, cols=10, style="Table Grid")
        hdr = table.rows[0].cells
        headers = [
            "Показатель",
            "K-W p",
            "BF₁₀",
            "ε²",
            "Лучшая пара (p_adj)",
            "ES пары",
            "Active vs Placebo p_adj",
            "ES (укрупнение)",
            "Δ (Abs)",
            "Δ (%)",
        ]
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        newbie_hits = []

        for ep_key, result in all_results.items():
            row = table.add_row().cells
            row[0].text = result.get("short", ep_key)

            vdata = result.get("by_visit", {}).get(v, {})
            kw = vdata.get("kruskal", {})
            bf = vdata.get("bf10", np.nan)

            row[1].text = f"{float(kw.get('p')):.4f}" if kw.get("p") is not None and np.isfinite(kw.get("p")) else "—"
            row[2].text = f"{float(bf):.2f}" if np.isfinite(bf) else "—"
            row[3].text = f"{float(kw.get('epsilon_sq')):.3f}" if kw.get("epsilon_sq") is not None and np.isfinite(kw.get("epsilon_sq")) else "—"

            best_pair_txt = "—"
            best_pair_es = "—"
            pv = result.get("pairwise", {}).get(v, {})
            best = None
            if isinstance(pv, dict) and pv:
                for pair_key, st in pv.items():
                    if not isinstance(st, dict) or "error" in st:
                        continue
                    p_adj = st.get("p_adj")
                    if p_adj is None or not np.isfinite(p_adj):
                        continue
                    if best is None or float(p_adj) < float(best["p_adj"]):
                        best = {"pair": pair_key, "p_adj": float(p_adj), "es": float(st.get("r", np.nan))}
            if best is not None and best["p_adj"] < 0.05:
                try:
                    g1, g2 = str(best["pair"]).split("_vs_")
                    best_pair_txt = f"Г{g1}–Г{g2} (p={best['p_adj']:.3f})"
                except Exception:
                    best_pair_txt = f"p={best['p_adj']:.3f}"
            row[4].text = best_pair_txt
            if best is not None and best.get("es") is not None and np.isfinite(best.get("es")) and best.get("p_adj") is not None and np.isfinite(best.get("p_adj")) and float(best.get("p_adj")) < 0.05:
                best_pair_es = f"{float(best.get('es')):.2f}"
            row[5].text = best_pair_es

            avp = active_vs_placebo.get(ep_key, {}).get("visits", {}).get(v, {})
            p_avp = avp.get("p_adj", avp.get("p_value", np.nan))
            row[6].text = f"{float(p_avp):.4f}" if p_avp is not None and np.isfinite(p_avp) else "—"
            es_avp = avp.get("effect_size", np.nan)
            row[7].text = f"{float(es_avp):.2f}" if es_avp is not None and np.isfinite(es_avp) else "—"
            diff = avp.get("diff_median", np.nan)
            diff_pct = avp.get("diff_pct", np.nan)
            row[8].text = f"{float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "—"
            row[9].text = f"{float(diff_pct):+.1f}%" if diff_pct is not None and np.isfinite(diff_pct) else "—"

            if avp.get("significant"):
                direction = "в пользу Active" if (result.get("direction") == "lower_is_better" and diff is not None and np.isfinite(diff) and float(diff) < 0) else ("в пользу Active" if (result.get("direction") != "lower_is_better" and diff is not None and np.isfinite(diff) and float(diff) > 0) else "направление неоднозначно")
                newbie_hits.append(f"{result.get('short', ep_key)}: значимо (p_adj={float(p_avp):.3f}, ES={float(es_avp):.2f}, {direction})")

        if newbie_hits:
            doc.add_paragraph("Интерпретация: на этом визите укрупнение выявило значимые различия по: " + "; ".join(newbie_hits) + ".")
        else:
            doc.add_paragraph("Интерпретация: на этом визите укрупнение не выявило значимых различий (p_adj ≥ 0.05) ни по одному показателю.")

        doc.add_paragraph(
            "Важно: отсутствие значимости не равно отсутствию эффекта; дополнительно оценивайте размер эффекта (ES) и абсолютную/процентную разницу (Δ)."
        )
        doc.add_page_break()

    # === 6. DISCUSSION & CONCLUSIONS ===
    doc.add_heading("6. ОБСУЖДЕНИЕ И ВЫВОДЫ", level=1)
    
    discussion_text = generate_discussion_text(exec_summary, all_results, active_vs_placebo, responders)
    for paragraph in discussion_text:
        doc.add_paragraph(paragraph)

    doc.add_heading("6.1 Цели и задачи", level=2)
    doc.add_paragraph("Цель: оценить эффективность аппарата ДИАМАГ по ключевым клиническим шкалам в сравнении с плацебо.")
    doc.add_paragraph("Задачи: (1) сравнить 4 рандомизированные группы на каждом визите; (2) выполнить post-hoc сравнения с коррекцией (Холм); (3) выполнить укрупнение Active vs Placebo; (4) оценить внутригрупповую динамику; (5) оценить долю респондеров и клиническую значимость (NNT); (6) подтвердить результаты mixed effects моделями повторных измерений.")

    doc.add_heading("6.2 Проверка гипотез", level=2)
    doc.add_paragraph("H₀: терапия не приводит к различиям между группами. H₁: терапия приводит к различиям (эффект есть).")
    for ep_key, result in all_results.items():
        if not result.get("primary"):
            continue
        v6_kw = result.get("by_visit", {}).get("V6", {}).get("kruskal", {})
        avp_v6 = active_vs_placebo.get(ep_key, {}).get("visits", {}).get("V6", {})
        kw_p = v6_kw.get("p", np.nan)
        avp_p = avp_v6.get("p_adj", avp_v6.get("p_value", np.nan))
        if kw_p is not None and np.isfinite(kw_p) and float(kw_p) < 0.05:
            doc.add_paragraph(f"{result.get('short', ep_key)}: H₀ отвергается на уровне 4 групп (V6, p={float(kw_p):.4f}).")
        elif avp_p is not None and np.isfinite(avp_p) and float(avp_p) < 0.05:
            doc.add_paragraph(f"{result.get('short', ep_key)}: H₀ отвергается в укрупнении (V6, p_adj={float(avp_p):.4f}, ES={float(avp_v6.get('effect_size', 0.0)):.2f}).")
        else:
            doc.add_paragraph(f"{result.get('short', ep_key)}: статистических оснований отвергать H₀ на V6 не получено (p≥0.05).")

    doc.add_heading("6.3 Итоговые выводы", level=2)
    doc.add_paragraph(f"Итог: {exec_summary.get('recommendation', '—')}")
    doc.add_paragraph("Рекомендация к интерпретации: опирайтесь на сочетание p-value (после коррекции), размера эффекта (ES) и клинических метрик (респондеры, NNT), а не на один показатель.")
        
    doc.save(output_path)
    print(f"   ✓ Saved: {output_path}")
    return str(output_path)


def generate_comprehensive_report_v2(all_results, responders, active_vs_placebo, duration_active_vs_placebo, exec_summary, figures, df, output_path):
    print(f"\nGenerating report: {output_path}")

    doc = Document()
    enable_word_update_fields_on_open(doc)
    try:
        doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    except Exception:
        pass
    num = Numbering()
    groups = sorted(df[GROUP_COL].dropna().unique())

    rewrite_items = [
        {
            "id": "s1_1",
            "text": "Пациенты рандомизированы в 4 группы: Г1 (ДИАМАГ 30 дней), Г2 (Плацебо 30 дней), Г3 (ДИАМАГ 20 дней), Г4 (Плацебо 20 дней). "
            "Основное сравнение эффективности терапии проводится как на уровне 4 групп (с учётом длительности), так и в укрупнённом анализе (укрупнение, pooled): Active (Г1+Г3) vs Placebo (Г2+Г4), "
            "который повышает статистическую мощность ценой потери детализации по длительности.",
        },
        {
            "id": "s1_2",
            "text": "В отчёте приводятся p-value и скорректированные p-value (p_adj, коррекция Холма для множественных сравнений), Bayes Factor (BF₁₀) "
            "как оценка силы доказательств в пользу H₁, размеры эффекта (ES) как величина различий, а также клинически интерпретируемые разницы: Δ (абсолютная) и Δ% (процентная), "
            "обычно по медианам. Для респондеров дополнительно приводятся 95% доверительные интервалы и NNT.",
        },
        {
            "id": "s2_intro",
            "text": "Цель раздела — проверить наличие различий между 4 рандомизированными группами на каждом визите и описать их величину (ES).",
        },
        {
            "id": "s3_intro",
            "text": "Цель раздела — проверить основную гипотезу эффективности укрупнённым сравнением Active (Г1+Г3) vs Placebo (Г2+Г4) на каждом визите с оценкой Δ и ES.",
        },
        {
            "id": "s4_intro",
            "text": "В разделе собраны post-hoc сравнения между 4 группами по всем визитам. "
            "Таблица 16 включает только статистически значимые различия после коррекции Холма (p_adj < 0.05). "
            "Если значимых строк немного, это обычно означает, что после учёта множественных сравнений сигнал устойчиво проявился лишь для части показателей, "
            "а для остальных наблюдаются тренды без достижения порога p_adj<0.05.",
        },
        {
            "id": "s5_intro",
            "text": "Модели Mixed Effects оценивают взаимодействие Визит×Группа с учётом повторных измерений у пациентов.",
        },
        {
            "id": "s6_intro",
            "text": "Респондер определяется как пациент с улучшением не менее чем на 20% относительно baseline (V2). Раздел отражает клиническую значимость эффекта по визитам.",
        },
        {
            "id": "s7_intro",
            "text": "Раздел агрегирует результаты по каждой временной точке (V2–V6): глобальный тест 4 групп (Kruskal–Wallis), "
            "наиболее выраженные post-hoc различия (p_adj), объединённый анализ Active vs Placebo (p_adj, ES, Δ) и сравнительные графики размеров эффекта (ε²).",
        },
    ]
    rewrites = generate_ai_rewrites(items=rewrite_items)

    def fallback_rewrite_text(text_value: str) -> str:
        s = str(text_value or "").strip()
        if not s:
            return ""
        re = __import__("re")
        parts = [p.strip() for p in re.split(r"(?<=[.!?])\s+", s) if p and p.strip()]
        if not parts:
            return ""
        out = parts[0].strip()
        out = " ".join(out.split()).replace("…", "").replace("...", "").strip()
        if not out:
            return ""
        if len(out) > 240:
            cut = out[:240]
            if " " in cut:
                cut = cut[: cut.rfind(" ")].rstrip()
            out = cut
        if out and out[-1] not in ".!?":
            out += "."
        return out

    def add_text_with_rewrite(text_id: str, text_value: str):
        doc.add_paragraph(text_value)
        rw = rewrites.get(text_id)
        short = str(rw).strip() if rw else ""
        short = " ".join(short.split()).replace("…", "").replace("...", "").strip()
        if short and short[-1] not in ".!?":
            short += "."
        if short and short != str(text_value or "").strip():
            p = doc.add_paragraph()
            r = p.add_run("Резюме: ")
            r.bold = True
            p.add_run(short)

    def fallback_artifact_explanation(payload: Dict[str, Any]) -> Dict[str, Any]:
        artifact_id = str(payload.get("id") or "").strip()
        label = str(payload.get("label") or payload.get("title") or "").strip()
        kind = str(payload.get("kind") or "").strip() or "artifact"

        title = label or ("Таблица" if kind == "table" else "График" if kind == "chart" else "Материал")
        plain = ""
        takeaways: List[str] = []
        caveats: List[str] = []

        if kind == "table":
            rows = payload.get("rows")
            if not isinstance(rows, list):
                rows = []

            p_candidates: List[tuple] = []
            p_adj_present = False
            for r in rows:
                if not isinstance(r, dict):
                    continue
                for k in [
                    "p_adj",
                    "p",
                    "p_value",
                    "kw_p",
                    "kw_min_p",
                    "min_p",
                    "min_p_adj",
                    "pooled_p_adj",
                    "avp_p_adj",
                ]:
                    v = r.get(k)
                    if v is None or not np.isfinite(v):
                        continue
                    is_adj = "adj" in str(k)
                    if is_adj:
                        p_adj_present = True
                    p_candidates.append((float(v), bool(is_adj)))

            pvals = [p for p, _ in p_candidates]
            sig_count = sum(1 for p in pvals if p < 0.05)
            best = min(p_candidates, key=lambda x: x[0]) if p_candidates else None
            best_p = best[0] if best is not None else None
            best_is_adj = bool(best[1]) if best is not None else False
            if best_p is not None and np.isfinite(best_p):
                key_txt = "p_adj" if best_is_adj else "p"
                p_txt = f"{key_txt} < 0.001" if best_p < 0.001 else f"{key_txt}≈{best_p:.3f}"
            else:
                p_txt = ""

            plain = f"Таблица суммирует результаты: {title}. По ней удобно быстро увидеть, где есть статистические сигналы и насколько они устойчивы после коррекций."
            if p_txt:
                plain += f" Наиболее выраженный статистический сигнал в таблице: {p_txt}."
            if pvals:
                takeaways.append(f"Статистически значимые строки: {sig_count} из {len(pvals)} (по доступным p-значениям).")
            if p_adj_present:
                takeaways.append("Для итоговой интерпретации ориентируйтесь на p_adj (после коррекции множественных сравнений).")
            caveats.append("Оценивайте вместе p, размер эффекта (ES) и Δ, а не только значимость.")

        elif kind == "chart":
            caption = str(payload.get("caption") or "").strip()
            kw_by_visit = payload.get("kw_by_visit")
            sig_visits = []
            if isinstance(kw_by_visit, list):
                for r in kw_by_visit:
                    if not isinstance(r, dict):
                        continue
                    p = r.get("p")
                    if p is not None and np.isfinite(p) and float(p) < 0.05:
                        sig_visits.append(str(r.get("visit") or "").strip())

            if caption:
                plain = f"График иллюстрирует результаты: {caption}"
            else:
                plain = f"График иллюстрирует результаты: {title}."
            if sig_visits:
                takeaways.append("Сигналы различий между группами заметны на визитах: " + ", ".join([v for v in sig_visits if v]) + ".")
            takeaways.append("Подписи под рисунками объясняют критерии, p-value и величину эффекта.")
            caveats.append("Визуальные различия нужно читать вместе со статистикой (p_adj, ES, Δ).")

        else:
            plain = f"Материал помогает интерпретировать результаты: {title}."
            takeaways.append("Смотрите подписи и значения в таблицах/графиках для конкретных выводов.")

        return {
            "id": artifact_id,
            "title": title,
            "plain": plain,
            "takeaways": takeaways,
            "caveats": caveats,
        }

    def add_ai_for_artifact(payload: Dict[str, Any]):
        if not isinstance(payload, dict):
            return
        artifact_id = str(payload.get("id") or "").strip()
        if not artifact_id:
            return
        expl = generate_ai_artifact_explanations(artifacts=[payload])
        got = expl.get(artifact_id) if isinstance(expl, dict) else None
        if not isinstance(got, dict):
            got = fallback_artifact_explanation(payload)
        _doc_add_ai_block(doc, got)

    title = doc.add_paragraph()
    title.add_run("ПОЛНЫЙ АНАЛИТИЧЕСКИЙ ОТЧЁТ ПО ИССЛЕДОВАНИЮ ДИАМАГ").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(22)

    doc.add_paragraph()
    doc.add_paragraph("Фокус: эффективность терапии, динамика ответа, сравнение с плацебо").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading("ОГЛАВЛЕНИЕ", level=1)
    add_toc(doc)

    doc.add_paragraph(
        "Если оглавление не обновилось автоматически: откройте документ в Microsoft Word и нажмите правой кнопкой мыши по оглавлению → «Обновить поле». "
        "Ниже приведён статический фолбэк структуры отчёта."
    )
    doc.add_paragraph("1. Дизайн и методология")
    doc.add_paragraph("2. Сравнение 4 групп")
    doc.add_paragraph("3. Укрупнение: Active (Г1+Г3) vs Placebo (Г2+Г4)")
    doc.add_paragraph("4. Попарные сравнения (post-hoc)")
    doc.add_paragraph("5. Смешанные эффекты (повторные измерения)")
    doc.add_paragraph("6. Анализ респондеров (улучшение ≥20%)")
    doc.add_paragraph("7. Итоговая сводка по визитам")
    doc.add_paragraph("8. Обсуждение")
    doc.add_paragraph("9. Выводы")
    doc.add_page_break()

    _doc_add_heading(doc, "1. ДИЗАЙН И МЕТОДОЛОГИЯ", level=1)
    doc.add_heading("1.1 Группы и сравнения", level=2)
    add_text_with_rewrite(
        "s1_1",
        "Пациенты рандомизированы в 4 группы: Г1 (ДИАМАГ 30 дней), Г2 (Плацебо 30 дней), Г3 (ДИАМАГ 20 дней), Г4 (Плацебо 20 дней). "
        "Основное сравнение эффективности терапии проводится как на уровне 4 групп (с учётом длительности), так и в укрупнённом анализе (укрупнение, pooled): Active (Г1+Г3) vs Placebo (Г2+Г4), "
        "который повышает статистическую мощность ценой потери детализации по длительности.",
    )
    doc.add_heading("1.2 Метрики значимости и интерпретации", level=2)
    add_text_with_rewrite(
        "s1_2",
        "В отчёте приводятся p-value и скорректированные p-value (p_adj, коррекция Холма для множественных сравнений), Bayes Factor (BF₁₀) "
        "как оценка силы доказательств в пользу H₁, размеры эффекта (ES) как величина различий, а также клинически интерпретируемые разницы: Δ (абсолютная) и Δ% (процентная), "
        "обычно по медианам. Для респондеров дополнительно приводятся 95% доверительные интервалы и NNT.",
    )

    doc.add_heading("1.3 Глоссарий ключевых статистических терминов", level=2)
    doc.add_paragraph(num.tab() + ". Глоссарий ключевых статистических терминов")
    table = doc.add_table(rows=1, cols=2, style="Table Grid")
    hdr = table.rows[0].cells
    hdr[0].text = "Термин"
    hdr[1].text = "Определение"
    hdr[0].paragraphs[0].runs[0].bold = True
    hdr[1].paragraphs[0].runs[0].bold = True

    terms = [
        ("H₀ (нулевая гипотеза)", "Гипотеза об отсутствии различий между группами."),
        ("H₁ (альтернативная гипотеза)", "Гипотеза о наличии различий (эффект терапии есть)."),
        ("p-value", "Вероятность получить такие или более экстремальные данные при верной H₀."),
        ("p_adj", "p-value после коррекции на множественные сравнения (Холм)."),
        ("Holm (коррекция Холма)", "Пошаговая коррекция p-value для контроля семейной ошибки I рода при множественных сравнениях."),
        ("Bayes Factor (BF₁₀)", "Отношение правдоподобий H₁ к H₀; BF₁₀>3 — умеренная поддержка H₁; BF₁₀<1/3 — поддержка H₀; 1/3–3 — нейтрально."),
        ("ES (размер эффекта)", "Величина различий, не сводимая к значимости: r≈0.1 малый, 0.3 средний, 0.5 большой; ε²≈0.01 малый, 0.06 средний, 0.14 большой."),
        ("ε² (epsilon-squared)", "Размер эффекта для Kruskal–Wallis (4 группы): доля вариативности, объясняемая группой."),
        ("r (rank-biserial)", "Размер эффекта для Mann–Whitney/Wilcoxon; знак отражает направление различий, модуль — величину."),
        ("Δ (дельта)", "Абсолютная разница между группами; в pooled-анализе рассчитывается как Med(Active) − Med(Placebo), знак интерпретируется с учётом направления шкалы."),
        ("Δ (%)", "Процентная разница, обычно относительно Med(Placebo): 100%·Δ/Med(Placebo)."),
        ("Медиана (Med)", "Центральное значение распределения, устойчивое к выбросам; обычно интерпретируется вместе с IQR (Q1–Q3)."),
        ("Доверительный интервал (95% ДИ)", "Диапазон значений параметра, совместимых с данными; для долей респондеров используется интервал Вилсона."),
        ("NNT", "Number Needed to Treat: сколько пациентов нужно лечить, чтобы получить 1 дополнительного респондера по сравнению с плацебо."),
        ("Укрупнение (pooled)", "Сведение 4 групп к 2: Active (Г1+Г3) vs Placebo (Г2+Г4) для повышения мощности; теряется детализация по длительности."),
        ("Kruskal–Wallis", "Непараметрический тест сравнения 3+ групп."),
        ("Mann–Whitney U", "Непараметрический тест сравнения 2 независимых групп."),
        ("Wilcoxon", "Непараметрический парный тест изменений относительно baseline."),
        ("Mixed Effects", "Модель повторных измерений (Визит×Группа) с учётом индивидуальных различий (рандом-эффект пациента)."),
    ]

    for term, desc in terms:
        row = table.add_row().cells
        row[0].text = term
        row[1].text = desc

    glossary_payload = {
        "id": "s1_glossary",
        "kind": "table",
        "label": "Глоссарий ключевых терминов",
        "terms": [t for t, _ in terms],
    }
    add_ai_for_artifact(glossary_payload)

    doc.add_page_break()

    _doc_add_heading(doc, "2. СРАВНЕНИЕ 4 ГРУПП", level=1)
    add_text_with_rewrite(
        "s2_intro",
        "Цель раздела — проверить наличие различий между 4 рандомизированными группами на каждом визите и описать их величину (ES).",
    )

    for idx, (key, result) in enumerate(all_results.items(), start=1):
        doc.add_heading(f"2.{idx} {result.get('name', key)}", level=2)

        doc.add_paragraph(num.tab() + f". Описательная статистика по визитам: {result.get('short', key)}")
        doc.add_paragraph("Показаны медиана и межквартильный размах (Q1–Q3) с указанием n.")
        desc_table = doc.add_table(rows=1, cols=5, style="Table Grid")
        dh = desc_table.rows[0].cells
        dh[0].text = "Визит"
        dh[0].paragraphs[0].runs[0].bold = True
        group_order = [g for g in ["1", "2", "3", "4"] if g in groups] + [str(g) for g in groups if str(g) not in {"1", "2", "3", "4"}]
        while len(group_order) < 4:
            group_order.append("")
        for j, g in enumerate(group_order[:4], start=1):
            dh[j].text = f"Г{g}" if g else "—"
            if dh[j].paragraphs and dh[j].paragraphs[0].runs:
                dh[j].paragraphs[0].runs[0].bold = True

        cfg = ENDPOINTS.get(key, {})
        for v in VISITS:
            row = desc_table.add_row().cells
            row[0].text = v
            col = (cfg.get("cols") or {}).get(v)
            for j, g in enumerate(group_order[:4], start=1):
                if not col or not g or col not in df.columns:
                    row[j].text = "—"
                    continue
                s = df[df[GROUP_COL].astype(str) == str(g)][col].dropna()
                if len(s) == 0:
                    row[j].text = "—"
                    continue
                med = float(s.median())
                q1 = float(s.quantile(0.25))
                q3 = float(s.quantile(0.75))
                row[j].text = f"n={len(s)}; {med:.1f} [{q1:.1f}; {q3:.1f}]"

        doc.add_paragraph(num.tab() + f". Изменение от baseline (Δ к V2) по визитам: {result.get('short', key)}")
        direction = str(cfg.get("direction") or "").strip()
        if direction == "lower_is_better":
            doc.add_paragraph("Для данного показателя снижение значения соответствует улучшению, поэтому отрицательная Δ означает улучшение относительно baseline.")
        elif direction == "higher_is_better":
            doc.add_paragraph("Для данного показателя повышение значения соответствует улучшению, поэтому положительная Δ означает улучшение относительно baseline.")
        doc.add_paragraph("Показаны Med(Δ) и IQR (Q1–Q3) по пациентам с доступными значениями на V2 и соответствующем визите; n может отличаться от n в уровнях показателя.")
        delta_table = doc.add_table(rows=1, cols=5, style="Table Grid")
        dth = delta_table.rows[0].cells
        dth[0].text = "Визит"
        dth[0].paragraphs[0].runs[0].bold = True
        for j, g in enumerate(group_order[:4], start=1):
            dth[j].text = f"Г{g}" if g else "—"
            if dth[j].paragraphs and dth[j].paragraphs[0].runs:
                dth[j].paragraphs[0].runs[0].bold = True

        baseline_col = (cfg.get("cols") or {}).get("V2")
        for v in VISITS[1:]:
            row = delta_table.add_row().cells
            row[0].text = v
            visit_col = (cfg.get("cols") or {}).get(v)
            for j, g in enumerate(group_order[:4], start=1):
                if not baseline_col or not visit_col or not g or baseline_col not in df.columns or visit_col not in df.columns:
                    row[j].text = "—"
                    continue
                gdf = df[df[GROUP_COL].astype(str) == str(g)][[baseline_col, visit_col]].dropna()
                if len(gdf) == 0:
                    row[j].text = "—"
                    continue
                delta = (gdf[visit_col] - gdf[baseline_col]).astype(float)
                med = float(delta.median())
                q1 = float(delta.quantile(0.25))
                q3 = float(delta.quantile(0.75))
                row[j].text = f"n={len(delta)}; {med:+.1f} [{q1:+.1f}; {q3:+.1f}]"

        doc.add_paragraph(num.tab() + f". Глобальный тест 4 групп по визитам: {result.get('short', key)}")
        table = doc.add_table(rows=1, cols=6, style="Table Grid")
        hdr = table.rows[0].cells
        for i, h in enumerate(["Визит", "H", "p", "BF₁₀", "ε²", "Вывод"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        kw_rows = []
        for v in VISITS:
            vdata = result.get("by_visit", {}).get(v, {})
            kw = vdata.get("kruskal", {})
            bf = vdata.get("bf10", np.nan)
            row = table.add_row().cells
            row[0].text = v
            row[1].text = f"{float(kw.get('H')):.2f}" if kw.get("H") is not None and np.isfinite(kw.get("H")) else "—"
            p = kw.get("p")
            row[2].text = f"{float(p):.4f}" if p is not None and np.isfinite(p) else "—"
            row[3].text = f"{float(bf):.2f}" if bf is not None and np.isfinite(bf) else "—"
            es = kw.get("epsilon_sq")
            row[4].text = f"{float(es):.3f}" if es is not None and np.isfinite(es) else "—"
            row[5].text = "значимо" if kw.get("significant") else "незначимо"

            kw_rows.append(
                {
                    "visit": v,
                    "H": kw.get("H"),
                    "p": p,
                    "bf10": bf,
                    "epsilon_sq": es,
                    "significant": bool(kw.get("significant")),
                }
            )

        doc.add_paragraph(
            "Интерпретация: p<0.05 означает различия между группами на данном визите; ε² показывает их величину; BF₁₀ отражает силу доказательств в пользу H₁."
        )

        add_ai_for_artifact(
            {
                "id": f"s2_{key}_kw_table",
                "kind": "table",
                "label": f"Глобальный тест 4 групп по визитам ({result.get('short', key)})",
                "rows": kw_rows,
            }
        )

        fig_key = f"{key}_spaghetti"
        if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
            doc.add_picture(figures[fig_key], width=Inches(6))
            caption = generate_figure_caption_text(all_results, key, visit_key=None, active_vs_placebo=None)
            fig_label = num.fig()
            p = doc.add_paragraph(fig_label + ". " + caption)
            p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None

            add_ai_for_artifact(
                {
                    "id": f"s2_{key}_spaghetti_fig",
                    "kind": "chart",
                    "label": f"{fig_label}. {result.get('short', key)} — динамика по визитам (4 группы)",
                    "caption": caption,
                    "kw_by_visit": kw_rows,
                }
            )

        doc.add_heading("Boxplot по визитам (4 группы)", level=3)

        boxplot_rows = []
        for v in VISITS:
            fig_key = f"{key}_boxplot_{v}"
            if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
                doc.add_picture(figures[fig_key], width=Inches(6))
                caption = generate_figure_caption_text(all_results, key, visit_key=v, active_vs_placebo=None)
                caption = "В контексте цели оценки эффективности терапии сравнение групп позволяет выявить, существует ли эффект вмешательства. " + caption
                fig_label = num.fig()
                p = doc.add_paragraph(fig_label + ". " + caption)
                p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None

                boxplot_rows.append({"visit": v, "caption": caption})

        if boxplot_rows:
            add_ai_for_artifact(
                {
                    "id": f"s2_{key}_boxplots_figset",
                    "kind": "chart",
                    "label": f"Boxplot по визитам (4 группы): {result.get('short', key)}",
                    "by_visit": boxplot_rows,
                    "kw_by_visit": kw_rows,
                }
            )

        doc.add_page_break()

    _doc_add_heading(doc, "3. УКРУПНЕНИЕ: ACTIVE (Г1+Г3) VS PLACEBO (Г2+Г4)", level=1)
    add_text_with_rewrite(
        "s3_intro",
        "Цель раздела — проверить основную гипотезу эффективности укрупнённым сравнением Active (Г1+Г3) vs Placebo (Г2+Г4) на каждом визите с оценкой Δ и ES.",
    )

    doc.add_heading("3.1 Сводка выигрыша от укрупнения", level=2)
    doc.add_paragraph(
        "Укрупнение обычно повышает мощность за счёт уменьшения числа сравниваемых уровней, но не отвечает на вопрос различий между длительностями (20 vs 30 дней)."
    )
    doc.add_paragraph(num.tab() + ". Сравнение чувствительности 4-группового и укрупнённого анализа")
    win_table = doc.add_table(rows=1, cols=5, style="Table Grid")
    wh = win_table.rows[0].cells
    for i, h in enumerate(["Показатель", "4 группы: значимые визиты", "4 группы: min p", "Укрупнение: значимые визиты", "Укрупнение: min p_adj"]):
        wh[i].text = h
        wh[i].paragraphs[0].runs[0].bold = True

    win_rows = []
    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)
        kw_sig = []
        kw_ps = []
        for v in VISITS:
            kw = (result.get("by_visit", {}).get(v, {}) or {}).get("kruskal", {})
            p = kw.get("p")
            if p is not None and np.isfinite(p):
                kw_ps.append(float(p))
                if float(p) < 0.05:
                    kw_sig.append(v)

        avp_sig = []
        avp_ps = []
        avp_visits = (active_vs_placebo.get(ep_key, {}) or {}).get("visits", {})
        if isinstance(avp_visits, dict):
            for v in VISITS:
                d = avp_visits.get(v, {}) if isinstance(avp_visits.get(v, {}), dict) else {}
                p = d.get("p_adj", d.get("p_value", np.nan))
                if p is not None and np.isfinite(p):
                    avp_ps.append(float(p))
                    if float(p) < 0.05:
                        avp_sig.append(v)

        row = win_table.add_row().cells
        row[0].text = short
        row[1].text = ", ".join(kw_sig) if kw_sig else "—"
        row[2].text = f"{min(kw_ps):.4f}" if kw_ps else "—"
        row[3].text = ", ".join(avp_sig) if avp_sig else "—"
        row[4].text = f"{min(avp_ps):.4f}" if avp_ps else "—"

        win_rows.append(
            {
                "endpoint": short,
                "kw_significant_visits": kw_sig,
                "kw_min_p": (min(kw_ps) if kw_ps else None),
                "pooled_significant_visits": avp_sig,
                "pooled_min_p_adj": (min(avp_ps) if avp_ps else None),
            }
        )

    add_ai_for_artifact(
        {
            "id": "s3_win_table",
            "kind": "table",
            "label": "Сравнение чувствительности 4-группового и укрупнённого анализа",
            "rows": win_rows,
        }
    )

    doc.add_page_break()

    pooled_idx = 1
    for key, res in active_vs_placebo.items():
        doc.add_heading(f"3.{pooled_idx} Показатель: {res.get('short', key)}", level=2)

        fig_key = f"{key}_pooled_spaghetti"
        if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
            doc.add_picture(figures[fig_key], width=Inches(6))
            fig_label = num.fig()
            p = doc.add_paragraph(fig_label + f". Динамика {res.get('short', key)} (Active vs Placebo); подписи p= приведены только для визитов со значимыми различиями.")
            p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None

            add_ai_for_artifact(
                {
                    "id": f"s3_{key}_pooled_spaghetti_fig",
                    "kind": "chart",
                    "label": f"{fig_label}. {res.get('short', key)} — динамика Active vs Placebo",
                    "by_visit": [
                        {
                            "visit": v,
                            "p_adj": (res.get("visits", {}).get(v, {}) or {}).get("p_adj"),
                            "effect_size": (res.get("visits", {}).get(v, {}) or {}).get("effect_size"),
                            "diff_median": (res.get("visits", {}).get(v, {}) or {}).get("diff_median"),
                            "diff_pct": (res.get("visits", {}).get(v, {}) or {}).get("diff_pct"),
                            "significant": bool((res.get("visits", {}).get(v, {}) or {}).get("significant")),
                        }
                        for v in VISITS
                    ],
                }
            )

        doc.add_paragraph(num.tab() + f". Active vs Placebo по визитам: {res.get('short', key)}")
        table = doc.add_table(rows=1, cols=9, style="Table Grid")
        hdr = table.rows[0].cells
        for i, h in enumerate(["Визит", "Active (Med)", "Placebo (Med)", "Δ (Abs)", "Δ (%)", "p_adj", "BF₁₀", "ES", "Вывод"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        pooled_rows = []
        visits_sorted = sorted(res.get("visits", {}).keys(), key=lambda v: VISITS.index(v) if v in VISITS else 99)
        for v in visits_sorted:
            d = res["visits"][v]
            row = table.add_row().cells
            row[0].text = v
            row[1].text = f"{float(d.get('active_median', np.nan)):.1f}" if np.isfinite(d.get('active_median', np.nan)) else "—"
            row[2].text = f"{float(d.get('placebo_median', np.nan)):.1f}" if np.isfinite(d.get('placebo_median', np.nan)) else "—"
            diff = d.get('diff_median', np.nan)
            row[3].text = f"{float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "—"
            diff_pct = d.get('diff_pct', np.nan)
            row[4].text = f"{float(diff_pct):+.1f}%" if diff_pct is not None and np.isfinite(diff_pct) else "—"
            p_val = d.get('p_adj', d.get('p_value', np.nan))
            row[5].text = f"{float(p_val):.4f}" if p_val is not None and np.isfinite(p_val) else "—"
            bf = d.get('bf10', np.nan)
            row[6].text = f"{float(bf):.2f}" if bf is not None and np.isfinite(bf) else "—"
            es = d.get('effect_size', np.nan)
            row[7].text = f"{float(es):.2f}" if es is not None and np.isfinite(es) else "—"
            row[8].text = "значимо" if d.get('significant') else "незначимо"

            pooled_rows.append(
                {
                    "visit": v,
                    "active_median": d.get("active_median"),
                    "placebo_median": d.get("placebo_median"),
                    "diff_median": diff,
                    "diff_pct": diff_pct,
                    "p_adj": p_val,
                    "bf10": bf,
                    "effect_size": es,
                    "significant": bool(d.get("significant")),
                }
            )

        add_ai_for_artifact(
            {
                "id": f"s3_{key}_avp_table",
                "kind": "table",
                "label": f"Active vs Placebo по визитам ({res.get('short', key)})",
                "rows": pooled_rows,
            }
        )

        doc.add_heading("Boxplot по визитам (Active vs Placebo)", level=3)
        pooled_box_rows = []
        for v in VISITS:
            fig_key = f"{key}_pooled_boxplot_{v}"
            if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
                doc.add_picture(figures[fig_key], width=Inches(6))
                vstats = res.get("visits", {}).get(v, {})
                p_used = vstats.get("p_adj", vstats.get("p_value", np.nan))
                p_txt = "p < 0.001" if p_used is not None and np.isfinite(p_used) and float(p_used) < 0.001 else (f"p = {float(p_used):.3f}" if p_used is not None and np.isfinite(p_used) else "p = —")
                diff = vstats.get("diff_median", np.nan)
                diff_pct = vstats.get("diff_pct", np.nan)
                es = vstats.get("effect_size", np.nan)
                diff_txt = f"Δ={float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "Δ=—"
                diffp_txt = f" ({float(diff_pct):+.1f}%)" if diff_pct is not None and np.isfinite(diff_pct) else ""
                es_txt = f"ES={float(es):.2f}" if es is not None and np.isfinite(es) else "ES=—"
                verdict = "значимо" if vstats.get("significant") else "незначимо"
                caption = f"В контексте основной цели исследования сравнение Active vs Placebo на визите {v} даёт: {p_txt} ({verdict}), {diff_txt}{diffp_txt}, {es_txt}."
                fig_label = num.fig()
                p = doc.add_paragraph(fig_label + ". " + caption)
                p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None

                pooled_box_rows.append({"visit": v, "caption": caption})

        if pooled_box_rows:
            add_ai_for_artifact(
                {
                    "id": f"s3_{key}_pooled_boxplots_figset",
                    "kind": "chart",
                    "label": f"Boxplot по визитам (Active vs Placebo): {res.get('short', key)}",
                    "by_visit": pooled_box_rows,
                    "by_visit_stats": [
                        {
                            "visit": v,
                            "p_adj": (res.get("visits", {}).get(v, {}) or {}).get("p_adj"),
                            "effect_size": (res.get("visits", {}).get(v, {}) or {}).get("effect_size"),
                            "diff_median": (res.get("visits", {}).get(v, {}) or {}).get("diff_median"),
                            "diff_pct": (res.get("visits", {}).get(v, {}) or {}).get("diff_pct"),
                            "significant": bool((res.get("visits", {}).get(v, {}) or {}).get("significant")),
                        }
                        for v in VISITS
                    ],
                }
            )

        pooled_idx += 1
        doc.add_page_break()

    _doc_add_heading(doc, "4. ПОПАРНЫЕ СРАВНЕНИЯ (POST-HOC)", level=1)
    add_text_with_rewrite(
        "s4_intro",
        "В разделе собраны post-hoc сравнения между 4 группами по всем визитам. "
        "Таблица 16 включает только статистически значимые различия после коррекции Холма (p_adj < 0.05). "
        "Если значимых строк немного, это обычно означает, что после учёта множественных сравнений сигнал устойчиво проявился лишь для части показателей, "
        "а для остальных наблюдаются тренды без достижения порога p_adj<0.05.",
    )

    doc.add_paragraph(num.tab() + ". Значимые попарные сравнения (4 группы, p_adj<0.05)")
    table = doc.add_table(rows=1, cols=9, style="Table Grid")
    hdr = table.rows[0].cells
    for i, h in enumerate(["Показатель", "Визит", "Пара", "p_adj", "BF₁₀", "Δ (Abs)", "Δ (%)", "ES (r)", "Вывод"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    any_rows = False
    sig_by_endpoint: Dict[str, int] = {}
    sig_payload_rows = []
    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)
        sig_by_endpoint[short] = 0
        for v in VISITS:
            pv = result.get("pairwise", {}).get(v, {})
            if not isinstance(pv, dict):
                continue
            for pair_key, st in pv.items():
                if not isinstance(st, dict) or st.get("error"):
                    continue
                p_adj = st.get("p_adj")
                if p_adj is None or not np.isfinite(p_adj) or float(p_adj) >= 0.05:
                    continue
                any_rows = True
                sig_by_endpoint[short] = int(sig_by_endpoint.get(short, 0)) + 1
                sig_payload_rows.append(
                    {
                        "endpoint": short,
                        "visit": v,
                        "pair": str(pair_key).replace("_vs_", " vs "),
                        "p_adj": float(p_adj),
                        "bf10": st.get("bf10"),
                        "diff_median": st.get("diff_median"),
                        "diff_pct": st.get("diff_pct"),
                        "effect_size_r": st.get("r"),
                    }
                )
                row = table.add_row().cells
                row[0].text = short
                row[1].text = v
                row[2].text = str(pair_key).replace("_vs_", " vs ")
                row[3].text = f"{float(p_adj):.4f}"
                bf = st.get("bf10", np.nan)
                row[4].text = f"{float(bf):.2f}" if bf is not None and np.isfinite(bf) else "—"
                diff = st.get("diff_median", np.nan)
                diff_pct = st.get("diff_pct", np.nan)
                row[5].text = f"{float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "—"
                row[6].text = f"{float(diff_pct):+.1f}%" if diff_pct is not None and np.isfinite(diff_pct) else "—"
                es = st.get("r", np.nan)
                row[7].text = f"{float(es):.2f}" if es is not None and np.isfinite(es) else "—"
                row[8].text = "значимо"

    if not any_rows:
        doc.add_paragraph("Статистически значимых post-hoc различий (p_adj<0.05) не выявлено.")
    else:
        kept = [(k, v) for k, v in sig_by_endpoint.items() if int(v) > 0]
        kept_sorted = sorted(kept, key=lambda x: x[1], reverse=True)
        summary_txt = "; ".join([f"{k}: {v}" for k, v in kept_sorted])
        doc.add_paragraph(
            "Пояснение по охвату показателей: значимые post-hoc после коррекции Холма обнаружены для следующих показателей (число значимых сравнений): "
            + summary_txt
            + "."
        )

    if sig_payload_rows:
        sig_payload_rows_sorted = sorted(sig_payload_rows, key=lambda r: float(r.get("p_adj", 1e9)))
        add_ai_for_artifact(
            {
                "id": "s4_sig_posthoc_table",
                "kind": "table",
                "label": "Значимые попарные сравнения (4 группы, p_adj<0.05)",
                "rows": sig_payload_rows_sorted[:40],
                "total_rows": len(sig_payload_rows_sorted),
            }
        )

    doc.add_paragraph(num.tab() + ". Наиболее выраженные post-hoc различия (топ-3 по p_adj на визит для каждого показателя)")
    table2 = doc.add_table(rows=1, cols=10, style="Table Grid")
    hdr2 = table2.rows[0].cells
    for i, h in enumerate(["Показатель", "Визит", "Пара", "p", "p_adj", "BF₁₀", "Δ (Abs)", "Δ (%)", "ES (r)", "Вывод"]):
        hdr2[i].text = h
        hdr2[i].paragraphs[0].runs[0].bold = True

    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)
        for v in VISITS:
            pv = result.get("pairwise", {}).get(v, {})
            if not isinstance(pv, dict) or not pv:
                continue
            rows = []
            for pair_key, st in pv.items():
                if not isinstance(st, dict) or st.get("error"):
                    continue
                p_adj = st.get("p_adj")
                if p_adj is None or not np.isfinite(p_adj):
                    continue
                rows.append((float(p_adj), pair_key, st))
            rows.sort(key=lambda x: x[0])
            for p_adj, pair_key, st in rows[:3]:
                row = table2.add_row().cells
                row[0].text = short
                row[1].text = v
                row[2].text = str(pair_key).replace("_vs_", " vs ")
                p_raw = st.get("p")
                row[3].text = f"{float(p_raw):.4f}" if p_raw is not None and np.isfinite(p_raw) else "—"
                row[4].text = f"{float(p_adj):.4f}"
                bf = st.get("bf10", np.nan)
                row[5].text = f"{float(bf):.2f}" if bf is not None and np.isfinite(bf) else "—"
                diff = st.get("diff_median", np.nan)
                diff_pct = st.get("diff_pct", np.nan)
                row[6].text = f"{float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "—"
                row[7].text = f"{float(diff_pct):+.1f}%" if diff_pct is not None and np.isfinite(diff_pct) else "—"
                es = st.get("r", np.nan)
                row[8].text = f"{float(es):.2f}" if es is not None and np.isfinite(es) else "—"
                row[9].text = "значимо" if (np.isfinite(p_adj) and float(p_adj) < 0.05) else "незначимо"

    top3_payload_rows = []
    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)
        for v in VISITS:
            pv = result.get("pairwise", {}).get(v, {})
            if not isinstance(pv, dict) or not pv:
                continue
            tmp = []
            for pair_key, st in pv.items():
                if not isinstance(st, dict) or st.get("error"):
                    continue
                p_adj = st.get("p_adj")
                if p_adj is None or not np.isfinite(p_adj):
                    continue
                tmp.append((float(p_adj), pair_key, st))
            tmp.sort(key=lambda x: x[0])
            for p_adj, pair_key, st in tmp[:3]:
                top3_payload_rows.append(
                    {
                        "endpoint": short,
                        "visit": v,
                        "pair": str(pair_key).replace("_vs_", " vs "),
                        "p": st.get("p"),
                        "p_adj": float(p_adj),
                        "bf10": st.get("bf10"),
                        "diff_median": st.get("diff_median"),
                        "diff_pct": st.get("diff_pct"),
                        "effect_size_r": st.get("r"),
                        "significant": bool(np.isfinite(p_adj) and float(p_adj) < 0.05),
                    }
                )

    if top3_payload_rows:
        top3_payload_rows_sorted = sorted(top3_payload_rows, key=lambda r: float(r.get("p_adj", 1e9)))
        add_ai_for_artifact(
            {
                "id": "s4_top3_posthoc_table",
                "kind": "table",
                "label": "Наиболее выраженные post-hoc различия (топ-3 по p_adj на визит)",
                "rows": top3_payload_rows_sorted[:40],
                "total_rows": len(top3_payload_rows_sorted),
            }
        )

    doc.add_page_break()

    _doc_add_heading(doc, "5. СМЕШАННЫЕ ЭФФЕКТЫ (ПОВТОРНЫЕ ИЗМЕРЕНИЯ)", level=1)
    add_text_with_rewrite(
        "s5_intro",
        "Модели Mixed Effects оценивают взаимодействие Визит×Группа с учётом повторных измерений у пациентов.",
    )

    doc.add_paragraph(num.tab() + ". Mixed Effects: взаимодействие Визит×Группа (4 группы и укрупнение)")
    table = doc.add_table(rows=1, cols=7, style="Table Grid")
    hdr = table.rows[0].cells
    for i, h in enumerate(["Показатель", "Модель", "N (obs)", "N (subjects)", "p (interaction)", "Значимость", "Интерпретация"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    mixed_rows = []
    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)
        for label, mm in [("4 группы", result.get("mixed_effects_4g", {})), ("Укрупнение (Active vs Placebo)", result.get("mixed_effects_pooled", {}))]:
            row = table.add_row().cells
            row[0].text = short
            row[1].text = label
            row[2].text = str(mm.get("n_observations", "—")) if isinstance(mm, dict) and "error" not in mm else "—"
            row[3].text = str(mm.get("n_subjects", "—")) if isinstance(mm, dict) and "error" not in mm else "—"
            p_int = mm.get("interaction_p_value") if isinstance(mm, dict) else None
            row[4].text = f"{float(p_int):.4f}" if p_int is not None and np.isfinite(p_int) else "—"
            sig = mm.get("interaction", {}).get("significant") if isinstance(mm, dict) else None
            row[5].text = "значимо" if sig is True else ("незначимо" if sig is False else "—")
            interp = mm.get("interaction", {}).get("interpretation") if isinstance(mm, dict) else None
            row[6].text = str(interp) if interp else (str(mm.get("error")) if isinstance(mm, dict) and mm.get("error") else "—")

            mixed_rows.append(
                {
                    "endpoint": short,
                    "model": label,
                    "n_observations": (mm.get("n_observations") if isinstance(mm, dict) else None),
                    "n_subjects": (mm.get("n_subjects") if isinstance(mm, dict) else None),
                    "interaction_p": p_int,
                    "interaction_significant": sig,
                    "interpretation": interp,
                    "error": (mm.get("error") if isinstance(mm, dict) else None),
                }
            )

    add_ai_for_artifact(
        {
            "id": "s5_mixed_overview_table",
            "kind": "table",
            "label": "Mixed Effects: взаимодействие Визит×Группа (4 группы и укрупнение)",
            "rows": mixed_rows,
        }
    )

    doc.add_paragraph(
        "Интерпретация: значимое взаимодействие Визит×Группа означает, что динамика показателя по визитам различается между группами. "
        "Это указывает на различия траекторий во времени, но само по себе не определяет, на каких визитах и между какими группами различия наиболее выражены; "
        "для этого приводится follow-up анализ."
    )
    any_mixed_sig = False
    detail_idx = 0
    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)
        mm4 = result.get("mixed_effects_4g", {})
        mmp = result.get("mixed_effects_pooled", {})
        sig4 = isinstance(mm4, dict) and (mm4.get("interaction", {}) or {}).get("significant") is True
        sigp = isinstance(mmp, dict) and (mmp.get("interaction", {}) or {}).get("significant") is True
        if not (sig4 or sigp):
            continue
        any_mixed_sig = True
        detail_idx += 1
        doc.add_heading(f"5.{detail_idx} Показатель: {short}", level=2)

        if sig4:
            p_int = mm4.get("interaction_p_value")
            p_txt = f"{float(p_int):.4f}" if p_int is not None and np.isfinite(p_int) else "—"
            interp = (mm4.get("interaction", {}) or {}).get("interpretation")
            doc.add_paragraph(f"4 группы: взаимодействие Визит×Группа: p={p_txt}. {interp}")
            direction = str((ENDPOINTS.get(ep_key, {}) or {}).get("direction") or "").strip()
            if direction == "lower_is_better":
                doc.add_paragraph("Направление шкалы: ниже — лучше; снижение по визитам интерпретируется как улучшение.")
            elif direction == "higher_is_better":
                doc.add_paragraph("Направление шкалы: выше — лучше; повышение по визитам интерпретируется как улучшение.")
            doc.add_paragraph(
                "Практический смысл: динамика по визитам неодинакова между группами; ожидается, что в отдельных временных точках будут различия уровней или темпов изменения. "
                "Если при этом post-hoc сравнения на отдельных визитах не достигают значимости после коррекции, это обычно означает, что эффект распределён по времени и/или недостаточно крупен для выявления на уровне отдельных визитов при множественных сравнениях."
            )

            doc.add_paragraph(num.tab() + f". Follow-up (4 группы): post-hoc сравнения для {short} (после коррекции Холма)")
            follow4_rows = []
            for v in VISITS:
                pv = result.get("pairwise", {}).get(v, {})
                if not isinstance(pv, dict):
                    continue
                for pair_key, st in pv.items():
                    if not isinstance(st, dict) or st.get("error"):
                        continue
                    p_adj = st.get("p_adj")
                    if p_adj is None or not np.isfinite(p_adj) or float(p_adj) >= 0.05:
                        continue
                    follow4_rows.append(
                        {
                            "visit": v,
                            "pair": str(pair_key).replace("_vs_", " vs "),
                            "p_adj": float(p_adj),
                            "bf10": st.get("bf10"),
                            "effect_size_r": st.get("r"),
                            "diff_median": st.get("diff_median"),
                            "diff_pct": st.get("diff_pct"),
                        }
                    )
            if not follow4_rows:
                doc.add_paragraph(
                    "Значимых post-hoc различий на отдельных визитах после коррекции Холма не выявлено. "
                    "В этом случае значимое взаимодействие Визит×Группа следует интерпретировать как различие траекторий во времени без устойчивых парных различий в отдельных временных точках при текущей мощности исследования."
                )
            else:
                t4 = doc.add_table(rows=1, cols=7, style="Table Grid")
                h4 = t4.rows[0].cells
                for i, h in enumerate(["Визит", "Пара", "p_adj", "BF₁₀", "ES (r)", "Δ (Med1−Med2)", "Δ (%)"]):
                    h4[i].text = h
                    h4[i].paragraphs[0].runs[0].bold = True
                for r in follow4_rows:
                    row = t4.add_row().cells
                    row[0].text = str(r.get("visit") or "—")
                    row[1].text = str(r.get("pair") or "—")
                    p_adj = r.get("p_adj")
                    row[2].text = f"{float(p_adj):.4f}" if p_adj is not None and np.isfinite(p_adj) else "—"
                    bf = r.get("bf10", np.nan)
                    row[3].text = f"{float(bf):.2f}" if bf is not None and np.isfinite(bf) else "—"
                    es = r.get("effect_size_r", np.nan)
                    row[4].text = f"{float(es):.2f}" if es is not None and np.isfinite(es) else "—"
                    diff = r.get("diff_median", np.nan)
                    row[5].text = f"{float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "—"
                    diff_pct = r.get("diff_pct", np.nan)
                    row[6].text = f"{float(diff_pct):+.1f}%" if diff_pct is not None and np.isfinite(diff_pct) else "—"
                add_ai_for_artifact(
                    {
                        "id": f"s5_{ep_key}_followup_4g_table",
                        "kind": "table",
                        "label": f"Follow-up (4 группы): значимые post-hoc сравнения для {short}",
                        "rows": follow4_rows,
                    }
                )

        if sigp:
            p_int = mmp.get("interaction_p_value")
            p_txt = f"{float(p_int):.4f}" if p_int is not None and np.isfinite(p_int) else "—"
            interp = (mmp.get("interaction", {}) or {}).get("interpretation")
            doc.add_paragraph(f"Укрупнение (Active vs Placebo): взаимодействие Визит×Группа: p={p_txt}. {interp}")
            direction = str((ENDPOINTS.get(ep_key, {}) or {}).get("direction") or "").strip()
            if direction == "lower_is_better":
                doc.add_paragraph("Интерпретация Δ: Δ=Med(Active)−Med(Placebo); отрицательная Δ означает преимущество Active (ниже — лучше).")
            elif direction == "higher_is_better":
                doc.add_paragraph("Интерпретация Δ: Δ=Med(Active)−Med(Placebo); положительная Δ означает преимущество Active (выше — лучше).")
            doc.add_paragraph(
                "Практический смысл: различие между Active и Placebo изменяется во времени. Для интерпретации приводятся оценки по визитам (p_adj, Δ и ES), "
                "которые показывают, на каких временных точках наблюдается наибольший разрыв и является ли он статистически устойчивым после коррекции."
            )

            doc.add_paragraph(num.tab() + f". Follow-up (укрупнение): контрасты Active vs Placebo по визитам для {short}")
            tp = doc.add_table(rows=1, cols=7, style="Table Grid")
            hp = tp.rows[0].cells
            for i, h in enumerate(["Визит", "p_adj", "BF₁₀", "Δ (Abs)", "Δ (%)", "ES", "Вывод"]):
                hp[i].text = h
                hp[i].paragraphs[0].runs[0].bold = True
            avp_visits = (active_vs_placebo.get(ep_key, {}) or {}).get("visits", {})
            followp_rows = []
            for v in VISITS:
                d = avp_visits.get(v, {}) if isinstance(avp_visits, dict) else {}
                row = tp.add_row().cells
                row[0].text = v
                p_val = d.get("p_adj", d.get("p_value", np.nan))
                row[1].text = f"{float(p_val):.4f}" if p_val is not None and np.isfinite(p_val) else "—"
                bf = d.get("bf10", np.nan)
                row[2].text = f"{float(bf):.2f}" if bf is not None and np.isfinite(bf) else "—"
                diff = d.get("diff_median", np.nan)
                diff_pct = d.get("diff_pct", np.nan)
                row[3].text = f"{float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "—"
                row[4].text = f"{float(diff_pct):+.1f}%" if diff_pct is not None and np.isfinite(diff_pct) else "—"
                es = d.get("effect_size", np.nan)
                row[5].text = f"{float(es):.2f}" if es is not None and np.isfinite(es) else "—"
                row[6].text = "значимо" if d.get("significant") else "незначимо"

                followp_rows.append(
                    {
                        "visit": v,
                        "p_adj": p_val,
                        "bf10": bf,
                        "diff_median": diff,
                        "diff_pct": diff_pct,
                        "effect_size": es,
                        "significant": bool(d.get("significant")),
                    }
                )

            add_ai_for_artifact(
                {
                    "id": f"s5_{ep_key}_followup_pooled_table",
                    "kind": "table",
                    "label": f"Follow-up (укрупнение): контрасты Active vs Placebo по визитам для {short}",
                    "rows": followp_rows,
                }
            )

    if not any_mixed_sig:
        doc.add_paragraph("Значимых взаимодействий Визит×Группа по Mixed Effects моделям не выявлено.")

    doc.add_page_break()

    _doc_add_heading(doc, "6. АНАЛИЗ РЕСПОНДЕРОВ (УЛУЧШЕНИЕ ≥20%)", level=1)
    add_text_with_rewrite(
        "s6_intro",
        "Респондер определяется как пациент с улучшением не менее чем на 20% относительно baseline (V2). Раздел отражает клиническую значимость эффекта по визитам.",
    )

    for idx, (ep_key, resp) in enumerate(responders.items(), start=1):
        short = resp.get("short", ep_key) if isinstance(resp, dict) else ep_key
        doc.add_heading(f"6.{idx} Показатель: {short}", level=2)

        resp_fig_rows = []

        for v in VISITS:
            chart_key = f"{ep_key}_resp_barplot_{v}"
            if chart_key in figures and figures[chart_key] and os.path.exists(figures[chart_key]):
                doc.add_picture(figures[chart_key], width=Inches(5))
                fig_label = num.fig()
                p = doc.add_paragraph(fig_label + f". Доля респондеров на визите {v} (Active vs Placebo показаны агрегированно; 95% ДИ).")
                p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None

                resp_fig_rows.append({"visit": v, "label": fig_label})

        if resp_fig_rows:
            add_ai_for_artifact(
                {
                    "id": f"s6_{ep_key}_resp_barplots_figset",
                    "kind": "chart",
                    "label": f"Доля респондеров по визитам: {short}",
                    "visits": resp_fig_rows,
                    "by_visit": [
                        {
                            "visit": v,
                            "chi2_p": ((resp.get("visits", {}).get(v, {}) or {}).get("test", {}) or {}).get("p") if isinstance(resp, dict) else None,
                            "groups": (resp.get("visits", {}).get(v, {}) or {}).get("groups") if isinstance(resp, dict) else None,
                        }
                        for v in VISITS
                    ],
                }
            )

        if isinstance(resp, dict) and resp.get("visits"):
            doc.add_paragraph(num.tab() + f". Респондеры по визитам: {short} (4 группы + укрупнение)")
            table = doc.add_table(rows=1, cols=9, style="Table Grid")
            hdr = table.rows[0].cells
            for i, h in enumerate(["Визит", "Г1 %", "Г2 %", "Г3 %", "Г4 %", "p (χ²)", "Active %", "Placebo %", "Интерпретация"]):
                hdr[i].text = h
                hdr[i].paragraphs[0].runs[0].bold = True

            resp_rows = []
            for v in VISITS:
                vres = resp.get("visits", {}).get(v, {})
                groups_list = vres.get("groups", []) if isinstance(vres, dict) else []
                by_g = {str(r.get("group")): r for r in groups_list if isinstance(r, dict)}

                row = table.add_row().cells
                row[0].text = v
                for col_i, g in enumerate(["1", "2", "3", "4"], start=1):
                    pct = by_g.get(g, {}).get("pct")
                    row[col_i].text = f"{float(pct):.0f}%" if pct is not None and np.isfinite(pct) else "—"

                p_val = vres.get("test", {}).get("p") if isinstance(vres, dict) else None
                row[5].text = f"{float(p_val):.4f}" if p_val is not None and np.isfinite(p_val) else "—"

                active_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["1", "3"])
                active_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["1", "3"])
                placebo_resp = sum(int(by_g.get(g, {}).get("responders", 0)) for g in ["2", "4"])
                placebo_n = sum(int(by_g.get(g, {}).get("n", 0)) for g in ["2", "4"])
                active_pct = (active_resp / active_n * 100.0) if active_n > 0 else np.nan
                placebo_pct = (placebo_resp / placebo_n * 100.0) if placebo_n > 0 else np.nan
                row[6].text = f"{active_pct:.0f}%" if np.isfinite(active_pct) else "—"
                row[7].text = f"{placebo_pct:.0f}%" if np.isfinite(placebo_pct) else "—"

                if v == "V2":
                    row[8].text = "baseline"
                    interp_txt = "baseline"
                else:
                    if np.isfinite(active_pct) and np.isfinite(placebo_pct):
                        arr = active_pct - placebo_pct
                        nnt, _ = calculate_nnt(float(active_pct), float(placebo_pct))
                        nnt_txt = f"NNT≈{nnt:.1f}" if nnt is not None else "NNT=—"
                        row[8].text = f"Δ={arr:+.1f} п.п.; {nnt_txt}"
                        interp_txt = f"Δ={arr:+.1f} п.п.; {nnt_txt}"
                    else:
                        row[8].text = "—"
                        interp_txt = "—"

                resp_rows.append(
                    {
                        "visit": v,
                        "g1_pct": by_g.get("1", {}).get("pct"),
                        "g2_pct": by_g.get("2", {}).get("pct"),
                        "g3_pct": by_g.get("3", {}).get("pct"),
                        "g4_pct": by_g.get("4", {}).get("pct"),
                        "chi2_p": p_val,
                        "active_pct": (float(active_pct) if np.isfinite(active_pct) else None),
                        "placebo_pct": (float(placebo_pct) if np.isfinite(placebo_pct) else None),
                        "interpretation": interp_txt,
                    }
                )

            add_ai_for_artifact(
                {
                    "id": f"s6_{ep_key}_resp_table",
                    "kind": "table",
                    "label": f"Респондеры по визитам ({short})",
                    "rows": resp_rows,
                }
            )

        doc.add_page_break()

    _doc_add_heading(doc, "7. ИТОГОВАЯ СВОДКА ПО ВИЗИТАМ И ОБСУЖДЕНИЕ", level=1)
    add_text_with_rewrite(
        "s7_intro",
        "Раздел агрегирует результаты по каждой временной точке (V2–V6): глобальный тест 4 групп (Kruskal–Wallis), "
        "наиболее выраженные post-hoc различия (p_adj), объединённый анализ Active vs Placebo (p_adj, ES, Δ) и сравнительные графики размеров эффекта (ε²).",
    )

    for idx, v in enumerate(VISITS, start=1):
        doc.add_heading(f"7.{idx} Визит {v}", level=2)

        fig_key = f"forest_{v}"
        if fig_key in figures and figures[fig_key] and os.path.exists(figures[fig_key]):
            doc.add_picture(figures[fig_key], width=Inches(6))
            fig_label = num.fig()
            p = doc.add_paragraph(fig_label + f". Сравнение размеров эффекта (ε²) на визите {v} по всем показателям.")
            p.style = doc.styles['Caption'] if 'Caption' in doc.styles else None

            eps_rows = []
            for ep_key, result in all_results.items():
                vdata = result.get("by_visit", {}).get(v, {})
                kw = vdata.get("kruskal", {})
                eps_rows.append(
                    {
                        "endpoint": result.get("short", ep_key),
                        "epsilon_sq": kw.get("epsilon_sq"),
                        "p": kw.get("p"),
                    }
                )
            add_ai_for_artifact(
                {
                    "id": f"s7_{v}_forest_fig",
                    "kind": "chart",
                    "label": f"{fig_label}. Сравнение размеров эффекта (ε²) на визите {v}",
                    "rows": eps_rows,
                }
            )

        doc.add_paragraph(num.tab() + f". Сводка по визиту {v} (все показатели)")
        table = doc.add_table(rows=1, cols=10, style="Table Grid")
        hdr = table.rows[0].cells
        headers = [
            "Показатель",
            "K-W p",
            "BF₁₀",
            "ε²",
            "Лучшая пара (p_adj)",
            "ES пары",
            "Active vs Placebo p_adj",
            "ES (укрупнение)",
            "Δ (Abs)",
            "Δ (%)",
        ]
        for i, h in enumerate(headers):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        newbie_hits = []
        visit_rows = []

        for ep_key, result in all_results.items():
            row = table.add_row().cells
            row[0].text = result.get("short", ep_key)

            vdata = result.get("by_visit", {}).get(v, {})
            kw = vdata.get("kruskal", {})
            bf = vdata.get("bf10", np.nan)

            row[1].text = f"{float(kw.get('p')):.4f}" if kw.get("p") is not None and np.isfinite(kw.get("p")) else "—"
            row[2].text = f"{float(bf):.2f}" if bf is not None and np.isfinite(bf) else "—"
            row[3].text = f"{float(kw.get('epsilon_sq')):.3f}" if kw.get("epsilon_sq") is not None and np.isfinite(kw.get("epsilon_sq")) else "—"

            best_pair_txt = "—"
            best_pair_es = "—"
            pv = result.get("pairwise", {}).get(v, {})
            best = None
            if isinstance(pv, dict) and pv:
                for pair_key, st in pv.items():
                    if not isinstance(st, dict) or "error" in st:
                        continue
                    p_adj = st.get("p_adj")
                    if p_adj is None or not np.isfinite(p_adj):
                        continue
                    if best is None or float(p_adj) < float(best["p_adj"]):
                        best = {"pair": pair_key, "p_adj": float(p_adj), "es": float(st.get("r", np.nan))}
            if best is not None and best["p_adj"] < 0.05:
                try:
                    g1, g2 = str(best["pair"]).split("_vs_")
                    best_pair_txt = f"Г{g1}–Г{g2} (p={best['p_adj']:.3f})"
                except Exception:
                    best_pair_txt = f"p={best['p_adj']:.3f}"
            row[4].text = best_pair_txt
            if best is not None and best.get("es") is not None and np.isfinite(best.get("es")) and best.get("p_adj") is not None and np.isfinite(best.get("p_adj")) and float(best.get("p_adj")) < 0.05:
                best_pair_es = f"{float(best.get('es')):.2f}"
            row[5].text = best_pair_es

            avp = active_vs_placebo.get(ep_key, {}).get("visits", {}).get(v, {})
            p_avp = avp.get("p_adj", avp.get("p_value", np.nan))
            row[6].text = f"{float(p_avp):.4f}" if p_avp is not None and np.isfinite(p_avp) else "—"
            es_avp = avp.get("effect_size", np.nan)
            row[7].text = f"{float(es_avp):.2f}" if es_avp is not None and np.isfinite(es_avp) else "—"
            diff = avp.get("diff_median", np.nan)
            diff_pct = avp.get("diff_pct", np.nan)
            row[8].text = f"{float(diff):+.1f}" if diff is not None and np.isfinite(diff) else "—"
            row[9].text = f"{float(diff_pct):+.1f}%" if diff_pct is not None and np.isfinite(diff_pct) else "—"

            if avp.get("significant"):
                newbie_hits.append(f"{result.get('short', ep_key)}: значимо (p_adj={float(p_avp):.3f}, ES={float(es_avp):.2f})")

            visit_rows.append(
                {
                    "endpoint": result.get("short", ep_key),
                    "kw_p": kw.get("p"),
                    "bf10": bf,
                    "epsilon_sq": kw.get("epsilon_sq"),
                    "best_pair": best_pair_txt,
                    "best_pair_es": (best.get("es") if (best is not None and best.get("es") is not None and np.isfinite(best.get("es")) and best.get("p_adj") is not None and np.isfinite(best.get("p_adj")) and float(best.get("p_adj")) < 0.05) else None),
                    "pooled_p_adj": p_avp,
                    "pooled_effect_size": es_avp,
                    "diff_median": diff,
                    "diff_pct": diff_pct,
                    "pooled_significant": bool(avp.get("significant")),
                }
            )

        if newbie_hits:
            doc.add_paragraph("Интерпретация: на этом визите укрупнение выявило значимые различия по: " + "; ".join(newbie_hits) + ".")
        else:
            doc.add_paragraph("Интерпретация: на этом визите укрупнение не выявило значимых различий (p_adj ≥ 0.05) ни по одному показателю.")

        add_ai_for_artifact(
            {
                "id": f"s7_{v}_summary_table",
                "kind": "table",
                "label": f"Сводка по визиту {v} (все показатели)",
                "rows": visit_rows,
            }
        )

        doc.add_page_break()

    _doc_add_heading(doc, "8. ОБСУЖДЕНИЕ", level=1)

    doc.add_heading("8.1 Цели и задачи", level=2)
    doc.add_paragraph("Цель: оценить эффективность аппарата ДИАМАГ по ключевым клиническим шкалам в сравнении с плацебо.")
    doc.add_paragraph(
        "Задачи: (1) сравнить 4 рандомизированные группы на каждом визите; (2) выполнить post-hoc сравнения с коррекцией (Холм); "
        "(3) выполнить укрупнение Active vs Placebo; (4) оценить внутригрупповую динамику; (5) оценить долю респондеров и клиническую значимость (NNT); "
        "(6) подтвердить результаты mixed effects моделями повторных измерений."
    )

    doc.add_heading("8.2 Проверка гипотез", level=2)
    doc.add_paragraph("H₀: терапия не приводит к различиям между группами. H₁: терапия приводит к различиям (эффект есть).")
    for ep_key, result in all_results.items():
        short = result.get("short", ep_key)
        kw_any = None
        for v in VISITS:
            kw = result.get("by_visit", {}).get(v, {}).get("kruskal", {})
            p = kw.get("p")
            if p is not None and np.isfinite(p) and float(p) < 0.05:
                kw_any = (v, float(p))
                break
        avp_any = None
        avp_visits = active_vs_placebo.get(ep_key, {}).get("visits", {})
        for v in VISITS:
            d = avp_visits.get(v, {}) if isinstance(avp_visits, dict) else {}
            p = d.get("p_adj", d.get("p_value", np.nan))
            if p is not None and np.isfinite(p) and float(p) < 0.05:
                avp_any = (v, float(p), float(d.get("effect_size", 0.0)))
                break
        if kw_any is not None and avp_any is not None:
            doc.add_paragraph(
                f"{short}: зарегистрированы статистические сигналы как на уровне 4 групп ({kw_any[0]}, p={kw_any[1]:.4f}), "
                f"так и в укрупнённом сравнении Active vs Placebo ({avp_any[0]}, p_adj={avp_any[1]:.4f}, ES={avp_any[2]:.2f})."
            )
        elif kw_any is not None:
            doc.add_paragraph(f"{short}: H₀ отклоняется на уровне 4 групп на отдельных визитах ({kw_any[0]}, p={kw_any[1]:.4f}).")
        elif avp_any is not None:
            doc.add_paragraph(f"{short}: H₀ отклоняется в укрупнённом сравнении Active vs Placebo на отдельных визитах ({avp_any[0]}, p_adj={avp_any[1]:.4f}, ES={avp_any[2]:.2f}).")
        else:
            doc.add_paragraph(f"{short}: статистических оснований отклонять H₀ не получено (p_adj≥0.05 по визитам).")

    doc.add_heading("8.3 Обсуждение результатов", level=2)
    ai_block = generate_ai_discussion_chunked(
        exec_summary=exec_summary,
        all_results=all_results,
        active_vs_placebo=active_vs_placebo,
        responders=responders,
    )
    if isinstance(ai_block, dict) and ai_block.get("discussion"):
        for paragraph in ai_block.get("discussion", []):
            doc.add_paragraph(str(paragraph))
    else:
        for paragraph in generate_discussion_text(exec_summary, all_results, active_vs_placebo, responders):
            doc.add_paragraph(paragraph)

    doc.add_page_break()

    _doc_add_heading(doc, "9. ВЫВОДЫ", level=1)
    if isinstance(ai_block, dict) and ai_block.get("conclusions"):
        for line in ai_block.get("conclusions", []):
            doc.add_paragraph(str(line))
    else:
        for line in generate_conclusions_text(exec_summary, all_results, active_vs_placebo, responders):
            doc.add_paragraph(str(line))

    doc.save(output_path)
    print(f"Saved: {output_path}")
    return str(output_path)


# ============================================================
# MAIN
# ============================================================

def main():
    print("=" * 60)
    print("  DIAMAG COMPREHENSIVE ANALYSIS + GRAPHS")
    print("=" * 60)
    
    OUTPUT_DIR.mkdir(exist_ok=True)
    
    # Load data
    print(f"\n📂 Loading: {EXCEL_PATH}")
    df = pd.read_excel(EXCEL_PATH, sheet_name="Лист1")
    df[GROUP_COL] = df[GROUP_COL].astype(str)
    print(f"   {len(df)} patients, {len(df.columns)} columns")
    
    # Analyze all endpoints
    all_results = {}
    for key in ENDPOINTS:
        all_results[key] = analyze_endpoint_full(df, key)
    
    # Responders (Longitudinal V3-V6)
    responders = analyze_responders(df, threshold=20)
    
    # Active vs Placebo comparison (Combined Longitudinal)
    active_vs_placebo = analyze_active_vs_placebo(df, all_results)

    # Active vs Placebo comparisons by duration (30d: G1 vs G2; 20d: G3 vs G4)
    duration_active_vs_placebo = analyze_active_vs_placebo_by_duration(df)
    
    # Generate graphs
    print("\n📊 Generating publication-quality figures...")
    figures = {}
    
    for key in ENDPOINTS:
        print(f"   📈 {ENDPOINTS[key]['short']}")
        figures[f"{key}_spaghetti"] = create_spaghetti_plot(
            df,
            key,
            OUTPUT_DIR / f"fig_{key}_spaghetti.png",
            pairwise_by_visit=all_results.get(key, {}).get("pairwise", {}),
        )
        
        # Pooled Spaghetti (Active vs Placebo)
        figures[f"{key}_pooled_spaghetti"] = create_pooled_spaghetti(
            active_vs_placebo, key, OUTPUT_DIR / f"fig_{key}_pooled_spaghetti.png"
        )

        # Duration-specific Spaghetti (G1 vs G2; G3 vs G4)
        figures[f"{key}_30d_spaghetti"] = create_duration_spaghetti(
            duration_active_vs_placebo, key, "30d", OUTPUT_DIR / f"fig_{key}_30d_spaghetti.png"
        )
        figures[f"{key}_20d_spaghetti"] = create_duration_spaghetti(
            duration_active_vs_placebo, key, "20d", OUTPUT_DIR / f"fig_{key}_20d_spaghetti.png"
        )
        
        # Boxplots per visit
        for visit in VISITS:
            # Standard 4-group boxplot
            figures[f"{key}_boxplot_{visit}"] = create_visit_boxplot(
                df,
                key,
                visit,
                OUTPUT_DIR / f"fig_{key}_boxplot_{visit}.png",
                pairwise_stats=all_results.get(key, {}).get("pairwise", {}).get(visit),
            )
            
            # Pooled Active vs Placebo boxplot
            figures[f"{key}_pooled_boxplot_{visit}"] = create_pooled_boxplot(
                df,
                key,
                visit,
                OUTPUT_DIR / f"fig_{key}_pooled_boxplot_{visit}.png",
                visit_stats=active_vs_placebo.get(key, {}).get("visits", {}).get(visit),
            )

            # Duration-specific boxplots
            figures[f"{key}_30d_boxplot_{visit}"] = create_duration_boxplot(
                df,
                key,
                visit,
                duration_active_vs_placebo.get(key, {}).get("comparisons", {}).get("30d", {}),
                OUTPUT_DIR / f"fig_{key}_30d_boxplot_{visit}.png",
                visit_stats=duration_active_vs_placebo.get(key, {}).get("comparisons", {}).get("30d", {}).get("visits", {}).get(visit),
            )
            figures[f"{key}_20d_boxplot_{visit}"] = create_duration_boxplot(
                df,
                key,
                visit,
                duration_active_vs_placebo.get(key, {}).get("comparisons", {}).get("20d", {}),
                OUTPUT_DIR / f"fig_{key}_20d_boxplot_{visit}.png",
                visit_stats=duration_active_vs_placebo.get(key, {}).get("comparisons", {}).get("20d", {}).get("visits", {}).get(visit),
            )
        
        # Responder dynamics (V2-V6)
        for visit in VISITS:
            figures[f"{key}_resp_barplot_{visit}"] = create_responder_visit_plot(
                responders, key, visit, OUTPUT_DIR / f"fig_{key}_resp_barplot_{visit}.png"
            )
    
    figures["forest"] = create_forest_plot(
        all_results, OUTPUT_DIR / "fig_forest_plot.png"
    )

    for visit in VISITS:
        figures[f"forest_{visit}"] = create_forest_plot_by_visit(
            all_results, visit, OUTPUT_DIR / f"fig_forest_plot_{visit}.png"
        )
    
    print(f"   ✓ Generated {len(figures)} figures")
    
    # Executive Summary
    exec_summary = generate_executive_summary(all_results, responders, active_vs_placebo, df)
    
    # Generate report
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_path = OUTPUT_DIR / f"diamag_SCIENTIFIC_{timestamp}.docx"
    
    generate_comprehensive_report_v2(
        all_results,
        responders,
        active_vs_placebo,
        duration_active_vs_placebo,
        exec_summary,
        figures,
        df,
        report_path,
    )
    
    print("\n" + "=" * 60)
    print(f"  ✅ DONE! Report: {report_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
