#!/usr/bin/env python3
"""
DIAMAG Clinical Trial Analysis Script

Автоматический анализ данных клинического исследования магнитотерапии
при болезни Паркинсона с генерацией Word-отчёта.

Использование:
    cd backend
    python3 scripts/run_diamag_analysis.py

Результат:
    output/diamag_report_YYYYMMDD_HHMMSS.docx
"""

import os
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple
import warnings

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import pandas as pd
import numpy as np
from scipy import stats
import pingouin as pg
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns

warnings.filterwarnings('ignore')

# ============================================================
# CONFIGURATION
# ============================================================

EXCEL_PATH = PROJECT_ROOT.parent / "docs" / "Первичка для анализа работа.xlsx"
OUTPUT_DIR = PROJECT_ROOT / "output"

# Column mappings for main endpoints
ENDPOINTS_CONFIG = {
    "updrs_part2": {
        "name": "УШОБП часть 2 (Повседневная активность)",
        "short_name": "UPDRS II",
        "columns": {
            "V2": "УШОБП часть 2 «Повседневная активность» баллы V2",
            "V3": "УШОБП часть 2 «Повседневная активность» баллы V3",
            "V4": "УШОБП часть 2 «Повседневная активность» баллы V4",
            "V5": "УШОБП часть 2 «Повседневная активность» баллы V5",
            "V6": "УШОБП часть 2 «Повседневная активность» баллы V6",
        },
        "primary": True,
    },
    "updrs_part3": {
        "name": "УШОБП часть 3 (Двигательные функции)",
        "short_name": "UPDRS III",
        "columns": {
            "V2": "УШОБП часть 3 «Двигательные функции» баллы V2",
            "V3": "УШОБП часть 3 «Двигательные функции» баллы V3",
            "V4": "УШОБП часть 3 «Двигательные функции» баллы V4",
            "V5": "УШОБП часть 3 «Двигательные функции» баллы V5",
            "V6": "УШОБП часть 3 «Двигательные функции» баллы V6",
        },
        "primary": True,
    },
    "dass21": {
        "name": "DASS-21 (Депрессия, тревога, стресс)",
        "short_name": "DASS-21",
        "columns": {
            "V2": "Шкала депрессии тревоги и стресса DASS-21 баллы V2",
            "V3": "Шкала депрессии тревоги и стресса DASS-21 баллы V3",
            "V4": "Шкала депрессии тревоги и стресса DASS-21 баллы V4",
            "V5": "Шкала депрессии тревоги и стресса DASS-21 баллы V5",
            "V6": "Шкала депрессии, тревоги и стресса DASS-21баллы V6",  # Note typo in original
        },
        "primary": False,
    },
    "epworth": {
        "name": "Шкала сонливости Эпуорта",
        "short_name": "Epworth",
        "columns": {
            "V2": "Шкала оценки сонливости Эпуорта баллы V2",
            "V3": "Шкала оценки сонливости Эпуорта баллы V3",
            "V4": "Шкала оценки сонливости Эпуорта баллы V4",
            "V5": "Шкала оценки сонливости Эпуорта баллы V5",
            "V6": "Шкала оценки сонливости Эпуорта баллы V6",
        },
        "primary": False,
    },
    "apathy": {
        "name": "Шкала апатии Старкстейна",
        "short_name": "Apathy",
        "columns": {
            "V2": "Шкала апатии Старкстейна баллы V2",
            "V3": "Шкала апатии Старкстейна баллы V3",
            "V4": "Шкала апатии Старкстейна баллы V4",
            "V5": "Шкала апатии Старкстейна баллы V5",
            "V6": "Шкала апатии Старкстейна баллы V6",
        },
        "primary": False,
    },
    "fatigue": {
        "name": "Шкала утомляемости при БП",
        "short_name": "Fatigue",
        "columns": {
            "V2": "Шкала оценки утомляемости при БП баллы V2",
            "V3": "Шкала оценки утомляемости при БП баллы V3",
            "V4": "Шкала оценки утомляемости при БП баллы V4",
            "V5": "Шкала оценки утомляемости при БП баллы V5",
            "V6": "Шкала оценки утомляемости при БП баллы V6",
        },
        "primary": False,
    },
    "pdq39": {
        "name": "PDQ-39 (Качество жизни при БП)",
        "short_name": "PDQ-39",
        "columns": {
            "V2": "Шкала оценки качества жизни при БП баллы V2",
            "V3": "Шкала оценки качества жизни при БП баллы V3", 
            "V4": "Шкала оценки качества жизни при БП баллы",  # V4 has different name
            "V5": "Шкала оценки качества жизни при БП баллы V5",
            "V6": "Шкала оценки качества жизни при БП баллы V6",
        },
        "primary": False,
    },
    "stroop": {
        "name": "Индекс интерференции Струпа",
        "short_name": "Stroop",
        "columns": {
            "V2": "Индекс интерференции V2",
            "V3": "Индекс интерференции V4",  # Note: V3 uses V4 label in original
            "V4": "Индекс интерференции V4.1",
            "V5": "Индекс интерференции V5",
            "V6": "Индекс интерференции V6",
        },
        "primary": False,
    },
}

GROUP_COL = "Группа"
ID_COL = "ID № участника исследования"

# ============================================================
# DATA LOADING
# ============================================================

def load_data() -> pd.DataFrame:
    """Load and preprocess the Excel file."""
    print(f"📂 Loading data from: {EXCEL_PATH}")
    
    if not EXCEL_PATH.exists():
        raise FileNotFoundError(f"Excel file not found: {EXCEL_PATH}")
    
    df = pd.read_excel(EXCEL_PATH, sheet_name="Лист1")
    print(f"   Shape: {df.shape[0]} patients × {df.shape[1]} columns")
    print(f"   Groups: {sorted(df[GROUP_COL].dropna().unique())}")
    
    # Convert group to string for consistent handling
    df[GROUP_COL] = df[GROUP_COL].astype(str)
    
    return df


def get_endpoint_data(df: pd.DataFrame, endpoint_key: str) -> pd.DataFrame:
    """Extract long-form data for a specific endpoint."""
    config = ENDPOINTS_CONFIG[endpoint_key]
    cols = config["columns"]
    
    records = []
    for _, row in df.iterrows():
        subject_id = row[ID_COL]
        group = row[GROUP_COL]
        
        for visit, col_name in cols.items():
            if col_name in df.columns:
                value = row[col_name]
                if pd.notna(value):
                    try:
                        records.append({
                            "subject_id": subject_id,
                            "group": group,
                            "visit": visit,
                            "value": float(value),
                        })
                    except (ValueError, TypeError):
                        pass
    
    return pd.DataFrame(records)


# ============================================================
# STATISTICS
# ============================================================

def compute_descriptive_stats(values: pd.Series) -> Dict[str, Any]:
    """Compute descriptive statistics for a series."""
    clean = values.dropna()
    n = len(clean)
    
    if n == 0:
        return {"n": 0, "mean": np.nan, "sd": np.nan, "median": np.nan, "q1": np.nan, "q3": np.nan}
    
    return {
        "n": n,
        "mean": float(clean.mean()),
        "sd": float(clean.std()),
        "se": float(clean.std() / np.sqrt(n)) if n > 1 else np.nan,
        "median": float(clean.median()),
        "q1": float(clean.quantile(0.25)),
        "q3": float(clean.quantile(0.75)),
        "min": float(clean.min()),
        "max": float(clean.max()),
    }


def run_kruskal_wallis(df: pd.DataFrame, value_col: str = "value", group_col: str = "group") -> Dict[str, Any]:
    """Run Kruskal-Wallis test and pairwise comparisons."""
    groups = df[group_col].unique()
    data_by_group = [df[df[group_col] == g][value_col].dropna() for g in groups]
    
    # Filter out empty groups
    valid_groups = [(g, d) for g, d in zip(groups, data_by_group) if len(d) > 0]
    if len(valid_groups) < 2:
        return {"error": "Not enough groups with data"}
    
    groups, data_by_group = zip(*valid_groups)
    
    # Kruskal-Wallis omnibus test
    stat, p_value = stats.kruskal(*data_by_group)
    
    # Effect size: epsilon-squared
    n_total = sum(len(d) for d in data_by_group)
    k = len(groups)
    epsilon_sq = (stat - k + 1) / (n_total - k) if n_total > k else np.nan
    
    # Pairwise comparisons with Dunn test
    pairwise = []
    try:
        posthoc = pg.pairwise_tests(data=df, dv=value_col, between=group_col, parametric=False)
        for _, row in posthoc.iterrows():
            pairwise.append({
                "A": str(row["A"]),
                "B": str(row["B"]),
                "p_value": float(row["p-unc"]) if pd.notna(row["p-unc"]) else np.nan,
                "p_adjusted": float(row["p-corr"]) if "p-corr" in row and pd.notna(row.get("p-corr")) else np.nan,
            })
    except Exception as e:
        print(f"   ⚠️ Pairwise tests failed: {e}")
    
    return {
        "test": "Kruskal-Wallis",
        "statistic": float(stat),
        "p_value": float(p_value),
        "effect_size": float(epsilon_sq) if np.isfinite(epsilon_sq) else None,
        "effect_size_name": "epsilon-squared",
        "pairwise": pairwise,
    }


def run_mixed_effects(df: pd.DataFrame, value_col: str = "value", group_col: str = "group", 
                      time_col: str = "visit", subject_col: str = "subject_id") -> Dict[str, Any]:
    """Run Linear Mixed Effects model using statsmodels: value ~ group * time + (1|subject)."""
    try:
        import statsmodels.formula.api as smf
        from statsmodels.regression.mixed_linear_model import MixedLMResults
        
        # Prepare data
        df = df.copy()
        df = df.dropna(subset=[value_col, group_col, time_col, subject_col])
        
        if len(df) < 10:
            return {"error": "Not enough data points"}
        
        # Ensure proper types
        df[group_col] = df[group_col].astype(str)
        df[time_col] = df[time_col].astype(str)
        df[subject_col] = df[subject_col].astype(str)
        
        # Formula: value ~ group * time (interaction)
        formula = f"{value_col} ~ C({group_col}) * C({time_col})"
        
        # Fit mixed model with random intercept per subject
        model = smf.mixedlm(formula, df, groups=df[subject_col])
        result = model.fit(reml=True, method='powell')
        
        # Extract key statistics
        effects = {}
        
        # Main effect of group (any coefficient starting with C(group))
        group_pvals = [result.pvalues[k] for k in result.pvalues.index if f"C({group_col})" in k and f"C({time_col})" not in k]
        if group_pvals:
            effects["Группа"] = {
                "p_value": float(min(group_pvals)),  # Most significant group effect
                "significant": any(p < 0.05 for p in group_pvals),
            }
        
        # Main effect of time
        time_pvals = [result.pvalues[k] for k in result.pvalues.index if f"C({time_col})" in k and f"C({group_col})" not in k]
        if time_pvals:
            effects["Время"] = {
                "p_value": float(min(time_pvals)),
                "significant": any(p < 0.05 for p in time_pvals),
            }
        
        # Interaction effect
        interaction_pvals = [result.pvalues[k] for k in result.pvalues.index if f"C({group_col})" in k and f"C({time_col})" in k]
        if interaction_pvals:
            effects["Взаимодействие"] = {
                "p_value": float(min(interaction_pvals)),
                "significant": any(p < 0.05 for p in interaction_pvals),
            }
        
        # Model fit statistics
        return {
            "test": "Linear Mixed Model (statsmodels)",
            "effects": effects,
            "interaction_p": effects.get("Взаимодействие", {}).get("p_value", np.nan),
            "aic": float(result.aic) if hasattr(result, 'aic') else None,
            "bic": float(result.bic) if hasattr(result, 'bic') else None,
            "converged": result.converged if hasattr(result, 'converged') else True,
        }
    
    except Exception as e:
        print(f"   ⚠️ Mixed effects failed: {e}")
        return {"error": str(e)}


def compute_bayes_factor(p_value: float, n: int = 30) -> float:
    """Estimate Bayes Factor (BF10) from p-value using Sellke et al. (2001) bound."""
    if p_value is None or np.isnan(p_value) or p_value <= 0 or p_value >= 1:
        return np.nan
    
    # Minimum bound: -1/(e * p * ln(p))
    try:
        bf10 = -1 / (np.e * p_value * np.log(p_value))
        return min(bf10, 1000)  # Cap at 1000
    except:
        return np.nan


def interpret_effect_size(effect_size: float, effect_type: str = "eta_squared") -> str:
    """Interpret effect size magnitude."""
    if effect_size is None or np.isnan(effect_size):
        return "—"
    
    abs_es = abs(effect_size)
    
    if effect_type in ("eta_squared", "epsilon_squared", "np2"):
        if abs_es < 0.01:
            return "незначительный"
        elif abs_es < 0.06:
            return "малый"
        elif abs_es < 0.14:
            return "средний"
        else:
            return "большой"
    else:  # Cohen's d
        if abs_es < 0.2:
            return "незначительный"
        elif abs_es < 0.5:
            return "малый"
        elif abs_es < 0.8:
            return "средний"
        else:
            return "большой"


# ============================================================
# ANALYSIS FUNCTIONS
# ============================================================

def analyze_endpoint(df: pd.DataFrame, endpoint_key: str) -> Dict[str, Any]:
    """Perform full analysis for one endpoint."""
    config = ENDPOINTS_CONFIG[endpoint_key]
    print(f"\n📊 Analyzing: {config['name']}")
    
    # Get long-form data
    long_df = get_endpoint_data(df, endpoint_key)
    
    if long_df.empty:
        print(f"   ⚠️ No data available")
        return {"error": "No data"}
    
    print(f"   Records: {len(long_df)}")
    
    # Descriptive statistics by group and visit
    descriptive = {}
    for group in long_df["group"].unique():
        descriptive[group] = {}
        for visit in sorted(long_df["visit"].unique()):
            values = long_df[(long_df["group"] == group) & (long_df["visit"] == visit)]["value"]
            descriptive[group][visit] = compute_descriptive_stats(values)
    
    # Cross-sectional analysis at V6 (endpoint)
    v6_data = long_df[long_df["visit"] == "V6"]
    kruskal_v6 = run_kruskal_wallis(v6_data) if not v6_data.empty else {"error": "No V6 data"}
    
    # Mixed effects analysis (longitudinal)
    mixed = run_mixed_effects(long_df)
    
    # Bayes Factor
    bf10_v6 = compute_bayes_factor(kruskal_v6.get("p_value", np.nan))
    
    result = {
        "name": config["name"],
        "short_name": config["short_name"],
        "primary": config["primary"],
        "descriptive": descriptive,
        "kruskal_v6": kruskal_v6,
        "mixed_effects": mixed,
        "bf10": bf10_v6,
        "long_data": long_df,
    }
    
    # Print summary
    p_v6 = kruskal_v6.get("p_value", np.nan)
    sig = "✓" if p_v6 < 0.05 else "✗"
    print(f"   V6 Kruskal-Wallis: H={kruskal_v6.get('statistic', 'N/A'):.2f}, p={p_v6:.4f} {sig}")
    if "effects" in mixed:
        int_p = mixed.get("interaction_p", np.nan)
        print(f"   Mixed ANOVA Interaction: p={int_p:.4f}")
    
    return result


def analyze_responders(df: pd.DataFrame, threshold_pct: float = 20.0) -> Dict[str, Any]:
    """Analyze responder rates (improvement >= threshold from baseline V2 to V6)."""
    print(f"\n📊 Analyzing responders (threshold: {threshold_pct}% improvement)")
    
    endpoint_key = "updrs_part3"  # Primary endpoint
    config = ENDPOINTS_CONFIG[endpoint_key]
    
    v2_col = config["columns"]["V2"]
    v6_col = config["columns"]["V6"]
    
    # Check columns exist
    if v2_col not in df.columns or v6_col not in df.columns:
        return {"error": "Missing columns"}
    
    # Calculate responders
    responder_data = []
    for group in df[GROUP_COL].unique():
        group_df = df[df[GROUP_COL] == group]
        
        valid = group_df[[v2_col, v6_col]].dropna()
        n_total = len(valid)
        
        if n_total == 0:
            continue
        
        # Responder = improvement >= threshold
        improvements = (valid[v2_col] - valid[v6_col]) / valid[v2_col] * 100
        n_responders = (improvements >= threshold_pct).sum()
        
        responder_data.append({
            "group": group,
            "n_total": n_total,
            "n_responders": int(n_responders),
            "pct_responders": float(n_responders / n_total * 100) if n_total > 0 else 0,
        })
    
    # Chi-square / Fisher test for responder rates
    if len(responder_data) >= 2:
        contingency = np.array([
            [r["n_responders"], r["n_total"] - r["n_responders"]] 
            for r in responder_data
        ])
        
        try:
            chi2, p_value, dof, expected = stats.chi2_contingency(contingency)
            test_result = {
                "test": "Chi-square",
                "statistic": float(chi2),
                "p_value": float(p_value),
                "dof": int(dof),
            }
        except:
            # Use Fisher for 2x2
            try:
                odds_ratio, p_value = stats.fisher_exact(contingency[:2])
                test_result = {
                    "test": "Fisher exact",
                    "odds_ratio": float(odds_ratio),
                    "p_value": float(p_value),
                }
            except:
                test_result = {"error": "Test failed"}
    else:
        test_result = {"error": "Not enough groups"}
    
    rates_str = [f"G{r['group']}: {r['pct_responders']:.1f}%" for r in responder_data]
    print(f"   Responder rates: {rates_str}")
    
    return {
        "threshold_pct": threshold_pct,
        "responders": responder_data,
        "test": test_result,
    }


def generate_table1(df: pd.DataFrame) -> Dict[str, Any]:
    """Generate Table 1: Baseline characteristics."""
    print("\n📊 Generating Table 1 (Baseline)")
    
    groups = sorted(df[GROUP_COL].unique())
    
    # Demographics
    age_col = "Возраст"
    sex_col = "Пол"
    
    table_data = []
    
    # Age
    row = {"variable": "Возраст, лет (M ± SD)"}
    for g in groups:
        values = df[df[GROUP_COL] == g][age_col].dropna()
        stats_dict = compute_descriptive_stats(values)
        row[f"group_{g}"] = f"{stats_dict['mean']:.1f} ± {stats_dict['sd']:.1f}" if stats_dict["n"] > 0 else "—"
    table_data.append(row)
    
    # Sex
    row = {"variable": "Пол (М/Ж)"}
    for g in groups:
        sex_counts = df[df[GROUP_COL] == g][sex_col].value_counts()
        m = sex_counts.get("М", sex_counts.get("м", sex_counts.get(1, 0)))
        f = sex_counts.get("Ж", sex_counts.get("ж", sex_counts.get(2, 0)))
        row[f"group_{g}"] = f"{m}/{f}"
    table_data.append(row)
    
    # Baseline UPDRS Part 3 (V2)
    updrs3_col = ENDPOINTS_CONFIG["updrs_part3"]["columns"]["V2"]
    if updrs3_col in df.columns:
        row = {"variable": "UPDRS III (V2), баллы"}
        for g in groups:
            values = df[df[GROUP_COL] == g][updrs3_col].dropna()
            stats_dict = compute_descriptive_stats(values)
            row[f"group_{g}"] = f"{stats_dict['median']:.1f} [{stats_dict['q1']:.1f}–{stats_dict['q3']:.1f}]" if stats_dict["n"] > 0 else "—"
        table_data.append(row)
    
    # N per group
    row = {"variable": "N"}
    for g in groups:
        row[f"group_{g}"] = str(len(df[df[GROUP_COL] == g]))
    table_data.insert(0, row)
    
    return {
        "groups": groups,
        "data": table_data,
    }


# ============================================================
# PLOTTING
# ============================================================

def create_boxplot(long_df: pd.DataFrame, title: str, output_path: Path) -> str:
    """Create boxplot for group comparison."""
    plt.figure(figsize=(10, 6))
    
    # Use last visit for cross-sectional view
    v6_data = long_df[long_df["visit"] == "V6"]
    
    if v6_data.empty:
        v6_data = long_df  # Fallback
    
    sns.boxplot(x="group", y="value", data=v6_data, palette="Set2", showfliers=False)
    sns.stripplot(x="group", y="value", data=v6_data, color="black", alpha=0.5, size=4)
    
    plt.title(title, fontsize=14, fontweight="bold")
    plt.xlabel("Группа", fontsize=12)
    plt.ylabel("Баллы", fontsize=12)
    plt.tight_layout()
    
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()
    
    return str(output_path)


def create_spaghetti_plot(long_df: pd.DataFrame, title: str, output_path: Path) -> str:
    """Create spaghetti plot for longitudinal data."""
    plt.figure(figsize=(12, 6))
    
    # Order visits
    visit_order = ["V2", "V3", "V4", "V5", "V6"]
    long_df = long_df.copy()
    long_df["visit"] = pd.Categorical(long_df["visit"], categories=visit_order, ordered=True)
    long_df = long_df.dropna(subset=["visit"])
    
    # Mean ± SE by group and visit
    summary = long_df.groupby(["group", "visit"]).agg(
        mean=("value", "mean"),
        se=("value", lambda x: x.std() / np.sqrt(len(x)) if len(x) > 1 else 0),
    ).reset_index()
    
    colors = {"1": "#1f77b4", "2": "#ff7f0e", "3": "#2ca02c", "4": "#d62728"}
    
    for group in summary["group"].unique():
        group_data = summary[summary["group"] == group].sort_values("visit")
        color = colors.get(str(group), "#333333")
        
        plt.errorbar(
            x=range(len(group_data)),
            y=group_data["mean"],
            yerr=group_data["se"],
            marker="o",
            linewidth=2,
            capsize=4,
            label=f"Группа {group}",
            color=color,
        )
    
    visits = sorted(summary["visit"].unique())
    plt.xticks(range(len(visits)), visits)
    
    plt.title(title, fontsize=14, fontweight="bold")
    plt.xlabel("Визит", fontsize=12)
    plt.ylabel("Баллы (M ± SE)", fontsize=12)
    plt.legend(frameon=False)
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    
    plt.savefig(output_path, dpi=300, bbox_inches="tight")
    plt.close()
    
    return str(output_path)


# ============================================================
# WORD REPORT GENERATION
# ============================================================

def add_heading(doc: Document, text: str, level: int = 1):
    """Add heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_table(doc: Document, headers: List[str], rows: List[List[str]], style: str = "Table Grid"):
    """Add formatted table to document."""
    table = doc.add_table(rows=1, cols=len(headers), style=style)
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
    
    return table


def format_p_value(p: float) -> str:
    """Format p-value for display."""
    if p is None or np.isnan(p):
        return "—"
    if p < 0.001:
        return "<0.001"
    elif p < 0.01:
        return f"{p:.3f}"
    else:
        return f"{p:.3f}"


def generate_word_report(
    table1: Dict,
    endpoint_results: Dict[str, Dict],
    responder_results: Dict,
    output_path: Path,
    figures: Dict[str, str],
) -> str:
    """Generate comprehensive Word report with all statistics."""
    print(f"\n📄 Generating Word report: {output_path}")
    
    doc = Document()
    
    # ============================================================
    # TITLE PAGE
    # ============================================================
    title = doc.add_paragraph()
    title.add_run("СТАТИСТИЧЕСКИЙ ОТЧЁТ").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    
    subtitle = doc.add_paragraph()
    subtitle.add_run("Клиническое исследование эффективности магнитотерапии\nпри болезни Паркинсона (ДИАМАГ)")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph()
    date_para.add_run(f"Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ============================================================
    # TABLE OF CONTENTS
    # ============================================================
    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    
    toc_items = [
        "1. Дизайн исследования",
        "   1.1. Цели и задачи",
        "   1.2. Группы и показатели",
        "   1.3. Визиты и временные точки",
        "2. Описание выборки (Table 1)",
        "3. Первичные конечные точки",
        "   3.1. UPDRS часть 2 (Повседневная активность)",
        "   3.2. UPDRS часть 3 (Двигательные функции)",
        "4. Вторичные конечные точки",
        "   4.1. DASS-21 (Депрессия, тревога)",
        "   4.2. Шкала сонливости Эпуорта",
        "   4.3. Шкала апатии Старкстейна",
        "   4.4. Шкала утомляемости при БП",
        "   4.5. PDQ-39 (Качество жизни)",
        "   4.6. Индекс интерференции Струпа",
        "5. Анализ респондеров",
        "6. Сводная таблица всех показателей",
        "7. Заключение",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 1: STUDY DESIGN
    # ============================================================
    add_heading(doc, "1. ДИЗАЙН ИССЛЕДОВАНИЯ", 1)
    
    add_heading(doc, "1.1. Цели и задачи", 2)
    doc.add_paragraph(
        "Цель исследования: оценить эффективность и безопасность применения "
        "медицинского изделия «Аппарат магнитотерапевтический «АЛМАГ-03» "
        "(торговый знак «ДИАМАГ») при болезни Паркинсона."
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Первичные конечные точки:").bold = True
    doc.add_paragraph("• UPDRS часть 3 — оценка двигательных функций")
    doc.add_paragraph("• UPDRS часть 2 — оценка повседневной активности")
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Вторичные конечные точки:").bold = True
    doc.add_paragraph("• DASS-21 — депрессия, тревога, стресс")
    doc.add_paragraph("• Шкала сонливости Эпуорта — дневная сонливость")
    doc.add_paragraph("• Шкала апатии Старкстейна")
    doc.add_paragraph("• Шкала утомляемости при БП")
    doc.add_paragraph("• PDQ-39 — качество жизни")
    doc.add_paragraph("• Тест Струпа — когнитивные функции")
    
    add_heading(doc, "1.2. Группы и показатели", 2)
    
    doc.add_paragraph("В исследовании участвовали 4 группы пациентов:")
    
    headers = ["Группа", "N", "Описание"]
    rows = []
    if table1 and table1.get("data"):
        groups = table1["groups"]
        for g in groups:
            n = table1['data'][0].get(f'group_{g}', '?')
            if g in ['1', '2']:
                desc = "ДИАМАГ + стандартная терапия"
            else:
                desc = "Плацебо + стандартная терапия"
            rows.append([f"Группа {g}", str(n), desc])
    else:
        rows = [["1", "11", "ДИАМАГ"], ["2", "10", "ДИАМАГ"], ["3", "11", "Плацебо"], ["4", "11", "Плацебо"]]
    add_table(doc, headers, rows)
    
    add_heading(doc, "1.3. Визиты и временные точки", 2)
    
    headers = ["Визит", "Описание", "День"]
    rows = [
        ["V1", "Скрининг", "−7 — 0"],
        ["V2", "Включение, рандомизация, начало терапии (baseline)", "0"],
        ["V3", "Оценка после окончания курса", "+10"],
        ["V4", "Оценка стабильности эффекта", "+20"],
        ["V5", "Оценка стабильности эффекта", "+30"],
        ["V6", "Финальная оценка (endpoint)", "+30 (±1)"],
    ]
    add_table(doc, headers, rows)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Статистические методы:").bold = True
    doc.add_paragraph("• Kruskal-Wallis — межгрупповое сравнение (непараметрический)")
    doc.add_paragraph("• Linear Mixed Model — анализ взаимодействия время × группа")
    doc.add_paragraph("• Bayes Factor — оценка силы доказательств")
    doc.add_paragraph("• Effect size (ε²) — размер эффекта")
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 2: TABLE 1
    # ============================================================
    add_heading(doc, "2. ОПИСАНИЕ ВЫБОРКИ (Table 1)", 1)
    
    if table1 and table1.get("data"):
        groups = table1["groups"]
        headers = ["Показатель"] + [f"Группа {g} (n={table1['data'][0].get(f'group_{g}', '?')})" for g in groups]
        
        rows = []
        for row_data in table1["data"][1:]:
            row = [row_data["variable"]]
            for g in groups:
                row.append(row_data.get(f"group_{g}", "—"))
            rows.append(row)
        
        add_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 3: PRIMARY ENDPOINTS (DETAILED)
    # ============================================================
    add_heading(doc, "3. ПЕРВИЧНЫЕ КОНЕЧНЫЕ ТОЧКИ", 1)
    
    for key, result in endpoint_results.items():
        if not result.get("primary"):
            continue
        
        add_heading(doc, result["name"], 2)
        
        # --- 2.1 Kruskal-Wallis ---
        p_para = doc.add_paragraph()
        p_para.add_run("2.1. Межгрупповое сравнение (Kruskal-Wallis)").bold = True
        
        kw = result.get("kruskal_v6", {})
        if "error" not in kw:
            p = kw.get("p_value", np.nan)
            sig = "статистически значимы" if p < 0.05 else "статистически не значимы"
            
            doc.add_paragraph(
                f"H = {kw.get('statistic', 0):.2f}, p = {format_p_value(p)}. "
                f"Различия между группами на визите V6 {sig}."
            )
            
            # Effect size with interpretation
            es = kw.get("effect_size")
            if es is not None:
                interp = interpret_effect_size(es, "epsilon_squared")
                doc.add_paragraph(f"Размер эффекта: ε² = {es:.3f} ({interp} эффект)")
        
        # --- Descriptive table with absolute values ---
        doc.add_paragraph()
        p_para = doc.add_paragraph()
        p_para.add_run("Описательные статистики по визитам (Median [Q1–Q3]):").bold = True
        
        desc = result.get("descriptive", {})
        if desc:
            groups_list = sorted(desc.keys())
            visits = ["V2", "V3", "V4", "V5", "V6"]
            
            headers = ["Визит"] + [f"Группа {g}" for g in groups_list]
            rows = []
            for v in visits:
                row = [v]
                for g in groups_list:
                    stats_d = desc.get(g, {}).get(v, {})
                    if stats_d.get("n", 0) > 0:
                        cell = f"{stats_d['median']:.1f} [{stats_d['q1']:.1f}–{stats_d['q3']:.1f}]"
                    else:
                        cell = "—"
                    row.append(cell)
                rows.append(row)
            add_table(doc, headers, rows)
        
        # --- Change from baseline (V2 → V6) ---
        doc.add_paragraph()
        p_para = doc.add_paragraph()
        p_para.add_run("Изменение от исходного уровня (V2 → V6):").bold = True
        
        if desc:
            headers = ["Группа", "V2 (baseline)", "V6 (endpoint)", "Δ абс.", "Δ %"]
            rows = []
            for g in groups_list:
                v2 = desc.get(g, {}).get("V2", {})
                v6 = desc.get(g, {}).get("V6", {})
                
                v2_med = v2.get("median", np.nan)
                v6_med = v6.get("median", np.nan)
                
                if np.isfinite(v2_med) and np.isfinite(v6_med):
                    delta_abs = v6_med - v2_med
                    delta_pct = (delta_abs / v2_med * 100) if v2_med != 0 else np.nan
                    rows.append([
                        f"Группа {g}",
                        f"{v2_med:.1f}",
                        f"{v6_med:.1f}",
                        f"{delta_abs:+.1f}",
                        f"{delta_pct:+.1f}%" if np.isfinite(delta_pct) else "—",
                    ])
                else:
                    rows.append([f"Группа {g}", "—", "—", "—", "—"])
            
            add_table(doc, headers, rows)
        
        # --- 2.2 Mixed Effects ---
        doc.add_paragraph()
        p_para = doc.add_paragraph()
        p_para.add_run("2.2. Смешанная модель (время × группа):").bold = True
        
        mixed = result.get("mixed_effects", {})
        if "effects" in mixed:
            doc.add_paragraph(f"Модель: {mixed.get('test', 'Linear Mixed Model')}")
            
            headers = ["Эффект", "p-value", "Значимо?"]
            rows = []
            for source, stats_d in mixed["effects"].items():
                p_val = stats_d.get("p_value", np.nan)
                sig_mark = "✓" if stats_d.get("significant", p_val < 0.05) else "✗"
                rows.append([
                    source,
                    format_p_value(p_val),
                    sig_mark,
                ])
            add_table(doc, headers, rows)
            
            int_p = mixed.get("interaction_p", np.nan)
            if np.isfinite(int_p):
                if int_p < 0.05:
                    doc.add_paragraph(
                        f"Взаимодействие время×группа значимо (p = {format_p_value(int_p)}): "
                        f"динамика показателя различается между группами."
                    )
                else:
                    doc.add_paragraph(
                        f"Взаимодействие время×группа не значимо (p = {format_p_value(int_p)})."
                    )
        elif "error" in mixed:
            doc.add_paragraph(f"Смешанная модель не рассчитана: {mixed['error']}")
        
        # --- 2.3 Bayes Factor ---
        doc.add_paragraph()
        p_para = doc.add_paragraph()
        p_para.add_run("2.3. Байесовский анализ:").bold = True
        
        bf = result.get("bf10")
        if bf and np.isfinite(bf):
            # Full Bayes Factor interpretation
            if bf > 100:
                bf_interp = "экстремально сильное свидетельство в пользу H₁"
                bf_meaning = "Данные очень убедительно свидетельствуют о наличии различий."
            elif bf > 30:
                bf_interp = "очень сильное свидетельство в пользу H₁"
                bf_meaning = "Данные убедительно поддерживают наличие различий."
            elif bf > 10:
                bf_interp = "сильное свидетельство в пользу H₁"
                bf_meaning = "Есть существенные основания полагать, что различия существуют."
            elif bf > 3:
                bf_interp = "умеренное свидетельство в пользу H₁"
                bf_meaning = "Данные склоняются к наличию различий, но требуется подтверждение."
            elif bf > 1:
                bf_interp = "слабое свидетельство в пользу H₁"
                bf_meaning = "Данные неопределённы, требуется больше наблюдений."
            elif bf > 1/3:
                bf_interp = "слабое свидетельство в пользу H₀"
                bf_meaning = "Данные слегка склоняются к отсутствию различий."
            elif bf > 1/10:
                bf_interp = "умеренное свидетельство в пользу H₀"
                bf_meaning = "Данные поддерживают отсутствие различий между группами."
            else:
                bf_interp = "сильное свидетельство в пользу H₀"
                bf_meaning = "Данные убедительно свидетельствуют об отсутствии различий."
            
            doc.add_paragraph(f"BF₁₀ = {bf:.2f} ({bf_interp})")
            doc.add_paragraph(f"Интерпретация: {bf_meaning}")
        else:
            doc.add_paragraph("Bayes Factor не рассчитан.")
        
        # --- Figures ---
        fig_key = f"{key}_boxplot"
        if fig_key in figures and os.path.exists(figures[fig_key]):
            doc.add_paragraph()
            doc.add_picture(figures[fig_key], width=Inches(5.5))
            doc.add_paragraph(f"Рисунок: Межгрупповое сравнение {result['short_name']} на визите V6")
        
        fig_key = f"{key}_spaghetti"
        if fig_key in figures and os.path.exists(figures[fig_key]):
            doc.add_paragraph()
            doc.add_picture(figures[fig_key], width=Inches(6))
            doc.add_paragraph(f"Рисунок: Динамика {result['short_name']} в ходе исследования (M ± SE)")
        
        doc.add_page_break()
    
    # ============================================================
    # SECTION 4: SECONDARY ENDPOINTS
    # ============================================================
    add_heading(doc, "4. ВТОРИЧНЫЕ КОНЕЧНЫЕ ТОЧКИ", 1)
    
    # Summary table
    p_para = doc.add_paragraph()
    p_para.add_run("Сводная таблица результатов:").bold = True
    
    headers = ["Показатель", "H (K-W)", "p", "BF₁₀", "Δ V2→V6 (Гр.1)", "Значимо?"]
    rows = []
    
    for key, result in endpoint_results.items():
        if result.get("primary"):
            continue
        
        kw = result.get("kruskal_v6", {})
        p_val = kw.get("p_value", np.nan)
        bf = result.get("bf10", np.nan)
        
        # Get delta for group 1
        desc = result.get("descriptive", {})
        g1_desc = desc.get("1", {})
        v2 = g1_desc.get("V2", {}).get("median", np.nan)
        v6 = g1_desc.get("V6", {}).get("median", np.nan)
        delta = f"{v6 - v2:+.1f}" if np.isfinite(v2) and np.isfinite(v6) else "—"
        
        sig_mark = "✓" if p_val < 0.05 else "✗"
        
        rows.append([
            result["short_name"],
            f"{kw.get('statistic', 0):.2f}" if "statistic" in kw else "—",
            format_p_value(p_val),
            f"{bf:.1f}" if np.isfinite(bf) else "—",
            delta,
            sig_mark,
        ])
    
    add_table(doc, headers, rows)
    
    # Detailed analysis for each secondary endpoint
    for key, result in endpoint_results.items():
        if result.get("primary"):
            continue
        
        add_heading(doc, result["name"], 2)
        
        kw = result.get("kruskal_v6", {})
        p_val = kw.get("p_value", np.nan)
        
        # Kruskal-Wallis
        doc.add_paragraph(
            f"Kruskal-Wallis: H = {kw.get('statistic', 0):.2f}, p = {format_p_value(p_val)}"
        )
        
        # Effect size
        es = kw.get("effect_size")
        if es is not None:
            interp = interpret_effect_size(es, "epsilon_squared")
            doc.add_paragraph(f"Размер эффекта: ε² = {es:.3f} ({interp})")
        
        # Bayes Factor
        bf = result.get("bf10")
        if bf and np.isfinite(bf):
            if bf > 10:
                bf_interp = "сильное свидетельство H₁"
            elif bf > 3:
                bf_interp = "умеренное H₁"
            elif bf > 1:
                bf_interp = "слабое H₁"
            else:
                bf_interp = "в пользу H₀"
            doc.add_paragraph(f"BF₁₀ = {bf:.2f} ({bf_interp})")
        
        # Mixed Effects interaction
        mixed = result.get("mixed_effects", {})
        int_p = mixed.get("interaction_p", np.nan)
        if np.isfinite(int_p):
            sig = "значимо" if int_p < 0.05 else "не значимо"
            doc.add_paragraph(f"Взаимодействие время×группа: p = {format_p_value(int_p)} ({sig})")
        
        # V2 → V6 change
        desc = result.get("descriptive", {})
        if desc:
            doc.add_paragraph()
            p_para = doc.add_paragraph()
            p_para.add_run("Изменение V2 → V6:").bold = True
            
            for g in sorted(desc.keys()):
                v2 = desc[g].get("V2", {}).get("median", np.nan)
                v6 = desc[g].get("V6", {}).get("median", np.nan)
                if np.isfinite(v2) and np.isfinite(v6):
                    delta = v6 - v2
                    pct = (delta / v2 * 100) if v2 != 0 else 0
                    doc.add_paragraph(f"  Группа {g}: {v2:.1f} → {v6:.1f} (Δ = {delta:+.1f}, {pct:+.1f}%)")
        
        # Figure
        fig_key = f"{key}_spaghetti"
        if fig_key in figures and os.path.exists(figures[fig_key]):
            doc.add_picture(figures[fig_key], width=Inches(5))
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 5: RESPONDERS
    # ============================================================
    add_heading(doc, "5. АНАЛИЗ РЕСПОНДЕРОВ", 1)
    
    if responder_results and "responders" in responder_results:
        threshold = responder_results.get("threshold_pct", 20)
        
        doc.add_paragraph(
            f"Критерий ответа: улучшение UPDRS III ≥ {threshold}% от исходного уровня (V2 → V6)"
        )
        doc.add_paragraph(
            "Респондер — пациент, у которого снижение балла по шкале UPDRS часть 3 "
            f"составило не менее {threshold}% от исходного значения."
        )
        
        doc.add_paragraph()
        p_para = doc.add_paragraph()
        p_para.add_run("Результаты:").bold = True
        
        headers = ["Группа", "N", "Респондеры", "%", "95% ДИ"]
        rows = []
        for r in responder_results["responders"]:
            n = r["n_total"]
            k = r["n_responders"]
            pct = r["pct_responders"]
            
            # Wilson confidence interval
            if n > 0:
                from scipy.stats import norm
                z = 1.96
                p_hat = k / n
                denom = 1 + z**2 / n
                center = (p_hat + z**2 / (2 * n)) / denom
                margin = z * np.sqrt((p_hat * (1 - p_hat) + z**2 / (4 * n)) / n) / denom
                ci_low = max(0, center - margin) * 100
                ci_high = min(1, center + margin) * 100
                ci_str = f"[{ci_low:.0f}–{ci_high:.0f}%]"
            else:
                ci_str = "—"
            
            rows.append([
                f"Группа {r['group']}",
                str(n),
                str(k),
                f"{pct:.1f}%",
                ci_str,
            ])
        
        add_table(doc, headers, rows)
        
        # Statistical test
        test = responder_results.get("test", {})
        if "p_value" in test:
            doc.add_paragraph()
            p_val = test["p_value"]
            test_name = test.get("test", "Chi-square")
            sig = "статистически значима" if p_val < 0.05 else "статистически не значима"
            
            doc.add_paragraph(
                f"{test_name}: p = {format_p_value(p_val)}. Разница в долях респондеров {sig}."
            )
            
            # Odds ratio if available
            if "odds_ratio" in test:
                odds = test["odds_ratio"]
                doc.add_paragraph(f"Отношение шансов (Odds Ratio) = {odds:.2f}")
        
        # Interpretation
        doc.add_paragraph()
        p_para = doc.add_paragraph()
        p_para.add_run("Интерпретация:").bold = True
        
        resp_data = responder_results["responders"]
        if resp_data:
            best_group = max(resp_data, key=lambda x: x["pct_responders"])
            worst_group = min(resp_data, key=lambda x: x["pct_responders"])
            doc.add_paragraph(
                f"Наибольшая доля респондеров наблюдается в группе {best_group['group']} "
                f"({best_group['pct_responders']:.0f}%), наименьшая — в группе {worst_group['group']} "
                f"({worst_group['pct_responders']:.0f}%)."
            )
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 6: SUMMARY TABLE
    # ============================================================
    add_heading(doc, "6. СВОДНАЯ ТАБЛИЦА ВСЕХ ПОКАЗАТЕЛЕЙ", 1)
    
    headers = ["Показатель", "Тип", "K-W p", "LMM p (инт.)", "BF₁₀", "ε²", "Вывод"]
    rows = []
    
    for key, result in endpoint_results.items():
        kw = result.get("kruskal_v6", {})
        mixed = result.get("mixed_effects", {})
        
        p_kw = kw.get("p_value", np.nan)
        p_int = mixed.get("interaction_p", np.nan)
        bf = result.get("bf10", np.nan)
        es = kw.get("effect_size")
        
        # Determine conclusion
        if p_kw < 0.05 and p_int < 0.05:
            conclusion = "Значимо оба"
        elif p_kw < 0.05:
            conclusion = "Знач. K-W"
        elif p_int < 0.05:
            conclusion = "Знач. LMM"
        else:
            conclusion = "Не значимо"
        
        rows.append([
            result["short_name"],
            "Первичный" if result.get("primary") else "Вторичный",
            format_p_value(p_kw),
            format_p_value(p_int),
            f"{bf:.1f}" if np.isfinite(bf) else "—",
            f"{es:.3f}" if es else "—",
            conclusion,
        ])
    
    add_table(doc, headers, rows)
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 7: CONCLUSIONS
    # ============================================================
    add_heading(doc, "7. ЗАКЛЮЧЕНИЕ", 1)
    
    doc.add_paragraph(
        "На основании проведённого статистического анализа можно сделать следующие выводы:"
    )
    
    # Primary endpoints conclusions
    p_para = doc.add_paragraph()
    p_para.add_run("Первичные конечные точки:").bold = True
    
    for key, result in endpoint_results.items():
        if not result.get("primary"):
            continue
        
        kw = result.get("kruskal_v6", {})
        mixed = result.get("mixed_effects", {})
        p_kw = kw.get("p_value", np.nan)
        p_int = mixed.get("interaction_p", np.nan)
        bf = result.get("bf10", np.nan)
        
        bullet = f"• {result['short_name']}: "
        
        if p_int < 0.05:
            bullet += f"выявлено значимое взаимодействие время×группа (p = {format_p_value(p_int)}), "
            bullet += "что свидетельствует о различной динамике показателя между группами. "
        else:
            bullet += f"взаимодействие время×группа не достигло значимости (p = {format_p_value(p_int)}). "
        
        if bf and np.isfinite(bf):
            if bf > 3:
                bullet += f"Байесовский анализ поддерживает наличие эффекта (BF₁₀ = {bf:.1f})."
            else:
                bullet += f"Байесовский анализ неопределён (BF₁₀ = {bf:.1f})."
        
        doc.add_paragraph(bullet)
    
    # Responders conclusion
    doc.add_paragraph()
    p_para = doc.add_paragraph()
    p_para.add_run("Анализ респондеров:").bold = True
    
    if responder_results and "responders" in responder_results:
        resp_data = responder_results["responders"]
        best = max(resp_data, key=lambda x: x["pct_responders"])
        test = responder_results.get("test", {})
        p_val = test.get("p_value", np.nan)
        
        doc.add_paragraph(
            f"• Наибольшая доля ответивших на терапию наблюдается в группе {best['group']} "
            f"({best['pct_responders']:.0f}%). "
            f"Различия между группами {'статистически значимы' if p_val < 0.05 else 'не достигли значимости'} "
            f"(p = {format_p_value(p_val)})."
        )
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(
        "Отчёт сгенерирован автоматически с использованием Clinimetria.\n"
        "Методы: Kruskal-Wallis, Linear Mixed Model (statsmodels), Bayes Factor (Sellke bound)."
    ).italic = True
    
    # Save
    doc.save(output_path)
    print(f"   ✓ Report saved: {output_path}")
    
    return str(output_path)


# ============================================================
# MAIN
# ============================================================

def main():
    print("=" * 60)
    print("  DIAMAG Clinical Trial Analysis")
    print("=" * 60)
    
    # Create output directory
    OUTPUT_DIR.mkdir(exist_ok=True)
    
    # Load data
    df = load_data()
    
    # Generate Table 1
    table1 = generate_table1(df)
    
    # Analyze all endpoints
    endpoint_results = {}
    figures = {}
    
    for key in ENDPOINTS_CONFIG:
        result = analyze_endpoint(df, key)
        endpoint_results[key] = result
        
        # Generate figures
        if "long_data" in result and not result["long_data"].empty:
            # Boxplot
            fig_path = OUTPUT_DIR / f"{key}_boxplot.png"
            create_boxplot(result["long_data"], result["name"], fig_path)
            figures[f"{key}_boxplot"] = str(fig_path)
            
            # Spaghetti plot
            fig_path = OUTPUT_DIR / f"{key}_spaghetti.png"
            create_spaghetti_plot(result["long_data"], result["name"], fig_path)
            figures[f"{key}_spaghetti"] = str(fig_path)
    
    # Responder analysis
    responder_results = analyze_responders(df, threshold_pct=20)
    
    # Generate Word report
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_path = OUTPUT_DIR / f"diamag_report_{timestamp}.docx"
    
    generate_word_report(
        table1=table1,
        endpoint_results=endpoint_results,
        responder_results=responder_results,
        output_path=report_path,
        figures=figures,
    )
    
    print("\n" + "=" * 60)
    print(f"  ✅ ANALYSIS COMPLETE!")
    print(f"  📄 Report: {report_path}")
    print("=" * 60)
    
    return str(report_path)


if __name__ == "__main__":
    main()
