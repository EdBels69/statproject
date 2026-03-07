import asyncio
import os
import sys
import uuid
import pandas as pd
from app.copilot.engine import CopilotEngine
from app.copilot.report import generate_report

sys.path.append(os.getcwd())

async def main():
    print("🧪 PHASE E: RESULTS PIPELINE VERIFICATION")
    print("=" * 50)
    
    engine = CopilotEngine()
    dataset_path = os.path.abspath("workspace/covid19_dataset.csv")
    
    if not os.path.exists(dataset_path):
        print(f"❌ Dataset not found: {dataset_path}")
        return

    # 1. Load Data and build RICH dataset_info
    df = pd.read_csv(dataset_path)
    print(f"📂 Loaded dataset: {df.shape[0]} rows, {df.shape[1]} columns")
    
    # Build detailed column info for LLM context
    column_details = []
    for col in df.columns[:50]:  # First 50 cols for context
        dtype = str(df[col].dtype)
        nunique = int(df[col].nunique())
        sample = df[col].dropna().head(3).tolist()
        col_info = f"{col} (dtype={dtype}, unique={nunique}, sample={sample})"
        column_details.append(col_info)

    dataset_info = {
        "filename": "covid19_dataset.csv",
        "n_rows": len(df),
        "n_cols": len(df.columns),
        "columns": list(df.columns),
        "dataset_meta": {
            "summary": {
                "n_rows": len(df),
                "columns_detail": column_details,
                "notes": [
                    "Column 'Исход' is the outcome variable: 'Выписан'=Discharged, 'Мертв'=Dead",
                    "Column 'возраст' is age (numeric)",
                    "Column 'пол' is sex: 'м'=male, 'ж'=female",
                    "Column 'SpO2 %' is oxygen saturation (numeric)",
                    "Column 'WBC V1 10*9' is white blood cell count (numeric)",
                    "Column 'Длительность госпитализации' is length of stay (numeric)",
                    "Many columns are in Russian (Cyrillic charset)"
                ]
            }
        }
    }
    
    # 2. Request using EXACT column names from dataset
    request = """
    Проведи анализ факторов риска летального исхода (колонка 'Исход', значения: 'Выписан' vs 'Мертв'):
    1. Описательная статистика: возраст, SpO2 %, WBC V1 10*9, Длительность госпитализации — по группам Исхода.
    2. Сравни группы (Выписан vs Мертв) по переменным: возраст, SpO2 %, WBC V1 10*9, температура. Используй критерий Манна-Уитни.
    3. Построй логистическую регрессию: Исход (закодируй: Выписан=0, Мертв=1) ~ возраст + SpO2 % + WBC V1 10*9.
    4. Визуализируй boxplot для возраста по группам Исхода.
    """
    
    session_id = str(uuid.uuid4())
    print(f"\n🚀 Creating Plan (Session: {session_id[:8]})...")
    
    # 3. Create Plan
    plan_res = await engine.create_plan(
        session_id=session_id,
        user_request=request,
        dataset_path=dataset_path,
        dataset_info=dataset_info,
        advanced=True
    )
    
    if not plan_res["success"]:
        print(f"❌ Plan failed: {plan_res.get('error')}")
        return

    print(f"✅ Plan Created:")
    print(f"   Goal: {plan_res['plan'].get('understood_goal', 'N/A')[:120]}...")
    print(f"   Analyses: {len(plan_res['plan'].get('analyses', []))}")
    for a in plan_res['plan'].get('analyses', []):
        print(f"   - {a.get('name', 'N/A')} [{a.get('type', '?')}]")
    
    # 4. Execute
    print(f"\n⚙️ Executing Analysis...")
    exec_res = await engine.execute_plan(session_id=session_id)
    
    if not exec_res["success"]:
        print(f"❌ Execution failed: {exec_res.get('error')}")
        return
    
    print(f"✅ Execution Success!")
    
    # 5. Report
    print("\n📄 Generating Report...")
    try:
        session = engine.sessions[session_id]
        results = session.get("results", {})
        
        # Count actual analysis items (exclude _errors and _raw_* keys)
        analysis_keys = [k for k in results.keys() if not k.startswith("_")]
        plot_count = sum(
            len(v.get("plots", [])) 
            for v in results.values() 
            if isinstance(v, dict) and "plots" in v
        )
        
        print(f"   Results: {len(analysis_keys)} analyses, {plot_count} plots")
        for k in analysis_keys:
            item = results[k]
            if isinstance(item, dict):
                table_rows = len(item.get("table", [])) - 1 if item.get("table") else 0
                print(f"   ✓ {k}: {item.get('title', 'N/A')} ({table_rows} data rows)")
        
        if results.get("_errors"):
            print(f"   ⚠️ Errors: {results['_errors']}")
        
        has_interpretation = session.get("interpretation") is not None
        print(f"   Interpretation: {'✅ present' if has_interpretation else '❌ missing'}")
        
        docx = generate_report(
            results=results,
            plan=session["plan"],
            code=session.get("code"),
            interpretation=session.get("interpretation"),
            dataset_info=session.get("dataset_info")
        )
        
        report_path = f"workspace/report_phase_e_{session_id[:8]}.docx"
        with open(report_path, "wb") as f:
            f.write(docx)
            
        # Verify report content
        from docx import Document
        doc = Document(report_path)
        table_count = len(doc.tables)
        para_count = len(doc.paragraphs)
        
        print(f"\n✅ Report saved: {report_path}")
        print(f"   Size: {len(docx)/1024:.1f} KB")
        print(f"   Paragraphs: {para_count}")
        print(f"   Tables: {table_count}")
        
        if table_count > 1:
            print("   🎉 SUCCESS: Report contains data tables!")
        else:
            print("   ⚠️ WARNING: Report may lack data tables")
            
    except Exception as e:
        print(f"❌ Report generation failed: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    asyncio.run(main())
