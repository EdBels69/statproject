"""
DiaMag-Level DOCX Report Generator
Enhanced report with TOC, BF10, Holm correction, responders, LMM results.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List, Optional
from datetime import datetime
import io


def _add_toc(doc: Document, title: str = "Содержание"):
    """Add Table of Contents to document."""
    doc.add_heading(title, level=1)
    
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    # Create TOC field
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    
    doc.add_paragraph("(Обновите оглавление после открытия в Word: Ctrl+A, F9)")


def _color_cell(cell, color: str):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _format_p(value: Optional[float], threshold: float = 0.05) -> tuple:
    """Format p-value and return (text, is_significant)."""
    if value is None or not isinstance(value, (int, float)):
        return ("—", False)
    if value < 0.001:
        return ("<0.001", value < threshold)
    return (f"{value:.4f}", value < threshold)


def _format_bf(value: Optional[float]) -> str:
    """Format Bayes Factor."""
    if value is None or not isinstance(value, (int, float)):
        return "—"
    if value > 100:
        return ">100"
    if value >= 10:
        return f"{value:.1f}"
    return f"{value:.2f}"


class DiaMagReportGenerator:
    """
    Enhanced DOCX generator for DiaMag-level analysis results.
    
    Features:
    - Table of Contents
    - BF10 interpretation column
    - Holm-corrected p-values
    - Responder analysis tables
    - LMM Time×Group results
    - Significance coloring
    """
    
    def __init__(self, results: Dict[str, Any], config: Dict[str, Any]):
        self.results = results
        self.config = config
    
    def generate(self) -> bytes:
        """Generate full DiaMag-style report."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Комплексный статистический отчёт', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        p = doc.add_paragraph()
        p.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
        p.add_run(f"Тип анализа: DiaMag-Level Comprehensive\n")
        p.add_run(f"Группировочная переменная: {self.config.get('group_col', 'N/A')}\n")
        p.add_run(f"N пациентов: {self.results.get('n_patients', 'N/A')}")
        
        # Table of Contents
        _add_toc(doc)
        doc.add_page_break()
        
        # Executive Summary
        self._add_executive_summary(doc)
        
        # Methodology
        self._add_methodology(doc)
        
        # Results by Endpoint
        self._add_endpoint_results(doc)
        
        # Responder Analysis Summary
        self._add_responder_summary(doc)
        
        # LMM Summary
        self._add_lmm_summary(doc)
        
        # AI Discussion (if available)
        self._add_ai_content(doc)
        
        # Limitations
        self._add_limitations(doc)
        
        # Save
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _add_executive_summary(self, doc: Document):
        """Add executive summary section."""
        doc.add_heading("1. Краткие результаты", level=1)
        
        endpoints = self.results.get("endpoints", {})
        groups = self.results.get("groups", [])
        
        summary_items = [
            f"Всего групп: {len(groups)}",
            f"Групп: {', '.join(str(g) for g in groups)}",
            f"Исследовано показателей: {len(endpoints)}",
        ]
        
        # Count significant results
        sig_kw = 0
        sig_pw = 0
        sig_lmm = 0
        
        for ep_data in endpoints.values():
            if ep_data.get("baseline_kw", {}).get("significant"):
                sig_kw += 1
            
            for visit_data in ep_data.get("by_visit", {}).values():
                if visit_data.get("kruskal", {}).get("significant"):
                    sig_kw += 1
                
                pw = visit_data.get("pairwise", {})
                sig_pw += sum(1 for p in pw.values() if isinstance(p, dict) and p.get("significant"))
            
            if ep_data.get("mixed_effects", {}).get("interaction", {}).get("significant"):
                sig_lmm += 1
        
        summary_items.append(f"Значимых Kruskal-Wallis: {sig_kw}")
        summary_items.append(f"Значимых pairwise (после Holm): {sig_pw}")
        summary_items.append(f"Значимых LMM Time×Group: {sig_lmm}")
        
        for item in summary_items:
            doc.add_paragraph(item, style='List Bullet')
    
    def _add_methodology(self, doc: Document):
        """Add methodology section."""
        doc.add_heading("2. Методология", level=1)
        
        methods = [
            "Kruskal-Wallis H-тест для сравнения групп (≥3 групп)",
            "Mann-Whitney U post-hoc с поправкой Holm-Bonferroni",
            "Байесовский фактор BF₁₀ (Sellke bound) для силы доказательств",
            "Effect size: ε² (Kruskal), r (Mann-Whitney)",
            "Респондеры: ≥20% улучшение от базелайна",
            "NNT (Number Needed to Treat) для лучшей группы",
            "Linear Mixed Model: Time × Group взаимодействие",
        ]
        
        for m in methods:
            doc.add_paragraph(m, style='List Bullet')
        
        # BF interpretation table
        doc.add_heading("Интерпретация BF₁₀", level=2)
        bf_table = doc.add_table(rows=1, cols=2)
        bf_table.style = 'Table Grid'
        bf_table.rows[0].cells[0].text = "BF₁₀"
        bf_table.rows[0].cells[1].text = "Интерпретация"
        
        bf_ranges = [
            (">100", "Экстремально сильно за H₁"),
            ("30-100", "Очень сильно за H₁"),
            ("10-30", "Сильно за H₁"),
            ("3-10", "Умеренно за H₁"),
            ("1-3", "Слабо за H₁"),
            ("0.33-1", "Неопределённо"),
            ("<0.33", "За H₀ (нет различий)"),
        ]
        for bf, interp in bf_ranges:
            row = bf_table.add_row().cells
            row[0].text = bf
            row[1].text = interp
    
    def _add_endpoint_results(self, doc: Document):
        """Add detailed results for each endpoint."""
        doc.add_heading("3. Результаты по показателям", level=1)
        
        for ep_name, ep_data in self.results.get("endpoints", {}).items():
            if "error" in ep_data:
                continue
            
            doc.add_heading(ep_name, level=2)
            
            # Baseline KW
            kw = ep_data.get("baseline_kw", {})
            if kw and "H" in kw:
                p_text, is_sig = _format_p(kw.get("p"))
                bf = _format_bf(kw.get("bf10"))
                sig_mark = " ✓" if is_sig else ""
                
                doc.add_paragraph(
                    f"Базелайн Kruskal-Wallis: H={kw.get('H', 0):.2f}, "
                    f"p={p_text}{sig_mark}, BF₁₀={bf}, "
                    f"ε²={kw.get('epsilon_sq', 0):.3f} ({kw.get('epsilon_interpretation', '')})"
                )
            
            # By-visit table
            by_visit = ep_data.get("by_visit", {})
            if by_visit:
                doc.add_heading("По визитам", level=3)
                
                # Kruskal-Wallis table
                kw_table = doc.add_table(rows=1, cols=5)
                kw_table.style = 'Table Grid'
                headers = ["Визит", "H", "p (raw)", "BF₁₀", "ε²"]
                for i, h in enumerate(headers):
                    kw_table.rows[0].cells[i].text = h
                
                for visit, v_data in by_visit.items():
                    kw = v_data.get("kruskal", {})
                    if not kw or "error" in kw:
                        continue
                    
                    row = kw_table.add_row().cells
                    row[0].text = visit
                    row[1].text = f"{kw.get('H', 0):.2f}"
                    p_text, is_sig = _format_p(kw.get("p"))
                    row[2].text = p_text
                    row[3].text = _format_bf(kw.get("bf10"))
                    row[4].text = f"{kw.get('epsilon_sq', 0):.3f}"
                    
                    if is_sig:
                        _color_cell(row[2], "90EE90")  # Light green
            
            # Responders table
            best_resp = None
            best_rate = 0
            for visit, v_data in by_visit.items():
                resp = v_data.get("responders", {})
                if "groups" in resp:
                    for g, g_data in resp.get("groups", {}).items():
                        rate = g_data.get("rate_pct", 0)
                        if rate > best_rate:
                            best_rate = rate
                            best_resp = {"visit": visit, "group": g, "rate": rate}
            
            if best_resp:
                doc.add_paragraph(
                    f"Лучший респонс: Группа {best_resp['group']} на визите {best_resp['visit']} "
                    f"({best_resp['rate']:.1f}% респондеров)"
                )
            
            # LMM result
            lmm = ep_data.get("mixed_effects", {})
            if lmm and "interaction" in lmm:
                interaction = lmm.get("interaction", {})
                p_val = interaction.get("min_p_value", lmm.get("interaction_p_value"))
                p_text, is_sig = _format_p(p_val)
                sig_text = "ЗНАЧИМО" if is_sig else "не значимо"
                
                doc.add_paragraph(
                    f"LMM Time×Group: p={p_text} — {sig_text}"
                )
                if interaction.get("interpretation"):
                    doc.add_paragraph(interaction.get("interpretation"), style='Quote')
    
    def _add_responder_summary(self, doc: Document):
        """Add responder analysis summary table."""
        doc.add_heading("4. Сводка по респондерам", level=1)
        
        # Collect all responder data
        resp_data = []
        for ep_name, ep_data in self.results.get("endpoints", {}).items():
            for visit, v_data in ep_data.get("by_visit", {}).items():
                resp = v_data.get("responders", {})
                if "groups" in resp:
                    for g, g_data in resp.get("groups", {}).items():
                        resp_data.append({
                            "endpoint": ep_name,
                            "visit": visit,
                            "group": g,
                            "n": g_data.get("n", 0),
                            "n_resp": g_data.get("n_responders", 0),
                            "rate": g_data.get("rate_pct", 0),
                        })
        
        if not resp_data:
            doc.add_paragraph("Данные о респондерах недоступны.")
            return
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ["Показатель", "Визит", "Группа", "Респонд.", "%"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        
        for r in resp_data:
            row = table.add_row().cells
            row[0].text = r["endpoint"][:20]
            row[1].text = r["visit"]
            row[2].text = str(r["group"])
            row[3].text = f"{r['n_resp']}/{r['n']}"
            row[4].text = f"{r['rate']:.1f}%"
            
            if r["rate"] >= 50:
                _color_cell(row[4], "90EE90")
    
    def _add_lmm_summary(self, doc: Document):
        """Add LMM summary table."""
        doc.add_heading("5. Mixed Effects Models (Time×Group)", level=1)
        
        lmm_results = []
        for ep_name, ep_data in self.results.get("endpoints", {}).items():
            lmm = ep_data.get("mixed_effects", {})
            if lmm and "interaction" in lmm:
                interaction = lmm.get("interaction", {})
                lmm_results.append({
                    "endpoint": ep_name,
                    "p": interaction.get("min_p_value", lmm.get("interaction_p_value")),
                    "significant": interaction.get("significant", False),
                })
        
        if not lmm_results:
            doc.add_paragraph("LMM не были вычислены (требуется subject_col в конфигурации).")
            return
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        headers = ["Показатель", "p (interaction)", "Значимость"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        
        for r in lmm_results:
            row = table.add_row().cells
            row[0].text = r["endpoint"][:30]
            p_text, is_sig = _format_p(r["p"])
            row[1].text = p_text
            row[2].text = "✓" if is_sig else "—"
            
            if is_sig:
                _color_cell(row[1], "90EE90")
                _color_cell(row[2], "90EE90")
    
    def _add_ai_content(self, doc: Document):
        """Add AI-generated discussion and conclusions."""
        ai_content = self.results.get("ai_content", {})
        
        if not ai_content or "error" in ai_content:
            return
        
        discussion = ai_content.get("discussion", [])
        conclusions = ai_content.get("conclusions", [])
        
        if discussion:
            doc.add_heading("6. Обсуждение", level=1)
            for para in discussion:
                doc.add_paragraph(para)
        
        if conclusions:
            doc.add_heading("7. Выводы", level=1)
            for item in conclusions:
                doc.add_paragraph(item, style='List Bullet')
    
    def _add_limitations(self, doc: Document):
        """Add limitations section."""
        doc.add_heading("8. Ограничения", level=1)
        
        limitations = [
            "Автоматический выбор тестов основан на общих правилах.",
            "BF₁₀ рассчитан из p-value (Sellke bound) — приближённый метод.",
            "LMM предполагает случайное отсутствие данных (MAR).",
            "Holm-коррекция консервативна; рассмотрите FDR при большом числе сравнений.",
        ]
        
        for lim in limitations:
            doc.add_paragraph(lim, style='List Bullet')
