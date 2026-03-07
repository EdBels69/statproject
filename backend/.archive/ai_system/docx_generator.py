"""
AI Analysis DOCX Report Generator
Generates transparent Word reports with methodology, audit trail, and reproducibility.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Dict, Any, List
from datetime import datetime
import io


class AIAnalysisReportGenerator:
    """Generates DOCX reports for AI-driven statistical analysis with full transparency."""
    
    def __init__(self, results: Dict[str, Any], config: Dict[str, Any], dataset_info: Dict[str, Any]):
        self.results = results
        self.config = config
        self.dataset_info = dataset_info
        
    def generate(self) -> bytes:
        """Generate DOCX report and return as bytes."""
        doc = Document()
        
        # Title
        title = doc.add_heading('Статистический анализ', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph(f"Датасет: {self.dataset_info.get('filename', 'N/A')}")
        doc.add_paragraph(f"Группировка: {self.config.get('group_col', 'N/A')}")
        
        # Executive Summary
        doc.add_heading('Краткие результаты', level=1)
        summary = self._generate_summary()
        for line in summary:
            doc.add_paragraph(line, style='List Bullet')
        
        # Methodology Section (TRANSPARENCY)
        doc.add_heading('Методология', level=1)
        doc.add_paragraph(
            "Данный отчёт содержит полностью воспроизводимый анализ. "
            "Все статистические тесты были выбраны автоматически на основе "
            "характеристик данных (нормальность, количество групп, тип переменной)."
        )
        
        # Test Selection Rationale Table
        doc.add_heading('Выбор статистических методов', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Переменная'
        hdr_cells[1].text = 'Тип'
        hdr_cells[2].text = 'Метод'
        hdr_cells[3].text = 'P-value'
        
        # Single variables
        for var, res in self.results.get('single_variables', {}).items():
            row = table.add_row().cells
            row[0].text = var
            row[1].text = res.get('type', 'unknown')
            test = res.get('test', {})
            row[2].text = test.get('method', '-')
            pval = test.get('p_value')
            row[3].text = f"{pval:.4f}" if pval is not None else '-'
            
            # Highlight significant results
            if test.get('significant'):
                for cell in row:
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                    cell.paragraphs[0].runs[0].font.bold = True
        
        # Longitudinal families
        doc.add_heading('Продольный анализ (семейства переменных)', level=2)
        for fam_name, fam_data in self.results.get('endpoints', {}).items():
            doc.add_heading(fam_name, level=3)
            
            # Table for each visit
            fam_table = doc.add_table(rows=1, cols=3)
            fam_table.style = 'Light Grid Accent 1'
            hdr = fam_table.rows[0].cells
            hdr[0].text = 'Визит'
            hdr[1].text = 'Метод'
            hdr[2].text = 'P-value'
            
            for visit, test_res in fam_data.get('tests', {}).items():
                row = fam_table.add_row().cells
                row[0].text = visit
                row[1].text = test_res.get('method', '-')
                pval = test_res.get('p_value')
                row[2].text = f"{pval:.4f}" if pval is not None else '-'
                
                if test_res.get('significant'):
                    for cell in row:
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                        cell.paragraphs[0].runs[0].font.bold = True
        
        # Interpretation
        doc.add_heading('Интерпретация', level=1)
        doc.add_paragraph(
            "P-value < 0.05 указывает на статистически значимое различие между группами. "
            "Значимые результаты выделены зелёным цветом в таблицах."
        )
        
        # Limitations
        doc.add_heading('Ограничения', level=2)
        doc.add_paragraph(
            "1. Автоматический выбор тестов основан на общих правилах и может не учитывать "
            "специфические особенности исследования."
        )
        doc.add_paragraph(
            "2. Множественные сравнения могут увеличить вероятность ошибки I рода. "
            "Рекомендуется применить поправку Bonferroni или FDR при необходимости."
        )
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _generate_summary(self) -> List[str]:
        """Generate executive summary bullet points."""
        summary = []
        
        # Count tests
        single_vars = self.results.get('single_variables', {})
        endpoints = self.results.get('endpoints', {})
        
        total_tests = len(single_vars)
        for fam in endpoints.values():
            total_tests += len(fam.get('tests', {}))
        
        sig_count = sum(1 for v in single_vars.values() if v.get('test', {}).get('significant'))
        for fam in endpoints.values():
            sig_count += sum(1 for t in fam.get('tests', {}).values() if t.get('significant'))
        
        summary.append(f"Всего выполнено {total_tests} статистических тестов")
        summary.append(f"Значимых результатов (p < 0.05): {sig_count}")
        summary.append(f"Группировочная переменная: {self.config.get('group_col', 'N/A')}")
        
        if endpoints:
            summary.append(f"Продольных семейств обнаружено: {len(endpoints)}")
        
        return summary
